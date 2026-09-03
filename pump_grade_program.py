# rebuild-marker: 2026-08-29 14:01:23
# ============================================================
# K-water tech
# 펌프 설비관리 통합 플랫폼
# QR 기반 설비관리 + 정밀진단 + CBM + 오버홀 + AI + KPI
#
# Streamlit Single File Application
# ============================================================

import os
import io
import urllib.request
import requests
import gspread
import smtplib
import time
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import urllib.parse
from datetime import datetime, date, timedelta
import random
import math

import streamlit as st
import streamlit.components.v1 as components

import pandas as pd
import numpy as np
import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, LineChart, Reference

import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

import qrcode
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import base64
import glob
import subprocess
import tempfile
import shutil
import zipfile
import re
import calendar
from filelock import FileLock

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 0. 기본 설정
# ============================================================

APP_TITLE = "K-water tech 설비관리 플랫폼"

DB_FILE_PATH = "Pump_Master_DB.xlsx"
OVERHAUL_DB_PATH = "Pump_Overhaul_DB.xlsx"
DOC_DB_PATH = "Pump_Docs_DB.xlsx"
KNOWHOW_DB_PATH = "Pump_Knowhow_DB.xlsx"
DAILY_LOG_DB_PATH = "Pump_DailyLog_DB.xlsx"
SAFETY_PERMIT_DB_PATH = "Pump_SafetyPermit_DB.xlsx"
KPI_DB_PATH = "Pump_KPI_DB.xlsx"
EQUIP_DB_PATH = "Pump_Equipment_DB.xlsx"

LOGO_FILE_PATH = "Logo.png"
SEED_FLAG_PATH = "seed_flag.txt"
EQUIP_SEED_FLAG_PATH = "equip_seed_flag.txt"

PHOTO_DIR = "overhaul_photos"
CONFIG_DIR = "app_config"
DRAFT_DIR = "diag_drafts"

# ------------------------------------------------------------
# 로그인(PIN) 게이트 켜고 끄기 스위치.
#
# 평가위원 시연 등으로 로그인 화면 없이 바로 앱을 보여줘야
# 할 때는 False로 두면 된다. 나중에 다시 PIN을 쓰려면
# True로만 바꾸면 되고, PIN 검증 로직 자체는 그대로 남아있다.
# ------------------------------------------------------------

LOGIN_GATE_ENABLED = True


# ------------------------------------------------------------
# PIN 번호: st.secrets에 넣어두면 그걸 우선 쓰고,
# 없으면 기본값(2580)을 쓴다. 기본값을 그대로 쓰는 경우
# 로그인 화면에 "st.secrets 설정을 권장한다"는 안내를 띄운다.
# (깃허브 저장소가 Public이면 코드에 박힌 PIN이 그대로
#  노출되기 때문 — Streamlit Cloud의 "Secrets" 설정에
#  AUTH_PIN = "원하는 번호" 를 넣으면 코드에는 안 남는다.)
# ------------------------------------------------------------

AUTH_PIN_IS_DEFAULT = False

try:

    AUTH_PIN = st.secrets["AUTH_PIN"]

except Exception:

    AUTH_PIN = "2580"

    AUTH_PIN_IS_DEFAULT = True


# ============================================================
# 1. 한글 폰트
# ============================================================

def init_korean_font():

    font_filename = "NanumGothic.ttf"

    if not os.path.exists(font_filename):

        font_url = (
            "https://github.com/google/fonts/raw/main/"
            "ofl/nanumgothic/NanumGothic-Regular.ttf"
        )

        try:
            urllib.request.urlretrieve(font_url, font_filename)
        except Exception:
            pass

    if os.path.exists(font_filename):

        try:
            fm.fontManager.addfont(font_filename)

            font_prop = fm.FontProperties(
                fname=font_filename
            )

            plt.rcParams["font.family"] = font_prop.get_name()

        except Exception:
            plt.rcParams["font.family"] = "DejaVu Sans"

    else:

        plt.rcParams["font.family"] = "DejaVu Sans"

    plt.rcParams["axes.unicode_minus"] = False


init_korean_font()


# ============================================================
# 2. 페이지 설정
# ============================================================

st.set_page_config(

    page_title=APP_TITLE,

    page_icon="💧",

    layout="wide",

    initial_sidebar_state="expanded"
)


# ============================================================
# 2-0. 긴급 차단 스위치
#
# 회사에서 "당장 내려라"는 지시가 오면, 코드를 다시 배포할
# 필요 없이 Streamlit Secrets에서 MAINTENANCE_MODE = true 만
# 켜면 즉시 이 화면으로 대체된다. 뒤에 나오는 로그인/데이터
# 로직은 아예 실행되지 않는다.
# ============================================================

try:

    _maintenance_mode = st.secrets.get(
        "MAINTENANCE_MODE",
        False
    )

except Exception:

    _maintenance_mode = False

if str(_maintenance_mode).lower() in ("true", "1", "yes"):

    st.markdown(

        """
        <div style="
        text-align:center; margin-top:15vh;
        color:#0f3552;">

        <div style="font-size:3rem;">🛠️</div>

        <div style="font-size:1.3rem; font-weight:800;
        margin-top:10px;">
        점검 중입니다
        </div>

        <div style="font-size:0.95rem; color:#64748b;
        margin-top:8px;">
        잠시 후 다시 접속해주세요.
        </div>

        </div>
        """,

        unsafe_allow_html=True

    )

    st.stop()


# ============================================================
# 2-1. HTML 렌더링 버그 수정 패치
#
# Streamlit의 st.markdown(unsafe_allow_html=True)은 내부적으로
# 마크다운 파서를 거치는데, 줄 앞에 공백이 4칸 이상 있으면
# 마크다운 문법상 "코드 블록"으로 인식되어 HTML 태그가
# 렌더링되지 않고 그대로 화면에 텍스트로 노출되는 문제가 있다.
#
# 이 프로젝트는 들여쓰기가 깊은 코드 스타일이라 여러 곳에서
# 이 문제가 발생하므로, st.markdown을 한 번만 감싸서
# unsafe_allow_html=True로 호출될 때 각 줄의 선행 공백을
# 자동으로 제거하도록 패치한다.
# ============================================================

_original_markdown = st.markdown


def _patched_markdown(body, *args, **kwargs):

    if kwargs.get("unsafe_allow_html") and isinstance(body, str):

        body = "\n".join(
            line.lstrip() for line in body.split("\n")
        )

    return _original_markdown(body, *args, **kwargs)


st.markdown = _patched_markdown


# ============================================================
# 3. 수자원 스타일 CSS
# ============================================================

st.markdown(
    """
<style>

/* ========================================================
   전체
======================================================== */

html,
body,
[data-testid="stAppViewContainer"] {

    background:
        linear-gradient(
            180deg,
            #f4f9fc 0%,
            #ffffff 55%
        );

}

/* 페이지 전환 시 스크롤이 맨 위로 안 가고 중간에 멈추는
   문제의 실제 원인은 브라우저의 "스크롤 앵커링" 기능이었다.
   콘텐츠(그래프 이미지 등)가 로딩되면서 화면 위쪽 내용이
   바뀌면, 브라우저가 "사용자가 보던 위치를 유지"하려고
   자동으로 스크롤을 다시 밀어낸다. 이 기능을 꺼서
   자바스크립트로 맨 위로 되돌린 위치가 유지되게 한다. */

section[data-testid="stMain"],
[data-testid="stAppViewContainer"] {

    overflow-anchor: none;

}

.block-container {

    max-width: 1450px;

    padding-top: 3.5rem !important;

    padding-bottom: 2rem !important;

    padding-left: 1.2rem !important;

    padding-right: 1.2rem !important;

}


/* ========================================================
   Streamlit 기본 상단 툴바가 커스텀 헤더 위에
   겹쳐서 그래픽 상단이 잘리는 것을 방지
======================================================== */

header[data-testid="stHeader"] {

    background: transparent;

}

.top-header {

    margin-top: 0.4rem;

}


/* ========================================================
   사이드바
======================================================== */

section[data-testid="stSidebar"] {

    background:
        linear-gradient(
            180deg,
            #063b63 0%,
            #075985 45%,
            #087ea4 100%
        );

}

section[data-testid="stSidebar"] > div {

    background: transparent;

}

section[data-testid="stSidebar"] * {

    color: white !important;

}

section[data-testid="stSidebar"] .stButton > button,
section[data-testid="stSidebar"] .stDownloadButton > button {

    background: rgba(255,255,255,0.08);

    color: white !important;

    border: 1px solid rgba(255,255,255,0.12);

    border-radius: 9px;

    text-align: left;

    padding: 9px 12px;

    margin-bottom: 4px;

}

section[data-testid="stSidebar"] .stButton > button:hover,
section[data-testid="stSidebar"] .stDownloadButton > button:hover {

    background: rgba(255,255,255,0.20);

    border-color: rgba(255,255,255,0.35);

}


/* ========================================================
   상단 헤더
======================================================== */

.top-header {

    background:
        linear-gradient(
            135deg,
            #063b63 0%,
            #087ea4 60%,
            #11a6c9 100%
        );

    border-radius: 18px;

    padding: 22px 26px;

    color: white;

    margin-bottom: 18px;

    box-shadow:
        0 8px 25px rgba(3, 65, 100, 0.12);

}

.top-title {

    font-size: 1.55rem;

    font-weight: 800;

    letter-spacing: -0.5px;

}

.top-sub {

    font-size: 0.85rem;

    opacity: 0.88;

    margin-top: 5px;

}


/* ========================================================
   섹션
======================================================== */

.section-title {

    font-size: 1.12rem;

    font-weight: 800;

    color: #0f3552;

    margin-top: 8px;

    margin-bottom: 10px;

}

.section-caption {

    font-size: 0.80rem;

    color: #64748b;

    margin-bottom: 12px;

}


/* ========================================================
   KPI
======================================================== */

.kpi-grid {

    display: grid;

    grid-template-columns:
        repeat(5, 1fr);

    gap: 12px;

    margin-bottom: 18px;

}

.kpi-card {

    background: white;

    border-radius: 14px;

    padding: 15px;

    border:
        1px solid #dce8ef;

    box-shadow:
        0 3px 12px rgba(10, 70, 100, 0.06);

}

.kpi-label {

    font-size: 0.74rem;

    color: #64748b;

    font-weight: 700;

}

.kpi-value {

    font-size: 1.45rem;

    font-weight: 800;

    color: #083b5c;

    margin-top: 4px;

}

.kpi-sub {

    font-size: 0.68rem;

    color: #94a3b8;

    margin-top: 3px;

}


/* ========================================================
   카드
======================================================== */

/* ========================================================
   등장 애니메이션 (카드/섹션이 딱딱하게 툭 나타나는 대신
   살짝 떠오르며 나타나서 화면 전환이 더 부드럽게 느껴지게)
======================================================== */

@keyframes fadeInUp {

    from {
        opacity: 0;
        transform: translateY(10px);
    }

    to {
        opacity: 1;
        transform: translateY(0);
    }

}

.platform-card,
.equip-card,
.kpi-card,
.feature-card,
.story-card {

    animation: fadeInUp 0.35s ease;

}


/* ========================================================
   설비 스토리 카드 (QR포털 "설비 스토리" 탭)
======================================================== */

.story-card {

    background:
        linear-gradient(135deg, #f4f9fc, #ffffff);

    border: 1px solid #dce8ef;

    border-radius: 14px;

    padding: 22px;

    box-shadow:
        0 3px 14px rgba(10, 70, 100, 0.06);

}


/* ========================================================
   페이지 제목 + 마스코트 (QR/정밀진단/CBM/오버홀 헤더)
   가로배치가 기본이지만, 모바일에서는 미디어쿼리로
   세로배치로 전환한다(히어로배너와 같은 원리).
======================================================== */

.page-header-mascot {

    display: flex;

    flex-direction: column;

    align-items: flex-start;

    gap: 12px;

}

@media (min-width: 700px) {

    .page-header-mascot {

        flex-direction: row;

        align-items: center;

        justify-content: space-between;

        gap: 16px;

        flex-wrap: wrap;

    }

}

.page-header-mascot > .page-header-text {

    width: 100%;

    min-width: 0;

}

@media (min-width: 700px) {

    .page-header-mascot > .page-header-text {

        flex: 1 1 260px;

        width: auto;

    }

}

.page-header-mascot > .page-header-mascot-img {

    flex-shrink: 0;

    align-self: center;

}


/* ========================================================
   통합검색 하이라이트 표
======================================================== */

.search-result-wrap {

    animation: fadeInUp 0.3s ease;

    overflow-x: auto;

    border-radius: 10px;

    border: 1px solid #dce8ef;

}

.search-result-wrap table {

    width: 100%;

    border-collapse: collapse;

}

.search-result-wrap th {

    padding: 8px 12px;

    text-align: left;

    background: #f4f9fc;

    font-size: 0.85rem;

    border-bottom: 2px solid #dce8ef;

}

.search-result-wrap td {

    padding: 8px 12px;

    font-size: 0.85rem;

    border-bottom: 1px solid #eef2f6;

}

.search-result-wrap mark {

    background: #ffe066;

    padding: 0 2px;

    border-radius: 3px;

}


.platform-card {

    background: white;

    border: 1px solid #dce8ef;

    border-radius: 15px;

    padding: 20px 22px;

    margin-bottom: 16px;

    box-shadow:
        0 3px 14px rgba(10, 70, 100, 0.055);

    transition:
        transform 0.15s ease,
        box-shadow 0.15s ease;

}

.platform-card:hover {

    transform: translateY(-2px);

    box-shadow:
        0 10px 26px rgba(10, 70, 100, 0.12);

}

.card-title {

    color: #0f3552;

    font-weight: 800;

    font-size: 0.98rem;

    margin-bottom: 8px;

}


/* ========================================================
   설비 카드 그리드 (홈)
======================================================== */

.equip-card-grid {

    display: grid;

    grid-template-columns:
        repeat(auto-fill, minmax(220px, 1fr));

    gap: 12px;

    margin-bottom: 10px;

}

.equip-card {

    background: white;

    border-radius: 13px;

    padding: 13px 15px;

    box-shadow:
        0 3px 12px rgba(10, 70, 100, 0.07);

    transition:
        transform 0.15s ease,
        box-shadow 0.15s ease;

}

.equip-card:hover {

    transform: translateY(-3px);

    box-shadow:
        0 12px 24px rgba(10, 70, 100, 0.16);

}

.equip-card-title {

    font-weight: 800;

    font-size: 0.92rem;

    color: #0f3552;

    margin-bottom: 4px;

}

.equip-card-meta {

    font-size: 0.72rem;

    color: #64748b;

    display: flex;

    justify-content: space-between;

    margin-top: 6px;

}

.cbm-bar-track {

    width: 100%;

    height: 8px;

    border-radius: 999px;

    background: #eef2f5;

    overflow: hidden;

    margin-top: 8px;

}

.cbm-bar-fill {

    height: 100%;

    border-radius: 999px;

}

.greeting-banner {

    background:
        linear-gradient(
            120deg,
            #063b63 0%,
            #087ea4 55%,
            #11a6c9 100%
        );

    border-radius: 16px;

    padding: 18px 22px;

    color: white;

    margin-bottom: 16px;

    box-shadow:
        0 8px 22px rgba(3, 65, 100, 0.15);

}


/* ========================================================
   피처 카드 (K-water Tech "부서안내" 카드 참고)

   실제 현장 사진 대신 그라데이션 배경 + 굵은 흰 글자 +
   "자세히 보기" 스타일의 화살표 링크로 비슷한 느낌을 냈다.
======================================================== */

.feature-card {

    position: relative;

    border-radius: 16px;

    padding: 22px 20px 18px 20px;

    color: white;

    min-height: 128px;

    margin-bottom: 4px;

    box-shadow:
        0 10px 24px rgba(10, 70, 100, 0.18);

    transition:
        transform 0.15s ease,
        box-shadow 0.15s ease;

}

.feature-card:hover {

    transform: translateY(-4px);

    box-shadow:
        0 16px 32px rgba(10, 70, 100, 0.26);

}

.feature-card-icon {

    font-size: 1.6rem;

}

.feature-card-title {

    font-size: 1.05rem;

    font-weight: 800;

    margin-top: 6px;

}

.feature-card-desc {

    font-size: 0.78rem;

    opacity: 0.92;

    margin-top: 4px;

    line-height: 1.4;

}

.feature-card-link {

    font-size: 0.78rem;

    font-weight: 700;

    margin-top: 10px;

    opacity: 0.95;

}


/* ========================================================
   홈 히어로 배너
   (K water Tech 사이트 참고 - 큰 그라데이션 배경 +
    로고 모양 아치 + 색상 혼합 대형 타이틀)
======================================================== */

.hero-banner {

    position: relative;

    overflow: hidden;

    background:
        linear-gradient(
            135deg,
            #052c4a 0%,
            #075985 45%,
            #0891b2 78%,
            #22c9d6 100%
        );

    border-radius: 22px;

    padding: 38px 32px 34px 32px;

    margin-bottom: 18px;

    box-shadow:
        0 14px 34px rgba(3, 65, 100, 0.22);

}

.hero-banner::before {

    content: "";

    position: absolute;

    top: -60px;

    right: -60px;

    width: 260px;

    height: 260px;

    border-radius: 50%;

    background:
        radial-gradient(
            circle,
            rgba(255, 255, 255, 0.14) 0%,
            rgba(255, 255, 255, 0) 70%
        );

}

.hero-banner::after {

    content: "";

    position: absolute;

    bottom: -80px;

    left: -40px;

    width: 220px;

    height: 220px;

    border-radius: 50%;

    background:
        radial-gradient(
            circle,
            rgba(255, 200, 87, 0.16) 0%,
            rgba(255, 200, 87, 0) 70%
        );

}

.hero-inner {

    position: relative;

    z-index: 1;

    display: block;

}

.hero-inner > * {

    display: block;

    width: 100%;

    margin-bottom: 14px;

}

.hero-logo-wrap {

    width: auto !important;

}

.hero-mascot-wrap {

    text-align: center;

}

.hero-text-wrap {

    width: 100%;

}

.hero-title {

    font-size: 1.55rem;

    font-weight: 900;

    letter-spacing: -0.5px;

    line-height: 1.35;

    color: white;

}

.hero-title .hl-cyan {

    color: #67e8f9;

}

.hero-title .hl-amber {

    color: #ffd166;

}

.hero-sub {

    font-size: 0.85rem;

    color: rgba(255, 255, 255, 0.82);

    margin-top: 6px;

    font-weight: 500;

}

.hero-greeting {

    font-size: 0.82rem;

    color: rgba(255, 255, 255, 0.7);

    margin-top: 10px;

}

.highlight-pill {

    display: inline-block;

    padding: 6px 12px;

    border-radius: 999px;

    font-size: 0.78rem;

    font-weight: 700;

    margin-right: 8px;

    margin-bottom: 6px;

}


/* ========================================================
   상태
======================================================== */

.status-normal {

    display: inline-block;

    background: #e7f8f1;

    color: #087f5b;

    border-radius: 999px;

    padding: 4px 10px;

    font-size: 0.72rem;

    font-weight: 800;

}

.status-watch {

    display: inline-block;

    background: #fff7df;

    color: #a16207;

    border-radius: 999px;

    padding: 4px 10px;

    font-size: 0.72rem;

    font-weight: 800;

}

.status-danger {

    display: inline-block;

    background: #fff0f0;

    color: #c62828;

    border-radius: 999px;

    padding: 4px 10px;

    font-size: 0.72rem;

    font-weight: 800;

}


/* ========================================================
   메뉴
======================================================== */

.menu-caption {

    font-size: 0.68rem;

    opacity: 0.65;

    margin-top: 18px;

    margin-bottom: 4px;

    letter-spacing: 0.5px;

}


/* ========================================================
   진행바 (KPI 실적 등) — 기본 빨간색 대신 브랜드 청록색으로
======================================================== */

[data-testid="stProgress"] > div > div > div {

    background-color: #087ea4 !important;

}


/* ========================================================
   버튼
======================================================== */

.stButton > button {

    border-radius: 9px;

    font-weight: 700;

    min-height: 38px;

    transition:
        transform 0.1s ease,
        box-shadow 0.15s ease,
        background 0.15s ease;

}

.stButton > button:active {

    transform: scale(0.97);

}


/* ========================================================
   탭
======================================================== */

button[data-baseweb="tab"] {

    font-weight: 700;

}


/* ========================================================
   모바일
======================================================== */

@media (max-width: 900px) {

    .kpi-grid {

        grid-template-columns:
            repeat(2, 1fr);

    }

    .top-title {

        font-size: 1.2rem;

    }

}

@media (max-width: 600px) {

    .block-container {

        padding-left: 0.65rem !important;

        padding-right: 0.65rem !important;

    }

    /* 히어로 배너: 좁은 화면에서는 가로 배치(로고-텍스트-
       마스코트)가 짓눌려서 텍스트가 세로로 길게 깨지므로,
       세로로 쌓이는 배치로 전환한다. */

    .hero-inner {

        flex-direction: column !important;

        align-items: flex-start !important;

    }

    .hero-inner > div {

        margin-left: 0 !important;

        width: 100%;

    }

    .hero-inner > div:last-child {

        display: flex !important;
        justify-content: center;
        margin-top: 4px;

    }

    .hero-title {

        font-size: 1.25rem !important;

    }

    /* QR/정밀진단/CBM/오버홀 페이지 헤더의 제목+마스코트도
       같은 원리로 좁은 화면에서 세로로 쌓이게 한다. */

    .page-header-mascot {

        flex-direction: column !important;

        align-items: flex-start !important;

    }

    .page-header-mascot > .page-header-mascot-img {

        align-self: center;

        margin-top: 4px;

    }

    /* 손가락 터치 기준 버튼 최소 높이 확보 (44px 권장) */

    .stButton > button,
    .stDownloadButton > button {

        min-height: 46px !important;

        font-size: 0.95rem !important;

    }

    .story-card {

        padding: 16px !important;

    }

    .kpi-grid {

        grid-template-columns:
            repeat(2, 1fr);

        gap: 7px;

    }

    .kpi-card {

        padding: 11px;

    }

    .kpi-value {

        font-size: 1.2rem;

    }

    /* ====================================================
       실외(햇빛 아래) 스마트폰 사용을 고려해
       글자를 더 굵고 크게, 명암을 더 진하게
    ==================================================== */

    body, p, span, div, label {

        font-size: 1.02rem !important;

    }

    .section-title {

        font-size: 1.25rem !important;

    }

    .card-title {

        font-size: 1.05rem !important;

    }

    .kpi-label {

        color: #334155 !important;

        font-weight: 800 !important;

    }

    .platform-card,
    .kpi-card {

        border-width: 1.5px !important;

    }

    /* ====================================================
       설비 카드 그리드 - 좁은 화면에서는 1열로,
       게이지바/스파크라인이 카드 폭에 맞게 줄어들도록
    ==================================================== */

    .equip-card-grid {

        grid-template-columns: 1fr !important;

    }

    .equip-card {

        padding: 12px 13px !important;

    }

    .equip-card svg {

        width: 70px !important;

    }

}


/* ========================================================
   다크모드 대응 (오늘 새로 추가한 컴포넌트 한정)

   Streamlit 시스템/브라우저가 다크모드일 때, 흰 배경을
   하드코딩해둔 카드들이 붕 떠 보이지 않도록 보정한다.
   (기존에 이미 검증된 platform-card/kpi-card 등은
    건드리지 않고, 오늘 새로 만든 것들만 대상으로 한다)
======================================================== */

@media (prefers-color-scheme: dark) {

    .equip-card {

        background: #1e293b;

        box-shadow:
            0 3px 12px rgba(0, 0, 0, 0.35);

    }

    .equip-card-title {

        color: #e2e8f0;

    }

    .equip-card-meta {

        color: #94a3b8;

    }

    .cbm-bar-track {

        background: #334155;

    }

    .highlight-pill {

        filter: brightness(0.92);

    }

}


/* ========================================================
   인쇄용 스타일

   결재 올릴 때 화면을 그대로 인쇄하면 사이드바·버튼·
   그림자 효과까지 다 찍혀서 지저분했다. 인쇄(또는
   브라우저의 "PDF로 저장")할 때는 꼭 필요한 내용만
   깔끔하게 나오도록 한다.
======================================================== */

/* ========================================================
   사이드바 열기 버튼 확대

   모바일에서 사이드바를 접었을 때 다시 여는 ">" 버튼이
   너무 작아서 누르기 어렵다는 피드백. 아이콘 자체와
   터치 영역을 눈에 띄게 키운다.
======================================================== */

button[data-testid="stExpandSidebarButton"],
button[data-testid="stSidebarCollapsedControl"],
div[data-testid="stSidebarCollapsedControl"] button {

    width: 52px !important;

    height: 52px !important;

    background: #087ea4 !important;

    border-radius: 12px !important;

    box-shadow: 0 4px 12px rgba(8, 126, 164, 0.35) !important;

    position: fixed !important;

    top: 10px !important;

    left: 10px !important;

    z-index: 999999 !important;

}

button[data-testid="stExpandSidebarButton"] svg,
button[data-testid="stSidebarCollapsedControl"] svg,
div[data-testid="stSidebarCollapsedControl"] svg {

    width: 28px !important;

    height: 28px !important;

    color: white !important;

    fill: white !important;

}

@media print {

    section[data-testid="stSidebar"],
    .stButton,
    div[data-testid="stStatusWidget"],
    header[data-testid="stHeader"] {

        display: none !important;

    }

    .platform-card,
    .equip-card,
    .kpi-card {

        box-shadow: none !important;

        border: 1px solid #ccc !important;

        break-inside: avoid;

    }

    .greeting-banner,
    .top-header {

        background: white !important;

        color: #0f3552 !important;

        border: 1px solid #ccc;

    }

    body {

        color: black !important;

    }

}

</style>
""",
    unsafe_allow_html=True
)



# ============================================================
# [pump_calc.py 내용 — 진단기준·등급계산·NPSH계산기·배점검증]
# ============================================================
# ============================================================
# pump_calc.py
#
# 진단 기준·등급 계산 로직 (Streamlit·엑셀과 무관한 순수 계산부).
# 원래 pump_grade_program.py 한 파일(6천줄+)에 전부 섞여 있던 것을
# 모듈로 분리했다. 다른 개발자가 인수인계 받을 때
# "등급 계산 규칙이 어디 있지?"는 이 파일 하나만 보면 된다.
#
# 참고 표준:
# - 진동(Overall Vibration) 판정 구간은 ISO 10816-3
#   (Mechanical vibration — Evaluation of machine vibration by
#   measurements on non-rotating parts — Part 3) Class II
#   (15kW~300kW, 강성 지지) 권고 구간을 참고해서
#   1.8 / 4.5 / 7.1 / 11.2 mm/s 경계값을 사용한다.
#   실제 심의 시에는 이 표준 번호를 근거로 제시할 것.
# ============================================================

# ============================================================
# 5. 진단 기준
# ============================================================

CATEGORIES = {

    "성능": 40,

    "내부상태": 30,

    "기계상태": 25,

    "정비이력": 5

}


def calc_eff(val):

    if val >= 98:
        return "A+"

    if val >= 96:
        return "A"

    if val >= 94:
        return "A-"

    if val >= 92:
        return "B+"

    if val >= 90:
        return "B"

    if val >= 88:
        return "B-"

    if val >= 86:
        return "C+"

    if val >= 84:
        return "C"

    if val >= 82:
        return "C-"

    if val >= 80:
        return "D+"

    if val >= 75:
        return "D"

    return "E"


def calc_reach(val):

    if val >= 98:
        return "A"

    if val >= 93:
        return "B"

    if val >= 88:
        return "C"

    if val >= 80:
        return "D"

    return "E"


def calc_bep(val):

    if 85 <= val <= 115:
        return "A"

    if 75 <= val <= 125:
        return "B"

    if 65 <= val <= 135:
        return "C"

    if 50 <= val <= 150:
        return "D"

    return "E"


def calc_ring_gap(val):

    if val < 1.5:
        return "A"

    if val < 2:
        return "B"

    if val < 2.5:
        return "C"

    if val < 3:
        return "D"

    return "E"


def calc_sleeve(val):

    if val < 1:
        return "A"

    if val < 1.8:
        return "B"

    if val < 2.5:
        return "C"

    if val < 3:
        return "D"

    return "E"


def calc_vib(val):

    if val < 1.8:
        return "A"

    if val < 4.5:
        return "B"

    if val < 7.1:
        return "C"

    if val < 11.2:
        return "D"

    return "E"


def calc_align(val):

    if val <= 0.05:
        return "A"

    if val <= 0.08:
        return "B"

    if val <= 0.12:
        return "C"

    if val <= 0.15:
        return "D"

    return "E"


def calc_overhaul(val):

    if val <= 10000:
        return "A"

    if val <= 12000:
        return "B"

    if val <= 15000:
        return "C"

    if val <= 20000:
        return "D"

    return "E"


# ============================================================
# 6. 17개 핵심 진단항목
# ============================================================

EVAL_ITEMS = [

    (
        "성능",
        "펌프 효율 유지율 (%)",
        25,
        "98% 이상",
        ["A+", "A", "A-", "B+", "B", "B-", "C+", "C", "C-", "D+", "D", "E"],
        {
            "A+": 25,
            "A": 24,
            "A-": 23,
            "B+": 22,
            "B": 21,
            "B-": 20,
            "C+": 18.5,
            "C": 17,
            "C-": 15.5,
            "D+": 14,
            "D": 11,
            "E": 7
        },
        calc_eff
    ),

    (
        "성능",
        "설계 양정/유량 도달률 (%)",
        8,
        "98% 이상",
        ["A", "B", "C", "D", "E"],
        {
            "A": 8,
            "B": 6.8,
            "C": 5.6,
            "D": 4.4,
            "E": 2.4
        },
        calc_reach
    ),

    (
        "성능",
        "BEP 운전점 적정성 (%)",
        7,
        "85~115%",
        ["A", "B", "C", "D", "E"],
        {
            "A": 7,
            "B": 5.95,
            "C": 4.9,
            "D": 3.85,
            "E": 2.1
        },
        calc_bep
    ),

    (
        "내부상태",
        "임펠러/케이싱 링 간극",
        10,
        "1.5 mm 미만",
        ["A", "B", "C", "D", "E"],
        {
            "A": 10,
            "B": 8,
            "C": 6.5,
            "D": 5,
            "E": 3
        },
        calc_ring_gap
    ),

    (
        "내부상태",
        "축슬리브 마모",
        5,
        "1.0 mm 미만",
        ["A", "B", "C", "D", "E"],
        {
            "A": 5,
            "B": 4.25,
            "C": 3.5,
            "D": 2.75,
            "E": 1.5
        },
        calc_sleeve
    ),

    (
        "내부상태",
        "임펠러 손상/침식",
        2,
        "균열 및 침식 없음",
        ["A", "B", "C", "D", "E"],
        {
            "A": 2,
            "B": 1.7,
            "C": 1.4,
            "D": 1.1,
            "E": 0.6
        },
        None
    ),

    (
        "내부상태",
        "임펠러 동적 밸런싱",
        2,
        "ISO G2.5 이하",
        ["A", "B", "C", "D", "E"],
        {
            "A": 2,
            "B": 1.7,
            "C": 1.4,
            "D": 1.1,
            "E": 0.6
        },
        None
    ),

    (
        "내부상태",
        "NPSH 여유율/캐비테이션",
        5,
        "NPSHa/NPSHr ≥ 1.3",
        ["A", "B", "C", "D", "E"],
        {
            "A": 5,
            "B": 4.25,
            "C": 3.5,
            "D": 2.75,
            "E": 1.5
        },
        None
    ),

    (
        "내부상태",
        "내부 코팅 상태",
        3,
        "박리 없음",
        ["A", "B", "C", "D", "E"],
        {
            "A": 3,
            "B": 2.55,
            "C": 2.1,
            "D": 1.65,
            "E": 0.9
        },
        None
    ),

    (
        "내부상태",
        "비금속 웨어링 개선",
        3,
        "복합소재 적용 및 간극 개선",
        ["A", "B", "C", "D", "E"],
        {
            "A": 3,
            "B": 2.55,
            "C": 2.1,
            "D": 1.65,
            "E": 0.9
        },
        None
    ),

    (
        "기계상태",
        "Overall 진동 (mm/s)",
        7,
        "1.8 mm/s 미만",
        ["A", "B", "C", "D", "E"],
        {
            "A": 7,
            "B": 5.95,
            "C": 4.9,
            "D": 3.85,
            "E": 2.1
        },
        calc_vib
    ),

    (
        "기계상태",
        "베어링 결함 진동",
        5,
        "결함 신호 없음",
        ["A", "B", "C", "D", "E"],
        {
            "A": 5,
            "B": 4.25,
            "C": 3.5,
            "D": 2.75,
            "E": 1.5
        },
        None
    ),

    (
        "기계상태",
        "주파수 성분 결함",
        3,
        "특이 피크 없음",
        ["A", "B", "C", "D", "E"],
        {
            "A": 3,
            "B": 2.55,
            "C": 2.1,
            "D": 1.65,
            "E": 0.9
        },
        None
    ),

    (
        "기계상태",
        "펌프-모터 센터링",
        7,
        "0.05 mm 이하",
        ["A", "B", "C", "D", "E"],
        {
            "A": 7,
            "B": 5.95,
            "C": 4.9,
            "D": 3.85,
            "E": 2.1
        },
        calc_align
    ),

    (
        "기계상태",
        "Soft Foot 및 배관 응력",
        3,
        "0.05 mm 이하",
        ["A", "B", "C", "D", "E"],
        {
            "A": 3,
            "B": 2.55,
            "C": 2.1,
            "D": 1.65,
            "E": 0.9
        },
        None
    ),

    (
        "정비이력",
        "오버홀 주기",
        3,
        "10,000시간 이내",
        ["A", "B", "C", "D", "E"],
        {
            "A": 3,
            "B": 2.55,
            "C": 2.1,
            "D": 1.65,
            "E": 0.9
        },
        calc_overhaul
    ),

    (
        "정비이력",
        "주요 소모품 교체이력",
        2,
        "주기 및 이력 양호",
        ["A", "B", "C", "D", "E"],
        {
            "A": 2,
            "B": 1.7,
            "C": 1.4,
            "D": 1.1,
            "E": 0.6
        },
        None
    )

]


# ============================================================
# 6-1. 자동판정 항목별 입력 설정
#
# EVAL_ITEMS 중 auto_fn(계산함수)이 지정된 항목은
# 등급을 직접 고르는 대신 측정값을 입력하면
# 자동으로 등급이 산출된다.
# (min, max, 기본값, step, 단위표시)
# ============================================================

AUTO_INPUT_CONFIG = {

    "펌프 효율 유지율 (%)": (
        0.0, 110.0, 95.0, 0.1, "%"
    ),

    "설계 양정/유량 도달률 (%)": (
        0.0, 150.0, 96.0, 0.1, "%"
    ),

    "BEP 운전점 적정성 (%)": (
        0.0, 200.0, 100.0, 0.1, "%"
    ),

    "임펠러/케이싱 링 간극": (
        0.0, 10.0, 1.2, 0.1, "mm"
    ),

    "축슬리브 마모": (
        0.0, 10.0, 0.8, 0.1, "mm"
    ),

    "Overall 진동 (mm/s)": (
        0.0, 20.0, 2.0, 0.1, "mm/s"
    ),

    "펌프-모터 센터링": (
        0.0, 1.0, 0.05, 0.01, "mm"
    ),

    "오버홀 주기": (
        0.0, 30000.0, 8000.0, 100.0, "시간"
    )

}


# ============================================================
# 6-2. 전문용어 사전
#
# 신입/관리부서 등 비전문가도 볼 수 있도록
# 항목명 옆에 간단한 설명을 붙인다. (도움말 아이콘)
# ============================================================

GLOSSARY = {

    "펌프 효율 유지율 (%)":
        "설계(신품) 효율 대비 현재 효율의 비율. "
        "100%에 가까울수록 신품 상태에 가깝다.",

    "설계 양정/유량 도달률 (%)":
        "설계상 목표한 양정(압력)·유량을 실제로 얼마나 "
        "달성하고 있는지의 비율.",

    "BEP 운전점 적정성 (%)":
        "BEP(Best Efficiency Point, 최고효율점) 대비 "
        "현재 운전점의 위치. 100%에 가까울수록 "
        "펌프가 설계상 가장 효율적인 지점에서 운전 중.",

    "임펠러/케이싱 링 간극":
        "회전체(임펠러)와 고정체(케이싱) 사이의 틈. "
        "마모되어 간극이 커지면 내부 누설이 늘어 효율이 떨어진다.",

    "축슬리브 마모":
        "축을 보호하는 슬리브(축 보호관)의 마모량. "
        "패킹·메커니컬씰과 맞닿는 부분이라 마모되면 누수로 이어진다.",

    "NPSH 여유율/캐비테이션":
        "NPSH(Net Positive Suction Head, 유효흡입양정). "
        "여유가 부족하면 공동현상(캐비테이션)이 발생해 "
        "임펠러가 손상될 수 있다.",

    "Overall 진동 (mm/s)":
        "펌프 전체적인 진동의 크기(속도 실효값, RMS). "
        "베어링·축정렬·불균형 등 여러 원인을 종합적으로 반영한다. "
        "판정구간은 ISO 10816-3 Class II(15~300kW, 강성 지지) "
        "권고 기준을 참고했다.",

    "펌프-모터 센터링":
        "펌프축과 모터축의 중심을 맞추는 작업(축정렬)의 오차. "
        "오차가 크면 진동·베어링 조기마모의 원인이 된다.",

    "Soft Foot 및 배관 응력":
        "Soft Foot: 설비 받침(Foot) 중 한 곳이 뜨거나 "
        "완전히 밀착되지 않은 상태. "
        "배관 응력: 배관이 설비를 당기거나 미는 힘으로, "
        "둘 다 진동·정렬불량의 숨은 원인이 된다.",

    "오버홀 주기":
        "직전 오버홀(분해정비) 이후 누적 운전시간."

}


def term_help(name):

    return GLOSSARY.get(
        name,
        None
    )


# ============================================================
# 6-3. 설비별 맞춤 진동 기준
#
# 원래는 모든 설비에 똑같은 진동 기준(1.8/4.5/7.1/11.2 mm/s)을
# 적용했다. 실제로는 펌프 마력·모델마다 정상범위가 다를 수 있어,
# 설비 마스터에 "기준진동" 값이 등록되어 있으면 그 값을 기준으로
# 등급을 재계산하고, 없으면 기존 공통 기준을 그대로 쓴다.
# ============================================================

def get_effective_auto_fn(
    name,
    default_fn,
    pump
):

    if (

        name == "Overall 진동 (mm/s)"

        and
        pump.get("기준진동")

    ):

        limit = pump["기준진동"]

        def custom_vib(
            val,
            limit=limit
        ):

            if val < limit * 0.6:
                return "A"

            if val < limit:
                return "B"

            if val < limit * 1.4:
                return "C"

            if val < limit * 1.8:
                return "D"

            return "E"

        return custom_vib

    return default_fn


# ============================================================
# 7. 종합등급
# ============================================================

def get_final_grade(score):

    if score >= 90:
        return "A"

    if score >= 80:
        return "B"

    if score >= 70:
        return "C"

    if score >= 60:
        return "D"

    return "E"


def get_grade_text(grade):

    data = {

        "A": "매우 우수 · 정상운전",

        "B": "우수 · 계획관리",

        "C": "관찰 · 추이관리",

        "D": "개선 필요 · 정비검토",

        "E": "위험 · 즉시정비"

    }

    return data.get(
        grade,
        "판정 필요"
    )


def status_icon(status):

    # 색맹(적록색약 등)인 사람도 색깔에만 의존하지 않고
    # 구분할 수 있도록, 상태 배지에 색과 함께 쓰는 기호.

    return {

        "정상": "✓",

        "관찰": "△",

        "정비검토": "✕"

    }.get(
        status,
        "?"
    )


# ============================================================
# 7-1. NPSH(유효흡입양정) 간이 계산기
#
# 기존에는 "NPSH 여유율/캐비테이션" 항목을 사람이 눈대중으로
# A~E 등급을 골랐다. 실제로는 배관 조건으로 NPSHa를 계산할 수
# 있으므로, 계산값이 있으면 그걸로 등급을 자동판정한다.
#
# NPSHa = 대기압수두 - 포화증기압수두 - 흡입양정 - 마찰손실수두
# 여유율(%) = NPSHa / NPSHr(요구값) x 100
# ============================================================

def calc_npsha(
    atm_pressure_head,
    vapor_pressure_head,
    suction_lift,
    friction_loss
):

    return (

        atm_pressure_head
        -
        vapor_pressure_head
        -
        suction_lift
        -
        friction_loss

    )


def calc_npsh_margin_grade(
    margin_ratio_percent
):

    # margin_ratio_percent = NPSHa / NPSHr x 100

    if margin_ratio_percent >= 150:
        return "A"

    if margin_ratio_percent >= 130:
        return "B"

    if margin_ratio_percent >= 115:
        return "C"

    if margin_ratio_percent >= 100:
        return "D"

    return "E"


# ============================================================
# 7-2. 다음 오버홀 예정 시점 추정
#
# 오버홀 주기 판정기준(10,000시간)을 그대로 활용해서
# "현재 운전시간 기준으로 몇 시간 뒤가 다음 오버홀 예정인지"를
# 계산한다. 설비마다 실제 직전 오버홀 시점을 기록해두면 더
# 정확해지지만, 우선은 누적운전시간을 10,000시간 단위로
# 나눈 근사치를 사용한다.
# ============================================================

OVERHAUL_INTERVAL_HOURS = 10000


def estimate_next_overhaul(op_hours):

    if op_hours is None:

        return None, None

    cycles_done = op_hours // OVERHAUL_INTERVAL_HOURS

    next_due_hours = (

        (cycles_done + 1)
        *
        OVERHAUL_INTERVAL_HOURS

    )

    remaining_hours = next_due_hours - op_hours

    return next_due_hours, remaining_hours


def estimate_next_overhaul_advanced(

    pump,
    current_op_hours,
    df_overhaul,
    df_vib

):

    # 기존 estimate_next_overhaul()은 "남은 시간"만 알려줘서
    # 실제 달력 날짜로 와닿지 않았다. 여기서는:
    # 1) 준공일 대비 실제 평균 가동률로 잔여시간을 날짜로 환산
    # 2) 오버홀이력에 실제 기록이 있으면 그 날짜를 기준으로
    #    "보통 2년 주기" 관행을 반영해 앵커로 삼음
    # 3) 최근 진동 추세가 뚜렷이 나빠지고 있으면 예정일을 앞당김
    # 세 가지를 종합해서 달력 날짜 하나로 정리한다.

    today = datetime.now().date()

    basis = []

    daily_rate = 24.0

    try:

        build_date = datetime.strptime(

            str(pump.get("build_date", ""))[:10],

            "%Y-%m-%d"

        ).date()

        days_since_build = (
            today - build_date
        ).days

        if (

            days_since_build > 0

            and current_op_hours

        ):

            daily_rate = (

                current_op_hours / days_since_build

            )

    except Exception:

        pass

    _, remaining_hours = estimate_next_overhaul(
        current_op_hours
    )

    if (

        remaining_hours is None

        or daily_rate <= 0

    ):

        return None, "예측 불가 (운전시간 정보 부족)", False

    remaining_days = remaining_hours / daily_rate

    est_date = today + timedelta(
        days=remaining_days
    )

    basis.append(

        f"누적운전시간 기준 (일평균 {daily_rate:.1f}h 가동 가정)"

    )

    if (

        df_overhaul is not None

        and not df_overhaul.empty

        and "설비명" in df_overhaul.columns

    ):

        equip_hist = df_overhaul[

            df_overhaul["설비명"] == pump["equip"]

        ].copy()

        if not equip_hist.empty:

            equip_hist["_date"] = pd.to_datetime(

                equip_hist["작업일자"],
                errors="coerce"

            )

            equip_hist = equip_hist.dropna(
                subset=["_date"]
            )

            if not equip_hist.empty:

                last_overhaul_date = equip_hist[

                    "_date"

                ].max().date()

                history_based_date = (

                    last_overhaul_date

                    +
                    timedelta(days=730)

                )

                if history_based_date < est_date:

                    est_date = history_based_date

                    basis = [

                        f"최근 오버홀일({last_overhaul_date}) "
                        "기준 2년 주기"

                    ]

                else:

                    basis.append(

                        f"최근 오버홀일({last_overhaul_date}) "
                        "기준 2년 주기와 비교 검토됨"

                    )

    vib_adjusted = False

    if (

        df_vib is not None

        and not df_vib.empty

        and "설비명" in df_vib.columns

    ):

        equip_vib = df_vib[

            df_vib["설비명"] == pump["equip"]

        ].copy()

        if not equip_vib.empty:

            equip_vib["_month"] = equip_vib[

                "측정일자"

            ].astype(str).str[:7]

            monthly_max = equip_vib.groupby(

                "_month"

            )["측정값"].max().sort_index()

            if len(monthly_max) >= 4:

                recent2 = monthly_max.iloc[-2:].mean()

                prev2 = monthly_max.iloc[-4:-2].mean()

                if (

                    prev2 > 0

                    and recent2 > prev2 * 1.15

                ):

                    days_left = (

                        est_date - today

                    ).days

                    if days_left > 0:

                        est_date = today + timedelta(

                            days=int(days_left * 0.8)

                        )

                        vib_adjusted = True

                        basis.append(

                            "최근 진동값 상승추세 감지 → "
                            "예정일 앞당김"

                        )

    return est_date, " · ".join(basis), vib_adjusted


def build_month_calendar_html(year, month, events_by_day):

    # events_by_day: {day(int): [설비명, ...]} 형태.
    # 한 달치 정비 예정을 달력 격자로 그린다(외부 라이브러리
    # 없이 순수 HTML/CSS로).

    import calendar as _cal

    cal = _cal.Calendar(firstweekday=6)  # 일요일 시작

    weeks = cal.monthdayscalendar(year, month)

    weekday_names = ["일", "월", "화", "수", "목", "금", "토"]

    today = datetime.now().date()

    header_cells = "".join(

        f'<th style="padding:6px; font-size:0.8rem; '
        f'color:{"#c62828" if i == 0 else ("#087ea4" if i == 6 else "#334155")};">'
        f'{name}</th>'

        for i, name in enumerate(weekday_names)

    )

    rows_html = []

    for week in weeks:

        cells = []

        for i, day in enumerate(week):

            if day == 0:

                cells.append(

                    '<td style="padding:4px; height:70px; '
                    'vertical-align:top; background:#fafbfc; '
                    'border:1px solid #eef2f6;"></td>'

                )

                continue

            is_today = (

                day == today.day

                and month == today.month

                and year == today.year

            )

            events = events_by_day.get(day, [])

            event_html = "".join(

                f'<div style="background:#fff3cd; '
                f'color:#7a5c00; font-size:0.68rem; '
                f'padding:1px 4px; border-radius:4px; '
                f'margin-top:2px; white-space:nowrap; '
                f'overflow:hidden; text-overflow:ellipsis;" '
                f'title="{e}">🔧 {e}</div>'

                for e in events[:3]

            )

            if len(events) > 3:

                event_html += (

                    f'<div style="font-size:0.65rem; '
                    f'color:#94a3b8;">+{len(events)-3}건 더</div>'

                )

            day_color = (

                "#c62828" if i == 0

                else ("#087ea4" if i == 6 else "#334155")

            )

            cell_bg = "#e6f7fb" if is_today else "white"

            cells.append(

                f'<td style="padding:4px; height:70px; '
                f'vertical-align:top; background:{cell_bg}; '
                f'border:1px solid #eef2f6;">'
                f'<div style="font-size:0.78rem; font-weight:700; '
                f'color:{day_color};">{day}</div>'
                f'{event_html}'
                f'</td>'

            )

        rows_html.append(
            f"<tr>{''.join(cells)}</tr>"
        )

    return (

        '<div style="overflow-x:auto;">'
        '<table style="width:100%; border-collapse:collapse; '
        'table-layout:fixed;">'
        f"<thead><tr>{header_cells}</tr></thead>"
        f"<tbody>{''.join(rows_html)}</tbody>"
        "</table></div>"

    )


def build_overhaul_ics_bytes(schedule_rows):

    # schedule_rows: [(설비명, 사업장, 날짜, 근거설명), ...]
    # 구글캘린더·아웃룩 등에 그대로 가져올 수 있는 표준 .ics 포맷.

    lines = [

        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Kwatertech//OverhaulSchedule//KO"

    ]

    for equip, site, date_obj, basis in schedule_rows:

        dt_str = date_obj.strftime("%Y%m%d")

        uid = (

            f"{equip}-{dt_str}@kwatertech"

        ).replace(" ", "_").replace("(", "").replace(")", "")

        lines += [

            "BEGIN:VEVENT",
            f"UID:{uid}",
            f"DTSTART;VALUE=DATE:{dt_str}",
            f"DTEND;VALUE=DATE:{dt_str}",
            f"SUMMARY:{equip} 오버홀 예정 ({site})",
            f"DESCRIPTION:{basis}",
            "END:VEVENT"

        ]

    lines.append("END:VCALENDAR")

    return "\r\n".join(lines).encode("utf-8")


def build_equipment_timeline(

    pump,
    df_history,
    df_overhaul,
    df_vib

):

    # 정밀진단·오버홀·진동측정 이력이 메뉴마다 따로 흩어져
    # 있어서, "이 설비한테 언제 무슨 일이 있었는지"를 한눈에
    # 보기 어려웠다. 세 이력을 시간순으로 합쳐서 하나의
    # 스트림으로 만든다.

    events = []

    if (

        df_history is not None

        and not df_history.empty

        and "설비명" in df_history.columns

    ):

        for _, r in df_history[

            df_history["설비명"] == pump["equip"]

        ].iterrows():

            events.append(

                {
                    "date": str(r.get("점검일", "")),
                    "type": "정밀진단",
                    "icon": "🔍",
                    "desc": (

                        f"{r.get('최종등급', '-')}등급 "
                        f"({r.get('종합점수', '-')}점)"

                    )
                }

            )

    if (

        df_overhaul is not None

        and not df_overhaul.empty

        and "설비명" in df_overhaul.columns

    ):

        for _, r in df_overhaul[

            df_overhaul["설비명"] == pump["equip"]

        ].iterrows():

            events.append(

                {
                    "date": str(r.get("작업일자", "")),
                    "type": "오버홀",
                    "icon": "🛠️",
                    "desc": (

                        f"{r.get('공정단계', '작업')} · "
                        f"{r.get('작업자', '')}"

                    )
                }

            )

    if (

        df_vib is not None

        and not df_vib.empty

        and "설비명" in df_vib.columns

    ):

        equip_vib = df_vib[

            df_vib["설비명"] == pump["equip"]

        ].copy()

        if not equip_vib.empty:

            equip_vib["_month"] = equip_vib[

                "측정일자"

            ].astype(str).str[:7]

            for month, g in equip_vib.groupby("_month"):

                events.append(

                    {
                        "date": f"{month}-01",
                        "type": "진동측정",
                        "icon": "📈",
                        "desc": (

                            f"최대 {g['측정값'].max():.1f} mm/s "
                            f"({len(g)}건 측정)"

                        )
                    }

                )

    events.sort(

        key=lambda e: e["date"],

        reverse=True

    )

    return events


# ============================================================
# 10-6. 소모품 재고관리
#
# 그랜드패킹·베어링·축슬리브 같은 소모품 교체이력을 지금까지
# 오버홀 작업내용에 텍스트로만 남겼는데, 실제 재고수량까지
# 관리하면 "다음 오버홀 예정은 있는데 재고가 부족하다" 같은
# 걸 미리 알 수 있다.
# ============================================================

CONSUMABLES_DB_PATH = "Pump_Consumables_DB.xlsx"


def ensure_consumables_db_exists():

    ensure_excel_file(

        CONSUMABLES_DB_PATH,

        "소모품재고",

        [
            "소모품명",
            "규격",
            "현재고",
            "안전재고",
            "최근입고일",
            "비고"
        ]

    )


def get_consumables():

    return read_excel(

        CONSUMABLES_DB_PATH,

        "소모품재고"

    )


def add_consumable(row):

    safe_append_row(

        CONSUMABLES_DB_PATH,

        "소모품재고",

        [
            row["소모품명"],
            row["규격"],
            row["현재고"],
            row["안전재고"],
            row["최근입고일"],
            row.get("비고", "")
        ]

    )

    log_audit(

        "소모품 등록",

        row["소모품명"],

        f"현재고={row['현재고']}, 안전재고={row['안전재고']}"

    )


def update_consumable_stock(name, new_qty, note=""):

    with get_lock(CONSUMABLES_DB_PATH):

        wb = load_workbook(
            CONSUMABLES_DB_PATH
        )

        ws = wb["소모품재고"]

        for r in ws.iter_rows(min_row=2):

            if r[0].value == name:

                r[2].value = new_qty

                if note:

                    r[4].value = note

                break

        wb.save(
            CONSUMABLES_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()

    log_audit(

        "소모품 재고 변경",

        name,

        f"현재고 → {new_qty}"

    )


def delete_consumable(name):

    with get_lock(CONSUMABLES_DB_PATH):

        wb = load_workbook(
            CONSUMABLES_DB_PATH
        )

        ws = wb["소모품재고"]

        for r in ws.iter_rows(min_row=2):

            if r[0].value == name:

                ws.delete_rows(r[0].row)

                break

        wb.save(
            CONSUMABLES_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()

    log_audit(

        "소모품 삭제",

        name

    )


def get_low_stock_consumables():

    df = get_consumables()

    if df.empty:

        return df

    return df[

        pd.to_numeric(df["현재고"], errors="coerce").fillna(0)

        <

        pd.to_numeric(df["안전재고"], errors="coerce").fillna(0)

    ]


# 7-3. 배점 합계 자동 검증
#
# 17개 항목의 가중치 합, 그리고 4개 분야(CATEGORIES) 배점 합이
# 각각 100점이어야 정상이다. 누군가 항목을 고치다가 실수로
# 합이 어긋나도 예전에는 아무 경고 없이 조용히 잘못된 점수를
# 계속 뱉어냈다. 이제 시작할 때 검증해서 어긋나면 바로 알린다.
# ============================================================

def validate_eval_weights():

    errors = []

    item_total = sum(
        item[2]
        for item in EVAL_ITEMS
    )

    if abs(item_total - 100) > 0.01:

        errors.append(

            f"EVAL_ITEMS 가중치 합계가 {item_total}점입니다 "
            f"(100점이어야 함)."

        )

    category_total = sum(
        CATEGORIES.values()
    )

    if abs(category_total - 100) > 0.01:

        errors.append(

            f"CATEGORIES 배점 합계가 {category_total}점입니다 "
            f"(100점이어야 함)."

        )

    category_item_sum = {}

    for item in EVAL_ITEMS:

        category_item_sum[item[0]] = (

            category_item_sum.get(item[0], 0)
            +
            item[2]

        )

    for cat, cat_weight in CATEGORIES.items():

        item_sum = category_item_sum.get(cat, 0)

        if abs(item_sum - cat_weight) > 0.01:

            errors.append(

                f"'{cat}' 분야: EVAL_ITEMS 합계 {item_sum}점 "
                f"vs CATEGORIES 배점 {cat_weight}점이 서로 다릅니다."

            )

    return errors


# ============================================================
# [pump_db.py 내용 — 설비마스터·진단이력 등 엑셀 DB 저장/조회]
# ============================================================
# ============================================================
# pump_db.py
#
# 엑셀 기반 데이터 저장/조회 로직 모음.
# (설비 마스터, 진단이력 DB 생성, 안전저장, 캐시 등)
#
# 이 모듈은 Streamlit 페이지를 그리지 않는다 — 오직
# "데이터를 어떻게 저장하고 읽어오는가"만 담당한다.
# 화면(UI)은 pump_grade_program.py 쪽에 있다.
# ============================================================

import os
import random
from datetime import datetime

import pandas as pd
import streamlit as st
from openpyxl import Workbook, load_workbook
from filelock import FileLock


# ============================================================
# 0. DB 경로 상수
# ============================================================

DB_FILE_PATH = "Pump_Master_DB.xlsx"
OVERHAUL_DB_PATH = "Pump_Overhaul_DB.xlsx"
DOC_DB_PATH = "Pump_Docs_DB.xlsx"
KNOWHOW_DB_PATH = "Pump_Knowhow_DB.xlsx"
DAILY_LOG_DB_PATH = "Pump_DailyLog_DB.xlsx"
SAFETY_PERMIT_DB_PATH = "Pump_SafetyPermit_DB.xlsx"
KPI_DB_PATH = "Pump_KPI_DB.xlsx"
EQUIP_DB_PATH = "Pump_Equipment_DB.xlsx"

SEED_FLAG_PATH = "seed_flag.txt"

EQUIP_SEED_FLAG_PATH = "equip_seed_flag.txt"




# ============================================================
# 4. 설비 기본 DB (최초 시딩용 — 실제 목록은 EQUIP_DB_PATH)
# ============================================================

DEFAULT_PUMPS = [

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #1",
        "maker": "효성펌프",
        "model": "DHP-1",
        "hp": 160,
        "head": 47,
        "flow": 1250,
        "rpm": 1780,
        "build_date": "2018-01-15",
        "op_hours": 10200
    },

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #2",
        "maker": "현대중공업",
        "model": "VTP-20",
        "hp": 170,
        "head": 49,
        "flow": 1300,
        "rpm": 1780,
        "build_date": "2018-02-15",
        "op_hours": 9200
    },

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #3",
        "maker": "효성펌프",
        "model": "DHP-3",
        "hp": 180,
        "head": 51,
        "flow": 1350,
        "rpm": 1780,
        "build_date": "2018-03-15",
        "op_hours": 7800
    },

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #4",
        "maker": "현대중공업",
        "model": "VTP-40",
        "hp": 190,
        "head": 53,
        "flow": 1400,
        "rpm": 1780,
        "build_date": "2018-04-15",
        "op_hours": 10500
    },

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #5",
        "maker": "효성펌프",
        "model": "DHP-5",
        "hp": 200,
        "head": 55,
        "flow": 1450,
        "rpm": 1780,
        "build_date": "2018-05-15",
        "op_hours": 9800
    },

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #6",
        "maker": "현대중공업",
        "model": "VTP-60",
        "hp": 210,
        "head": 57,
        "flow": 1500,
        "rpm": 1780,
        "build_date": "2018-06-15",
        "op_hours": 8700
    },

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #7",
        "maker": "효성펌프",
        "model": "DHP-7",
        "hp": 220,
        "head": 59,
        "flow": 1550,
        "rpm": 1780,
        "build_date": "2018-07-15",
        "op_hours": 11400
    },

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #8",
        "maker": "현대중공업",
        "model": "VTP-80",
        "hp": 230,
        "head": 61,
        "flow": 1600,
        "rpm": 1780,
        "build_date": "2018-08-15",
        "op_hours": 7200
    },

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #9",
        "maker": "효성펌프",
        "model": "DHP-9",
        "hp": 240,
        "head": 63,
        "flow": 1650,
        "rpm": 1780,
        "build_date": "2018-09-15",
        "op_hours": 12300
    },

    {
        "site": "밀양정수장",
        "equip": "가압펌프 #10",
        "maker": "현대중공업",
        "model": "VTP-100",
        "hp": 250,
        "head": 65,
        "flow": 1700,
        "rpm": 1780,
        "build_date": "2018-10-15",
        "op_hours": 9100
    }

]


# ============================================================
# 7-4. 앱 설정 파일 (배포 주소 등 코드 수정 없이 바꿀 수 있는 값)
#
# 예전에는 QR코드가 가리키는 배포 주소가 코드에 박혀 있어서
# 실제 배포 주소가 바뀌면 코드를 다시 고쳐서 올려야 했다.
# 이제 관리자 화면(데이터관리 페이지)에서 입력하면
# 텍스트 파일로 저장되고, 다음부터는 그 값을 쓴다.
# ============================================================

APP_CONFIG_FILE = os.path.join(
    CONFIG_DIR,
    "app_base_url.txt"
)

DEFAULT_APP_BASE_URL = "https://kwatertech-pump.streamlit.app"


def get_app_base_url():

    if os.path.exists(APP_CONFIG_FILE):

        try:

            with open(
                APP_CONFIG_FILE,
                "r",
                encoding="utf-8"
            ) as f:

                saved = f.read().strip()

            if saved:

                return saved

        except Exception:

            pass

    return DEFAULT_APP_BASE_URL


def set_app_base_url(url):

    os.makedirs(
        CONFIG_DIR,
        exist_ok=True
    )

    with open(
        APP_CONFIG_FILE,
        "w",
        encoding="utf-8"
    ) as f:

        f.write(
            url.strip()
        )


# ============================================================
# 7-4-1. 알림 임계값 (관리자가 화면에서 직접 조정)
#
# 주의: 정밀진단 17개 항목의 등급(A~E) 계산 기준은
# ISO 10816-3 등 공식 기준을 따르는 채점 엔진이라 바꾸지 않는다.
# 여기서 조정 가능한 건 AI 이상징후·홈 알림·정비권고사항에서
# 쓰는 "관찰/주의" 보조 알림 기준이다.
# ============================================================

THRESHOLDS_CONFIG_FILE = os.path.join(
    CONFIG_DIR,
    "alert_thresholds.json"
)

DEFAULT_ALERT_THRESHOLDS = {

    "vib_watch": 4.5,
    "vib_danger": 7.1,
    "eff_watch": 80.0,
    "eff_danger": 70.0,
    "temp_watch": 50.0,
    "temp_danger": 55.0

}


def get_alert_thresholds():

    if os.path.exists(THRESHOLDS_CONFIG_FILE):

        try:

            with open(

                THRESHOLDS_CONFIG_FILE,

                "r",

                encoding="utf-8"

            ) as f:

                saved = json.load(f)

            merged = dict(
                DEFAULT_ALERT_THRESHOLDS
            )

            merged.update(
                saved
            )

            return merged

        except Exception:

            pass

    return dict(
        DEFAULT_ALERT_THRESHOLDS
    )


def save_alert_thresholds(values):

    os.makedirs(
        CONFIG_DIR,
        exist_ok=True
    )

    with open(

        THRESHOLDS_CONFIG_FILE,

        "w",

        encoding="utf-8"

    ) as f:

        json.dump(
            values,
            f,
            ensure_ascii=False
        )


# ============================================================
# 7-4-2. 알림 웹훅(Slack 등) 설정
#
# 위험설비가 감지되면(정밀진단 D/E등급 저장 시) 설정된 웹훅
# 주소로 메시지를 보낸다. Slack의 "Incoming Webhook" 형식과
# 호환되는 형태({"text": "..."})로 보낸다.
# ============================================================

WEBHOOK_CONFIG_FILE = os.path.join(
    CONFIG_DIR,
    "webhook_url.txt"
)


def get_webhook_url():

    if os.path.exists(WEBHOOK_CONFIG_FILE):

        try:

            with open(

                WEBHOOK_CONFIG_FILE,

                "r",

                encoding="utf-8"

            ) as f:

                saved = f.read().strip()

            return saved

        except Exception:

            pass

    return ""


def set_webhook_url(url):

    os.makedirs(
        CONFIG_DIR,
        exist_ok=True
    )

    with open(

        WEBHOOK_CONFIG_FILE,

        "w",

        encoding="utf-8"

    ) as f:

        f.write(
            url.strip()
        )


def send_webhook_notification(message):

    # 웹훅 주소가 설정 안 돼 있으면 조용히 건너뛴다.
    # 네트워크 문제 등으로 실패해도 앱이 멈추면 안 되므로
    # 예외를 전부 흡수한다.

    url = get_webhook_url()

    if not url:

        return False

    try:

        req = urllib.request.Request(

            url,

            data=json.dumps(

                {"text": message}

            ).encode("utf-8"),

            headers={
                "Content-Type": "application/json"
            }

        )

        urllib.request.urlopen(
            req,
            timeout=5
        )

        return True

    except Exception:

        return False


# ============================================================
# 7-4-3. 감사 로그 (누가 언제 무엇을 삭제/추가했는지)
# ============================================================

AUDIT_LOG_PATH = "Pump_AuditLog_DB.xlsx"


def ensure_audit_log_exists():

    ensure_excel_file(

        AUDIT_LOG_PATH,

        "감사로그",

        [
            "일시",
            "사용자",
            "동작",
            "대상",
            "상세"
        ]

    )


def log_audit(action, target, detail=""):

    try:

        safe_append_row(

            AUDIT_LOG_PATH,

            "감사로그",

            [
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                st.session_state.get("user_name", ""),
                action,
                target,
                detail
            ]

        )

    except Exception:

        pass


# ============================================================
# 7-4-4. 도면 이력 관리
#
# 예전에는 도면을 새로 올리면 이전 도면이 세션에서만 잠깐
# 저장됐다가 그냥 덮어써졌다(디스크 저장조차 안 됨). 이제
# 디스크에 저장하고 이력을 DB로 남겨서, 예전 도면과 비교해
# 볼 수 있게 한다.
# ============================================================

DRAWING_DIR = "drawings"

DRAWING_DB_PATH = "Pump_Drawing_DB.xlsx"


def ensure_drawing_db_exists():

    ensure_excel_file(

        DRAWING_DB_PATH,

        "도면이력",

        [
            "등록일자",
            "설비명",
            "파일명",
            "원본파일명"
        ]

    )


def save_drawing(equip_name, file_bytes, original_filename):

    os.makedirs(
        DRAWING_DIR,
        exist_ok=True
    )

    ext = (

        original_filename.rsplit(".", 1)[-1]

        if "." in original_filename

        else "png"

    )

    safe_name = "".join(

        c if c.isalnum() else "_"

        for c in equip_name

    )

    saved_filename = (

        f"{safe_name}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}."
        f"{ext}"

    )

    with open(

        os.path.join(
            DRAWING_DIR,
            saved_filename
        ),

        "wb"

    ) as f:

        f.write(
            file_bytes
        )

    safe_append_row(

        DRAWING_DB_PATH,

        "도면이력",

        [
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            equip_name,
            saved_filename,
            original_filename
        ]

    )

    log_audit(

        "도면 업로드",

        equip_name,

        original_filename

    )


# ============================================================
# 7-5. 정밀진단 임시저장(초안) — 디스크 백업
#
# 세션 안에서 페이지 이동은 diag_draft(세션 상태)로 해결했지만,
# 브라우저를 새로고침하거나 와이파이가 끊겼다 붙으면
# 세션 자체가 새로 시작되어 입력하던 값이 날아간다.
# on_change 콜백이 일어날 때마다 디스크에도 같이 저장해두고,
# 세션이 새로 시작될 때 디스크에서 복구한다.
# ============================================================

import json


def _draft_file_path(equip_name):

    safe_name = "".join(

        c if c.isalnum() else "_"

        for c in equip_name

    )

    return os.path.join(
        DRAFT_DIR,
        f"{safe_name}.json"
    )


def load_draft_from_disk(equip_name):

    path = _draft_file_path(
        equip_name
    )

    if not os.path.exists(path):

        return {}

    try:

        with open(
            path,
            "r",
            encoding="utf-8"
        ) as f:

            return json.load(f)

    except Exception:

        return {}


def save_draft_to_disk(equip_name, draft):

    os.makedirs(
        DRAFT_DIR,
        exist_ok=True
    )

    path = _draft_file_path(
        equip_name
    )

    try:

        with open(
            path,
            "w",
            encoding="utf-8"
        ) as f:

            json.dump(
                draft,
                f,
                ensure_ascii=False
            )

    except Exception:

        pass


def clear_draft(equip_name):

    if equip_name in st.session_state.get(
        "diag_draft",
        {}
    ):

        st.session_state.diag_draft[equip_name] = {}

    path = _draft_file_path(
        equip_name
    )

    if os.path.exists(path):

        try:

            os.remove(path)

        except Exception:

            pass


# ============================================================
# 7-6. 예쁜 엑셀 생성 (스타일 + 자동 차트)
#
# 예전에는 df.to_excel()로 서식 하나 없이 그냥 표만 덤프해서
# "보고서 내용이 너무 간단하다"는 지적을 받았다. 이제 제목행,
# 색깔이 들어간 헤더, 지브라 줄무늬, 열너비 자동조정, 헤더 고정에
# 더해 숫자 데이터가 있으면 엑셀 자체 차트(그림이 아니라 진짜
# 엑셀 차트 객체)까지 자동으로 넣는다.
# ============================================================

BRAND_HEADER_COLOR = "0F3552"
BRAND_ACCENT_COLOR = "087EA4"
BRAND_ZEBRA_COLOR = "F4F9FC"


def _write_excel_sheet(
    ws,
    df,
    title,
    chart_kind
):

    n_cols = max(
        len(df.columns),
        1
    )

    thin = Side(
        style="thin",
        color="CCCCCC"
    )

    border = Border(
        left=thin,
        right=thin,
        top=thin,
        bottom=thin
    )

    # ---- 제목행 ----

    ws.merge_cells(

        start_row=1,
        start_column=1,
        end_row=1,
        end_column=n_cols

    )

    title_cell = ws.cell(
        row=1,
        column=1,
        value=f"💧 {title}"
    )

    title_cell.font = Font(
        size=16,
        bold=True,
        color="FFFFFF"
    )

    title_cell.fill = PatternFill(
        "solid",
        fgColor=BRAND_ACCENT_COLOR
    )

    title_cell.alignment = Alignment(
        horizontal="center",
        vertical="center"
    )

    ws.row_dimensions[1].height = 28

    # ---- 부제(생성시각) ----

    ws.merge_cells(

        start_row=2,
        start_column=1,
        end_row=2,
        end_column=n_cols

    )

    sub_cell = ws.cell(

        row=2,

        column=1,

        value=(

            "생성일시 : "
            +
            datetime.now().strftime("%Y-%m-%d %H:%M")
            +
            "  ·  K-water tech 설비관리 통합 플랫폼"

        )

    )

    sub_cell.font = Font(
        size=9,
        italic=True,
        color="666666"
    )

    sub_cell.alignment = Alignment(
        horizontal="center"
    )

    header_row = 4

    # ---- 헤더 ----

    header_fill = PatternFill(
        "solid",
        fgColor=BRAND_HEADER_COLOR
    )

    header_font = Font(
        bold=True,
        color="FFFFFF"
    )

    for j, col in enumerate(
        df.columns,
        start=1
    ):

        c = ws.cell(
            row=header_row,
            column=j,
            value=str(col)
        )

        c.font = header_font

        c.fill = header_fill

        c.alignment = Alignment(
            horizontal="center",
            vertical="center"
        )

        c.border = border

    # ---- 데이터(지브라 줄무늬) ----

    zebra_fill = PatternFill(
        "solid",
        fgColor=BRAND_ZEBRA_COLOR
    )

    for i, row in enumerate(

        df.itertuples(index=False),

        start=header_row + 1

    ):

        for j, val in enumerate(
            row,
            start=1
        ):

            c = ws.cell(
                row=i,
                column=j,
                value=val
            )

            c.border = border

            if (i - header_row) % 2 == 0:

                c.fill = zebra_fill

    # ---- 열 너비 자동조정 ----

    for j, col in enumerate(
        df.columns,
        start=1
    ):

        if len(df) > 0:

            max_len = max(

                [len(str(col))]
                +
                [
                    len(str(v))
                    for v in df[col].astype(str)
                ]

            )

        else:

            max_len = len(str(col))

        ws.column_dimensions[

            get_column_letter(j)

        ].width = min(
            max(max_len + 4, 10),
            40
        )

    ws.freeze_panes = ws.cell(

        row=header_row + 1,

        column=1

    ).coordinate

    # ---- 그래프 삽입 ----
    #
    # chart_kind:
    #   "auto" - 숫자열이 있으면 전부 묶어서 막대차트 (기존 방식)
    #   "bar"  - 첫 숫자열 하나로 막대차트 (화면의 막대그래프와 매칭)
    #   "line" - 첫 숫자열 하나로 꺾은선차트 (화면의 선/영역그래프와 매칭)
    #   None   - 차트 없음

    if chart_kind and len(df) >= 2:

        numeric_cols = [

            j

            for j, col in enumerate(df.columns, start=1)

            if pd.api.types.is_numeric_dtype(df[col])

        ]

        if numeric_cols:

            if chart_kind == "line":

                chart = LineChart()

            else:

                chart = BarChart()

            chart.title = title

            chart.height = 9

            chart.width = 18

            chart.style = 10

            if chart_kind == "auto":

                min_c, max_c = numeric_cols[0], numeric_cols[-1]

            else:

                min_c = max_c = numeric_cols[0]

            data_ref = Reference(

                ws,

                min_col=min_c,

                max_col=max_c,

                min_row=header_row,

                max_row=header_row + len(df)

            )

            cats_ref = Reference(

                ws,

                min_col=1,

                min_row=header_row + 1,

                max_row=header_row + len(df)

            )

            chart.add_data(

                data_ref,

                titles_from_data=True

            )

            chart.set_categories(
                cats_ref
            )

            chart_anchor = (

                f"{get_column_letter(n_cols + 2)}"
                f"{header_row}"

            )

            ws.add_chart(
                chart,
                chart_anchor
            )


def build_pretty_excel_bytes(
    sheets,
    title
):

    # sheets: pandas DataFrame 하나를 그대로 넘기면 단일 시트로
    # 처리하고(예전 방식과 호환), 여러 그래프를 함께 넣고 싶으면
    # [{"name":.., "df":.., "chart": "bar"/"line"/"auto"/None}, ...]
    # 형태의 리스트를 넘긴다.

    if isinstance(sheets, pd.DataFrame):

        sheets = [

            {
                "name": "데이터",
                "df": sheets,
                "chart": "auto"
            }

        ]

    wb = Workbook()

    wb.remove(
        wb.active
    )

    for spec in sheets:

        sheet_title = spec.get(
            "name",
            "데이터"
        )[:31]

        ws = wb.create_sheet(
            sheet_title
        )

        _write_excel_sheet(

            ws,

            spec["df"],

            spec.get("title", title),

            spec.get("chart", "auto")

        )

    buffer = io.BytesIO()

    wb.save(buffer)

    return buffer.getvalue()


# ============================================================
# 7-6-1. 설비별 월간 보고서 (Word)
#
# 월간보고서는 사내에서 쓰는 다른 화면들과 달리 회사 밖으로
# 나가는(혹은 상위 결재를 받는) 유일한 산출물이라 별도로
# 공을 들인다. 화면에 보이는 데이터를 최대한 그대로 담고,
# 그래프도 이미지가 아니라 화면과 동일한 그래프를 그대로 삽입한다.
# ============================================================

def find_logo_file():

    for pattern in (
        "[Ll]ogo.png",
        "[Ll]ogo.jpg",
        "[Ll]ogo.jpeg"
    ):

        matches = glob.glob(pattern)

        if matches:

            return matches[0]

    return None


def find_manual_pdf():

    # 사용자 매뉴얼 PDF를 앱 폴더에 같이 올려두면(파일명에
    # "매뉴얼"이 들어가면) 자동으로 찾아서 다운로드 버튼을
    # 보여준다. 없으면 그냥 조용히 버튼을 숨긴다.

    for pattern in (
        "*매뉴얼*.pdf",
        "*manual*.pdf",
        "*Manual*.pdf"
    ):

        matches = glob.glob(pattern)

        if matches:

            return matches[0]

    return None


def render_manual_pdf_viewer_html(pdf_path, height_px=700):

    # 다운로드 없이 앱 화면 안에서 바로 매뉴얼을 넘겨볼 수
    # 있도록, PDF를 base64로 인코딩해서 iframe에 그대로
    # 끼워 넣는다(브라우저 내장 PDF 뷰어가 그려준다).

    try:

        with open(
            pdf_path,
            "rb"
        ) as f:

            b64 = base64.b64encode(
                f.read()
            ).decode("utf-8")

    except Exception:

        return None

    return (

        f'<iframe src="data:application/pdf;base64,{b64}" '
        f'width="100%" height="{height_px}" '
        f'style="border:1px solid #dce8ef; border-radius:10px;">'
        f'</iframe>'

    )


_KOREAN_PDF_FONT_NAME = "Helvetica"

_korean_font_registered = False


def ensure_korean_pdf_font():

    # reportlab 기본 폰트(Helvetica 등)는 한글을 못 그린다.
    # 한글 폰트 파일(NanumGothic.ttf 등)을 앱 폴더에 같이
    # 올려두면 그걸 등록해서 쓰고, 없으면 어쩔 수 없이
    # 기본 폰트로 폴백한다(그러면 한글이 깨져 보일 수 있음).

    global _KOREAN_PDF_FONT_NAME, _korean_font_registered

    if _korean_font_registered:

        return _KOREAN_PDF_FONT_NAME

    for pattern in (
        "NanumGothic.ttf",
        "[Nn]anum*.ttf",
        "*고딕*.ttf",
        "*Gothic*.ttf"
    ):

        matches = glob.glob(pattern)

        if matches:

            try:

                pdfmetrics.registerFont(

                    TTFont(
                        "KoreanFont",
                        matches[0]
                    )

                )

                _KOREAN_PDF_FONT_NAME = "KoreanFont"

                break

            except Exception:

                continue

    _korean_font_registered = True

    return _KOREAN_PDF_FONT_NAME


def build_dashboard_summary_pdf(

    all_pumps,
    df_history,
    site_list

):

    # 설비 하나씩 보고서를 뽑는 게 아니라, 이번 달 전체 현황을
    # 한 장으로 정리한 경영보고용 요약 PDF.

    font_name = ensure_korean_pdf_font()

    buffer = io.BytesIO()

    c = rl_canvas.Canvas(
        buffer,
        pagesize=A4
    )

    page_w, page_h = A4

    margin = 18 * mm

    y = page_h - margin

    logo_path = find_logo_file()

    if logo_path:

        try:

            c.drawImage(

                logo_path,

                margin,

                y - 16 * mm,

                width=40 * mm,

                height=12 * mm,

                preserveAspectRatio=True,

                mask="auto"

            )

        except Exception:

            pass

    c.setFont(
        font_name,
        18
    )

    c.setFillColorRGB(
        0.06, 0.21, 0.33
    )

    c.drawString(

        margin,

        y - 26 * mm,

        f"설비관리 통합 플랫폼 · "
        f"{datetime.now().strftime('%Y년 %m월')} 종합 대시보드"

    )

    y -= 34 * mm

    # 전체 요약

    normal = watch = repair = 0

    repair_items = []

    site_summary = {

        site: {"정상": 0, "관찰": 0, "정비검토": 0}

        for site in site_list

    }

    for pump in all_pumps:

        result = pump_status(

            pump,
            df_history

        )

        site_summary.setdefault(

            pump["site"],
            {"정상": 0, "관찰": 0, "정비검토": 0}

        )

        if result["상태"] == "정상":

            normal += 1

            site_summary[pump["site"]]["정상"] += 1

        elif result["상태"] == "관찰":

            watch += 1

            site_summary[pump["site"]]["관찰"] += 1

        else:

            repair += 1

            site_summary[pump["site"]]["정비검토"] += 1

            repair_items.append(

                (pump["equip"], pump["site"], result["점수"])

            )

    c.setFont(
        font_name,
        13
    )

    c.setFillColorRGB(
        0.06, 0.21, 0.33
    )

    c.drawString(

        margin,

        y,

        f"전체 설비: {len(all_pumps)}대   "
        f"정상: {normal}대   관찰: {watch}대   "
        f"정비검토: {repair}대"

    )

    y -= 12 * mm

    # 사업장별 표

    c.setFont(
        font_name,
        12
    )

    c.drawString(
        margin,
        y,
        "사업장별 현황"
    )

    y -= 8 * mm

    col_x = [margin, margin + 55 * mm, margin + 85 * mm, margin + 115 * mm]

    c.setFont(
        font_name,
        10
    )

    for label, x in zip(
        ["사업장", "정상", "관찰", "정비검토"],
        col_x

    ):

        c.drawString(
            x,
            y,
            label
        )

    y -= 6 * mm

    c.line(
        margin,
        y + 3 * mm,
        page_w - margin,
        y + 3 * mm

    )

    for site, counts in site_summary.items():

        c.drawString(
            col_x[0],
            y,
            site
        )

        c.drawString(
            col_x[1],
            y,
            str(counts["정상"])
        )

        c.drawString(
            col_x[2],
            y,
            str(counts["관찰"])
        )

        c.drawString(
            col_x[3],
            y,
            str(counts["정비검토"])
        )

        y -= 6 * mm

    y -= 8 * mm

    # 정비검토 필요 설비 목록

    c.setFont(
        font_name,
        12
    )

    c.setFillColorRGB(
        0.78, 0.16, 0.16
    )

    c.drawString(

        margin,

        y,

        f"⚠ 정비검토 필요 설비 ({len(repair_items)}건)"

    )

    y -= 8 * mm

    c.setFont(
        font_name,
        10
    )

    c.setFillColorRGB(
        0.06, 0.21, 0.33
    )

    if repair_items:

        for equip, site, score in repair_items:

            c.drawString(

                margin,

                y,

                f"- {equip} ({site}) · CBM Score {score}점"

            )

            y -= 6 * mm

    else:

        c.drawString(

            margin,

            y,

            "- 없음 (전 설비 정상/관찰 수준)"

        )

        y -= 6 * mm

    c.setFont(
        font_name,
        8
    )

    c.setFillColorRGB(
        0.5, 0.5, 0.5
    )

    c.drawString(

        margin,

        margin,

        f"생성일시: "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}"

    )

    c.save()

    buffer.seek(0)

    return buffer.getvalue()


def build_qr_label_sheet_pdf(

    pumps,
    cols=3,
    rows=6

):

    # 현장에 붙일 QR 라벨을 A4 한 장에 여러 개씩(기본
    # 3열x6행=18개) 격자로 배치해서 인쇄용 PDF로 만든다.
    # 자르는 선까지 같이 그려서 가위로 자르기 편하게 한다.

    font_name = ensure_korean_pdf_font()

    buffer = io.BytesIO()

    c = rl_canvas.Canvas(
        buffer,
        pagesize=A4
    )

    page_w, page_h = A4

    margin = 10 * mm

    usable_w = page_w - 2 * margin

    usable_h = page_h - 2 * margin

    cell_w = usable_w / cols

    cell_h = usable_h / rows

    per_page = cols * rows

    app_base_url = get_app_base_url()

    for i, pump in enumerate(pumps):

        page_pos = i % per_page

        if (

            i > 0

            and page_pos == 0

        ):

            c.showPage()

        col = page_pos % cols

        row = page_pos // cols

        x0 = margin + col * cell_w

        y0 = page_h - margin - (row + 1) * cell_h

        # 자르는 선 (점선 테두리)

        c.setDash(
            2,
            2
        )

        c.setStrokeColorRGB(
            0.7,
            0.7,
            0.7
        )

        c.rect(
            x0,
            y0,
            cell_w,
            cell_h
        )

        c.setDash()

        # QR 이미지

        qr_target_url = (

            f"{app_base_url}/?page=QR&equip="
            +
            urllib.parse.quote(
                pump["equip"]
            )

        )

        qr_img = qrcode.make(
            qr_target_url
        )

        qr_buf = io.BytesIO()

        qr_img.save(
            qr_buf,
            format="PNG"
        )

        qr_buf.seek(0)

        qr_size = min(
            cell_w * 0.55,
            cell_h * 0.62
        )

        qr_x = x0 + 6

        qr_y = y0 + (cell_h - qr_size) / 2

        c.drawImage(

            ImageReader(qr_buf),

            qr_x,

            qr_y,

            width=qr_size,

            height=qr_size

        )

        # 설비명 / 사업장 텍스트

        text_x = qr_x + qr_size + 8

        text_w = cell_w - (qr_size + 20)

        c.setFont(
            font_name,
            9
        )

        c.setFillColorRGB(
            0.06,
            0.21,
            0.33
        )

        equip_display = pump["equip"]

        if (

            font_name != "Helvetica"

            and c.stringWidth(
                equip_display,
                font_name,
                9

            ) > text_w

        ):

            while (

                len(equip_display) > 3

                and c.stringWidth(
                    equip_display + "…",
                    font_name,
                    9

                ) > text_w

            ):

                equip_display = equip_display[:-1]

            equip_display += "…"

        c.drawString(

            text_x,

            y0 + cell_h / 2 + 4,

            equip_display

        )

        c.setFont(
            font_name,
            7.5
        )

        c.setFillColorRGB(
            0.4,
            0.45,
            0.5
        )

        c.drawString(

            text_x,

            y0 + cell_h / 2 - 9,

            pump["site"]

        )

    c.save()

    buffer.seek(0)

    return buffer.getvalue()


def convert_docx_bytes_to_pdf_bytes(docx_bytes):

    # Word 보고서를 인쇄용 PDF로도 바로 받을 수 있게 한다.
    # 서버에 LibreOffice(soffice)가 설치돼 있어야 변환이
    # 되는데, 없으면 조용히 None을 돌려주고 호출부에서
    # "PDF 변환은 서버에 LibreOffice가 필요합니다" 안내를
    # 보여준다. (packages.txt에 libreoffice-writer 추가 필요)

    tmp_dir = tempfile.mkdtemp()

    try:

        docx_path = os.path.join(
            tmp_dir,
            "report.docx"
        )

        with open(
            docx_path,
            "wb"
        ) as f:

            f.write(
                docx_bytes
            )

        result = subprocess.run(

            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                tmp_dir,
                docx_path
            ],

            capture_output=True,

            timeout=60

        )

        pdf_path = os.path.join(
            tmp_dir,
            "report.pdf"
        )

        if (

            result.returncode == 0

            and os.path.exists(pdf_path)

        ):

            with open(
                pdf_path,
                "rb"
            ) as f:

                return f.read()

        return None

    except Exception:

        return None

    finally:

        shutil.rmtree(
            tmp_dir,
            ignore_errors=True
        )


# ============================================================
# 10-7. 설계적산 (일위대가 라이브러리 + 공사)
#
# 실제 설계·시공 산출내역서 엑셀을 올리면 "일위대가표" 시트를
# 파싱해서 라이브러리에 쌓는다. 라이브러리는 자주 쓰는 순서로
# 정렬되고, 신입사원이 새 공사를 만들 때 자유롭게 검색해서
# 담을 수 있다. 공종별 "가이드라인"(추천 세트)도 별도로 두어,
# 처음 시작할 때 참고할 수 있게 한다.
# ============================================================

DESIGN_UNITPRICE_DB_PATH = "Pump_DesignUnitPrice_DB.xlsx"

DESIGN_GUIDELINE_DB_PATH = "Pump_DesignGuideline_DB.xlsx"

DESIGN_PROJECT_DB_PATH = "Pump_DesignProject_DB.xlsx"

DESIGN_PROJECT_ITEM_DB_PATH = "Pump_DesignProjectItem_DB.xlsx"

DESIGN_GONGJONG_LIST = [

    "배관 교체",
    "펌프 오버홀",
    "배수펌프 교체",
    "잡철물 설치",
    "전동밸브 교체"

]


def ensure_design_dbs_exist():

    ensure_excel_file(

        DESIGN_UNITPRICE_DB_PATH,

        "일위대가라이브러리",

        [
            "항목명",
            "규격",
            "단위",
            "노무비단가",
            "재료비단가",
            "기계경비단가",
            "외주경비단가",
            "합계단가",
            "사용횟수",
            "등록일",
            "출처공사",
            "상세내역JSON"
        ]

    )

    ensure_excel_file(

        DESIGN_GUIDELINE_DB_PATH,

        "공종가이드라인",

        [
            "공사유형",
            "추천항목명"
        ]

    )

    ensure_excel_file(

        DESIGN_PROJECT_DB_PATH,

        "설계공사",

        [
            "공사ID",
            "공사명",
            "공사유형",
            "사업장",
            "작성자",
            "작성일"
        ]

    )

    ensure_excel_file(

        DESIGN_PROJECT_ITEM_DB_PATH,

        "설계공사상세",

        [
            "공사ID",
            "항목명",
            "규격",
            "수량",
            "단위",
            "적용단가",
            "금액"
        ]

    )


def repair_broken_xlsx_bytes(file_bytes):

    # 오래된 실무 엑셀(특히 여러 번 다른 파일을 참조하며
    # 편집된 산출내역서)에는 손상된 definedNames(#REF!,
    # #N/A 등이 수천 개 쌓인 경우)가 많아서 openpyxl/pandas가
    # 아예 못 여는 경우가 잦다. definedNames 섹션을 통째로
    # 제거하면 시트 내용 자체는 멀쩡하게 읽을 수 있다.

    import zipfile
    import re as _re

    src = io.BytesIO(file_bytes)

    out_buf = io.BytesIO()

    with zipfile.ZipFile(src, "r") as zin:

        with zipfile.ZipFile(

            out_buf,
            "w",
            zipfile.ZIP_DEFLATED

        ) as zout:

            for item in zin.namelist():

                data = zin.read(item)

                if item == "xl/workbook.xml":

                    text = data.decode(
                        "utf-8",
                        errors="ignore"
                    )

                    text = _re.sub(

                        r"<definedNames>.*?</definedNames>",

                        "",

                        text,

                        flags=_re.DOTALL

                    )

                    data = text.encode("utf-8")

                zout.writestr(item, data)

    out_buf.seek(0)

    return out_buf.getvalue()


def safe_read_excel_any(file_bytes, sheet_name=None):

    # 정상 파일이면 그냥 읽고, 깨진 파일이면 자동 복구 후
    # 다시 읽는다.

    try:

        return pd.read_excel(

            io.BytesIO(file_bytes),

            sheet_name=sheet_name,

            header=None

        )

    except Exception:

        fixed_bytes = repair_broken_xlsx_bytes(
            file_bytes
        )

        return pd.read_excel(

            io.BytesIO(fixed_bytes),

            sheet_name=sheet_name,

            header=None

        )


def parse_unitprice_sheet(df):

    # "일위대가표" 시트 하나를 파싱해서, 항목(호표)별로
    # [항목명, 단위, 노무비, 재료비, 기계경비, 외주경비, 합계,
    #  상세내역(종별/규격/수량/단위/단가/금액 전체 구조)] 를
    # 뽑아낸다. 상세내역을 완전한 구조로 저장해두면 나중에
    # 원본과 똑같은 엑셀 양식으로 다시 뽑을 수 있다.
    #
    # 열 구조(고정): 1=종별 2=규격 3=수량 4=단위
    # 5=노무비단가 6=노무비금액 7=재료비단가 8=재료비금액
    # 9=기계경비단가 10=기계경비금액 11=외주경비단가
    # 12=외주경비금액 13=비고

    items = []

    n = len(df)

    i = 0

    while i < n:

        row = df.iloc[i]

        row_vals = [

            str(v) if pd.notna(v) else ""

            for v in row.tolist()

        ]

        if any(

            "항목번호" in v

            for v in row_vals

        ):

            gongjong_row = df.iloc[i + 1] if i + 1 < n else None

            unit_row = df.iloc[i + 2] if i + 2 < n else None

            total_row = df.iloc[i + 3] if i + 3 < n else None

            if (

                gongjong_row is None

                or unit_row is None

                or total_row is None

            ):

                i += 1

                continue

            gongjong_vals = [

                v for v in gongjong_row.tolist()

                if pd.notna(v)

            ]

            unit_vals = [

                v for v in unit_row.tolist()

                if pd.notna(v)

            ]

            total_vals = [

                v for v in total_row.tolist()

                if pd.notna(v)

            ]

            if (

                len(gongjong_vals) < 2

                or len(unit_vals) < 2

                or len(total_vals) < 4

            ):

                i += 1

                continue

            item_name = str(
                gongjong_vals[-1]
            ).strip()

            unit = str(
                unit_vals[1]
            ).strip()

            def _num(v):

                try:

                    return float(v)

                except (TypeError, ValueError):

                    return 0.0

            labor = _num(
                total_vals[0]
            )

            material = _num(
                total_vals[1]

                if len(total_vals) > 1

                else 0

            )

            machine = _num(
                total_vals[2]

                if len(total_vals) > 2

                else 0

            )

            outsourcing = _num(
                total_vals[3]

                if len(total_vals) > 3

                else 0

            )

            total = (

                total_vals[-1]

                if len(total_vals) > 4

                else labor + material + machine + outsourcing

            )

            total = _num(total)

            # 상세내역: 헤더(종별/규격/수량/단위...) 2줄을 건너
            # 뛰고, "합   계" 행이나 다음 "일위대가표" 제목이
            # 나올 때까지 상세 라인을 하나씩 읽는다.

            detail_lines = []

            current_section = ""

            j = i + 6

            while j < n:

                drow = df.iloc[j]

                dvals = [

                    str(v) if pd.notna(v) else ""

                    for v in drow.tolist()

                ]

                col1 = dvals[1] if len(dvals) > 1 else ""

                if "합   계" in col1:

                    break

                if "일   위   대   가   표" in "".join(dvals):

                    break

                if (

                    "노 무 비" in col1

                    or "재 료 비" in col1

                    or "기계 경비" in col1

                    or "외주 경비" in col1

                ):

                    current_section = col1.strip()

                    j += 1

                    continue

                jongbyeol = drow.iloc[1] if pd.notna(drow.iloc[1]) else ""

                if str(jongbyeol).strip():

                    detail_lines.append(

                        {
                            "구분": current_section,
                            "종별": str(jongbyeol).strip(),
                            "규격": (

                                str(drow.iloc[2]).strip()

                                if pd.notna(drow.iloc[2])

                                else ""

                            ),
                            "수량": (

                                float(drow.iloc[3])

                                if pd.notna(drow.iloc[3])

                                else 0

                            ),
                            "단위": (

                                str(drow.iloc[4]).strip()

                                if pd.notna(drow.iloc[4])

                                else ""

                            ),
                            "노무비단가": (

                                float(drow.iloc[5])

                                if pd.notna(drow.iloc[5])

                                else 0

                            ),
                            "노무비금액": (

                                float(drow.iloc[6])

                                if pd.notna(drow.iloc[6])

                                else 0

                            ),
                            "재료비단가": (

                                float(drow.iloc[7])

                                if len(drow) > 7 and pd.notna(drow.iloc[7])

                                else 0

                            ),
                            "재료비금액": (

                                float(drow.iloc[8])

                                if len(drow) > 8 and pd.notna(drow.iloc[8])

                                else 0

                            ),
                            "기계경비단가": (

                                float(drow.iloc[9])

                                if len(drow) > 9 and pd.notna(drow.iloc[9])

                                else 0

                            ),
                            "기계경비금액": (

                                float(drow.iloc[10])

                                if len(drow) > 10 and pd.notna(drow.iloc[10])

                                else 0

                            ),
                            "외주경비단가": (

                                float(drow.iloc[11])

                                if len(drow) > 11 and pd.notna(drow.iloc[11])

                                else 0

                            ),
                            "외주경비금액": (

                                float(drow.iloc[12])

                                if len(drow) > 12 and pd.notna(drow.iloc[12])

                                else 0

                            )
                        }

                    )

                j += 1

                if j - i > 30:

                    break

            # 원본 항목명(공종)에 관경 등 규격 구분이 안 붙어
            # 있는 경우가 있어서(예: "배관 해체"가 32A용/50A용
            # 두 개인데 이름이 똑같음), 상세내역 첫 줄의 규격
            # 문자열에서 "(32A)" 같은 관경 표기를 찾아 항목명에
            # 자동으로 붙여준다. 그래야 라이브러리에서 서로
            # 다른 항목으로 구분되어 등록된다.

            spec_suffix = ""

            for d in detail_lines:

                m = re.search(

                    r"\(([0-9]+\s*A)\)",

                    d.get("규격", "")

                )

                if m:

                    spec_suffix = f" ({m.group(1).replace(' ', '')})"

                    break

            if spec_suffix and spec_suffix not in item_name:

                item_name = item_name + spec_suffix

            items.append(

                {
                    "항목명": item_name,
                    "규격": "",
                    "단위": unit,
                    "노무비단가": labor,
                    "재료비단가": material,
                    "기계경비단가": machine,
                    "외주경비단가": outsourcing,
                    "합계단가": total,
                    "상세내역": detail_lines
                }

            )

            i = j

        else:

            i += 1

    return items




# 기계설비 자주쓰는 일위대가 기본 탑재 데이터
# (실제 산출내역서 2건에서 뽑아온 값. 신입사원도 처음부터
#  바로 검색해서 쓸 수 있도록, 앱 처음 실행 시 라이브러리가
#  비어있으면 자동으로 채워 넣는다.)

DEFAULT_UNITPRICE_ITEMS = [{'항목명': '터파기', '규격': '', '단위': '㎥', '노무비단가': 16064.0, '재료비단가': 626.0, '기계경비단가': 823.0, '외주경비단가': 0.0, '합계단가': 17513.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '터파기', '규격': '보통토사', '수량': 0.05555555555555556, '단위': '인', '노무비단가': 226122.0, '노무비금액': 12562.0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '터파기 및 상차', '규격': '', '수량': 0.05934718100890207, '단위': '인', '노무비단가': 59020.0, '노무비금액': 3502.0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '2. 재 료 비', '종별': '터파기 및 상차', '규격': '', '수량': 0.05934718100890207, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 10559.0, '재료비금액': 626.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '3. 기계 경비', '종별': '굴착기', '규격': '', '수량': 0.05934718100890207, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 13873.0, '기계경비금액': 823.0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '_붙임1_산출내역서__부곡_가__인라인펌프장_배수펌프_배관_재설치.xlsx'}, {'항목명': '되메우기', '규격': '', '단위': '㎥', '노무비단가': 11625.0, '재료비단가': 626.0, '기계경비단가': 619.0, '외주경비단가': 0.0, '합계단가': 12870.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '부설', '규격': '보통인부', '수량': 0.020000000000000004, '단위': '인', '노무비단가': 172068.0, '노무비금액': 3441.0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0.0}, {'구분': '1. 노 무 비', '종별': '기계부설', '규격': '', '수량': 0.03848003848003848, '단위': '인', '노무비단가': 59020.0, '노무비금액': 2271.0918710918713, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '다짐', '규격': '', '수량': 0.1631321370309951, '단위': '인', '노무비단가': 36248.0, '노무비금액': 5913.0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '2. 재 료 비', '종별': '기계부설', '규격': '', '수량': 0.03848003848003848, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 10559.0, '재료비금액': 406.31072631072635, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '2. 재 료 비', '종별': '다짐', '규격': '', '수량': 0.1631321370309951, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 1352.0, '재료비금액': 220.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '3. 기계 경비', '종별': '기계부설', '규격': '', '수량': 0.03848003848003848, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 13873.0, '기계경비금액': 533.8335738335738, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '3. 기계 경비', '종별': '다짐', '규격': '', '수량': 0.1631321370309951, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 525.0, '기계경비금액': 85.0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '_붙임1_산출내역서__부곡_가__인라인펌프장_배수펌프_배관_재설치.xlsx'}, {'항목명': '장비대기시간', '규격': '', '단위': '시간당', '노무비단가': 59020.0, '재료비단가': 10559.0, '기계경비단가': 13873.0, '외주경비단가': 0.0, '합계단가': 83452.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '건설기계운전사', '규격': '', '수량': 1.0, '단위': '인', '노무비단가': 59020.2, '노무비금액': 59020.2, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '2. 재 료 비', '종별': '경유', '규격': '', '수량': 5.0, '단위': 'L', '노무비단가': 0, '노무비금액': 0, '재료비단가': 1745.455, '재료비금액': 8727.275, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '2. 재 료 비', '종별': '잡품', '규격': '경유값의 21%', '수량': 21.0, '단위': '%', '노무비단가': 0, '노무비금액': 0, '재료비단가': 0, '재료비금액': 1831.72775, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '3. 기계 경비', '종별': '굴착기', '규격': '0.2 M3', '수량': 0.2085, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 66538.0, '기계경비금액': 13873.172999999999, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '_붙임1_산출내역서__부곡_가__인라인펌프장_배수펌프_배관_재설치.xlsx'}, {'항목명': '엘보 설치', '규격': '', '단위': '개당', '노무비단가': 27700.0, '재료비단가': 8536.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 36236.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '엘보 설치', '규격': '배관공', '수량': 0.11173999999999999, '단위': '인', '노무비단가': 247897.0, '노무비금액': 27700.0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '2. 재 료 비', '종별': '엘보', '규격': '', '수량': 1.0, '단위': '개', '노무비단가': 0, '노무비금액': 0, '재료비단가': 8536.0, '재료비금액': 8536.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '_붙임1_산출내역서__부곡_가__인라인펌프장_배수펌프_배관_재설치.xlsx'}, {'항목명': '배관 설치', '규격': '', '단위': '일당', '노무비단가': 915759.0, '재료비단가': 0.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 915759.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '용접 배관', '규격': '배관공', '수량': 3.0, '단위': '인', '노무비단가': 247897.0, '노무비금액': 743691.0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '_붙임1_산출내역서__부곡_가__인라인펌프장_배수펌프_배관_재설치.xlsx'}, {'항목명': '모래 채움', '규격': '', '단위': '식', '노무비단가': 0.0, '재료비단가': 177000.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 177000.0, '상세내역': [{'구분': '1. 재 료 비', '종별': '모래 자갈 채움', '규격': '', '수량': 3.0, '단위': '㎥', '노무비단가': 0, '노무비금액': 0, '재료비단가': 59000.0, '재료비금액': 177000.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '_붙임1_산출내역서__부곡_가__인라인펌프장_배수펌프_배관_재설치.xlsx'}, {'항목명': '배관 및 보온재', '규격': '', '단위': '식', '노무비단가': 0.0, '재료비단가': 606200.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 606200.0, '상세내역': [{'구분': '1. 재 료 비', '종별': '배관 설치', '규격': '', '수량': 3.0, '단위': '본', '노무비단가': 0, '노무비금액': 0, '재료비단가': 180000.0, '재료비금액': 540000.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 재 료 비', '종별': '보온재', '규격': '', '수량': 10.0, '단위': '매', '노무비단가': 0, '노무비금액': 0, '재료비단가': 6620.0, '재료비금액': 66200.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '_붙임1_산출내역서__부곡_가__인라인펌프장_배수펌프_배관_재설치.xlsx'}, {'항목명': '배관해체', '규격': '', '단위': '일당', '노무비단가': 272184.0, '재료비단가': 0.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 272184.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '배관공', '규격': '', '수량': 0.8119999999999999, '단위': '인', '노무비단가': 247897.0, '노무비금액': 201292.0, '재료비단가': 0.0, '재료비금액': 0.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '보통인부', '규격': '', '수량': 0.41200000000000003, '단위': '인', '노무비단가': 172068.0, '노무비금액': 70892.0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '_붙임1_산출내역서__부곡_가__인라인펌프장_배수펌프_배관_재설치.xlsx'}, {'항목명': '배관보온', '규격': '', '단위': '일당', '노무비단가': 606092.0, '재료비단가': 0.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 606092.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '보온공', '규격': '', '수량': 2.0, '단위': '인', '노무비단가': 217012.0, '노무비금액': 434024.0, '재료비단가': 0.0, '재료비금액': 0.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '보통인부', '규격': '', '수량': 1.0, '단위': '인', '노무비단가': 172068.0, '노무비금액': 172068.0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '_붙임1_산출내역서__부곡_가__인라인펌프장_배수펌프_배관_재설치.xlsx'}, {'항목명': '배관 해체 (32A)', '규격': '', '단위': 'M', '노무비단가': 6947.0, '재료비단가': 139.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 7086.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '배관 철거', '규격': '배관공(32A)', '수량': 0.019, '단위': '인', '노무비단가': 247897.0, '노무비금액': 4710.1, '재료비단가': 0, '재료비금액': 0.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '2. 공구손료', '규격': '', '수량': 0, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '노무비의', '규격': '0.02', '수량': 2.0, '단위': '%', '노무비단가': 0, '노무비금액': 0, '재료비단가': 139.0, '재료비금액': 139.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '산출내역서__밀양_정__방류수_샘플링_배관_교체.xlsx'}, {'항목명': '배관 해체 (50A)', '규격': '', '단위': 'M', '노무비단가': 9791.0, '재료비단가': 196.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 9987.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '배관 철거', '규격': '배관공(50A)', '수량': 0.027, '단위': '인', '노무비단가': 247897.0, '노무비금액': 6693.3, '재료비단가': 0, '재료비금액': 0.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '2. 공구손료', '규격': '', '수량': 0, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '노무비의', '규격': '0.02', '수량': 2.0, '단위': '%', '노무비단가': 0, '노무비금액': 0, '재료비단가': 196.0, '재료비금액': 196.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '산출내역서__밀양_정__방류수_샘플링_배관_교체.xlsx'}, {'항목명': '배관 설치 (32A)', '규격': '', '단위': '일당', '노무비단가': 1226298.0, '재료비단가': 24526.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 1250824.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '강관 설치', '규격': '배관공(32A) / 0.67배 반영', '수량': 2.0100000000000002, '단위': '인', '노무비단가': 247897.0, '노무비금액': 498273.0, '재료비단가': 0, '재료비금액': 0.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '2. 공구손료', '규격': '', '수량': 0, '단위': '', '노무비단가': 0, '노무비금액': 0, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '노무비의', '규격': '0.02', '수량': 2.0, '단위': '%', '노무비단가': 0, '노무비금액': 0, '재료비단가': 24526.0, '재료비금액': 24526.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '산출내역서__밀양_정__방류수_샘플링_배관_교체.xlsx'}, {'항목명': '배관 보온', '규격': '', '단위': '일당', '노무비단가': 454569.0, '재료비단가': 0.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': 454569.0, '상세내역': [{'구분': '1. 노 무 비', '종별': '보온 두께 25mm이하', '규격': '보온공(32A 이하) / 0.32배 반영', '수량': 0.64, '단위': '인', '노무비단가': 217012.0, '노무비금액': 138887.7, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}, {'구분': '1. 노 무 비', '종별': '(발포 폴리에틸렌)', '규격': '보통인부(32A 이하) / 0.32배 반영', '수량': 0.32, '단위': '인', '노무비단가': 172068.0, '노무비금액': 55061.799999999996, '재료비단가': 0, '재료비금액': 0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '산출내역서__밀양_정__방류수_샘플링_배관_교체.xlsx'}, {'항목명': '발생품 공제(고비철)', '규격': '', '단위': '식', '노무비단가': 0.0, '재료비단가': -1700.0, '기계경비단가': 0.0, '외주경비단가': 0.0, '합계단가': -1700.0, '상세내역': [{'구분': '1. 재 료 비', '종별': '발생품 고비철', '규격': '스테인리스', '수량': 1.0, '단위': 'kg', '노무비단가': 0, '노무비금액': 0.0, '재료비단가': -1700.0, '재료비금액': -1700.0, '기계경비단가': 0, '기계경비금액': 0, '외주경비단가': 0, '외주경비금액': 0}], '출처공사': '산출내역서__밀양_정__방류수_샘플링_배관_교체.xlsx'}]



def find_similar_item_names(new_name, existing_names, threshold=0.82):

    # 완전히 똑같은 이름은 자동으로 걸러지지만("배관해체" 두
    # 번 등록 안 됨), "배관해체"와 "배관 해체 (32A)"처럼
    # 사람이 봐도 헷갈리는 비슷한 이름은 그냥 다른 항목으로
    # 등록되어 라이브러리가 지저분해질 수 있다. 등록 직전에
    # 미리 알려준다.

    import difflib

    similar = []

    for name in existing_names:

        if name == new_name:

            continue

        ratio = difflib.SequenceMatcher(

            None,
            new_name,
            name

        ).ratio()

        if ratio >= threshold:

            similar.append(
                (name, ratio)
            )

    return sorted(

        similar,

        key=lambda x: x[1],

        reverse=True

    )


def import_unitprice_items_to_library(items, source_name=""):

    added = 0

    existing = get_unitprice_library()

    existing_names = set(

        existing["항목명"].tolist()

        if not existing.empty

        else []

    )

    with get_lock(DESIGN_UNITPRICE_DB_PATH):

        wb = load_workbook(
            DESIGN_UNITPRICE_DB_PATH
        )

        ws = wb["일위대가라이브러리"]

        for item in items:

            if item["항목명"] in existing_names:

                continue

            ws.append(

                [
                    item["항목명"],
                    item.get("규격", ""),
                    item["단위"],
                    item["노무비단가"],
                    item["재료비단가"],
                    item["기계경비단가"],
                    item["외주경비단가"],
                    item["합계단가"],
                    0,
                    datetime.now().strftime("%Y-%m-%d"),
                    source_name,
                    json.dumps(

                        item.get("상세내역", []),

                        ensure_ascii=False

                    )
                ]

            )

            existing_names.add(
                item["항목명"]
            )

            added += 1

        wb.save(
            DESIGN_UNITPRICE_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()

    log_audit(

        "일위대가 라이브러리 추가",

        source_name,

        f"{added}건 추가"

    )

    return added


def seed_default_unitprice_library():

    # 예전엔 "라이브러리가 완전히 비어있을 때만" 기본 데이터를
    # 심었는데, 이미 사용자가 뭔가(예: 직접 업로드한 항목
    # 몇 개)를 넣어둔 상태면 "이미 있네" 하고 통째로 건너뛰어
    # 버려서, 기본 탑재 항목이 하나도 안 들어가는 문제가
    # 있었다. import_unitprice_items_to_library() 자체가
    # 이름별로 이미 있으면 건너뛰고 없으면만 추가하므로,
    # 매번 그냥 실행해도 안전하다(중복 추가 걱정 없음).

    import_unitprice_items_to_library(

        DEFAULT_UNITPRICE_ITEMS,

        source_name="기본 탑재 데이터"

    )


def get_unitprice_library():

    df = read_excel(

        DESIGN_UNITPRICE_DB_PATH,

        "일위대가라이브러리"

    )

    if df.empty:

        return df

    return df.sort_values(

        "사용횟수",

        ascending=False

    ).reset_index(drop=True)


def bump_unitprice_usage(item_name):

    with get_lock(DESIGN_UNITPRICE_DB_PATH):

        wb = load_workbook(
            DESIGN_UNITPRICE_DB_PATH
        )

        ws = wb["일위대가라이브러리"]

        for r in ws.iter_rows(min_row=2):

            if r[0].value == item_name:

                cur = r[8].value or 0

                r[8].value = cur + 1

                break

        wb.save(
            DESIGN_UNITPRICE_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()


def create_design_project(name, gongjong, site, author):

    project_id = f"P{datetime.now().strftime('%Y%m%d%H%M%S')}"

    safe_append_row(

        DESIGN_PROJECT_DB_PATH,

        "설계공사",

        [
            project_id,
            name,
            gongjong,
            site,
            author,
            datetime.now().strftime("%Y-%m-%d")
        ]

    )

    log_audit(

        "설계공사 생성",

        name,

        f"공종={gongjong}, 사업장={site}"

    )

    return project_id


def add_item_to_project(

    project_id,
    item_name,
    spec,
    qty,
    unit,
    unit_price

):

    amount = qty * unit_price

    safe_append_row(

        DESIGN_PROJECT_ITEM_DB_PATH,

        "설계공사상세",

        [
            project_id,
            item_name,
            spec,
            qty,
            unit,
            unit_price,
            amount
        ]

    )

    bump_unitprice_usage(
        item_name
    )


def get_project_items(project_id):

    df = read_excel(

        DESIGN_PROJECT_ITEM_DB_PATH,

        "설계공사상세"

    )

    if df.empty:

        return df

    return df[

        df["공사ID"] == project_id

    ]


def get_guideline_items(gongjong):

    df = read_excel(

        DESIGN_GUIDELINE_DB_PATH,

        "공종가이드라인"

    )

    if df.empty:

        return []

    matched = df[

        df["공사유형"] == gongjong

    ]

    return matched["추천항목명"].tolist()


def set_guideline_items(gongjong, item_names):

    with get_lock(DESIGN_GUIDELINE_DB_PATH):

        wb = load_workbook(
            DESIGN_GUIDELINE_DB_PATH
        )

        ws = wb["공종가이드라인"]

        # 기존 이 공종의 가이드라인은 지우고 새로 쓴다

        rows_to_delete = [

            r[0].row

            for r in ws.iter_rows(min_row=2)

            if r[0].value == gongjong

        ]

        for row_num in sorted(

            rows_to_delete,

            reverse=True

        ):

            ws.delete_rows(row_num)

        for name in item_names:

            ws.append(
                [gongjong, name]
            )

        wb.save(
            DESIGN_GUIDELINE_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()


def build_unitprice_excel_bytes(items_with_detail):

    # 원본 산출내역서의 "일위대가표" 시트와 똑같은 양식으로
    # 다시 만든다. items_with_detail: 라이브러리 DataFrame 행
    # (상세내역JSON 컬럼 포함) 리스트.
    #
    # 원본 파일을 직접 열어서 확인한 실제 구조:
    # - 종별(B)/규격(C)/수량(D)/단위(E)는 세로로 2행 병합
    # - 노무비/재료비/기계경비/외주경비는 각각 2열씩 병합
    #   (F:G=노무비, H:I=재료비, J:K=기계경비, L:M=외주경비)
    # - "라벨"은 병합영역의 앞쪽 열(F,H,J,L)에, "합계 숫자"는
    #   뒤쪽 열(G,I,K,M)에 들어간다(같은 열이 아님)
    # - 비고는 N열(종별헤더 줄 기준 세로 2행 병합)

    from openpyxl.styles import Font, Alignment, Border, Side

    wb = Workbook()

    ws = wb.active

    ws.title = "일위대가표"

    # 원본과 동일하게 가로방향 + 페이지맞춤(폭 1페이지)로
    # 인쇄되도록 설정한다. 이게 없으면 우측 열들이 인쇄/PDF
    # 변환 시 페이지 밖으로 잘려나간다.

    ws.page_setup.orientation = "landscape"

    ws.page_setup.fitToPage = True

    ws.page_setup.fitToWidth = 1

    ws.page_setup.fitToHeight = 0

    ws.sheet_properties.pageSetUpPr.fitToPage = True

    thin = Side(
        style="thin",
        color="000000"
    )

    box = Border(

        left=thin,
        right=thin,
        top=thin,
        bottom=thin

    )

    center = Alignment(

        horizontal="center",
        vertical="center"

    )

    bold = Font(
        bold=True
    )

    # 원본 파일 실측값 그대로 반영

    ws.column_dimensions["A"].width = 2.22

    ws.column_dimensions["B"].width = 19.11

    ws.column_dimensions["C"].width = 21.78

    ws.column_dimensions["D"].width = 11.22

    ws.column_dimensions["E"].width = 5.78

    ws.column_dimensions["F"].width = 11.78

    for col_letter in ["G", "H", "I", "J", "K", "L", "M", "N"]:

        ws.column_dimensions[col_letter].width = 8.89

    def _box_range(r1, c1, r2, c2):

        for rr in range(r1, r2 + 1):

            for cc in range(c1, c2 + 1):

                ws.cell(rr, cc).border = box

    row = 1

    for item_row in items_with_detail:

        item_name = item_row["항목명"]

        unit = item_row["단위"]

        labor = item_row["노무비단가"]

        material = item_row["재료비단가"]

        machine = item_row["기계경비단가"]

        outsourcing = item_row["외주경비단가"]

        total = item_row["합계단가"]

        try:

            details = json.loads(

                item_row.get("상세내역JSON", "[]")

                if isinstance(item_row, dict)

                else item_row["상세내역JSON"]

            )

        except Exception:

            details = []

        # 1행: 제목

        ws.cell(
            row, 2, "일   위   대   가   표"
        ).font = Font(
            bold=True,
            size=14
        )

        ws.merge_cells(

            start_row=row, start_column=2,

            end_row=row, end_column=14

        )

        ws.cell(row, 2).alignment = center

        ws.row_dimensions[row].height = 27.0

        title_row = row

        row += 2

        # 항목번호 / 공종 / 단위

        ws.cell(row, 2, "   항목번호 :")

        ws.cell(row, 3, item_name)

        ws.row_dimensions[row].height = 27.0

        row += 1

        ws.cell(row, 2, "   공    종 : ")

        ws.cell(row, 3, item_name)

        ws.row_dimensions[row].height = 27.0

        row += 1

        unit_row = row

        ws.cell(row, 2, "   단    위 :")

        ws.cell(row, 3, unit)

        # 노무비/재료비/기계경비/외주경비/합계 라벨
        # (각 2열 병합: F:G, H:I, J:K, L:M / 합계는 N열 단독)

        label_spans = [

            (6, 7, "노 무 비"),
            (8, 9, "재 료 비"),
            (10, 11, "기계경비"),
            (12, 13, "외주경비")

        ]

        for c1, c2, label in label_spans:

            ws.merge_cells(

                start_row=row, start_column=c1,

                end_row=row, end_column=c2

            )

            ws.cell(row, c1, label).font = bold

            ws.cell(row, c1).alignment = center

        ws.cell(row, 14, "합   계").font = bold

        ws.cell(row, 14).alignment = center

        _box_range(row, 6, row, 14)

        ws.row_dimensions[row].height = 27.0

        row += 1

        # 합계 숫자 (라벨의 "뒤쪽 열"에 위치: G,I,K,M / 총합계 N)

        val_positions = [

            (7, labor),
            (9, material),
            (11, machine),
            (13, outsourcing)

        ]

        for col_pos, v in val_positions:

            ws.cell(row, col_pos, v)

        ws.cell(row, 14, total)

        _box_range(row, 6, row, 14)

        ws.row_dimensions[row].height = 27.0

        total_val_row = row

        row += 1

        # 종별/규격/수량/단위(세로 2행 병합) +
        # 노무비/재료비/기계경비/외주경비(가로 2열 병합) + 비고

        header_row1 = row

        header_row2 = row + 1

        for col_idx, label in [

            (2, "종   별"),
            (3, "규   격"),
            (4, "수   량"),
            (5, "단  위")

        ]:

            ws.merge_cells(

                start_row=header_row1, start_column=col_idx,

                end_row=header_row2, end_column=col_idx

            )

            ws.cell(header_row1, col_idx, label).font = bold

            ws.cell(header_row1, col_idx).alignment = center

        for c1, c2, label in label_spans:

            ws.merge_cells(

                start_row=header_row1, start_column=c1,

                end_row=header_row1, end_column=c2

            )

            ws.cell(header_row1, c1, label).font = bold

            ws.cell(header_row1, c1).alignment = center

        ws.merge_cells(

            start_row=header_row1, start_column=14,

            end_row=header_row2, end_column=14

        )

        ws.cell(header_row1, 14, "비   고").font = bold

        ws.cell(header_row1, 14).alignment = center

        _box_range(header_row1, 2, header_row2, 14)

        ws.row_dimensions[header_row1].height = 20.1

        # 단가/금액 서브헤더 (2번째 줄)

        sub_cols = [6, 7, 8, 9, 10, 11, 12, 13]

        sub_labels = [

            "단 가", "금 액", "단 가", "금 액",
            "단 가", "금 액", "단 가", "금 액"

        ]

        for c, label in zip(sub_cols, sub_labels):

            ws.cell(header_row2, c, label)

            ws.cell(header_row2, c).alignment = center

        ws.row_dimensions[header_row2].height = 20.1

        row = header_row2 + 1

        # 상세내역

        current_section = None

        for d in details:

            if d.get("구분") != current_section:

                current_section = d.get("구분")

                ws.cell(
                    row,
                    2,
                    current_section
                ).font = bold

                ws.row_dimensions[row].height = 27.95

                row += 1

            ws.cell(row, 2, d.get("종별", ""))

            ws.cell(row, 3, d.get("규격", ""))

            ws.cell(row, 4, d.get("수량", ""))

            ws.cell(row, 5, d.get("단위", ""))

            ws.cell(row, 6, d.get("노무비단가") or "")

            ws.cell(row, 7, d.get("노무비금액") or "")

            ws.cell(row, 8, d.get("재료비단가") or "")

            ws.cell(row, 9, d.get("재료비금액") or "")

            ws.cell(row, 10, d.get("기계경비단가") or "")

            ws.cell(row, 11, d.get("기계경비금액") or "")

            ws.cell(row, 12, d.get("외주경비단가") or "")

            ws.cell(row, 13, d.get("외주경비금액") or "")

            _box_range(row, 2, row, 14)

            ws.row_dimensions[row].height = 27.95

            row += 1

        # 합계 행

        ws.cell(row, 2, "합   계").font = bold

        ws.cell(row, 7, labor)

        ws.cell(row, 9, material)

        ws.cell(row, 11, machine)

        ws.cell(row, 13, outsourcing)

        _box_range(row, 2, row, 14)

        ws.row_dimensions[row].height = 27.95

        row += 3

    buffer = io.BytesIO()

    wb.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()


def get_logo_base64_html(

    max_height_px=40

):

    # 사이드바·홈 히어로 배너처럼 화면(HTML) 안에 로고를 넣는
    # 자리에서 쓴다. Streamlit의 markdown은 로컬 파일 경로를
    # <img src="...">로 바로 못 읽으므로 base64로 인라인 삽입한다.
    # 로고 파일이 없으면 빈 문자열을 돌려주고, 호출부에서
    # 기존 이모지/아이콘으로 대체하면 된다.

    logo_path = find_logo_file()

    if not logo_path:

        return ""

    try:

        with open(

            logo_path,

            "rb"

        ) as f:

            b64 = base64.b64encode(

                f.read()

            ).decode("utf-8")

        ext = logo_path.rsplit(

            ".",
            1

        )[-1].lower()

        mime = (

            "image/png"

            if ext == "png"

            else "image/jpeg"

        )

        return (

            f'<img src="data:{mime};base64,{b64}" '
            f'style="max-height:{max_height_px}px; '
            f'width:auto;" />'

        )

    except Exception:

        return ""


def find_mascot_file(name="mascot"):

    for pattern in (
        f"{name}.png",
        f"{name}.jpg",
        f"[{name[0].upper()}{name[0].lower()}]{name[1:]}.png"
    ):

        matches = glob.glob(pattern)

        if matches:

            return matches[0]

    return None


def get_mascot_base64_html(

    max_height_px=90,
    name="mascot"

):

    mascot_path = find_mascot_file(name)

    if not mascot_path:

        return ""

    try:

        with open(

            mascot_path,

            "rb"

        ) as f:

            b64 = base64.b64encode(

                f.read()

            ).decode("utf-8")

        return (

            f'<img src="data:image/png;base64,{b64}" '
            f'style="max-height:{max_height_px}px; '
            f'width:auto; filter:drop-shadow(0 6px 12px '
            f'rgba(0,0,0,0.25));" />'

        )

    except Exception:

        return ""


def build_report_qr_image_bytes(
    pump,
    all_pumps
):

    # 종이로 인쇄한 보고서에서도 QR을 찍으면 바로 그 설비의
    # 실시간 화면으로 연결되도록 표지에 작은 QR을 넣는다.
    # (QR 포털의 QR 생성 로직과 동일한 방식)

    equip_no = (

        all_pumps.index(pump) + 1

        if pump in all_pumps

        else 0

    )

    app_base_url = get_app_base_url()

    qr_target_url = (

        f"{app_base_url}/?page=QR&equip="

        +
        urllib.parse.quote(
            pump["equip"]
        )

    )

    qr_img = qrcode.make(
        qr_target_url
    )

    buf = io.BytesIO()

    qr_img.save(
        buf,
        format="PNG"
    )

    buf.seek(0)

    return buf


def check_data_sufficiency(
    result,
    month_history
):

    # 신규 설비이거나 실측 이력이 적으면 보고서 여러 섹션이
    # 예시데이터/이력없음으로 휑하게 보일 수 있다. 이런 경우
    # 맨 앞에 안내 문구를 넣어, 보는 사람이 "왜 이렇게
    # 비어있지?"라고 오해하지 않도록 한다.

    if result.get("실측이력있음") and not month_history.empty:

        return None

    return (

        "※ 본 설비는 아직 축적된 진단 이력이 충분하지 않아 "
        "일부 항목(추세 그래프 등)이 실측이 아닌 예시 데이터로 "
        "표시됩니다. 정밀진단을 반복해서 저장할수록 이후 "
        "보고서의 신뢰도가 높아집니다."

    )


def _clean_cell_value(
    val,
    default=""
):

    # 엑셀에서 읽어온 값이 NaN(빈칸)이면 워드 표에
    # "nan"이라고 그대로 찍히던 문제를 막는다.

    if val is None:

        return default

    try:

        if pd.isna(val):

            return default

    except (TypeError, ValueError):

        pass

    return val


def _shade_cell(
    cell,
    color_hex
):

    tc_pr = cell._tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")

    shd.set(qn("w:fill"), color_hex)

    tc_pr.append(shd)


def _set_korean_font(
    doc
):

    style = doc.styles["Normal"]

    style.font.name = "맑은 고딕"

    style.font.size = Pt(10.5)

    r_pr = style.element.get_or_add_rPr()

    r_fonts = r_pr.find(qn("w:rFonts"))

    if r_fonts is None:

        r_fonts = OxmlElement("w:rFonts")

        r_pr.append(r_fonts)

    r_fonts.set(qn("w:eastAsia"), "맑은 고딕")

    # 줄간격·문단간격을 공공기관 보고서 스타일(빽빽하게,
    # 문단 사이 여백 최소화)에 맞춘다. Word 기본 스타일은
    # 문단 뒤 간격이 8pt씩 붙어서 원본보다 훨씬 헐렁하게
    # 보이는 원인이 된다.

    style.paragraph_format.line_spacing = 1.15

    style.paragraph_format.space_before = Pt(0)

    style.paragraph_format.space_after = Pt(4)

    for _level in range(1, 4):

        try:

            _heading_style = doc.styles[
                f"Heading {_level}"
            ]

            _heading_style.font.name = "맑은 고딕"

            _heading_style.font.color.rgb = RGBColor(
                0x0F, 0x35, 0x52
            )

            _h_rpr = _heading_style.element.get_or_add_rPr()

            _h_fonts = _h_rpr.find(qn("w:rFonts"))

            if _h_fonts is None:

                _h_fonts = OxmlElement("w:rFonts")

                _h_rpr.append(_h_fonts)

            _h_fonts.set(qn("w:eastAsia"), "맑은 고딕")

            _heading_style.paragraph_format.space_before = (
                Pt(10)
            )

            _heading_style.paragraph_format.space_after = (
                Pt(4)
            )

        except Exception:

            continue


def _add_styled_table(
    doc,
    headers,
    rows
):

    table = doc.add_table(
        rows=1,
        cols=len(headers)
    )

    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells

    for j, h in enumerate(headers):

        hdr_cells[j].text = str(h)

        _shade_cell(
            hdr_cells[j],
            BRAND_HEADER_COLOR
        )

        for p in hdr_cells[j].paragraphs:

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in p.runs:

                run.font.color.rgb = RGBColor(
                    0xFF, 0xFF, 0xFF
                )

                run.font.bold = True

    for i, row in enumerate(rows):

        cells = table.add_row().cells

        for j, val in enumerate(row):

            cells[j].text = str(val)

            if i % 2 == 1:

                _shade_cell(
                    cells[j],
                    BRAND_ZEBRA_COLOR
                )

    return table


def _add_fig_to_doc(
    doc,
    fig,
    width_inches=6.0,
    caption=None
):

    buf = io.BytesIO()

    fig.savefig(

        buf,

        format="png",

        dpi=150,

        bbox_inches="tight"

    )

    plt.close(fig)

    buf.seek(0)

    doc.add_picture(
        buf,
        width=Inches(width_inches)
    )

    if caption:

        cap = doc.add_paragraph()

        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = cap.add_run(
            caption
        )

        run.italic = True

        run.font.size = Pt(9)

        run.font.color.rgb = RGBColor(
            0x66, 0x66, 0x66
        )


def add_page_number_field(
    paragraph
):

    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")

    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")

    instr.set(qn("xml:space"), "preserve")

    instr.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")

    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)

    run._r.append(instr)

    run._r.append(fld_end)


def add_report_header_footer(
    doc,
    doc_no
):

    section = doc.sections[0]

    header_p = section.header.paragraphs[0]

    header_p.text = (

        "K-water tech 설비관리 통합 플랫폼"

    )

    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in header_p.runs:

        run.font.size = Pt(8)

        run.font.color.rgb = RGBColor(
            0x66, 0x66, 0x66
        )

    footer_p = section.footer.paragraphs[0]

    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_run = footer_p.add_run(
        f"{doc_no}   ·   "
    )

    footer_run.font.size = Pt(8)

    footer_run.font.color.rgb = RGBColor(
        0x66, 0x66, 0x66
    )

    add_page_number_field(
        footer_p
    )


def add_table_of_contents(
    doc
):

    # Word 자체 목차 필드를 삽입한다. 문서를 열면 Word가
    # "필드 업데이트"를 묻거나, 설정에 따라 자동으로
    # 목차를 채워준다. (LibreOffice 등에서는 안내 문구만
    # 보일 수 있어, 우클릭 -> 필드 업데이트를 안내한다.)

    doc.add_heading(
        "목차",
        level=1
    )

    paragraph = doc.add_paragraph()

    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")

    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")

    instr.set(qn("xml:space"), "preserve")

    instr.text = 'TOC \\o "1-2" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")

    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")

    fld_end.set(qn("w:fldCharType"), "end")

    r_elem = run._r

    r_elem.append(fld_begin)

    r_elem.append(instr)

    r_elem.append(fld_separate)

    r_elem.append(fld_end)

    hint_p = doc.add_paragraph()

    hint_run = hint_p.add_run(

        "※ 목차가 비어 보이면, 목차 위에서 마우스 오른쪽 버튼 → "
        "'필드 업데이트'를 눌러주세요 (Word에서 자동 채워집니다)."

    )

    hint_run.italic = True

    hint_run.font.size = Pt(8)

    hint_run.font.color.rgb = RGBColor(
        0x99, 0x99, 0x99
    )

    # Word가 문서를 열 때 필드를 자동으로 갱신하도록 설정.

    settings_elem = doc.settings.element

    update_fields = OxmlElement("w:updateFields")

    update_fields.set(qn("w:val"), "true")

    settings_elem.append(
        update_fields
    )


def get_month_latest_record(
    df_history,
    equip,
    month_label
):

    if (

        df_history.empty

        or
        "설비명" not in df_history.columns

    ):

        return None

    rows = df_history[

        (df_history["설비명"] == equip)

        &
        (

            df_history["점검일"].astype(str).str.startswith(
                month_label
            )

        )

    ]

    if rows.empty:

        return None

    return rows.sort_values("점검일").iloc[-1]


def get_previous_month_label(
    month_label
):

    year, month = [
        int(x)
        for x in month_label.split("-")
    ]

    if month == 1:

        return f"{year - 1}-12"

    return f"{year}-{month - 1:02d}"


def generate_executive_summary(
    pump,
    result,
    vib_values
):

    lines = []

    grade_text = get_grade_text(
        result["등급"]
    )

    lines.append(

        f"본 설비({pump['equip']})는 {result['점수']}점"
        f"({result['등급']}등급)으로 판정되어 "
        f"'{grade_text}' 상태입니다."

    )

    if len(vib_values) >= 2:

        diff = vib_values[-1] - vib_values[-2]

        if diff > 0.3:

            lines.append(

                f"진동값이 직전 시점 대비 {diff:+.1f} mm/s "
                f"상승하는 추세를 보이고 있어 지속적인 관찰이 필요합니다."

            )

        elif diff < -0.3:

            lines.append(

                f"진동값이 직전 시점 대비 {diff:+.1f} mm/s "
                f"개선되는 추세를 보이고 있습니다."

            )

        else:

            lines.append(

                "진동값은 직전 시점과 비교해 큰 변화 없이 "
                "안정적인 수준을 유지하고 있습니다."

            )

    if result["상태"] == "정비검토":

        lines.append(

            "현재 정비검토 등급으로 판정되어 "
            "조속한 정밀점검 및 정비 조치가 필요한 상태입니다."

        )

    elif result["상태"] == "관찰":

        lines.append(

            "현재 관찰 등급으로 판정되어 "
            "다음 정기점검 시 추이를 중점적으로 확인할 필요가 있습니다."

        )

    else:

        lines.append(

            "현재 정상 범위에서 안정적으로 운전되고 있습니다."

        )

    return " ".join(
        lines
    )


def generate_action_items(
    pump,
    result
):

    items = []

    _th = get_alert_thresholds()

    if result["등급"] in ["D", "E"]:

        items.append(

            "종합등급이 D등급 이하로 판정되어 "
            "조속한 정밀진단 및 정비계획 수립이 필요합니다."

        )

    if result["진동"] >= _th["vib_danger"]:

        items.append(

            f"진동이 주의 기준({_th['vib_danger']} mm/s)을 초과하여 "
            "베어링·축정렬 상태에 대한 정밀점검이 필요합니다."

        )

    elif result["진동"] >= _th["vib_watch"]:

        items.append(

            f"진동이 관찰 기준({_th['vib_watch']} mm/s)을 초과하여 "
            "다음 점검 시 원인 분석이 권장됩니다."

        )

    if result["효율"] < _th["eff_watch"]:

        items.append(

            f"효율이 {_th['eff_watch']}% 미만으로 저하되어 "
            "내부 마모(링간극·축슬리브 등) 점검이 권장됩니다."

        )

    if result["온도"] >= _th["temp_danger"]:

        items.append(

            f"온도가 주의 기준({_th['temp_danger']}°C)을 초과하여 "
            "냉각·윤활 상태 점검이 필요합니다."

        )

    remaining = result.get(
        "다음오버홀까지남은시간"
    )

    if remaining is not None and remaining <= 1000:

        items.append(

            f"다음 오버홀 예정시점까지 {remaining:,}시간 남아 "
            f"오버홀 일정을 사전에 계획할 것을 권장합니다."

        )

    if not items:

        items.append(

            "현재 특별한 조치가 필요한 사항은 없으며, "
            "정기 점검 주기를 유지하시기 바랍니다."

        )

    return items


def build_pump_monthly_report_docx(
    pump,
    result,
    month_label,
    df_history,
    all_pumps,
    use_ai_commentary=False
):

    doc = Document()

    _set_korean_font(
        doc
    )

    equip_no = (

        all_pumps.index(pump) + 1

        if pump in all_pumps

        else 0

    )

    doc_no = (

        f"KWT-{month_label.replace('-', '')}"
        f"-{equip_no:03d}"

    )

    add_report_header_footer(

        doc,

        doc_no

    )

    # ---------------- 표지 ----------------

    logo_path = find_logo_file()

    if logo_path:

        try:

            logo_p = doc.add_paragraph()

            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            logo_p.add_run().add_picture(

                logo_path,

                width=Inches(1.6)

            )

        except Exception:

            pass

    title = doc.add_heading(
        "월간 설비 진단 보고서",
        level=0
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_p = doc.add_paragraph()

    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle_p.add_run(

        f"{pump['equip']}  ·  {month_label}"

    )

    subtitle_run.font.size = Pt(16)

    subtitle_run.bold = True

    doc.add_paragraph()

    # QR코드 삽입 (인쇄본에서도 스캔하면 실시간 화면으로 연결)

    try:

        qr_buf = build_report_qr_image_bytes(

            pump,

            all_pumps

        )

        qr_p = doc.add_paragraph()

        qr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        qr_p.add_run().add_picture(

            qr_buf,

            width=Inches(1.2)

        )

        qr_cap = doc.add_paragraph()

        qr_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        qr_cap_run = qr_cap.add_run(

            "스캔하면 이 설비의 실시간 화면으로 연결됩니다"

        )

        qr_cap_run.italic = True

        qr_cap_run.font.size = Pt(8)

        qr_cap_run.font.color.rgb = RGBColor(
            0x99, 0x99, 0x99
        )

    except Exception:

        pass

    _add_styled_table(

        doc,

        ["항목", "내용"],

        [
            ["문서번호", doc_no],
            ["사업장", pump["site"]],
            ["설비명", pump["equip"]],
            ["보고월", month_label],
            [
                "작성자",
                st.session_state.get(
                    "user_name",
                    ""
                )
            ],
            [
                "작성일",
                datetime.now().strftime("%Y-%m-%d")
            ]
        ]

    )

    doc.add_page_break()

    # ---------------- 목차 ----------------

    add_table_of_contents(
        doc
    )

    doc.add_page_break()

    # ---------------- 0. 종합 소견 ----------------

    doc.add_heading(
        "종합 소견",
        level=1
    )

    month_history = pd.DataFrame()

    if (

        not df_history.empty

        and
        "설비명" in df_history.columns

    ):

        pump_hist = df_history[

            df_history["설비명"] == pump["equip"]

        ].copy()

        if not pump_hist.empty:

            month_history = pump_hist[

                pump_hist["점검일"].astype(str).str.startswith(
                    month_label
                )

            ]

    data_note = check_data_sufficiency(

        result,

        month_history

    )

    if data_note:

        note_p = doc.add_paragraph()

        note_run = note_p.add_run(
            data_note
        )

        note_run.italic = True

        note_run.font.size = Pt(9)

        note_run.font.color.rgb = RGBColor(
            0xB0, 0x60, 0x00
        )

        note_p.paragraph_format.space_after = Pt(8)

    vib_dates, vib_values, _ = get_vibration_trend_data(

        pump,

        result,

        df_history

    )

    summary_text = generate_executive_summary(

        pump,

        result,

        vib_values

    )

    ai_used = False

    if use_ai_commentary:

        _df_vib_full = read_excel(

            VIBRATION_DB_PATH,

            "진동측정이력"

        )

        _pred_for_report = predict_failure_date_regression(

            pump,

            _df_vib_full

        )

        _ai_text, _ai_error = generate_ai_commentary_for_equipment(

            pump,

            result,

            _pred_for_report

        )

        if _ai_text:

            summary_text = _ai_text

            ai_used = True

    summary_p = doc.add_paragraph(
        summary_text
    )

    summary_p.paragraph_format.space_after = Pt(12)

    if use_ai_commentary:

        ai_note_p = doc.add_paragraph()

        ai_note_run = ai_note_p.add_run(

            "🤖 AI(Claude)가 데이터를 보고 직접 작성한 "
            "종합소견입니다."

            if ai_used

            else "※ AI 종합소견 생성에 실패해 기본 요약으로 "
            "대체되었습니다."

        )

        ai_note_run.italic = True

        ai_note_run.font.size = Pt(8.5)

        ai_note_run.font.color.rgb = RGBColor(

            0x08, 0x7E, 0xA4

        )

    doc.add_page_break()

    # 그림/표 번호 매기기용 카운터
    # (사내 결재·회의에서 "표 2번 다시 보여주세요" 식으로
    # 참조할 수 있도록 순번을 붙인다)

    _report_counters = {"fig": 0, "table": 0}

    def _next_num(kind):

        _report_counters[kind] += 1

        return _report_counters[kind]

    def _add_table_caption(
        text
    ):

        num = _next_num("table")

        cap_p = doc.add_paragraph()

        cap_run = cap_p.add_run(

            f"표 {num}. {text}"

        )

        cap_run.bold = True

        cap_run.font.size = Pt(9.5)

        return num

    # ---------------- 1. 설비 스펙 ----------------

    doc.add_heading(
        "1. 설비 스펙",
        level=1
    )

    _add_table_caption(
        "설비 스펙"
    )

    _add_styled_table(

        doc,

        ["항목", "값"],

        [
            ["사업장", pump["site"]],
            ["설비명", pump["equip"]],
            ["제조사", pump["maker"]],
            ["모델명", pump["model"]],
            ["정격출력(HP)", pump["hp"]],
            ["정격양정(m)", pump["head"]],
            ["정격유량(m³/h)", pump["flow"]],
            ["회전수(RPM)", pump["rpm"]],
            ["준공일", pump["build_date"]],
            ["누적 운전시간(h)", f'{pump["op_hours"]:,}'],
            [
                "기준진동(mm/s)",
                pump.get("기준진동") or "공통기준 적용"
            ],
            [
                "기준효율(%)",
                pump.get("기준효율") or "공통기준 적용"
            ]
        ]

    )

    # ---------------- 2. 펌프 점수 및 현재 상태 ----------------

    doc.add_heading(
        "2. 펌프 점수 및 현재 상태",
        level=1
    )

    status_p = doc.add_paragraph()

    status_run = status_p.add_run(

        f"CBM Score {result['점수']}점 · "
        f"{result['등급']}등급 · "
        f"상태 : {result['상태']}"

    )

    status_run.bold = True

    status_run.font.size = Pt(13)

    if result.get("다음오버홀까지남은시간") is not None:

        doc.add_paragraph(

            f"다음 오버홀 예정까지 약 "
            f"{result['다음오버홀까지남은시간']:,}시간 남음 "
            f"(기준 {OVERHAUL_INTERVAL_HOURS:,}시간 주기)"

        )

    gauge_fig = build_score_gauge_fig(
        pump,
        result
    )

    _add_fig_to_doc(

        doc,

        gauge_fig,

        width_inches=6.0,

        caption=f"그림 {_next_num('fig')}. CBM Score 게이지 (구간별 등급 기준)"

    )

    # ---------------- 2-1. 전월 대비 변화 ----------------

    doc.add_heading(

        "전월 대비 변화",

        level=2

    )

    prev_month_label = get_previous_month_label(
        month_label
    )

    this_month_rec = get_month_latest_record(

        df_history,

        pump["equip"],

        month_label

    )

    prev_month_rec = get_month_latest_record(

        df_history,

        pump["equip"],

        prev_month_label

    )

    if this_month_rec is not None and prev_month_rec is not None:

        def _diff_row(
            label,
            col,
            unit
        ):

            cur = this_month_rec.get(col)

            prev = prev_month_rec.get(col)

            if pd.isna(cur) or pd.isna(prev):

                return [label, "-", "-", "-"]

            cur = float(cur)

            prev = float(prev)

            return [

                label,

                f"{prev:.1f}{unit}",

                f"{cur:.1f}{unit}",

                f"{cur - prev:+.1f}{unit}"

            ]

        _add_table_caption(
            "전월 대비 변화"
        )

        _add_styled_table(

            doc,

            ["지표", f"{prev_month_label}", f"{month_label}", "변화"],

            [
                _diff_row("종합점수", "종합점수", "점"),
                _diff_row("효율", "효율측정값(%)", "%"),
                _diff_row("진동", "진동측정값(mm/s)", "mm/s"),
                _diff_row("온도", "온도측정값(°C)", "°C")
            ]

        )

    else:

        doc.add_paragraph(

            f"{prev_month_label} 또는 {month_label}에 "
            "점검 이력이 없어 전월 대비 비교를 산출할 수 없습니다."

        )

    # ---------------- 3. 점검 이력 ----------------

    doc.add_heading(
        "3. 점검 이력",
        level=1
    )

    # (month_history는 종합소견 섹션에서 이미 계산해뒀다)

    if not month_history.empty:

        rows = []

        for _, r in month_history.iterrows():

            rows.append(

                [
                    _clean_cell_value(r.get("점검일")),
                    _clean_cell_value(r.get("점검자")),
                    _clean_cell_value(r.get("종합점수")),
                    _clean_cell_value(r.get("최종등급")),
                    _clean_cell_value(r.get("효율측정값(%)")),
                    _clean_cell_value(r.get("진동측정값(mm/s)")),
                    _clean_cell_value(r.get("온도측정값(°C)"))
                ]

            )

        _add_table_caption(
            f"{month_label} 점검 이력"
        )

        _add_styled_table(

            doc,

            [
                "점검일", "점검자", "종합점수", "등급",
                "효율(%)", "진동(mm/s)", "온도(°C)"
            ],

            rows

        )

    else:

        doc.add_paragraph(

            f"{month_label}에 저장된 점검 이력이 없습니다."

        )

    # ---------------- 4. 오버홀 내역 ----------------

    doc.add_heading(
        "4. 오버홀 내역",
        level=1
    )

    df_overhaul_all = read_excel(
        OVERHAUL_DB_PATH,
        "오버홀이력"
    )

    month_overhaul = pd.DataFrame()

    if (

        not df_overhaul_all.empty

        and
        "설비명" in df_overhaul_all.columns

    ):

        pump_overhaul = df_overhaul_all[

            df_overhaul_all["설비명"] == pump["equip"]

        ].copy()

        if not pump_overhaul.empty:

            month_overhaul = pump_overhaul[

                pump_overhaul["작업일자"].astype(str).str.startswith(
                    month_label
                )

            ]

    if not month_overhaul.empty:

        rows = []

        for _, r in month_overhaul.iterrows():

            rows.append(

                [
                    _clean_cell_value(r.get("작업일자")),
                    _clean_cell_value(r.get("공정단계")),
                    _clean_cell_value(r.get("작업자")),
                    str(_clean_cell_value(r.get("작업내용")))[:60]
                ]

            )

        _add_table_caption(
            f"{month_label} 오버홀 내역"
        )

        _add_styled_table(

            doc,

            ["작업일자", "공정단계", "작업자", "작업내용"],

            rows

        )

        # 오버홀 작업사진 첨부

        photo_rows = month_overhaul[

            month_overhaul["사진파일명"].astype(str) != ""

        ] if "사진파일명" in month_overhaul.columns else pd.DataFrame()

        for _, prow in photo_rows.iterrows():

            photo_path = os.path.join(

                PHOTO_DIR,

                str(prow["사진파일명"])

            )

            if os.path.exists(photo_path):

                try:

                    doc.add_picture(

                        photo_path,

                        width=Inches(5.0)

                    )

                    cap = doc.add_paragraph()

                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cap_run = cap.add_run(

                        f"{prow.get('작업일자', '')} 작업사진"

                    )

                    cap_run.italic = True

                    cap_run.font.size = Pt(9)

                    cap_run.font.color.rgb = RGBColor(
                        0x66, 0x66, 0x66
                    )

                except Exception:

                    pass

    else:

        doc.add_paragraph(

            f"{month_label}에 저장된 오버홀·작업 이력이 없습니다."

        )

    # ---------------- 5. 정비 권고사항 ----------------

    doc.add_heading(
        "5. 정비 권고사항",
        level=1
    )

    for item in generate_action_items(
        pump,
        result
    ):

        bullet_p = doc.add_paragraph(

            item,

            style="List Bullet"

        )

    # ---------------- 6. 등급 판정 기준 ----------------

    doc.add_heading(
        "6. 등급 판정 기준",
        level=1
    )

    _add_table_caption(
        "등급 판정 기준"
    )

    _add_styled_table(

        doc,

        ["등급", "종합점수 기준", "판정"],

        [
            ["A", "90점 이상", get_grade_text("A")],
            ["B", "80점 이상 ~ 90점 미만", get_grade_text("B")],
            ["C", "70점 이상 ~ 80점 미만", get_grade_text("C")],
            ["D", "60점 이상 ~ 70점 미만", get_grade_text("D")],
            ["E", "60점 미만", get_grade_text("E")]
        ]

    )

    doc.add_paragraph(

        "진동 판정구간은 ISO 10816-3 Class II"
        "(15~300kW, 강성 지지) 권고 기준을 참고했습니다. "
        "17개 세부 진단항목의 배점·기준은 정밀진단 화면의 "
        "각 항목별 판정기준을 따릅니다."

    )

    # ---------------- 6. 추세 그래프 ----------------

    doc.add_heading(
        "7. 추세 그래프",
        level=1
    )

    vib_fig = build_vibration_trend_fig(

        pump,

        result,

        df_history=df_history,

        figsize=(7, 3.2)

    )

    _add_fig_to_doc(
        doc,
        vib_fig,
        caption=f"그림 {_next_num('fig')}. 진동 추세"
    )

    eff_fig = build_efficiency_trend_fig(

        pump,

        result,

        df_history=df_history,

        figsize=(7, 3.2)

    )

    _add_fig_to_doc(
        doc,
        eff_fig,
        caption=f"그림 {_next_num('fig')}. 효율 추세"
    )

    temp_fig = build_temperature_trend_fig(

        pump,

        result,

        df_history=df_history,

        figsize=(7, 3.2)

    )

    _add_fig_to_doc(
        doc,
        temp_fig,
        caption=f"그림 {_next_num('fig')}. 온도 추세"
    )

    hours_fig = build_op_hours_trend_fig(

        pump,

        figsize=(7, 3.2)

    )

    _add_fig_to_doc(
        doc,
        hours_fig,
        caption=f"그림 {_next_num('fig')}. 누적 운전시간 추세 (추정치)"
    )

    compare_fig = build_fleet_compare_fig(

        pump,

        result,

        all_pumps,

        df_history,

        figsize=(7.5, 2.8)

    )

    _add_fig_to_doc(
        doc,
        compare_fig,
        caption=f"그림 {_next_num('fig')}. {pump['equip']} vs 전체 {len(all_pumps)}대 평균 비교"
    )

    # ---------------- 8. 부록: 용어 설명 ----------------

    doc.add_page_break()

    doc.add_heading(

        "8. 부록 : 용어 설명",

        level=1

    )

    for term, desc in GLOSSARY.items():

        term_p = doc.add_paragraph()

        term_run = term_p.add_run(
            term
        )

        term_run.bold = True

        desc_p = doc.add_paragraph(
            desc
        )

        desc_p.paragraph_format.space_after = Pt(8)

    buffer = io.BytesIO()

    doc.save(
        buffer
    )

    return buffer.getvalue()


# ============================================================
# 7-7. 메뉴 공용 "엑셀로 저장" 버튼 + 확인 팝업
#
# 모든 메뉴 하단에 "엑셀로 저장" 버튼을 놓고, 누르면
# 바로 저장하지 않고 예/아니오를 묻는 팝업(st.dialog)을
# 띄운다.
#
# "예" 버튼 자체가 실제 st.download_button이다. (처음엔
# 숨겨진 링크를 자바스크립트로 자동 클릭시키는 방식을 썼는데,
# Streamlit 컴포넌트는 sandbox가 걸린 iframe 안에서 렌더링되어
# 브라우저가 다운로드를 막아버리는 경우가 있어 실제로 저장이
# 안 되는 문제가 있었다. st.download_button은 Streamlit이
# 공식적으로 지원하는 방식이라 브라우저·환경에 관계없이
# 확실하게 동작한다. 예를 누르는 즉시 파일을 미리 만들어
# "예" 버튼 자체에 담아두는 방식으로, 클릭 한 번에 바로
# 다운로드되도록 했다.)
# ============================================================

EXCEL_MIME = (
    "application/vnd.openxmlformats-"
    "officedocument.spreadsheetml.sheet"
)


@st.dialog("엑셀로 저장")
def confirm_excel_export_dialog(
    page_key,
    filename,
    build_df_fn,
    title
):

    st.write(
        "현재 화면의 데이터를 엑셀 파일로 저장하시겠습니까?"
    )

    cache_key = f"_dialog_excel_bytes_{page_key}"

    if cache_key not in st.session_state:

        sheets = build_df_fn()

        st.session_state[cache_key] = build_pretty_excel_bytes(

            sheets,

            title

        )

    c1, c2 = st.columns(2)

    with c1:

        clicked_yes = st.download_button(

            "예",

            data=st.session_state[cache_key],

            file_name=filename,

            mime=EXCEL_MIME,

            key=f"confirm_yes_{page_key}",

            type="primary",

            use_container_width=True

        )

        if clicked_yes:

            del st.session_state[cache_key]

            st.session_state[
                f"_export_done_{page_key}"
            ] = True

            # 팝업을 닫는다. (엑셀로 저장 플래그를 꺼서
            # 다음 rerun부터 다이얼로그를 다시 열지 않게 함)

            st.session_state[
                f"_show_export_confirm_{page_key}"
            ] = False

            st.rerun()

    with c2:

        if st.button(
            "아니오",
            key=f"confirm_no_{page_key}",
            use_container_width=True
        ):

            if cache_key in st.session_state:

                del st.session_state[cache_key]

            st.session_state[
                f"_show_export_confirm_{page_key}"
            ] = False

            st.rerun()


def render_excel_export_section(
    page_key,
    filename,
    build_df_fn,
    title=None
):

    st.write("---")

    if st.button(

        "📥 엑셀로 저장",

        key=f"export_btn_{page_key}",

        use_container_width=True

    ):

        # 버튼을 누른 그 순간뿐 아니라, 이후 rerun에서도
        # 팝업이 계속 열려 있도록 세션에 플래그를 남긴다.

        st.session_state[
            f"_show_export_confirm_{page_key}"
        ] = True

        # 새로 저장 버튼을 누르면 이전 완료 표시는 지운다.

        st.session_state[
            f"_export_done_{page_key}"
        ] = False

    if st.session_state.get(
        f"_show_export_confirm_{page_key}"
    ):

        confirm_excel_export_dialog(

            page_key,

            filename,

            build_df_fn,

            title or filename.rsplit(".", 1)[0]

        )

    if st.session_state.get(
        f"_export_done_{page_key}"
    ):

        st.success(
            "✅ 엑셀 다운로드가 시작되었습니다. "
            "브라우저 하단/다운로드 폴더를 확인하세요."
        )



# ============================================================
# 8. Excel DB 생성
# ============================================================

def get_lock(
    path
):

    # 여러 사람이 거의 동시에 저장할 때 파일이 깨지는 것을
    # 막기 위한 파일 잠금. 실제 다중 사용자 동시편집을 완벽히
    # 해결하지는 못하지만(마지막 저장이 우선), 최소한
    # 두 저장 작업이 겹쳐서 엑셀 파일 자체가 손상되는 것은 막아준다.

    return FileLock(
        f"{path}.lock",
        timeout=10
    )


def safe_append_row(
    path,
    sheet,
    row
):

    with get_lock(path):

        wb = load_workbook(
            path
        )

        ws = wb[
            sheet
        ]

        ws.append(
            row
        )

        wb.save(
            path
        )

        wb.close()

    # 캐시된 read_excel() 결과가 mtime 기준이라, 같은 초 안에
    # 여러 번 쓰기가 일어나면 mtime이 안 바뀌어서 옛날 데이터를
    # 계속 돌려주는 문제가 있었다. 쓰기 직후엔 무조건
    # 캐시를 비워서 다음 읽기가 최신 내용을 가져오게 한다.

    _read_excel_cached.clear()


def migrate_sheet_headers(
    path,
    sheet,
    required_headers
):

    # 이미 배포되어 있는 엑셀 파일에 컬럼이 부족하면
    # (예전 버전으로 만들어진 파일) 뒤쪽에 새 컬럼을 추가한다.
    # 기존 행 데이터는 그대로 두고 새 컬럼만 빈 값으로 늘어난다.

    with get_lock(path):

        wb = load_workbook(
            path
        )

        ws = wb[
            sheet
        ]

        existing = [

            cell.value

            for cell in ws[1]

        ]

        changed = False

        for header in required_headers:

            if header not in existing:

                ws.cell(

                    row=1,

                    column=len(existing) + 1,

                    value=header

                )

                existing.append(
                    header
                )

                changed = True

        if changed:

            wb.save(
                path
            )

            _read_excel_cached.clear()

        wb.close()


def ensure_excel_file(
    path,
    sheet,
    headers
):

    if not os.path.exists(path):

        wb = Workbook()

        ws = wb.active

        ws.title = sheet

        ws.append(headers)

        wb.save(path)

        wb.close()


def _migrate_overhaul_cost_column():

    # 오버홀이력에 "비용" 컬럼이 없던 예전 버전으로 이미 파일이
    # 만들어져 있으면(ensure_excel_file은 파일이 없을 때만
    # 헤더를 새로 만들기 때문에), 여기서 한 번 열어서 "비용"
    # 컬럼이 없으면 뒤에 추가해준다. TCO(총소유비용) 계산에
    # 이 컬럼이 필요하다.

    if not os.path.exists(OVERHAUL_DB_PATH):

        return

    try:

        wb = load_workbook(
            OVERHAUL_DB_PATH
        )

        ws = wb["오버홀이력"]

        headers = [

            cell.value

            for cell in ws[1]

        ]

        if "비용" not in headers:

            ws.cell(
                1,
                len(headers) + 1,
                "비용"
            )

            wb.save(
                OVERHAUL_DB_PATH
            )

        wb.close()

    except Exception:

        pass


def ensure_db_exists():

    ensure_excel_file(

        DB_FILE_PATH,

        "진단이력",

        [
            "점검일",
            "사업장",
            "설비명",
            "제조사",
            "모델명",
            "마력",
            "양정",
            "준공일",
            "점검자",
            "종합점수",
            "최종등급"
        ]
        +
        [
            item[1]
            for item in EVAL_ITEMS
        ]
        +
        [
            "효율측정값(%)",
            "진동측정값(mm/s)",
            "온도측정값(°C)",
            "전류측정값(A)"
        ]

    )

    migrate_sheet_headers(

        DB_FILE_PATH,

        "진단이력",

        [
            "효율측정값(%)",
            "진동측정값(mm/s)",
            "온도측정값(°C)",
            "전류측정값(A)"
        ]

    )

    ensure_excel_file(

        EQUIP_DB_PATH,

        "설비마스터",

        [
            "사업장",
            "설비명",
            "제조사",
            "모델명",
            "정격출력(HP)",
            "정격양정(m)",
            "정격유량(m3/h)",
            "회전수(RPM)",
            "준공일",
            "누적운전시간",
            "기준진동(mm/s)",
            "기준효율(%)"
        ]

    )

    ensure_excel_file(

        OVERHAUL_DB_PATH,

        "오버홀이력",

        [
            "작업일자",
            "사업장",
            "설비명",
            "공정단계",
            "작업자",
            "작업내용",
            "사진파일명",
            "전후효율",
            "전후진동",
            "비용"
        ]

    )

    _migrate_overhaul_cost_column()

    ensure_excel_file(

        DOC_DB_PATH,

        "전문보고서",

        [
            "등록일자",
            "사업장",
            "설비명",
            "보고서구분",
            "수행기관",
            "요약소견",
            "저장파일명"
        ]

    )

    ensure_excel_file(

        KNOWHOW_DB_PATH,

        "노하우DB",

        [
            "등록일자",
            "분류",
            "관련모델",
            "현상및원인",
            "해결노하우",
            "작성자"
        ]

    )

    ensure_excel_file(

        DAILY_LOG_DB_PATH,

        "점검일지",

        [
            "점검일자",
            "점검자",
            "설비명",
            "누수",
            "진동",
            "온도",
            "전류",
            "특이사항"
        ]

    )

    ensure_excel_file(

        SAFETY_PERMIT_DB_PATH,

        "위험작업허가",

        [
            "신청일자",
            "작업명",
            "대상설비",
            "위험유형",
            "작업기간",
            "신청자",
            "승인상태",
            "안전조치"
        ]

    )

    ensure_excel_file(

        KPI_DB_PATH,

        "KPI실적",

        [
            "등록일자",
            "지표",
            "목표",
            "실적",
            "단위",
            "비고"
        ]

    )


ensure_db_exists()

ensure_audit_log_exists()

ensure_consumables_db_exists()

ensure_design_dbs_exist()



# ============================================================
# 9. 샘플 데이터
# ============================================================

def seed_sample_data():

    if os.path.exists(SEED_FLAG_PATH):

        return

    wb = load_workbook(DB_FILE_PATH)

    ws = wb["진단이력"]

    dates = [

        "2024-03-15",
        "2024-09-10",
        "2025-03-20",
        "2025-08-15",
        "2026-02-10",
        "2026-08-15"

    ]

    for pump in DEFAULT_PUMPS:

        for dt in dates:

            score = round(
                random.uniform(
                    72,
                    96
                ),
                1
            )

            grade = get_final_grade(
                score
            )

            grades = [

                random.choice(
                    [
                        "A",
                        "B",
                        "B",
                        "C"
                    ]
                )

                for _ in EVAL_ITEMS
            ]

            ws.append(

                [
                    dt,
                    pump["site"],
                    pump["equip"],
                    pump["maker"],
                    pump["model"],
                    pump["hp"],
                    pump["head"],
                    pump["build_date"],
                    "최진욱",
                    score,
                    grade
                ]
                +
                grades

            )

    wb.save(DB_FILE_PATH)

    wb.close()

    with open(
        SEED_FLAG_PATH,
        "w",
        encoding="utf-8"
    ) as f:

        f.write(
            datetime.now().isoformat()
        )


# seed_sample_data()

# 예전엔 앱을 처음 켰을 때 코드에 박힌 가짜 설비 10대와,
# 그에 딸린 가짜 진단점수 이력을 자동으로 채워 넣었다.
# 실제 회사 데이터를 쓰기 시작하면서 더는 필요 없어져서
# 자동 실행을 껐다(함수 자체는 남겨뒀으니 나중에 데모용으로
# 다시 필요하면 이 줄의 주석만 풀면 된다).



# ============================================================
# 10. 데이터 읽기
#
# - 파일 수정시각(mtime)을 캐시 키에 포함시켜서, 저장이 일어나면
#   자동으로 캐시가 무효화되도록 했다. (매번 통째로 새로 읽던
#   문제 -> 저장 전까지는 캐시된 결과를 재사용해서 빨라진다)
# - "파일이 아예 없음"과 "파일은 있는데 읽다가 깨짐"을 구분해서,
#   후자는 화면에 경고를 띄운다. 예전에는 둘 다 그냥
#   빈 데이터로 조용히 처리해서 실제 파일이 손상돼도
#   사용자가 알아챌 방법이 없었다.
# ============================================================

read_errors = []


@st.cache_data(show_spinner=False)
def _read_excel_cached(path, sheet, mtime, size):

    return pd.read_excel(
        path,
        sheet_name=sheet
    )


def read_excel(
    path,
    sheet
):

    if not os.path.exists(path):

        return pd.DataFrame()

    try:

        # mtime만 캐시 키로 쓰면, 짧은 시간 안에 여러 번 쓰기가
        # 일어날 때(예: 설비 10개를 연달아 빠르게 삭제) 파일
        # 시스템의 mtime 해상도(보통 1초 단위)에 걸려서 mtime이
        # 그대로인 채로 내용만 바뀌는 경우가 있었다. 그러면
        # 캐시 키가 안 바뀌어서 옛날 데이터를 계속 돌려주는
        # 버그가 생긴다. 파일 크기도 같이 캐시 키에 넣으면
        # 행이 추가/삭제될 때 크기가 달라지므로 훨씬 안전하다.

        mtime = os.path.getmtime(
            path
        )

        size = os.path.getsize(
            path
        )

        return _read_excel_cached(
            path,
            sheet,
            mtime,
            size
        )

    except Exception as e:

        read_errors.append(
            f"{os.path.basename(path)} ({sheet}) 읽기 실패: {e}"
        )

        return pd.DataFrame()


seed_default_unitprice_library()


# ============================================================
# 10-1. 설비 마스터 (추가·삭제 가능한 설비 목록)
#
# 예전에는 설비 10대가 코드(DEFAULT_PUMPS)에 고정되어 있어서
# 현장에서 설비를 추가·삭제할 수 없었다.
# 이제 설비 목록을 엑셀 DB(EQUIP_DB_PATH)로 관리하고,
# 비어 있으면 기존 10대로 최초 1회 시딩한다.
# ============================================================

def seed_equipment_if_empty():

    # 예전엔 "설비마스터가 비어있으면 무조건 다시 채운다"는
    # 로직이라, 사용자가 샘플(가짜) 설비를 일부러 다 지워도
    # 다음 rerun 때 자동으로 부활하는 버그가 있었다.
    # 진단이력 샘플 시딩(seed_sample_data)과 같은 방식으로
    # "앱을 통틀어 딱 한 번만" 시딩하되, 서로 다른 시점에
    # 실행되는 두 시딩 함수가 플래그를 공유하면 순서 문제가
    # 생길 수 있어 설비마스터 전용 플래그를 따로 쓴다.

    if os.path.exists(EQUIP_SEED_FLAG_PATH):

        return

    df_equip = read_excel(
        EQUIP_DB_PATH,
        "설비마스터"
    )

    if not df_equip.empty:

        with open(

            EQUIP_SEED_FLAG_PATH,

            "w",

            encoding="utf-8"

        ) as f:

            f.write(
                datetime.now().isoformat()
            )

        return

    with get_lock(EQUIP_DB_PATH):

        wb = load_workbook(
            EQUIP_DB_PATH
        )

        ws = wb[
            "설비마스터"
        ]

        for pump in DEFAULT_PUMPS:

            ws.append(

                [
                    pump["site"],
                    pump["equip"],
                    pump["maker"],
                    pump["model"],
                    pump["hp"],
                    pump["head"],
                    pump["flow"],
                    pump["rpm"],
                    pump["build_date"],
                    pump["op_hours"],
                    None,
                    None
                ]

            )

        wb.save(
            EQUIP_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()

    with open(

        EQUIP_SEED_FLAG_PATH,

        "w",

        encoding="utf-8"

    ) as f:

        f.write(
            datetime.now().isoformat()
        )


# seed_equipment_if_empty()

# 위와 같은 이유로 가짜 설비 10대 자동 채우기도 껐다.


def get_all_pumps():

    df_equip = read_excel(
        EQUIP_DB_PATH,
        "설비마스터"
    )

    if df_equip.empty:

        # 예전엔 여기서 DEFAULT_PUMPS(코드에 박힌 가짜 10대)를
        # 그대로 돌려줬다. 그래서 사용자가 진짜로 설비를 전부
        # 지워서 0대로 만들어도, 화면에는 계속 가짜 10대가
        # 나타나는 버그가 있었다. 이제는 정직하게 빈 리스트를
        # 돌려주고, 화면 쪽에서 "등록된 설비가 없습니다"를
        # 보여주게 한다.

        return []

    pumps = []

    for _, row in df_equip.iterrows():

        pumps.append(

            {
                "site": row["사업장"],
                "equip": row["설비명"],
                "maker": row["제조사"],
                "model": row["모델명"],
                "hp": row["정격출력(HP)"],
                "head": row["정격양정(m)"],
                "flow": row["정격유량(m3/h)"],
                "rpm": row["회전수(RPM)"],
                "build_date": str(
                    row["준공일"]
                )[:10],
                "op_hours": row["누적운전시간"],

                "기준진동": (

                    float(row["기준진동(mm/s)"])

                    if pd.notna(
                        row.get("기준진동(mm/s)")
                    )

                    else None

                ),

                "기준효율": (

                    float(row["기준효율(%)"])

                    if pd.notna(
                        row.get("기준효율(%)")
                    )

                    else None

                )

            }

        )

    return pumps


def add_equipment(new_pump):

    # 수동 등록 화면에서는 중복이름을 미리 걸러주지만, 통합
    # 업로드·일괄교체 등 다른 경로로 들어올 때도 안전하도록
    # 함수 자체에도 마지막 방어선을 둔다. 안 그러면 같은
    # 이름의 설비가 두 번 들어가서, 화면에서 그 설비 이름으로
    # 만든 버튼 key가 중복돼 앱이 통째로 죽는 문제가 생긴다.

    existing_names = {

        p["equip"].strip()

        for p in get_all_pumps()

    }

    if new_pump["equip"].strip() in existing_names:

        return False

    safe_append_row(

        EQUIP_DB_PATH,

        "설비마스터",

        [
            new_pump["site"],
            new_pump["equip"],
            new_pump["maker"],
            new_pump["model"],
            new_pump["hp"],
            new_pump["head"],
            new_pump["flow"],
            new_pump["rpm"],
            new_pump["build_date"],
            new_pump["op_hours"],
            new_pump.get("기준진동"),
            new_pump.get("기준효율")
        ]

    )

    log_audit(

        "설비 추가",

        new_pump["equip"],

        f"사업장={new_pump['site']}, 모델={new_pump['model']}"

    )

    return True


def delete_equipment(equip_name):

    with get_lock(EQUIP_DB_PATH):

        wb = load_workbook(
            EQUIP_DB_PATH
        )

        ws = wb[
            "설비마스터"
        ]

        headers = [
            c.value for c in ws[1]
        ]

        rows = list(
            ws.iter_rows(
                min_row=2
            )
        )

        for r in rows:

            if r[1].value == equip_name:

                # 지우기 전에 휴지통에 먼저 옮겨 담는다
                # (30일간 복원 가능)

                row_dict = {

                    headers[i]: cell.value

                    for i, cell in enumerate(r)

                    if i < len(headers)

                }

                move_to_trash(

                    "설비",

                    equip_name,

                    "설비마스터",

                    row_dict

                )

                ws.delete_rows(
                    r[0].row
                )

                break

        wb.save(
            EQUIP_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()

    log_audit(

        "설비 삭제",

        equip_name

    )


def delete_sample_default_equipment():

    # DEFAULT_PUMPS(코드에 예시로 박혀있던 가짜 10대)와
    # 이름이 정확히 일치하는 설비만 골라 한 번에 지운다.
    # 실제 데이터를 업로드해서 새로 등록한 설비는
    # 이름이 다르므로 건드리지 않는다.

    sample_names = {

        p["equip"]

        for p in DEFAULT_PUMPS

    }

    current_names = {

        p["equip"]

        for p in get_all_pumps()

    }

    to_delete = sample_names & current_names

    for name in to_delete:

        delete_equipment(
            name
        )

    log_audit(

        "샘플 설비 일괄 삭제",

        "전체",

        f"{len(to_delete)}건 삭제"

    )

    return len(to_delete)


# ============================================================
# 10-2. 실제 회사 데이터(엑셀) 일괄 업로드
#
# 사업소에서 받은 "설비 기본정보" 엑셀과 "오버홀 실적" 엑셀을
# 그대로 올리면 설비마스터·오버홀이력·QR까지 한 번에 채워진다.
# 원본 파일에 없는 값(제조사·모델·정격출력 등)은 빈칸으로 두고
# 억지로 채워넣지 않는다.
# ============================================================

def track_recent_view(equip_name):

    # 홈 화면 "최근 본 설비"에 쓸 방문 기록. 세션 동안만
    # 유지되는 간단한 목록이라 별도 DB 없이 session_state로
    # 처리한다. 이미 있던 항목이면 맨 앞으로 올리고, 최대
    # 5개까지만 남긴다.

    recent = st.session_state.get(
        "_recent_viewed",
        []
    )

    recent = [

        e for e in recent

        if e != equip_name

    ]

    recent.insert(
        0,
        equip_name
    )

    st.session_state["_recent_viewed"] = recent[:5]


def anonymize_name(name):

    # 정밀진단·오버홀 이력에 실제 담당자 실명이 그대로
    # 노출되던 문제. "최진욱"(이 앱 관리자)만 예외로 두고,
    # 나머지는 성만 남기고 나머지 글자를 ○로 가린다.
    # "김광일 외02명"처럼 뒤에 붙는 표현은 그대로 유지한다.

    if not name:

        return name

    name = str(name).strip()

    m = re.match(
        r"^([가-힣]+)(.*)$",
        name
    )

    if not m:

        return name

    korean_part = m.group(1)

    rest = m.group(2)

    if korean_part == "최진욱":

        return name

    if len(korean_part) <= 1:

        return name

    masked = (

        korean_part[0]
        +
        "○" * (len(korean_part) - 1)

    )

    return masked + rest


def parse_equipment_import_excel(df1):

    site_pattern = re.compile(
        r'^(.*?)\s*가압펌프동'
    )

    equip_rows = []

    equip_label_by_no = {}

    site_by_no = {}

    # 평촌가압장처럼 스펙이 같은 설비 여러 대가 "설비내역"에
    # 아무 번호 없이 "원심펌프"로만 똑같이 적혀있는 경우가
    # 있다. 이러면 나중에 "평촌#1/#2/#3" 같은 진동그래프
    # 시트명과 매칭할 방법이 없어서, 같은 사업장+설명이
    # 중복되면 순서대로 "#1,#2,#3..."을 자동으로 붙여준다.

    dup_counter = {}

    for _, r in df1.iterrows():

        equip_no = r.get("설비번호")

        if pd.isna(equip_no):

            continue

        equip_no = int(equip_no)

        m = site_pattern.match(

            str(r.get("기능위치내역", ""))

        )

        site_name = (

            m.group(1)

            if m

            else str(r.get("기능위치내역", "")).strip()

        )

        equip_desc = str(

            r.get("설비내역", "")

        ).replace(",", "").strip()

        equip_desc = re.sub(
            r"\s+",
            " ",
            equip_desc
        )

        if "#" not in equip_desc:

            dup_key = (
                site_name,
                equip_desc
            )

            dup_counter[dup_key] = dup_counter.get(
                dup_key,
                0
            ) + 1

            equip_desc = (

                f"{equip_desc} #{dup_counter[dup_key]}"

            )

        equip_label = f"{equip_desc} ({equip_no})"

        equip_label_by_no[equip_no] = equip_label

        site_by_no[equip_no] = site_name

        equip_rows.append(

            {
                "site": site_name,
                "equip": equip_label,
                "maker": "",
                "model": "",
                "hp": 0,
                "head": 0,
                "flow": 0,
                "rpm": 0,
                "build_date": "",
                "op_hours": 0,
                "기준진동": None,
                "기준효율": None
            }

        )

    return equip_rows, equip_label_by_no, site_by_no


def parse_overhaul_import_excel(

    df2,

    equip_label_by_no,

    site_by_no

):

    overhaul_rows = []

    skipped = 0

    for _, r in df2.iterrows():

        eq_no = r.get("설비")

        if pd.isna(eq_no):

            skipped += 1

            continue

        eq_no = int(eq_no)

        if eq_no not in equip_label_by_no:

            skipped += 1

            continue

        raw_date = r.get("측정일자")

        if pd.isna(raw_date):

            skipped += 1

            continue

        try:

            work_date = datetime.strptime(

                str(int(raw_date)),

                "%Y%m%d"

            ).strftime("%Y-%m-%d")

        except Exception:

            skipped += 1

            continue

        parts = []

        if pd.notna(r.get("평가코드내역")):

            parts.append(
                str(r["평가코드내역"])
            )

        vib_value = None

        if pd.notna(r.get("측정값")):

            # 확인 결과 이 mm 측정값은 진동 측정값이다.
            # 문장 안에 묻어두지 않고 전후진동 컬럼에
            # 직접 넣어서 나중에 데이터로 활용할 수 있게 한다.

            vib_value = float(
                r["측정값"]
            )

            parts.append(

                f"진동 측정값 {vib_value}mm/s"

            )

        if pd.notna(r.get("오더내역")):

            parts.append(
                str(r["오더내역"])
            )

        content = " · ".join(
            parts
        )

        raw_worker = (

            str(r.get("측정인", ""))

            if pd.notna(r.get("측정인"))

            else ""

        )

        overhaul_rows.append(

            {
                "작업일자": work_date,
                "site": site_by_no[eq_no],
                "equip": equip_label_by_no[eq_no],
                "공정단계": (

                    str(r.get("측정지점내역", ""))

                    if pd.notna(r.get("측정지점내역"))

                    else ""

                ),
                "작업자": anonymize_name(
                    raw_worker
                ),
                "작업내용": content,
                "전후진동": vib_value
            }

        )

    return overhaul_rows, skipped


def replace_all_equipment(new_equip_rows):

    # 기존 설비마스터를 전부 지우고(헤더는 유지) 새 목록으로
    # 통째로 교체한다.

    with get_lock(EQUIP_DB_PATH):

        wb = load_workbook(
            EQUIP_DB_PATH
        )

        ws = wb["설비마스터"]

        max_row = ws.max_row

        if max_row > 1:

            ws.delete_rows(
                2,
                max_row - 1
            )

        for row in new_equip_rows:

            ws.append(

                [
                    row["site"],
                    row["equip"],
                    row["maker"],
                    row["model"],
                    row["hp"],
                    row["head"],
                    row["flow"],
                    row["rpm"],
                    row["build_date"],
                    row["op_hours"],
                    row.get("기준진동"),
                    row.get("기준효율")
                ]

            )

        wb.save(
            EQUIP_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()

    log_audit(

        "설비 일괄 교체",

        "전체",

        f"{len(new_equip_rows)}건으로 교체"

    )


def bulk_append_overhaul(overhaul_rows):

    for row in overhaul_rows:

        safe_append_row(

            OVERHAUL_DB_PATH,

            "오버홀이력",

            [
                row["작업일자"],
                row["site"],
                row["equip"],
                row["공정단계"],
                row["작업자"],
                row["작업내용"],
                "",
                "",
                row.get("전후진동", ""),
                row.get("비용", 0)
            ]

        )

    log_audit(

        "오버홀 일괄 업로드",

        "전체",

        f"{len(overhaul_rows)}건 추가"

    )


# ============================================================
# 10-3. 통합 업로드 양식 (설비정보 + 정밀진단 + 오버홀이력)
#
# 엑셀 파일 하나만 채워서 올리면 설비등록 · QR생성(자동) ·
# 정밀진단 저장 · 오버홀이력까지 한 번에 반영되도록 만든
# 사용자 작성용 양식이다.
# ============================================================

DIAG_TEMPLATE_QUAL_ITEMS = [

    ("임펠러손상등급", "임펠러 손상/침식"),
    ("임펠러밸런싱등급", "임펠러 동적 밸런싱"),
    ("NPSH등급", "NPSH 여유율/캐비테이션"),
    ("코팅상태등급", "내부 코팅 상태"),
    ("비금속웨어링등급", "비금속 웨어링 개선"),
    ("베어링결함등급", "베어링 결함 진동"),
    ("주파수성분등급", "주파수 성분 결함"),
    ("softfoot등급", "Soft Foot 및 배관 응력"),
    ("소모품이력등급", "주요 소모품 교체이력")

]


def _style_template_header(ws, headers):

    header_fill = PatternFill(
        "solid",
        fgColor=BRAND_HEADER_COLOR
    )

    header_font = Font(
        bold=True,
        color="FFFFFF"
    )

    for j, h in enumerate(headers, start=1):

        c = ws.cell(
            row=1,
            column=j,
            value=h
        )

        c.font = header_font

        c.fill = header_fill

        c.alignment = Alignment(
            horizontal="center",
            vertical="center",
            wrap_text=True
        )

        ws.column_dimensions[

            get_column_letter(j)

        ].width = max(
            len(h) + 4,
            12
        )

    ws.freeze_panes = "A2"


def build_unified_import_template_bytes():

    wb = Workbook()

    wb.remove(
        wb.active
    )

    example_fill = PatternFill(
        "solid",
        fgColor="FFF9DB"
    )

    # ---- 시트1. 설비정보 ----

    ws1 = wb.create_sheet(
        "설비정보"
    )

    headers1 = [
        "사업장", "설비명", "제조사", "모델명",
        "정격출력(HP)", "정격양정(m)", "정격유량(m3/h)",
        "회전수(RPM)", "준공일", "누적운전시간(h)",
        "기준진동(mm/s)", "기준효율(%)"
    ]

    _style_template_header(
        ws1,
        headers1
    )

    example1 = [
        "부곡가압장", "원심펌프 #1 (33003471)", "효성펌프", "DHP-1",
        160, 47, 1250, 1780, "2018-01-15", 10200, 4.5, 90
    ]

    for j, v in enumerate(example1, start=1):

        c = ws1.cell(
            row=2,
            column=j,
            value=v
        )

        c.fill = example_fill

    # ---- 시트2. 정밀진단 ----

    ws2 = wb.create_sheet(
        "정밀진단"
    )

    headers2 = (

        ["설비명(설비정보 시트와 동일해야 함)", "점검일(YYYY-MM-DD)", "점검자"]

        +
        ["효율유지율(%)", "양정유량도달률(%)", "BEP적정성(%)"]

        +
        ["링간극(mm)", "축슬리브마모(mm)"]

        +
        [label for _, label in [
            ("", "임펠러손상등급(A~E)"),
            ("", "임펠러밸런싱등급(A~E)"),
            ("", "NPSH등급(A~E)"),
            ("", "코팅상태등급(A~E)"),
            ("", "비금속웨어링등급(A~E)")
        ]]

        +
        ["펌프부하진동(mm/s)", "펌프반부하진동(mm/s)",
         "모터부하진동(mm/s)", "모터반부하진동(mm/s)"]

        +
        ["베어링결함등급(A~E)", "주파수성분등급(A~E)"]

        +
        ["센터링(mm)"]

        +
        ["softfoot등급(A~E)"]

        +
        ["오버홀주기(누적운전시간h)"]

        +
        ["소모품이력등급(A~E)"]

        +
        ["측정온도(°C)", "측정전류(A)"]

    )

    _style_template_header(
        ws2,
        headers2
    )

    example2 = [
        "원심펌프 #1 (33003471)", "2026-08-23", "홍길동",
        95.0, 96.0, 100.0,
        1.2, 0.8,
        "A", "A", "A", "A", "A",
        2.0, 1.5, 2.2, 1.8,
        "A", "A",
        0.05,
        "A",
        8000,
        "A",
        45.0, 30.0
    ]

    for j, v in enumerate(example2, start=1):

        c = ws2.cell(
            row=2,
            column=j,
            value=v
        )

        c.fill = example_fill

    # ---- 시트3. 오버홀이력 ----

    ws3 = wb.create_sheet(
        "오버홀이력"
    )

    headers3 = [
        "작업일자(YYYY-MM-DD)", "설비명(설비정보 시트와 동일해야 함)",
        "공정단계", "작업자", "작업내용", "진동측정값(mm/s, 선택)"
    ]

    _style_template_header(
        ws3,
        headers3
    )

    example3 = [
        "2026-08-23", "원심펌프 #1 (33003471)",
        "펌프모터 OVERHAUL", "홍길동", "축슬리브 교체 · 센터링 완료", 1.0
    ]

    for j, v in enumerate(example3, start=1):

        c = ws3.cell(
            row=2,
            column=j,
            value=v
        )

        c.fill = example_fill

    # ---- 시트0. 작성안내 (맨 앞에 오도록 마지막에 삽입) ----

    ws0 = wb.create_sheet(
        "작성안내",
        0
    )

    guide_lines = [

        "① 이 파일 하나만 업로드하면 설비등록·QR생성·정밀진단·"
        "오버홀이력이 한 번에 반영됩니다.",

        "② 노란색으로 칠해진 2행은 작성 예시입니다. "
        "실제 데이터를 2행부터 이어서 입력하고, 예시행은 지우거나 "
        "덮어써도 됩니다.",

        "③ '설비명'은 세 시트 모두에서 반드시 똑같은 문자열이어야 "
        "매칭됩니다 (예: '원심펌프 #1 (33003471)').",

        "④ 값을 모르는 칸은 비워두세요. 빈칸은 억지로 채우지 않고 "
        "그대로 반영됩니다.",

        "⑤ 등급 칸(A~E)은 정밀진단 화면에서 사람이 직접 판정하는 "
        "항목입니다. 대문자 A/B/C/D/E 중 하나로만 입력하세요.",

        "⑥ 진동은 '펌프 부하/반부하'·'모터 부하/반부하' 4개 값을 "
        "각각 입력하면, 그중 가장 큰 값(worst case)을 종합 진동값으로 "
        "사용해 등급을 계산합니다.",

        "⑦ '정밀진단' 시트만 있고 '설비정보' 시트가 비어있어도 "
        "괜찮습니다 (이미 등록된 설비에 진단결과만 추가하는 경우).",

        "⑧ 오버홀이력의 작업자 이름은 업로드 시 자동으로 "
        "이니셜 처리됩니다 (예: 홍길동 → 홍○○)."

    ]

    for i, line in enumerate(guide_lines, start=1):

        p = ws0.cell(
            row=i,
            column=1,
            value=line
        )

        p.alignment = Alignment(
            wrap_text=True,
            vertical="top"
        )

        p.font = Font(
            size=11
        )

    ws0.column_dimensions["A"].width = 100

    for i in range(
        1,
        len(guide_lines) + 1
    ):

        ws0.row_dimensions[i].height = 32

    buffer = io.BytesIO()

    wb.save(
        buffer
    )

    return buffer.getvalue()


def upsert_equipment_rows(equip_rows):

    # 기존 설비는 그대로 두고, 새 설비명만 추가한다.
    # (통합 업로드는 반복해서 여러 번 올릴 걸 가정하므로,
    #  이전 것을 지우는 "전체 교체"가 아니라 "이미 있으면
    #  건너뛰고 없으면 추가"로 동작한다.)
    #
    # 예전엔 existing_names를 루프 시작 전에 한 번만 계산해서,
    # 같은 업로드 파일 "안에" 같은 이름이 두 번 있으면 그
    # 배치 안에서의 중복은 못 걸러내고 둘 다 추가되는 문제가
    # 있었다(그러면 나중에 그 이름으로 만든 화면 버튼 key가
    # 겹쳐서 앱이 죽는다). 추가할 때마다 그 이름도 바로
    # existing_names에 넣어서 같은 배치 안의 중복도 막는다.

    existing_names = {

        p["equip"]

        for p in get_all_pumps()

    }

    added = 0

    for row in equip_rows:

        if row["equip"] in existing_names:

            continue

        if add_equipment(row):

            existing_names.add(
                row["equip"]
            )

            added += 1

    return added


def parse_unified_import_workbook(file_bytes):

    xl = pd.ExcelFile(
        io.BytesIO(file_bytes)
    )

    result = {

        "equip_rows": [],
        "diag_rows": [],
        "overhaul_rows": [],
        "diag_skipped": [],
        "overhaul_skipped": 0

    }

    # ---- 설비정보 ----

    if "설비정보" in xl.sheet_names:

        df_eq = xl.parse(
            "설비정보"
        )

        for _, r in df_eq.iterrows():

            equip_name = r.get("설비명")

            if pd.isna(equip_name) or not str(equip_name).strip():

                continue

            def _num_or_zero(val):

                return (

                    float(val)

                    if pd.notna(val)

                    else 0

                )

            result["equip_rows"].append(

                {
                    "site": (

                        str(r.get("사업장", ""))

                        if pd.notna(r.get("사업장"))

                        else ""

                    ),
                    "equip": str(equip_name).strip(),
                    "maker": (

                        str(r.get("제조사", ""))

                        if pd.notna(r.get("제조사"))

                        else ""

                    ),
                    "model": (

                        str(r.get("모델명", ""))

                        if pd.notna(r.get("모델명"))

                        else ""

                    ),
                    "hp": _num_or_zero(
                        r.get("정격출력(HP)")
                    ),
                    "head": _num_or_zero(
                        r.get("정격양정(m)")
                    ),
                    "flow": _num_or_zero(
                        r.get("정격유량(m3/h)")
                    ),
                    "rpm": _num_or_zero(
                        r.get("회전수(RPM)")
                    ),
                    "build_date": (

                        str(r.get("준공일", ""))

                        if pd.notna(r.get("준공일"))

                        else ""

                    ),
                    "op_hours": _num_or_zero(
                        r.get("누적운전시간(h)")
                    ),
                    "기준진동": (

                        float(r.get("기준진동(mm/s)"))

                        if pd.notna(r.get("기준진동(mm/s)"))

                        else None

                    ),
                    "기준효율": (

                        float(r.get("기준효율(%)"))

                        if pd.notna(r.get("기준효율(%)"))

                        else None

                    )
                }

            )

    # 설비명 -> pump dict 매핑 (정밀진단/오버홀에서 참조)

    all_pumps_by_name = {

        p["equip"]: p

        for p in get_all_pumps()

    }

    for row in result["equip_rows"]:

        all_pumps_by_name.setdefault(

            row["equip"],

            {
                "site": row["site"],
                "equip": row["equip"],
                "maker": row["maker"],
                "model": row["model"],
                "hp": row["hp"],
                "head": row["head"],
                "build_date": row["build_date"]
            }

        )

    # ---- 정밀진단 ----

    diag_col = "설비명(설비정보 시트와 동일해야 함)"

    if "정밀진단" in xl.sheet_names:

        df_diag = xl.parse(
            "정밀진단"
        )

        for _, r in df_diag.iterrows():

            equip_name = r.get(diag_col)

            if pd.isna(equip_name) or not str(equip_name).strip():

                continue

            equip_name = str(equip_name).strip()

            pump = all_pumps_by_name.get(
                equip_name
            )

            if pump is None:

                result["diag_skipped"].append(

                    f"{equip_name} (설비정보에 없음)"

                )

                continue

            def _f(colname, default=None):

                v = r.get(colname)

                return (

                    float(v)

                    if pd.notna(v)

                    else default

                )

            def _grade(colname, default="B"):

                v = r.get(colname)

                if pd.isna(v):

                    return default

                v = str(v).strip().upper()

                return v if v in ("A", "B", "C", "D", "E") else default

            eff_val = _f(
                "효율유지율(%)",
                None
            )

            reach_val = _f(
                "양정유량도달률(%)",
                None
            )

            bep_val = _f(
                "BEP적정성(%)",
                None
            )

            ring_val = _f(
                "링간극(mm)",
                None
            )

            sleeve_val = _f(
                "축슬리브마모(mm)",
                None
            )

            vib_candidates = [

                v for v in [

                    _f("펌프부하진동(mm/s)"),
                    _f("펌프반부하진동(mm/s)"),
                    _f("모터부하진동(mm/s)"),
                    _f("모터반부하진동(mm/s)")

                ]

                if v is not None

            ]

            vib_val = (

                max(vib_candidates)

                if vib_candidates

                else None

            )

            align_val = _f(
                "센터링(mm)",
                None
            )

            overhaul_hours_val = _f(
                "오버홀주기(누적운전시간h)",
                None
            )

            temp_val = _f(
                "측정온도(°C)",
                45.0
            )

            current_val = _f(
                "측정전류(A)",
                None
            )

            grade_by_item_name = {}

            if eff_val is not None:
                grade_by_item_name["펌프 효율 유지율 (%)"] = calc_eff(eff_val)

            if reach_val is not None:
                grade_by_item_name["설계 양정/유량 도달률 (%)"] = calc_reach(reach_val)

            if bep_val is not None:
                grade_by_item_name["BEP 운전점 적정성 (%)"] = calc_bep(bep_val)

            if ring_val is not None:
                grade_by_item_name["임펠러/케이싱 링 간극"] = calc_ring_gap(ring_val)

            if sleeve_val is not None:
                grade_by_item_name["축슬리브 마모"] = calc_sleeve(sleeve_val)

            if vib_val is not None:

                effective_vib_fn = get_effective_auto_fn(

                    "Overall 진동 (mm/s)",

                    calc_vib,

                    pump

                )

                grade_by_item_name["Overall 진동 (mm/s)"] = effective_vib_fn(
                    vib_val
                )

            if align_val is not None:
                grade_by_item_name["펌프-모터 센터링"] = calc_align(align_val)

            if overhaul_hours_val is not None:
                grade_by_item_name["오버홀 주기"] = calc_overhaul(overhaul_hours_val)

            qual_col_lookup = {

                "임펠러손상등급": "임펠러손상등급(A~E)",
                "임펠러밸런싱등급": "임펠러밸런싱등급(A~E)",
                "NPSH등급": "NPSH등급(A~E)",
                "코팅상태등급": "코팅상태등급(A~E)",
                "비금속웨어링등급": "비금속웨어링등급(A~E)",
                "베어링결함등급": "베어링결함등급(A~E)",
                "주파수성분등급": "주파수성분등급(A~E)",
                "softfoot등급": "softfoot등급(A~E)",
                "소모품이력등급": "소모품이력등급(A~E)"

            }

            for col_key, item_name in DIAG_TEMPLATE_QUAL_ITEMS:

                grade_by_item_name[item_name] = _grade(

                    qual_col_lookup[col_key]

                )

            details_grades = []

            total_score = 0

            for item in EVAL_ITEMS:

                name = item[1]

                grade = grade_by_item_name.get(
                    name,
                    "B"
                )

                score = item[5].get(
                    grade,
                    0
                )

                details_grades.append(
                    grade
                )

                total_score += score

            total_score = round(
                total_score,
                2
            )

            final_grade = get_final_grade(
                total_score
            )

            checker = (

                str(r.get("점검자", ""))

                if pd.notna(r.get("점검자"))

                else ""

            )

            checker = anonymize_name(
                checker
            )

            raw_date = r.get("점검일(YYYY-MM-DD)")

            try:

                if pd.isna(raw_date):

                    check_date = datetime.now().strftime("%Y-%m-%d")

                elif isinstance(raw_date, datetime):

                    check_date = raw_date.strftime("%Y-%m-%d")

                else:

                    check_date = str(raw_date).strip()[:10]

            except Exception:

                check_date = datetime.now().strftime("%Y-%m-%d")

            db_row = (

                [
                    check_date,
                    pump.get("site", ""),
                    equip_name,
                    pump.get("maker", ""),
                    pump.get("model", ""),
                    pump.get("hp", ""),
                    pump.get("head", ""),
                    pump.get("build_date", ""),
                    checker,
                    total_score,
                    final_grade
                ]

                +
                details_grades

                +
                [

                    eff_val if eff_val is not None else "",
                    vib_val if vib_val is not None else "",
                    temp_val if temp_val is not None else "",
                    current_val if current_val is not None else ""

                ]

            )

            result["diag_rows"].append(
                db_row
            )

    # ---- 오버홀이력 ----

    if "오버홀이력" in xl.sheet_names:

        df_oh = xl.parse(
            "오버홀이력"
        )

        for _, r in df_oh.iterrows():

            equip_name = r.get(

                "설비명(설비정보 시트와 동일해야 함)"

            )

            if pd.isna(equip_name) or not str(equip_name).strip():

                result["overhaul_skipped"] += 1

                continue

            equip_name = str(equip_name).strip()

            pump = all_pumps_by_name.get(
                equip_name
            )

            if pump is None:

                result["overhaul_skipped"] += 1

                continue

            raw_date = r.get(
                "작업일자(YYYY-MM-DD)"
            )

            try:

                if pd.isna(raw_date):

                    work_date = datetime.now().strftime("%Y-%m-%d")

                elif isinstance(raw_date, datetime):

                    work_date = raw_date.strftime("%Y-%m-%d")

                else:

                    work_date = str(raw_date).strip()[:10]

            except Exception:

                work_date = datetime.now().strftime("%Y-%m-%d")

            worker = (

                str(r.get("작업자", ""))

                if pd.notna(r.get("작업자"))

                else ""

            )

            worker = anonymize_name(
                worker
            )

            vib_note = r.get(
                "진동측정값(mm/s, 선택)"
            )

            vib_note = (

                float(vib_note)

                if pd.notna(vib_note)

                else ""

            )

            result["overhaul_rows"].append(

                [
                    work_date,
                    pump.get("site", ""),
                    equip_name,
                    (

                        str(r.get("공정단계", ""))

                        if pd.notna(r.get("공정단계"))

                        else ""

                    ),
                    worker,
                    (

                        str(r.get("작업내용", ""))

                        if pd.notna(r.get("작업내용"))

                        else ""

                    ),
                    "",
                    "",
                    vib_note,
                    0
                ]

            )

    return result


def apply_unified_import(parsed):

    added_equip = upsert_equipment_rows(

        parsed["equip_rows"]

    )

    for db_row in parsed["diag_rows"]:

        safe_append_row(

            DB_FILE_PATH,

            "진단이력",

            db_row

        )

    for db_row in parsed["overhaul_rows"]:

        safe_append_row(

            OVERHAUL_DB_PATH,

            "오버홀이력",

            db_row

        )

    log_audit(

        "통합 엑셀 업로드",

        "전체",

        f"설비 {added_equip}건 추가, "
        f"정밀진단 {len(parsed['diag_rows'])}건, "
        f"오버홀 {len(parsed['overhaul_rows'])}건"

    )

    return added_equip


# ============================================================
# 10-4. 월간 진동측정 보고서
#
# 사업소에서 실제로 매달 쓰는 "펌프모터 진동 측정 분석 보고서"
# 양식을 그대로 재현한다. SAP에서 뽑은 "정기점검 실적목록" 원본
# 엑셀(측정점 10개 x 설비 수만큼 행)을 매달 업로드하면, 누적
# DB에 쌓이고 그걸로 붙임2(기록대장)·붙임3(추세분석)까지 자동
# 생성한다.
# ============================================================

VIBRATION_DB_PATH = "Pump_VibrationMeasure_DB.xlsx"

VIB_JUDGE_GOOD = 3.2
VIB_JUDGE_FAIR = 5.1
VIB_JUDGE_BAD = 8.5

VIB_POINT_COLS = [

    ("모터", "반부하", "수직"),
    ("모터", "반부하", "수평"),
    ("모터", "반부하", "축"),
    ("모터", "부하", "수직"),
    ("모터", "부하", "수평"),
    ("펌프", "반부하", "수직"),
    ("펌프", "반부하", "수평"),
    ("펌프", "반부하", "축"),
    ("펌프", "부하", "수직"),
    ("펌프", "부하", "수평")

]


def ensure_vibration_db_exists():

    ensure_excel_file(

        VIBRATION_DB_PATH,

        "진동측정이력",

        [
            "측정일자",
            "사업장",
            "설비명",
            "펌프모터구분",
            "부하구분",
            "측정방향",
            "측정값",
            "평가코드내역",
            "측정인"
        ]

    )


ensure_vibration_db_exists()


def _extract_site_from_raw(raw_site):

    # "양산정수장 여과지동" -> "양산정수장",
    # "부곡가압장 가압펌프동" -> "부곡가압장" 처럼
    # 뒤에 붙는 "동" 이름을 떼어내 사업장명만 남긴다.

    if not raw_site:

        return ""

    m = re.match(
        r"^(.*?)\s*(가압펌프동|여과지동)?$",
        str(raw_site).strip()
    )

    return m.group(1) if m else str(raw_site).strip()


def _extract_ho_label(raw_equip_desc):

    # "원심펌프, #3" -> "#3", "원심펌프, 인라인 #1" -> "인라인#1"

    text = str(raw_equip_desc)

    if "인라인" in text:

        m = re.search(r"인라인\s*#?\s*(\d+)", text)

        if m:

            return f"인라인#{m.group(1)}"

    m = re.search(r"#\s*(\d+)", text)

    if m:

        return f"#{m.group(1)}"

    return text.strip()


def match_equip_by_site_and_ho(

    site,
    ho_label,
    all_pumps

):

    candidates = [

        p for p in all_pumps

        if p["site"] == site

    ]

    # 설비명에 공백이 들어가는 방식이 제각각이라("인라인 #1" vs
    # "인라인#1") 공백을 다 지우고 비교해야 정확히 매칭된다.

    ho_label_norm = (

        re.sub(r"\s+", "", ho_label)

        if ho_label

        else ""

    )

    is_inline_query = "인라인" in ho_label_norm

    # 1차: "인라인" 여부까지 정확히 맞는 후보 우선.
    # (안 그러면 "#1"을 찾을 때 "인라인 #1"도 "#1"을 포함하고
    #  있어서 엉뚱한 인라인펌프로 잘못 매칭되는 문제가 있었다)

    for p in candidates:

        equip_norm = re.sub(

            r"\s+",
            "",
            p["equip"]

        )

        equip_is_inline = "인라인" in equip_norm

        if (

            ho_label_norm

            and ho_label_norm in equip_norm

            and is_inline_query == equip_is_inline

        ):

            return p

    # 2차: 인라인 여부는 무시하고 그냥 포함되는지만 확인

    for p in candidates:

        equip_norm = re.sub(

            r"\s+",
            "",
            p["equip"]

        )

        if ho_label_norm and ho_label_norm in equip_norm:

            return p

    # 정확히 안 맞으면, 사업장이 같고 설비가 1개뿐이면 그거라도 매칭

    if len(candidates) == 1:

        return candidates[0]

    return None


def parse_vibration_measure_upload(

    file_bytes,
    all_pumps

):

    df = pd.read_excel(
        io.BytesIO(file_bytes)
    )

    rows = []

    skipped = 0

    unmatched_labels = set()

    for _, r in df.iterrows():

        raw_site = r.get("기능위치내역")

        raw_equip = r.get("설비내역")

        raw_point = r.get("측정지점내역")

        raw_val = r.get("측정값")

        raw_date = r.get("측정일자")

        if (

            pd.isna(raw_site)

            or pd.isna(raw_equip)

            or pd.isna(raw_point)

            or pd.isna(raw_val)

            or pd.isna(raw_date)

        ):

            skipped += 1

            continue

        site = _extract_site_from_raw(
            raw_site
        )

        ho_label = _extract_ho_label(
            raw_equip
        )

        pump = match_equip_by_site_and_ho(

            site,
            ho_label,
            all_pumps

        )

        if pump is None:

            unmatched_labels.add(
                f"{site} {ho_label}"
            )

            skipped += 1

            continue

        m = re.search(

            r"진동측정\((펌프|모터)(부하|반부하)\s*(수직|수평|축)\)",

            str(raw_point)

        )

        if not m:

            skipped += 1

            continue

        pm_type, load_type, direction = m.groups()

        try:

            work_date = datetime.strptime(

                str(int(raw_date)),

                "%Y%m%d"

            ).strftime("%Y-%m-%d")

        except Exception:

            skipped += 1

            continue

        worker = (

            str(r.get("측정인", ""))

            if pd.notna(r.get("측정인"))

            else ""

        )

        rows.append(

            [
                work_date,
                site,
                pump["equip"],
                pm_type,
                load_type,
                direction,
                float(raw_val),
                (

                    str(r.get("평가코드내역", ""))

                    if pd.notna(r.get("평가코드내역"))

                    else ""

                ),
                anonymize_name(worker)
            ]

        )

    return rows, skipped, unmatched_labels


def apply_vibration_measure_upload(rows):

    # 예전엔 행 하나하나마다 파일을 열고 닫아서(safe_append_row
    # 반복호출), 그래프 엑셀처럼 한 번에 수천 행이 들어오면
    # 매우 느려서 타임아웃이 나는 문제가 있었다. 파일을 딱
    # 한 번만 열고 다 쓴 다음 한 번만 저장한다.

    if not rows:

        return

    with get_lock(VIBRATION_DB_PATH):

        wb = load_workbook(
            VIBRATION_DB_PATH
        )

        ws = wb["진동측정이력"]

        for row in rows:

            ws.append(
                row
            )

        wb.save(
            VIBRATION_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()

    log_audit(

        "월간 진동측정 업로드",

        "전체",

        f"{len(rows)}건 추가"

    )


# ============================================================
# 10-5. 진동 그래프 엑셀(과거 이력 일괄) 업로드
#
# "정기점검 실적목록"은 그달치만 들어있어서, 이전 몇 년치
# 이력을 한 번에 채우려면 매달 것을 일일이 올려야 하는
# 문제가 있었다. 사업소에서 별도로 관리하는 "진동 그래프"
# 엑셀(설비 1대당 시트 1개, 월별로 이미 집계된 값)을 올리면
# 과거 이력을 한 번에 채울 수 있게 한다.
# ============================================================

VIB_GRAPH_SITE_PREFIX = {

    "평촌": "평촌가압장",
    "팔도": "팔도가압장",
    "부곡": "부곡가압장",
    "밀양": "밀양정수장",
    "양산": "양산정수장",
    "나노": "나노가압장"

}

VIB_GRAPH_COL_ORDER = [

    ("모터", "반부하", "수직"),
    ("모터", "반부하", "수평"),
    ("모터", "반부하", "축"),
    ("모터", "부하", "수직"),
    ("모터", "부하", "수평"),
    ("펌프", "반부하", "수직"),
    ("펌프", "반부하", "수평"),
    ("펌프", "반부하", "축"),
    ("펌프", "부하", "수직"),
    ("펌프", "부하", "수평")

]


def parse_vib_graph_sheet_name(sheet_name):

    name = sheet_name.strip()

    # "평촌#1 (대체)" 같은 꼬리표는 떼어낸다

    name = re.sub(

        r"\s*\(.*?\)\s*$",

        "",

        name

    )

    site = None

    rest = name

    for prefix, site_full in VIB_GRAPH_SITE_PREFIX.items():

        if name.startswith(prefix):

            site = site_full

            rest = name[len(prefix):]

            break

    if site is None:

        return None, None

    ho_label = _extract_ho_label(
        rest
    )

    return site, ho_label


def parse_vibration_graph_workbook(

    file_bytes,
    all_pumps

):

    xl = pd.ExcelFile(
        io.BytesIO(file_bytes)
    )

    rows = []

    unmatched = set()

    sheet_month_counts = {}

    for sheet_name in xl.sheet_names:

        site, ho_label = parse_vib_graph_sheet_name(
            sheet_name
        )

        if site is None:

            unmatched.add(
                f"{sheet_name} (사업장 인식 실패)"
            )

            continue

        pump = match_equip_by_site_and_ho(

            site,
            ho_label,
            all_pumps

        )

        if pump is None:

            unmatched.add(

                f"{sheet_name} ({site} {ho_label})"

            )

            continue

        df = xl.parse(
            sheet_name,
            header=0
        )

        current_year = None

        month_count = 0

        for _, r in df.iterrows():

            raw_month = r.iloc[0]

            if pd.isna(raw_month):

                continue

            raw_month = str(raw_month).strip()

            m = re.match(

                r"^(?:(\d{2})\.)?(\d{2})월",

                raw_month

            )

            if not m:

                continue

            if m.group(1):

                current_year = 2000 + int(
                    m.group(1)
                )

            if current_year is None:

                # 첫 행부터 연도가 안 붙어있으면 매칭 불가라 건너뜀

                continue

            month_num = int(
                m.group(2)
            )

            work_date = (

                f"{current_year}-{month_num:02d}-01"

            )

            got_any_value = False

            for col_i, (pm, load, direction) in enumerate(

                VIB_GRAPH_COL_ORDER,

                start=1

            ):

                if col_i >= len(r):

                    continue

                val = r.iloc[col_i]

                if pd.isna(val):

                    continue

                try:

                    val = float(val)

                except (TypeError, ValueError):

                    continue

                got_any_value = True

                rows.append(

                    [
                        work_date,
                        pump["site"],
                        pump["equip"],
                        pm,
                        load,
                        direction,
                        val,
                        "",
                        ""
                    ]

                )

            if got_any_value:

                month_count += 1

        sheet_month_counts[sheet_name] = month_count

    return rows, unmatched, sheet_month_counts


# ============================================================
# 10-8. AI 고장예측 (회귀분석)
#
# 축적된 진동 이력에 실제 선형회귀를 적용해서, "이 추세대로
# 가면 위험 기준(8.5mm/s)에 언제쯤 도달할지"를 통계적으로
# 예측한다. 임계값 비교 같은 규칙기반이 아니라 실제 회귀식을
# 적합시키고 결정계수(R²)로 신뢰도까지 같이 보여준다.
# ============================================================

def predict_failure_date_regression(pump, df_vib):

    if (

        df_vib is None

        or df_vib.empty

        or "설비명" not in df_vib.columns

    ):

        return None

    equip_hist = df_vib[

        df_vib["설비명"] == pump["equip"]

    ].copy()

    if equip_hist.empty:

        return None

    equip_hist["_month"] = equip_hist[

        "측정일자"

    ].astype(str).str[:7]

    monthly_max = equip_hist.groupby(

        "_month"

    )["측정값"].max().sort_index()

    if len(monthly_max) < 3:

        return {

            "status": "insufficient_data",

            "months_count": len(monthly_max)

        }

    x = np.arange(
        len(monthly_max)
    )

    y = monthly_max.values.astype(float)

    # 실제 선형회귀 적합 (최소자승법)

    slope, intercept = np.polyfit(x, y, 1)

    y_pred = slope * x + intercept

    ss_res = np.sum(
        (y - y_pred) ** 2
    )

    ss_tot = np.sum(
        (y - np.mean(y)) ** 2
    )

    r_squared = (

        1 - (ss_res / ss_tot)

        if ss_tot > 0

        else 0.0

    )

    threshold = VIB_JUDGE_BAD

    result = {

        "status": "stable",
        "months": monthly_max.index.tolist(),
        "values": y.tolist(),
        "slope": float(slope),
        "intercept": float(intercept),
        "r_squared": float(r_squared),
        "trend_line": y_pred.tolist(),
        "seasonal_applied": False

    }

    # 계절성 반영: 데이터가 6개월 이상 쌓이면, 추세선만으로
    # 밀어보는 게 아니라 "이 달(예: 여름철)엔 원래 부하가
    # 몰려서 진동이 더 오르더라"는 월별 편차(잔차)까지 반영
    # 한다. 데이터가 적으면(6개월 미만) 그냥 순수 추세선으로
    # 폴백한다.

    month_numbers = [

        int(m.split("-")[1])

        for m in monthly_max.index

    ]

    seasonal_avg = {}

    if len(monthly_max) >= 6:

        residuals = y - y_pred

        by_month = {}

        for mn, r in zip(month_numbers, residuals):

            by_month.setdefault(
                mn, []
            ).append(r)

        raw_seasonal = {

            mn: float(np.mean(vals))

            for mn, vals in by_month.items()

        }

        # 잔차 평균이 0 근처가 되도록 중심을 맞춘다
        # (계절성분만 남기고, 추세는 이미 회귀식이 담당)

        center = np.mean(
            list(raw_seasonal.values())
        )

        seasonal_avg = {

            mn: v - center

            for mn, v in raw_seasonal.items()

        }

        result["seasonal_applied"] = True

        result["seasonal_by_month"] = seasonal_avg

    if slope <= 0.001 and not seasonal_avg:

        return result

    if seasonal_avg:

        # 한 달씩 미래로 걸어가면서, "추세값 + 그 달의 계절
        # 편차"가 위험기준을 넘는 첫 시점을 찾는다(대수적으로
        # 한 번에 풀 수 없어서 반복 탐색).

        last_year, last_month = (

            int(v)

            for v in monthly_max.index[-1].split("-")

        )

        cur_x = len(monthly_max) - 1

        cur_year, cur_month = last_year, last_month

        crossed_x = None

        for _ in range(60):

            cur_x += 1

            cur_month += 1

            if cur_month > 12:

                cur_month = 1

                cur_year += 1

            trend_val = slope * cur_x + intercept

            adj_val = trend_val + seasonal_avg.get(

                cur_month,
                0.0

            )

            if adj_val >= threshold:

                crossed_x = cur_x

                crossed_year_month = (
                    cur_year,
                    cur_month
                )

                break

        if crossed_x is None:

            return result

        result["status"] = "predicted"

        result["predicted_date"] = date(

            crossed_year_month[0],

            crossed_year_month[1],

            1

        )

        result["months_ahead"] = float(

            crossed_x - (len(monthly_max) - 1)

        )

        return result

    # 계절성 반영 불가(데이터 부족) → 기존처럼 순수 추세선으로

    x_threshold = (
        threshold - intercept
    ) / slope

    months_ahead = x_threshold - (
        len(monthly_max) - 1
    )

    result["status"] = "predicted"

    if months_ahead < 0:

        result["predicted_date"] = datetime.now().date()

        result["months_ahead"] = 0.0

    else:

        result["predicted_date"] = (

            datetime.now().date()

            +
            timedelta(
                days=int(months_ahead * 30.4)
            )

        )

        result["months_ahead"] = float(
            months_ahead
        )

    return result


FAVORITE_DB_PATH = "Pump_Favorite_DB.xlsx"

RECENT_VIEW_DB_PATH = "Pump_RecentView_DB.xlsx"


def ensure_favorite_recent_dbs_exist():

    ensure_excel_file(

        FAVORITE_DB_PATH,

        "즐겨찾기",

        ["설비명", "등록일시"]

    )

    ensure_excel_file(

        RECENT_VIEW_DB_PATH,

        "최근조회",

        ["설비명", "조회일시"]

    )


CHECKLIST_TEMPLATE_DB_PATH = "Pump_ChecklistTemplate_DB.xlsx"

DEFAULT_CHECKLIST_TEMPLATES = {

    "원심펌프용": [

        "베어링 이상 소음 여부",
        "축봉부(메커니컬씰) 누유 여부",
        "커플링 정렬 상태",
        "임펠러 마모·부식 상태",
        "케이싱 부식·크랙 여부",
        "흡입·토출 배관 이상진동",
        "모터 온도(과열) 확인",
        "윤활유·그리스 상태",
        "기초볼트 풀림 여부"

    ],

    "제어반용": [

        "외함 부식·손상 여부",
        "단자대 접속 상태(풀림·변색)",
        "차단기·계전기 동작 이상",
        "절연저항 측정값",
        "접지선 연결 상태",
        "내부 발열·이취 여부",
        "표시등·경보 정상 동작"

    ],

    "배수펌프용": [

        "임펠러 이물질 끼임 여부",
        "케이블·전선 피복 손상",
        "부력스위치(플로트) 동작",
        "펌프 하우징 부식 상태",
        "체크밸브 역류 여부",
        "설치 앵커 고정 상태"

    ]

}


TRASH_DB_PATH = "Pump_Trash_DB.xlsx"

TRASH_RETENTION_DAYS = 30


ALIGNMENT_DB_PATH = "Pump_Alignment_DB.xlsx"

VIB_ANALYSIS_DB_PATH = "Pump_VibAnalysis_DB.xlsx"

EFFICIENCY_REPORT_DB_PATH = "Pump_EfficiencyReport_DB.xlsx"


def ensure_specialized_report_dbs_exist():

    # 연간 축정렬(얼라인먼트) 보고서용 데이터.
    # 실제 회사 양식을 받으면 컬럼을 그에 맞춰 조정할 예정이라,
    # 지금은 업계에서 일반적으로 쓰는 항목(오프셋·각도·TIR)
    # 기준으로 임시 구성해뒀다.

    ensure_excel_file(

        ALIGNMENT_DB_PATH,

        "축정렬이력",

        [
            "측정일자",
            "사업장",
            "설비명",
            "측정자",
            "커플링구분",
            "수직오프셋(mm)",
            "수평오프셋(mm)",
            "수직각도(mm/100mm)",
            "수평각도(mm/100mm)",
            "판정",
            "조치사항"
        ]

    )

    # 연간 진동분석 보고서(정밀 스펙트럼 분석)용 데이터.
    # 월간 진동측정(overall 값 위주)과 달리, 방향별·베어링
    # 결함주파수 성분까지 다루는 정밀진단 성격이다.

    ensure_excel_file(

        VIB_ANALYSIS_DB_PATH,

        "진동분석이력",

        [
            "측정일자",
            "사업장",
            "설비명",
            "측정자",
            "측정위치",
            "Overall(mm/s)",
            "1X성분(mm/s)",
            "2X성분(mm/s)",
            "베어링결함성분",
            "종합판정",
            "권고사항"
        ]

    )

    # 효율진단 보고서용 데이터.

    ensure_excel_file(

        EFFICIENCY_REPORT_DB_PATH,

        "효율진단이력",

        [
            "측정일자",
            "사업장",
            "설비명",
            "측정자",
            "정격효율(%)",
            "실측효율(%)",
            "정격유량(m3/h)",
            "실측유량(m3/h)",
            "정격양정(m)",
            "실측양정(m)",
            "전동기효율(%)",
            "원인분석",
            "개선권고"
        ]

    )


def ensure_trash_db_exists():

    ensure_excel_file(

        TRASH_DB_PATH,

        "휴지통",

        [
            "삭제일시",
            "삭제자",
            "종류",
            "이름",
            "원본시트",
            "데이터JSON"
        ]

    )


def move_to_trash(category, item_name, sheet_name, row_dict):

    # 바로 지우지 않고, 나중에 되돌릴 수 있게 휴지통에 먼저
    # 옮겨 담는다. 30일이 지나면 자동으로 완전 삭제된다.

    import json as _json

    safe_append_row(

        TRASH_DB_PATH,

        "휴지통",

        [
            datetime.now().strftime("%Y-%m-%d %H:%M"),
            st.session_state.get("user_name", "알수없음"),
            category,
            item_name,
            sheet_name,
            _json.dumps(row_dict, ensure_ascii=False, default=str)
        ]

    )


def purge_old_trash():

    # 30일 지난 휴지통 항목은 화면에서도 안 보이고 실제로도
    # 지운다(공간 관리).

    df = read_excel(
        TRASH_DB_PATH,
        "휴지통"
    )

    if df.empty:

        return

    cutoff = datetime.now() - timedelta(
        days=TRASH_RETENTION_DAYS
    )

    with get_lock(TRASH_DB_PATH):

        wb = load_workbook(
            TRASH_DB_PATH
        )

        ws = wb["휴지통"]

        rows_to_delete = []

        for r in ws.iter_rows(min_row=2):

            try:

                deleted_at = datetime.strptime(

                    str(r[0].value),

                    "%Y-%m-%d %H:%M"

                )

                if deleted_at < cutoff:

                    rows_to_delete.append(
                        r[0].row
                    )

            except Exception:

                continue

        for row_num in sorted(

            rows_to_delete,
            reverse=True

        ):

            ws.delete_rows(
                row_num
            )

        if rows_to_delete:

            wb.save(
                TRASH_DB_PATH
            )

        wb.close()

    if rows_to_delete:

        _read_excel_cached.clear()


def get_trash_items():

    purge_old_trash()

    df = read_excel(
        TRASH_DB_PATH,
        "휴지통"
    )

    return df


def restore_from_trash(row_index, category, item_name, sheet_name, row_dict):

    import json as _json

    if category == "설비":

        add_equipment({

            "site": row_dict.get("사업장"),
            "equip": row_dict.get("설비명"),
            "maker": row_dict.get("제조사"),
            "model": row_dict.get("모델명"),
            "hp": row_dict.get("정격출력(HP)"),
            "head": row_dict.get("정격양정(m)"),
            "flow": row_dict.get("정격유량(m3/h)"),
            "rpm": row_dict.get("회전수(RPM)"),
            "build_date": row_dict.get("준공일"),
            "op_hours": row_dict.get("누적운전시간"),
            "기준진동": row_dict.get("기준진동"),
            "기준효율": row_dict.get("기준효율")

        })

    elif category == "노하우":

        safe_append_row(

            KNOWHOW_DB_PATH,

            sheet_name,

            list(row_dict.values())

        )

    elif category == "체크리스트템플릿":

        # 체크리스트는 템플릿명+항목리스트 형태로 저장

        save_checklist_template(

            row_dict.get("템플릿명"),

            row_dict.get("항목목록", [])

        )

    # 복원했으면 휴지통에서는 제거

    with get_lock(TRASH_DB_PATH):

        wb = load_workbook(
            TRASH_DB_PATH
        )

        ws = wb["휴지통"]

        for r in list(

            ws.iter_rows(min_row=2)

        ):

            if (

                r[3].value == item_name

                and r[2].value == category

            ):

                ws.delete_rows(
                    r[0].row
                )

                break

        wb.save(
            TRASH_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()

    log_audit(

        "휴지통에서 복원",

        item_name,

        category

    )


def permanently_delete_from_trash(category, item_name):

    with get_lock(TRASH_DB_PATH):

        wb = load_workbook(
            TRASH_DB_PATH
        )

        ws = wb["휴지통"]

        for r in list(

            ws.iter_rows(min_row=2)

        ):

            if (

                r[3].value == item_name

                and r[2].value == category

            ):

                ws.delete_rows(
                    r[0].row
                )

                break

        wb.save(
            TRASH_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()


def ensure_checklist_template_db_exists():

    ensure_excel_file(

        CHECKLIST_TEMPLATE_DB_PATH,

        "체크리스트템플릿",

        ["템플릿명", "순번", "점검항목"]

    )


def seed_default_checklist_templates():

    df = read_excel(

        CHECKLIST_TEMPLATE_DB_PATH,

        "체크리스트템플릿"

    )

    if not df.empty:

        return

    with get_lock(CHECKLIST_TEMPLATE_DB_PATH):

        wb = load_workbook(
            CHECKLIST_TEMPLATE_DB_PATH
        )

        ws = wb["체크리스트템플릿"]

        for template_name, items in (

            DEFAULT_CHECKLIST_TEMPLATES.items()

        ):

            for idx, item in enumerate(items, start=1):

                ws.append(

                    [template_name, idx, item]

                )

        wb.save(
            CHECKLIST_TEMPLATE_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()


def get_checklist_template_names():

    df = read_excel(

        CHECKLIST_TEMPLATE_DB_PATH,

        "체크리스트템플릿"

    )

    if df.empty:

        return []

    return df["템플릿명"].drop_duplicates().tolist()


def get_checklist_template_items(template_name):

    df = read_excel(

        CHECKLIST_TEMPLATE_DB_PATH,

        "체크리스트템플릿"

    )

    if df.empty:

        return []

    matched = df[

        df["템플릿명"] == template_name

    ].sort_values("순번")

    return matched["점검항목"].tolist()


def save_checklist_template(template_name, items):

    with get_lock(CHECKLIST_TEMPLATE_DB_PATH):

        wb = load_workbook(
            CHECKLIST_TEMPLATE_DB_PATH
        )

        ws = wb["체크리스트템플릿"]

        # 같은 이름의 기존 템플릿은 지우고 새로 저장

        for r in list(

            ws.iter_rows(min_row=2)

        ):

            if r[0].value == template_name:

                ws.delete_rows(
                    r[0].row
                )

        for idx, item in enumerate(items, start=1):

            ws.append(

                [template_name, idx, item]

            )

        wb.save(
            CHECKLIST_TEMPLATE_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()


def delete_checklist_template(template_name):

    with get_lock(CHECKLIST_TEMPLATE_DB_PATH):

        wb = load_workbook(
            CHECKLIST_TEMPLATE_DB_PATH
        )

        ws = wb["체크리스트템플릿"]

        for r in list(

            ws.iter_rows(min_row=2)

        ):

            if r[0].value == template_name:

                ws.delete_rows(
                    r[0].row
                )

        wb.save(
            CHECKLIST_TEMPLATE_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()


def is_favorite_equipment(equip_name):

    df = read_excel(
        FAVORITE_DB_PATH,
        "즐겨찾기"
    )

    if df.empty:

        return False

    return equip_name in df["설비명"].values


def toggle_favorite_equipment(equip_name):

    if is_favorite_equipment(equip_name):

        with get_lock(FAVORITE_DB_PATH):

            wb = load_workbook(
                FAVORITE_DB_PATH
            )

            ws = wb["즐겨찾기"]

            for r in list(

                ws.iter_rows(min_row=2)

            ):

                if r[0].value == equip_name:

                    ws.delete_rows(
                        r[0].row
                    )

                    break

            wb.save(
                FAVORITE_DB_PATH
            )

            wb.close()

    else:

        safe_append_row(

            FAVORITE_DB_PATH,

            "즐겨찾기",

            [
                equip_name,
                datetime.now().strftime(
                    "%Y-%m-%d %H:%M"
                )
            ]

        )

    _read_excel_cached.clear()


def get_favorite_equipment_list():

    df = read_excel(
        FAVORITE_DB_PATH,
        "즐겨찾기"
    )

    if df.empty:

        return []

    return df["설비명"].tolist()


def log_recent_view(equip_name):

    df = read_excel(
        RECENT_VIEW_DB_PATH,
        "최근조회"
    )

    with get_lock(RECENT_VIEW_DB_PATH):

        wb = load_workbook(
            RECENT_VIEW_DB_PATH
        )

        ws = wb["최근조회"]

        # 같은 설비의 예전 기록은 지우고 맨 위로(최신으로) 갱신

        for r in list(

            ws.iter_rows(min_row=2)

        ):

            if r[0].value == equip_name:

                ws.delete_rows(
                    r[0].row
                )

        ws.append(

            [
                equip_name,
                datetime.now().strftime(
                    "%Y-%m-%d %H:%M"
                )
            ]

        )

        # 최근 20개만 유지(파일이 무한정 커지지 않게)

        all_rows = list(
            ws.iter_rows(min_row=2)
        )

        if len(all_rows) > 20:

            for r in all_rows[:-20]:

                ws.delete_rows(
                    r[0].row
                )

        wb.save(
            RECENT_VIEW_DB_PATH
        )

        wb.close()

    _read_excel_cached.clear()


def get_recent_view_list(limit=5):

    df = read_excel(
        RECENT_VIEW_DB_PATH,
        "최근조회"
    )

    if df.empty:

        return []

    return df["설비명"].tolist()[::-1][:limit]


PREDICTION_LOG_DB_PATH = "Pump_PredictionLog_DB.xlsx"


def ensure_prediction_log_exists():

    ensure_excel_file(

        PREDICTION_LOG_DB_PATH,

        "예측이력",

        [
            "기록일",
            "설비명",
            "예측위험도달일",
            "예측시점진동값",
            "R2",
            "계절성반영"
        ]

    )


ensure_prediction_log_exists()

ensure_favorite_recent_dbs_exist()

ensure_specialized_report_dbs_exist()

ensure_trash_db_exists()

ensure_checklist_template_db_exists()

seed_default_checklist_templates()


def log_prediction(pump, prediction):

    # 예측이 얼마나 맞았는지 나중에 검증하려면, "그때 뭐라고
    # 예측했었는지"가 남아있어야 한다. 같은 날 같은 설비로
    # 중복 기록하지 않도록 확인 후 남긴다.

    if (

        prediction is None

        or prediction.get("status") != "predicted"

    ):

        return

    today_str = datetime.now().strftime("%Y-%m-%d")

    df_log = read_excel(

        PREDICTION_LOG_DB_PATH,

        "예측이력"

    )

    if not df_log.empty:

        already = df_log[

            (df_log["설비명"] == pump["equip"])

            &
            (df_log["기록일"].astype(str) == today_str)

        ]

        if not already.empty:

            return

    safe_append_row(

        PREDICTION_LOG_DB_PATH,

        "예측이력",

        [
            today_str,
            pump["equip"],
            str(prediction["predicted_date"]),
            prediction["values"][-1],
            round(prediction["r_squared"], 3),
            "예" if prediction.get("seasonal_applied") else "아니오"
        ]

    )


def verify_past_predictions(all_pumps, df_vib):

    # 예측일이 이미 지난 과거 예측들을 모아서, 실제로 그
    # 시점에 위험기준(8.5)에 도달했었는지 확인한다.

    df_log = read_excel(

        PREDICTION_LOG_DB_PATH,

        "예측이력"

    )

    if df_log.empty:

        return pd.DataFrame()

    today = datetime.now().date()

    rows = []

    for _, log_row in df_log.iterrows():

        try:

            pred_date = datetime.strptime(

                str(log_row["예측위험도달일"]),

                "%Y-%m-%d"

            ).date()

        except Exception:

            continue

        if pred_date > today:

            continue

        equip_name = log_row["설비명"]

        if (

            df_vib is None

            or df_vib.empty

            or "설비명" not in df_vib.columns

        ):

            actual_status = "데이터없음"

        else:

            equip_hist = df_vib[

                df_vib["설비명"] == equip_name

            ]

            window_start = (

                pred_date - timedelta(days=20)

            ).strftime("%Y-%m")

            window_end = (

                pred_date + timedelta(days=20)

            ).strftime("%Y-%m")

            near = equip_hist[

                equip_hist["측정일자"].astype(
                    str
                ).str[:7].between(

                    window_start,
                    window_end

                )

            ]

            if near.empty:

                actual_status = "데이터없음"

            else:

                max_val = near["측정값"].max()

                actual_status = (

                    f"실측 {max_val:.1f} "
                    f"({'적중' if max_val >= VIB_JUDGE_BAD else '미도달'})"

                )

        rows.append(

            {
                "설비명": equip_name,
                "예측시점": log_row["기록일"],
                "예측위험도달일": log_row["예측위험도달일"],
                "당시 R²": log_row["R2"],
                "검증결과": actual_status
            }

        )

    return pd.DataFrame(rows)


def build_failure_prediction_chart_fig(pump, pred):

    fig, ax = plt.subplots(
        figsize=(7, 3.6)
    )

    months = pred["months"]

    values = pred["values"]

    ax.plot(

        months,
        values,
        marker="o",
        linewidth=2,
        color="#087ea4",
        label="실측 진동값(월별 최댓값)"

    )

    ax.plot(

        months,
        pred["trend_line"],
        linestyle="--",
        color="#e8590c",
        linewidth=2,
        label="회귀 추세선"

    )

    ax.axhline(

        VIB_JUDGE_BAD,

        color="#c62828",

        linestyle=":",

        linewidth=1.5,

        label=f"위험 기준({VIB_JUDGE_BAD})"

    )

    if pred["status"] == "predicted":

        ax.set_title(

            f"{pump['equip']} — 예측 위험도달일: "
            f"{pred['predicted_date']} "
            f"(R²={pred['r_squared']:.2f})"

        )

    else:

        ax.set_title(

            f"{pump['equip']} — 추세 안정적 "
            f"(R²={pred['r_squared']:.2f})"

        )

    ax.set_ylabel(
        "진동값 (mm/s, rms)"
    )

    ax.legend(
        fontsize=8
    )

    ax.grid(
        alpha=0.2
    )

    fig.tight_layout()

    return fig


# ============================================================
# 10-9. AI(LLM) 종합의견 자동 생성
#
# 미리 정해둔 문장 틀에 숫자만 끼워넣는 방식이 아니라, 실제
# Claude API를 호출해서 데이터를 보고 코멘트를 새로 작성하게
# 한다. Streamlit Cloud의 Secrets에 ANTHROPIC_API_KEY를
# 등록해야 동작한다.
# ============================================================

def call_claude_for_commentary(prompt_text):

    try:

        api_key = st.secrets.get(
            "ANTHROPIC_API_KEY"
        )

    except Exception:

        api_key = None

    if not api_key:

        return None, (

            "ANTHROPIC_API_KEY가 설정되지 않았습니다. "
            "Streamlit Cloud 앱 설정의 Secrets에 "
            "ANTHROPIC_API_KEY를 추가해주세요."

        )

    try:

        resp = requests.post(

            "https://api.anthropic.com/v1/messages",

            headers={

                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json"

            },

            json={

                "model": "claude-haiku-4-5-20251001",
                "max_tokens": 500,
                "messages": [

                    {
                        "role": "user",
                        "content": prompt_text
                    }

                ]

            },

            timeout=30

        )

        resp.raise_for_status()

        data = resp.json()

        text = "".join(

            block.get("text", "")

            for block in data.get("content", [])

            if block.get("type") == "text"

        )

        return text, None

    except Exception as e:

        return None, f"AI 호출 중 오류가 발생했습니다: {e}"


def build_align_excel_template():

    wb = Workbook()

    ws = wb.active

    ws.title = "축정렬입력"

    headers = [

        "각도오차_수평", "각도오차_수직",
        "옵셋오차_수평", "옵셋오차_수직",
        "전면변위_수평", "전면변위_수직",
        "후면변위_수평", "후면변위_수직",
        "판정_수평", "판정_수직",
        "측정장비", "종합의견"

    ]

    ws.append(headers)

    ws.append([

        "-0.01/100", "+0.02/100",
        "-0.05", "-0.03",
        "-0.09", "+0.04",
        "-0.14", "+0.11",
        "양호", "양호",
        "Fixurlsser EVO",
        "전체 정렬 상태는 ISO/ANSI 회전기계 정렬 허용 기준 "
        "내에 있으며 즉각적인 재정렬이 필요한 수준은 아님"

    ])

    buf = io.BytesIO()

    wb.save(buf)

    return buf.getvalue()


def parse_align_excel(file_bytes):

    df = pd.read_excel(
        io.BytesIO(file_bytes)
    )

    df = df.fillna("")

    row = df.iloc[0]

    return {

        "각도오차_수평": row.get("각도오차_수평", ""),
        "각도오차_수직": row.get("각도오차_수직", ""),
        "옵셋오차_수평": row.get("옵셋오차_수평", ""),
        "옵셋오차_수직": row.get("옵셋오차_수직", ""),
        "전면변위_수평": row.get("전면변위_수평", ""),
        "전면변위_수직": row.get("전면변위_수직", ""),
        "후면변위_수평": row.get("후면변위_수평", ""),
        "후면변위_수직": row.get("후면변위_수직", ""),
        "판정_수평": row.get("판정_수평", ""),
        "판정_수직": row.get("판정_수직", ""),
        "측정장비": row.get("측정장비", ""),
        "종합의견": row.get("종합의견", "")

    }


def build_va_excel_template():

    wb = Workbook()

    ws = wb.active

    ws.title = "진동분석입력"

    ws.append(

        ["측정점", "Overall(mm/s)", "판정", "pk-pk(µm)", "주파수", "키워드"]

    )

    for _point in [

        "모터 부하 수직", "모터 부하 수평",
        "모터 반부하 수직", "모터 반부하 수평", "모터 반부하 축",
        "펌프 부하 수직", "펌프 부하 수평",
        "펌프 반부하 수직", "펌프 반부하 수평", "펌프 반부하 축"

    ]:

        ws.append([_point, 0.0, "A(양호)", 0.0, "", ""])

    buf = io.BytesIO()

    wb.save(buf)

    return buf.getvalue()


def parse_va_excel(file_bytes):

    df = pd.read_excel(
        io.BytesIO(file_bytes)
    )

    df = df.fillna("")

    points = {}

    for _, row in df.iterrows():

        points[row.get("측정점", "")] = {

            "overall_rms": row.get("Overall(mm/s)", ""),
            "판정": row.get("판정", ""),
            "pk_pk": row.get("pk-pk(µm)", ""),
            "주파수": row.get("주파수", ""),
            "키워드": row.get("키워드", "")

        }

    return points


def build_eff_excel_template():

    wb = Workbook()

    ws = wb.active

    ws.title = "효율진단입력"

    ws.append(

        ["정격효율(%)", "실측효율(%)", "정격유량(m3/h)",
         "실측유량(m3/h)", "원인분석및개선권고"]

    )

    ws.append([

        82, 73, 1250, 1180,
        "정격 대비 실측효율이 9%p 낮음. 임펠러 마모가 의심되며 "
        "정밀점검 및 교체를 권고함."

    ])

    buf = io.BytesIO()

    wb.save(buf)

    return buf.getvalue()


def parse_eff_excel(file_bytes):

    df = pd.read_excel(
        io.BytesIO(file_bytes)
    )

    df = df.fillna("")

    row = df.iloc[0]

    return (

        row.get("정격효율(%)", ""),
        row.get("실측효율(%)", ""),
        row.get("정격유량(m3/h)", ""),
        row.get("실측유량(m3/h)", ""),
        row.get("원인분석및개선권고", "")

    )


def call_claude_vision_analysis(images_with_labels, instruction):

    # images_with_labels: [(라벨, 이미지bytes, mime타입), ...]
    # 사진(진동 스펙트럼, 축정렬 측정기 화면, 효율진단 데이터
    # 등)을 실제 Claude에게 보여주고 직접 읽어서 분석하게
    # 한다. 문자로 값을 일일이 입력할 필요 없이, 현장에서
    # 찍은 사진 그대로 올리면 된다.

    try:

        api_key = st.secrets.get(
            "ANTHROPIC_API_KEY"
        )

    except Exception:

        api_key = None

    if not api_key:

        return None, (

            "ANTHROPIC_API_KEY가 설정되지 않았습니다. "
            "Streamlit Secrets에 등록해주세요."

        )

    content_blocks = []

    for label, img_bytes, mime_type in images_with_labels:

        content_blocks.append(

            {
                "type": "text",
                "text": f"[{label}]"
            }

        )

        content_blocks.append(

            {
                "type": "image",
                "source": {

                    "type": "base64",

                    "media_type": mime_type,

                    "data": base64.b64encode(
                        img_bytes
                    ).decode("utf-8")

                }
            }

        )

    content_blocks.append(

        {
            "type": "text",
            "text": instruction
        }

    )

    try:

        resp = requests.post(

            "https://api.anthropic.com/v1/messages",

            headers={

                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json"

            },

            json={

                "model": "claude-sonnet-5",
                "max_tokens": 8000,
                "messages": [

                    {
                        "role": "user",
                        "content": content_blocks
                    }

                ]

            },

            timeout=90

        )

        resp.raise_for_status()

        data = resp.json()

        text = "".join(

            block.get("text", "")

            for block in data.get("content", [])

            if block.get("type") == "text"

        )

        return text, None

    except Exception as e:

        return None, f"AI 이미지 분석 중 오류가 발생했습니다: {e}"


def extract_json_from_ai_response(text):

    # AI가 ```json ... ``` 코드블록으로 감싸서 답하거나,
    # 앞뒤에 설명 문장을 붙이는 경우가 있어서, 그 안에서
    # 실제 JSON 객체만 안전하게 뽑아낸다.

    import json as _json
    import re as _re

    _match = _re.search(

        r"\{.*\}",

        text,

        _re.DOTALL

    )

    if not _match:

        return None

    try:

        return _json.loads(
            _match.group(0)
        )

    except Exception:

        return None


def ocr_extract_alignment_values(image_bytes):

    # API 없이도 쓸 수 있도록, 축정렬 측정기 화면(숫자가
    # 또렷하게 찍힌 디지털 표시판)에서 글자만 읽어내는
    # 방식이다. 스펙트럼 그래프처럼 "곡선을 해석"하는 게
    # 아니라 "찍혀있는 숫자를 읽는" 거라 API 없이도 가능하다.
    # 다만 100% 정확하진 않으니, 읽어낸 값은 사람이 한번
    # 확인하고 필요하면 고치도록 폼에 미리 채워 넣는 용도다.

    try:

        import pytesseract

        from PIL import Image as _PILImage

        img = _PILImage.open(
            io.BytesIO(image_bytes)
        )

        # 흑백 변환 + 확대하면 작은 숫자 인식률이 올라간다

        img = img.convert("L")

        _scale = 2

        img = img.resize(

            (img.width * _scale, img.height * _scale),

            _PILImage.LANCZOS

        )

        raw_text = pytesseract.image_to_string(

            img,

            lang="eng",

            config="--psm 6"

        )

    except Exception as e:

        return None, f"OCR 실행 중 오류: {e}"

    # "+0.02/100", "-0.01 /100" 같은 각도오차 패턴 (보통
    # 화면에 2개, 위가 수직 아래가 수평인 경우가 많다)

    angle_matches = re.findall(

        r"([+\-]?\d+\.\d+)\s*/\s*100",

        raw_text

    )

    # 나머지 소수점 값들 (옵셋오차·전면변위·후면변위 후보)

    all_decimals = re.findall(

        r"[+\-]\d+\.\d+",

        raw_text

    )

    # angle_matches에 이미 들어간 값은 all_decimals에서도
    # 중복으로 잡히므로, 부호+숫자 기준으로 걸러낸다

    angle_set = set(

        a.replace(" ", "")

        for a in angle_matches

    )

    remaining = [

        d for d in all_decimals

        if d not in angle_set

    ]

    def _get(lst, idx):

        return lst[idx] if idx < len(lst) else ""

    result = {

        "raw_text": raw_text,

        "각도오차_수직": _get(angle_matches, 0),
        "각도오차_수평": _get(angle_matches, 1),
        "옵셋오차_수직": _get(remaining, 0),
        "전면변위_수직": _get(remaining, 1),
        "후면변위_수직": _get(remaining, 2),
        "옵셋오차_수평": _get(remaining, 3),
        "전면변위_수평": _get(remaining, 4),
        "후면변위_수평": _get(remaining, 5)

    }

    return result, None


def ocr_extract_vibration_numbers(image_bytes):

    # 진동측정 소프트웨어 화면 상단엔 보통 "총진동:2.41mm/s
    # rms"나 "곡선:0mSec(0.000) 최대진폭:2.4" 같은 숫자가
    # 텍스트로 찍혀있다. 스펙트럼 곡선의 모양(불평형이냐
    # 베어링결함이냐)은 해석 못 해도, 이 찍힌 숫자만큼은
    # OCR로 읽어낼 수 있다. 사진 한 장에 수직/수평/축 여러
    # 방향이 같이 있으면 값이 여러 개 나올 수 있어서 리스트로
    # 반환한다.

    try:

        import pytesseract

        from PIL import Image as _PILImage

        img = _PILImage.open(
            io.BytesIO(image_bytes)
        )

        img = img.convert("L")

        _scale = 2

        img = img.resize(

            (img.width * _scale, img.height * _scale),

            _PILImage.LANCZOS

        )

        raw_text = pytesseract.image_to_string(

            img,

            lang="kor+eng",

            config="--psm 6"

        )

    except Exception as e:

        return None, f"OCR 실행 중 오류: {e}"

    overall_matches = re.findall(

        r"총\s*진동\S*\s*[:：]?\s*([\d.]+)",

        raw_text

    )

    peak_matches = re.findall(

        r"최대\s*진폭\S*\s*[:：]?\s*([\d.]+)",

        raw_text

    )

    freq_matches = re.findall(

        r"주파수\S*\s*[:：]?\s*([\d.eE+\-]+)\s*Hz",

        raw_text

    )

    return {

        "raw_text": raw_text,
        "총진동_후보": overall_matches,
        "최대진폭_후보": peak_matches,
        "주파수_후보": freq_matches

    }, None


def guess_image_mime(filename):

    _lower = filename.lower()

    if _lower.endswith(".png"):

        return "image/png"

    if _lower.endswith(".webp"):

        return "image/webp"

    return "image/jpeg"


def generate_ai_commentary_for_equipment(

    pump,
    result,
    prediction

):

    pred_text = "예측 데이터 부족"

    if prediction:

        if prediction["status"] == "predicted":

            pred_text = (

                f"회귀분석 결과 약 {prediction['predicted_date']}"
                f"경 위험기준(8.5mm/s) 도달 예상 "
                f"(신뢰도 R²={prediction['r_squared']:.2f})"

            )

        elif prediction["status"] == "stable":

            pred_text = (

                f"진동 추세 안정적 "
                f"(R²={prediction['r_squared']:.2f})"

            )

    prompt = (

        "당신은 수도설비 정비 전문가입니다. 아래 데이터를 "
        "보고, 정비 담당자에게 전달할 종합의견을 한국어로 "
        "2~3문장, 실무적인 어투로 작성해주세요. 과장하지 "
        "말고 데이터에 근거해서만 작성하세요.\n\n"
        f"- 설비명: {pump['equip']} ({pump['site']})\n"
        f"- CBM Score: {result['점수']}점 ({result['등급']}등급)\n"
        f"- 현재 상태: {result['상태']}\n"
        f"- 현재 진동값: {result['진동']:.1f} mm/s\n"
        f"- 현재 효율: {result['효율']:.1f}%\n"
        f"- AI 회귀분석 예측: {pred_text}\n"

    )

    return call_claude_for_commentary(
        prompt
    )


# ============================================================
# 10-10. 구글시트 백업/복원
#
# 로컬 엑셀 파일은 서버가 재배포되면 통째로 사라질 수 있다.
# 완전히 구글시트로 갈아엎으면(모든 읽기/쓰기를 API로) 위험
# 부담이 크므로, 대신 "지금 상태를 구글시트에 통째로 복사"
# (백업)하고 "구글시트에서 다시 불러오기"(복원)하는 안전망만
# 둔다. 평소엔 그대로 빠른 로컬 엑셀로 동작한다.
# ============================================================

GSHEET_BACKUP_TARGETS = [

    (DB_FILE_PATH, "진단이력"),
    (OVERHAUL_DB_PATH, "오버홀이력"),
    (EQUIP_DB_PATH, "설비마스터"),
    (CONSUMABLES_DB_PATH, "소모품재고"),
    (VIBRATION_DB_PATH, "진동측정이력"),
    (AUDIT_LOG_PATH, "감사로그"),
    (KNOWHOW_DB_PATH, "노하우DB"),
    (KPI_DB_PATH, "KPI실적"),
    (DESIGN_UNITPRICE_DB_PATH, "일위대가라이브러리"),
    (DESIGN_GUIDELINE_DB_PATH, "공종가이드라인"),
    (DESIGN_PROJECT_DB_PATH, "설계공사"),
    (DESIGN_PROJECT_ITEM_DB_PATH, "설계공사상세")

]


def get_gsheet_client():

    try:

        creds_info = st.secrets.get(
            "gcp_service_account"
        )

    except Exception:

        creds_info = None

    if not creds_info:

        return None, (

            "구글 서비스계정 정보가 없습니다. Streamlit "
            "Secrets에 [gcp_service_account]를 등록해주세요."

        )

    try:

        from google.oauth2.service_account import Credentials

        scopes = [

            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive"

        ]

        creds = Credentials.from_service_account_info(

            dict(creds_info),

            scopes=scopes

        )

        client = gspread.authorize(
            creds
        )

        return client, None

    except Exception as e:

        return None, f"구글 인증 실패: {e}"


def get_gsheet_url():

    try:

        return st.secrets.get(
            "GSHEET_BACKUP_URL"
        )

    except Exception:

        return None


LAST_BACKUP_FLAG_PATH = "Pump_LastBackup.txt"


def _record_last_backup_time():

    try:

        with open(

            LAST_BACKUP_FLAG_PATH,

            "w",

            encoding="utf-8"

        ) as f:

            f.write(
                datetime.now().isoformat()
            )

    except Exception:

        pass


def get_days_since_last_backup():

    if not os.path.exists(
        LAST_BACKUP_FLAG_PATH
    ):

        return None

    try:

        with open(

            LAST_BACKUP_FLAG_PATH,

            encoding="utf-8"

        ) as f:

            last_time = datetime.fromisoformat(
                f.read().strip()
            )

        return (

            datetime.now() - last_time

        ).days

    except Exception:

        return None


def backup_all_to_gsheet():

    client, error = get_gsheet_client()

    if error:

        return False, error

    sheet_url = get_gsheet_url()

    if not sheet_url:

        return False, (

            "백업할 구글시트 URL이 없습니다. Streamlit "
            "Secrets에 GSHEET_BACKUP_URL을 등록해주세요."

        )

    try:

        doc = client.open_by_url(
            sheet_url
        )

        backed_up = 0

        verification_issues = []

        for local_path, sheet_name in GSHEET_BACKUP_TARGETS:

            df = read_excel(

                local_path,
                sheet_name

            )

            try:

                ws = doc.worksheet(
                    sheet_name
                )

                ws.clear()

            except gspread.WorksheetNotFound:

                ws = doc.add_worksheet(

                    title=sheet_name,

                    rows=max(len(df) + 10, 100),

                    cols=max(len(df.columns) + 2, 10)

                        if not df.empty

                        else 10

                )

            if df.empty:

                ws.update(
                    [list(df.columns)]

                    if len(df.columns) > 0

                    else [["(빈 시트)"]]

                )

            else:

                data = (

                    [df.columns.tolist()]

                    +
                    df.astype(str).values.tolist()

                )

                ws.update(
                    data
                )

            backed_up += 1

            # 백업 직후 실제로 구글시트에서 다시 읽어서
            # 행 개수가 일치하는지 검증한다("눌렀는데 실제로
            # 다 옮겨졌는지" 확신할 수 있게).

            uploaded_rows = len(
                ws.get_all_values()
            ) - 1

            local_rows = len(df)

            if uploaded_rows != local_rows:

                verification_issues.append(

                    f"{sheet_name}: 로컬 {local_rows}행 vs "
                    f"업로드확인 {uploaded_rows}행 불일치"

                )

        if verification_issues:

            _record_last_backup_time()

            return True, (

                f"{backed_up}개 시트를 백업했지만, 일부 검증에서 "
                "행 개수가 안 맞았습니다:\n"
                +
                "\n".join(verification_issues)

            )

        _record_last_backup_time()

        return True, (

            f"{backed_up}개 시트를 백업했고, 전부 행 개수까지 "
            "확인했습니다."

        )

    except Exception as e:

        return False, f"백업 중 오류: {e}"


def restore_all_from_gsheet():

    client, error = get_gsheet_client()

    if error:

        return False, error

    sheet_url = get_gsheet_url()

    if not sheet_url:

        return False, "복원할 구글시트 URL이 없습니다."

    try:

        doc = client.open_by_url(
            sheet_url
        )

        restored = 0

        for local_path, sheet_name in GSHEET_BACKUP_TARGETS:

            try:

                ws = doc.worksheet(
                    sheet_name
                )

            except gspread.WorksheetNotFound:

                continue

            values = ws.get_all_values()

            if not values:

                continue

            headers = values[0]

            rows = values[1:]

            wb = Workbook()

            out_ws = wb.active

            out_ws.title = sheet_name

            out_ws.append(
                headers
            )

            for row in rows:

                out_ws.append(
                    row
                )

            wb.save(
                local_path
            )

            wb.close()

            restored += 1

        _read_excel_cached.clear()

        return True, f"{restored}개 시트를 복원했습니다."

    except Exception as e:

        return False, f"복원 중 오류: {e}"


# ============================================================
# 10-11. 이메일 자동발송
#
# Streamlit 앱 자체는 "매달 1일 자동으로" 실행되는 백그라운드
# 스케줄러가 없다(사람이 접속해야 코드가 돈다). 그래서 여기선
# "버튼 누르면 지금 바로 발송"까지만 만든다. 진짜 매달 자동
# 발송을 원하면, 외부 스케줄러(예: GitHub Actions)가 이
# 버튼과 동일한 함수를 주기적으로 호출해줘야 한다.
# ============================================================

def send_monthly_summary_email(

    recipient_list,
    subject,
    body_text

):

    try:

        smtp_host = st.secrets.get(
            "SMTP_HOST",
            "smtp.gmail.com"
        )

        smtp_port = int(

            st.secrets.get(
                "SMTP_PORT",
                587
            )

        )

        smtp_user = st.secrets.get(
            "SMTP_USER"
        )

        smtp_password = st.secrets.get(
            "SMTP_PASSWORD"
        )

    except Exception:

        smtp_user = None

        smtp_password = None

    if not smtp_user or not smtp_password:

        return False, (

            "이메일 발송 계정이 설정되지 않았습니다. "
            "Streamlit Secrets에 SMTP_USER, SMTP_PASSWORD를 "
            "등록해주세요 (Gmail이면 앱 비밀번호 필요)."

        )

    msg = MIMEMultipart()

    msg["From"] = smtp_user

    msg["To"] = ", ".join(
        recipient_list
    )

    msg["Subject"] = subject

    msg.attach(

        MIMEText(
            body_text,
            "plain",
            "utf-8"
        )

    )

    # 일시적인 네트워크 문제로 실패할 수 있어서, 바로 포기하지
    # 않고 짧은 간격을 두고 한 번 더 시도한다.

    last_error = None

    for attempt in range(2):

        try:

            with smtplib.SMTP(

                smtp_host,
                smtp_port

            ) as server:

                server.starttls()

                server.login(
                    smtp_user,
                    smtp_password
                )

                server.sendmail(

                    smtp_user,

                    recipient_list,

                    msg.as_string()

                )

            return True, (

                f"{len(recipient_list)}명에게 발송했습니다."
                +
                (
                    " (1차 시도 실패 후 재시도로 성공)"

                    if attempt > 0

                    else ""
                )

            )

        except Exception as e:

            last_error = e

            if attempt == 0:

                time.sleep(3)

    return False, (

        f"이메일 발송에 2회 시도했지만 실패했습니다: "
        f"{last_error}\n"
        "SMTP 계정 정보(SMTP_USER/SMTP_PASSWORD)나 네트워크 "
        "상태를 확인해주세요."

    )


def calculate_equipment_tco(equip_name, df_overhaul):

    # 총소유비용(TCO): 이 설비에 지금까지 실제로 들어간
    # 오버홀 비용을 전부 더한다. "비용" 컬럼이 비어있는 예전
    # 기록은 0으로 취급한다.

    if (

        df_overhaul is None

        or df_overhaul.empty

        or "설비명" not in df_overhaul.columns

    ):

        return 0, 0

    equip_hist = df_overhaul[

        df_overhaul["설비명"] == equip_name

    ]

    if equip_hist.empty:

        return 0, 0

    costs = pd.to_numeric(

        equip_hist.get("비용", 0),

        errors="coerce"

    ).fillna(0)

    return float(costs.sum()), len(equip_hist)


def build_monthly_summary_email_body(all_pumps, df_history):

    normal = watch = repair = 0

    repair_names = []

    for pump in all_pumps:

        result = pump_status(

            pump,
            df_history

        )

        if result["상태"] == "정상":

            normal += 1

        elif result["상태"] == "관찰":

            watch += 1

        else:

            repair += 1

            repair_names.append(
                pump["equip"]
            )

    lines = [

        f"[K-water tech 설비관리] "
        f"{datetime.now().strftime('%Y년 %m월')} 월간 요약",
        "",
        f"전체 설비: {len(all_pumps)}대",
        f"정상: {normal}대 / 관찰: {watch}대 / 정비검토: {repair}대",
        ""

    ]

    if repair_names:

        lines.append(

            "정비검토 필요 설비: "
            +
            ", ".join(repair_names)

        )

    lines.append(

        "\n자세한 내용은 K-water tech 설비관리 플랫폼에서 "
        "확인하세요."

    )

    return "\n".join(lines)


# ============================================================
# 10-12. 통합검색
#
# 설비마스터·진단이력·오버홀이력·소모품재고·일위대가·노하우가
# 다 따로따로 있어서, 뭔가 찾으려면 메뉴를 여러 개 돌아다녀야
# 했다. 검색어 하나로 전부 한 번에 훑어서 어디에 뭐가
# 있는지 보여준다.
# ============================================================

def find_similar_by_tfidf(

    keyword,
    df,
    text_columns,
    exclude_index=None,
    top_n=5,
    min_similarity=0.12

):

    # 정확한 단어가 하나도 안 겹쳐도, TF-IDF 코사인 유사도로
    # "의미가 비슷한" 과거 사례를 찾아준다(진짜 의미 이해는
    # 아니지만, 단어 조합의 패턴이 비슷한 문서를 순위 매겨
    # 보여주는 가벼운 통계기반 검색).

    if df is None or df.empty:

        return pd.DataFrame()

    valid_cols = [

        c for c in text_columns

        if c in df.columns

    ]

    if not valid_cols:

        return pd.DataFrame()

    corpus_df = df

    if exclude_index is not None:

        corpus_df = df.drop(

            index=exclude_index,

            errors="ignore"

        )

    if corpus_df.empty:

        return pd.DataFrame()

    combined_text = corpus_df[valid_cols].astype(

        str

    ).agg(" ".join, axis=1)

    try:

        vectorizer = TfidfVectorizer(

            analyzer="char_wb",

            ngram_range=(2, 3)

        )

        tfidf_matrix = vectorizer.fit_transform(

            combined_text.tolist() + [keyword]

        )

        sims = cosine_similarity(

            tfidf_matrix[-1],

            tfidf_matrix[:-1]

        )[0]

    except Exception:

        return pd.DataFrame()

    ranked_idx = np.argsort(sims)[::-1]

    picked = []

    for i in ranked_idx[:top_n]:

        if sims[i] >= min_similarity:

            picked.append(

                corpus_df.index[i]

            )

    if not picked:

        return pd.DataFrame()

    return corpus_df.loc[picked]


def render_highlighted_table_html(df, keyword):

    # st.dataframe은 셀 안에 HTML을 못 넣어서 검색어 강조가
    # 안 된다. 검색결과만큼은 표를 직접 HTML로 그려서, 찾은
    # 키워드 부분을 노란색으로 강조 표시한다. 스타일은 CSS
    # 클래스(.search-result-wrap)로 분리해서 다크모드에서도
    # 색이 자연스럽게 바뀌도록 한다.

    keyword_escaped = re.escape(
        keyword.strip()
    )

    def _highlight(val):

        text = str(val)

        try:

            return re.sub(

                f"({keyword_escaped})",

                r"<mark>\1</mark>",

                text,

                flags=re.IGNORECASE

            )

        except Exception:

            return text

    header_html = "".join(

        f"<th>{col}</th>"

        for col in df.columns

    )

    rows_html = []

    for _, row in df.iterrows():

        cells = "".join(

            f"<td>{_highlight(v)}</td>"

            for v in row.tolist()

        )

        rows_html.append(
            f"<tr>{cells}</tr>"
        )

    return (

        '<div class="search-result-wrap">'
        "<table>"
        f"<thead><tr>{header_html}</tr></thead>"
        f"<tbody>{''.join(rows_html)}</tbody>"
        "</table></div>"

    )


def run_unified_search(

    keyword,
    all_pumps,
    similarity_threshold=0.12

):

    keyword = keyword.strip()

    if not keyword:

        return {}

    results = {}

    # 1) 설비마스터

    equip_matches = [

        p for p in all_pumps

        if keyword.lower() in p["equip"].lower()

        or keyword.lower() in p["site"].lower()

    ]

    if equip_matches:

        results["설비"] = equip_matches

    # 2) 정밀진단 이력

    df_diag = read_excel(
        DB_FILE_PATH,
        "진단이력"
    )

    if not df_diag.empty:

        mask = df_diag.apply(

            lambda row: keyword.lower() in str(

                row.get("설비명", "")

            ).lower()

            or keyword.lower() in str(

                row.get("점검자", "")

            ).lower(),

            axis=1

        )

        matched = df_diag[mask]

        if not matched.empty:

            results["정밀진단 이력"] = matched

    # 3) 오버홀 이력

    df_overhaul = read_excel(
        OVERHAUL_DB_PATH,
        "오버홀이력"
    )

    if not df_overhaul.empty:

        mask = df_overhaul.apply(

            lambda row: any(

                keyword.lower() in str(v).lower()

                for v in row.values

            ),

            axis=1

        )

        matched = df_overhaul[mask]

        if not matched.empty:

            results["오버홀 이력"] = matched

    # 4) 소모품 재고

    df_cons = get_consumables()

    if not df_cons.empty:

        mask = df_cons["소모품명"].str.contains(

            keyword,
            case=False,
            na=False

        )

        matched = df_cons[mask]

        if not matched.empty:

            results["소모품 재고"] = matched

    # 5) 일위대가 라이브러리

    df_unitprice = get_unitprice_library()

    if not df_unitprice.empty:

        mask = df_unitprice["항목명"].str.contains(

            keyword,
            case=False,
            na=False

        )

        matched = df_unitprice[mask]

        if not matched.empty:

            results["일위대가 라이브러리"] = matched

    # 6) 기술 노하우 — 정확한 단어가 없어도 "의미가 비슷한"
    # 사례를 찾을 수 있도록, 정확 매칭에 더해 TF-IDF 기반
    # 유사도 검색도 같이 돌린다.

    df_knowhow = read_excel(
        KNOWHOW_DB_PATH,
        "노하우DB"
    )

    if not df_knowhow.empty:

        mask = df_knowhow.apply(

            lambda row: any(

                keyword.lower() in str(v).lower()

                for v in row.values

            ),

            axis=1

        )

        matched = df_knowhow[mask]

        similar = find_similar_by_tfidf(

            keyword,

            df_knowhow,

            text_columns=["현상및원인", "해결노하우"],

            exclude_index=matched.index,

            min_similarity=similarity_threshold

        )

        if not matched.empty:

            results["기술 노하우 (정확일치)"] = matched

        if not similar.empty:

            results["기술 노하우 (비슷한 사례·AI 추천)"] = similar

    return results


def get_vib_judgement(value):

    if value is None:

        return "-"

    if value <= VIB_JUDGE_GOOD:

        return "A"

    if value <= VIB_JUDGE_FAIR:

        return "B"

    if value <= VIB_JUDGE_BAD:

        return "C"

    return "D"


def format_vib_month_label(

    month_str,
    prev_year

):

    # 이미지의 그래프처럼 "25.08월, 09월, 10월... 26.01월, 02월"
    # 식으로 연도가 바뀔 때만 앞에 연도를 붙인다.

    year, month = month_str.split("-")

    year_short = year[2:]

    if year_short != prev_year:

        return f"{year_short}.{month}월", year_short

    return f"{month}월", year_short


def build_vibration_trend_chart_fig(

    equip_name,
    months_avail,
    trend_rows

):

    # 보내주신 "펌프모터진동 그래프 SHEET" 이미지와 동일하게
    # 재현한다. trend_rows는 [월, 모터반부하V, 모터반부하H,
    # 모터반부하축, 모터부하V, 모터부하H, 펌프반부하V,
    # 펌프반부하H, 펌프반부하축, 펌프부하V, 펌프부하H] 순서로
    # 붙임3 표를 만들 때 쓴 것과 동일한 데이터를 재사용한다.

    if not trend_rows:

        return None

    series_names = [

        "모터반부하V", "모터반부하H", "모터반부하축",
        "모터부하V", "모터부하H",
        "펌프반부하V", "펌프반부하H", "펌프반부하축",
        "펌프부하V", "펌프부하H"

    ]

    fig, ax = plt.subplots(
        figsize=(9, 4)
    )

    for col_idx, name in enumerate(
        series_names,
        start=1
    ):

        values = []

        for row in trend_rows:

            v = row[col_idx]

            values.append(

                v if isinstance(v, (int, float)) else None

            )

        # 값이 하나도 없는 계열은 그래프에서 생략

        if all(v is None for v in values):

            continue

        ax.plot(

            months_avail,

            values,

            marker="o",

            markersize=3,

            linewidth=1.4,

            label=name

        )

    # 이미지와 동일하게 양호(3.2)=파랑, 보통(5.1)=초록

    ax.axhline(

        VIB_JUDGE_FAIR,

        linestyle="-",

        color="#2f9e44",

        linewidth=1.6,

        label=f"보통({VIB_JUDGE_FAIR})"

    )

    ax.axhline(

        VIB_JUDGE_GOOD,

        linestyle="-",

        color="#1c7ed6",

        linewidth=1.6,

        label=f"양호({VIB_JUDGE_GOOD})"

    )

    ax.set_ylabel(
        "진동값 (mm/s, rms)"
    )

    ax.legend(

        loc="upper left",

        bbox_to_anchor=(0, -0.16, 1, 0.1),

        mode="expand",

        ncol=6,

        fontsize=10.5,

        columnspacing=1.0,

        handletextpad=0.4

    )

    ax.grid(
        alpha=0.2
    )

    fig.tight_layout()

    return fig


def build_vibration_monthly_report_docx(

    site_list,
    month_label,
    all_pumps,
    work_period_text,
    work_members_text

):

    df_vib = read_excel(

        VIBRATION_DB_PATH,

        "진동측정이력"

    )

    month_df = pd.DataFrame()

    if not df_vib.empty:

        month_df = df_vib[

            df_vib["측정일자"].astype(str).str.startswith(
                month_label
            )

        ]

    doc = Document()

    _set_korean_font(doc)

    # 표가 많은 문서라 기본 여백(보통 상하좌우 1인치)이면
    # 표 열이 옹색하게 나온다. 여백을 줄여서 표가 넉넉하게
    # 나오도록 한다.

    for section in doc.sections:

        section.left_margin = Inches(0.6)

        section.right_margin = Inches(0.6)

        section.top_margin = Inches(0.5)

        section.bottom_margin = Inches(0.5)

    logo_path = find_logo_file()

    if logo_path:

        try:

            logo_p = doc.add_paragraph()

            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            logo_p.add_run().add_picture(
                logo_path,
                width=Inches(1.3)
            )

        except Exception:

            pass

    title = doc.add_heading(
        "펌프모터 진동 측정 분석 보고서",
        level=0
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()

    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_run = sub.add_run(

        f"('{month_label[2:4]}. "
        f"{int(month_label[5:7])}월 밀양권사업소)"

    )

    sub_run.bold = True

    doc.add_paragraph()

    # ---------------- 1. 개요 ----------------

    doc.add_heading(
        "1. 개요",
        level=1
    )

    _add_styled_table(

        doc,

        ["항목", "내용"],

        [
            [
                "작업 목적",
                "펌프모터의 진동을 주기적으로 측정·분석하여 "
                "이상 징후를 조기에 감지하고, 측정값을 기존 "
                "데이터 및 기준값과 비교해 상태를 진단한다. "
                "필요시 적절한 조치를 수행하여 설비의 안정적 "
                "운영과 효율적 유지관리를 도모하고자 한다."
            ],
            ["작업 기간", work_period_text],
            ["작업 인원", work_members_text],
            ["작업 내용", "휴대용 진동측정기를 활용하여 수행하는 진동측정"],
            ["측정 장비", "진동계(ACO 3116)"]
        ]

    )

    # ---------------- 계획 대비 실적 (공란) ----------------

    doc.add_heading(
        "계획 대비 실적",
        level=1
    )

    doc.add_paragraph(

        "※ 계약수량·주기횟수 등은 별도 계약관리 자료 확인 후 "
        "기입하시기 바랍니다 (표만 제공)."

    )

    _add_styled_table(

        doc,

        [
            "구분/분류", "주기", "계약수량", "주기횟수", "총 횟수",
            "전월", "금월", "누계", "잔여물량", "비고"
        ],

        [
            ["2. 측정 및 시험", "", "", "", "", "", "", "", "", ""],
            ["2.1.기계설비", "", "", "", "", "", "", "", "", ""],
            ["2.1.2.진동측정_A", "월간", "", "", "", "", "", "", "", ""],
            ["2.1.3.진동측정_B", "월간", "", "", "", "", "", "", "", ""]
        ]

    )

    # ---------------- 대상 시설(작업범위) ----------------

    doc.add_heading(
        "대상 시설(작업 범위)",
        level=1
    )

    facility_rows = []

    for pump in all_pumps:

        if pump["site"] not in site_list:

            continue

        flow_v = pump.get("flow")

        head_v = pump.get("head")

        hp_v = pump.get("hp")

        spec_parts = []

        if flow_v and head_v:

            spec_parts.append(

                f"(펌프) {flow_v}m³/h × {head_v}m"

            )

        if hp_v:

            spec_parts.append(

                f"(모터) {hp_v}HP"

            )

        spec_text = (

            "\n".join(spec_parts)

            if spec_parts

            else "제원 미등록"

        )

        facility_rows.append(

            [
                pump["site"],
                pump["equip"],
                spec_text,
                "○"
            ]

        )

    if facility_rows:

        _add_styled_table(

            doc,

            ["사업장", "설비명", "제원", "작업대상"],

            facility_rows

        )

    else:

        doc.add_paragraph(
            "등록된 대상 설비가 없습니다."
        )

    # ---------------- 2. 측정‧시험 결과 ----------------

    doc.add_heading(
        "2. 측정‧시험 결과",
        level=1
    )

    summary_rows = []

    equip_worst = {}

    if not month_df.empty:

        for equip_name, g in month_df.groupby("설비명"):

            idx_max = g["측정값"].idxmax()

            max_row = g.loc[idx_max]

            max_val = max_row["측정값"]

            judgement = get_vib_judgement(
                max_val
            )

            site_name = g["사업장"].iloc[0]

            point_desc = (

                f"{max_row['펌프모터구분']} "
                f"{max_row['부하구분']} 부분 "
                f"{max_row['측정방향']}"

            )

            comment = (

                "특이사항 없음"

                if judgement == "A"

                else f"{point_desc} 진동값 상승, 원인 분석 필요"

            )

            equip_worst[equip_name] = (

                site_name,
                max_val,
                judgement,
                point_desc,
                comment

            )

    for equip_name, (site_name, max_val, judgement, point_desc, comment) in sorted(

        equip_worst.items(),

        key=lambda x: (x[1][0], x[0])

    ):

        summary_rows.append(

            [
                site_name,
                equip_name,
                round(max_val, 1),
                judgement,
                comment
            ]

        )

    if summary_rows:

        _add_styled_table(

            doc,

            ["사업장", "설비명", "진동값(mm/s,rms)", "판정", "측정·시험 결과"],

            summary_rows

        )

    else:

        doc.add_paragraph(

            f"{month_label}에 업로드된 진동측정 데이터가 없습니다."

        )

    # ---------------- 3. 종합의견 ----------------

    doc.add_heading(
        "3. 종합의견",
        level=1
    )

    # A영역이라도 이 값 이상이면 "특이사항 없음"으로 뭉개지 않고
    # 추이관찰 문구를 붙인다. (예시 보고서의 "2.2㎜/s(A영역)이지만
    # 지속관찰 중" 같은 표현을 재현)

    VIB_WATCH_IN_A = 2.0

    def _build_vib_comment(

        equip_name,
        point_desc,
        value,
        judgement

    ):

        if judgement == "A":

            if value >= VIB_WATCH_IN_A:

                return (

                    f"○ {equip_name} : {point_desc} 진동값이 "
                    f"최대 {value:.1f}㎜/s(A영역)로 측정되어 "
                    "그리스 주입 후 추이 관찰을 진행하고 있음."

                )

            return None

        if judgement == "B":

            return (

                f"○ {equip_name} : {point_desc} 진동값이 "
                f"{value:.1f}㎜/s(B영역, 보통)로 측정되어 "
                "정기점검 주기를 단축하고 원인 분석이 필요함."

            )

        if judgement == "C":

            return (

                f"○ {equip_name} : {point_desc} 진동값이 "
                f"{value:.1f}㎜/s(C영역, 주의)로 측정되어 "
                "정밀진동분석 및 정비계획 수립이 필요함."

            )

        return (

            f"○ {equip_name} : {point_desc} 진동값이 "
            f"{value:.1f}㎜/s(D영역, 불량)로 측정되어 "
            "즉시 정밀진동분석 및 정비 조치가 필요함 "
            "(주관부서 정밀진동분석 의뢰 요망)."

        )

    if equip_worst:

        by_site = {}

        for equip_name, (site_name, max_val, judgement, point_desc, comment) in equip_worst.items():

            by_site.setdefault(
                site_name,
                []
            ).append(

                (equip_name, max_val, judgement, point_desc)

            )

        for site_name, items in sorted(by_site.items()):

            watch_comments = []

            for equip_name, value, judgement, point_desc in sorted(

                items,

                key=lambda x: x[0]

            ):

                c = _build_vib_comment(

                    equip_name,
                    point_desc,
                    value,
                    judgement

                )

                if c:

                    watch_comments.append(c)

            if watch_comments:

                doc.add_paragraph(

                    f"☐ {site_name} 펌프모터 설비의 진동값은 "
                    "대체로 양호(A영역)하였으나, 일부 설비에서 "
                    "진동값 상승이 확인되어 아래와 같이 지속적인 "
                    "관리가 필요함."

                )

                for c in watch_comments:

                    doc.add_paragraph(
                        c,
                        style="List Bullet"
                    )

            else:

                doc.add_paragraph(

                    f"☐ {site_name} 펌프모터 설비의 진동값은 "
                    "전 항목 양호(A영역)로 특이사항 없음."

                )

    else:

        doc.add_paragraph(
            "이번 달 측정 데이터가 없어 종합의견을 생성할 수 없습니다."
        )

    # ---------------- 붙임1. 관리상태 판단기준 ----------------

    doc.add_page_break()

    doc.add_heading(
        "붙임1. 관리상태 판단기준",
        level=1
    )

    doc.add_paragraph(

        "◦ 진동판정기준 (본 보고서 적용 기준)\n"
        f"- 양호(A) : 총진동값 {VIB_JUDGE_GOOD}(mm/s, rms) 이하\n"
        f"- 보통(B) : 총진동값 {VIB_JUDGE_FAIR}(mm/s, rms) 이하\n"
        f"- 주의(C) : 총진동값 {VIB_JUDGE_BAD}(mm/s, rms) 이하\n"
        f"- 불량(D) : 총진동값 {VIB_JUDGE_BAD}(mm/s, rms) 초과 "
        "(☞ 불량판정의 경우 주관부서에 정밀진동분석 의뢰)"

    )

    doc.add_paragraph(

        "◦ ISO10816-3/7 등 공식 규격은 설비 종류·용량·회전수에 "
        "따라 세분화된 영역(A/B/C/D)을 정의하며, 본 보고서의 "
        "판정기준(위)은 이를 참고해 자체적으로 단순화한 관리기준"
        "입니다. 공식 규격 전문은 사내 규정집을 참고하십시오."

    )

    # ---------------- 붙임2. 측정‧시험 DATA(결과값) ----------------

    doc.add_page_break()

    doc.add_heading(
        "붙임2. 측정‧시험 DATA(결과값)",
        level=1
    )

    for site_name in site_list:

        site_pumps = [

            p for p in all_pumps

            if p["site"] == site_name

        ]

        if not site_pumps:

            continue

        doc.add_heading(
            f"{site_name} 펌프모터 총진동값 기록대장",
            level=2
        )

        rec_rows = []

        for pump in site_pumps:

            equip_df = (

                month_df[month_df["설비명"] == pump["equip"]]

                if not month_df.empty

                else pd.DataFrame()

            )

            if equip_df.empty:

                rec_rows.append(

                    [pump["equip"], "", "", "", "", "", "", ""]

                )

                continue

            for load_type in ["부하", "반부하"]:

                sub = equip_df[

                    equip_df["부하구분"] == load_type

                ]

                if sub.empty:

                    continue

                def _v(pm, d):

                    m = sub[

                        (sub["펌프모터구분"] == pm)

                        &
                        (sub["측정방향"] == d)

                    ]

                    return (

                        round(m["측정값"].iloc[0], 1)

                        if not m.empty

                        else ""

                    )

                worst = sub["측정값"].max()

                rec_rows.append(

                    [
                        f"{pump['equip']} ({load_type})",
                        _v("펌프", "수직"),
                        _v("펌프", "수평"),
                        _v("펌프", "축"),
                        _v("모터", "수직"),
                        _v("모터", "수평"),
                        _v("모터", "축"),
                        get_vib_judgement(worst)
                    ]

                )

        _add_styled_table(

            doc,

            [
                "호기(측정위치)", "펌프 V", "펌프 H", "펌프 축(A)",
                "모터 V", "모터 H", "모터 축(A)", "판정"
            ],

            rec_rows

        )

    # ---------------- 붙임3. 상태변화 및 추이분석 ----------------

    doc.add_page_break()

    doc.add_heading(
        "붙임3. 상태변화 및 추이분석",
        level=1
    )

    if not df_vib.empty:

        for pump in all_pumps:

            if pump["site"] not in site_list:

                continue

            equip_hist = df_vib[

                df_vib["설비명"] == pump["equip"]

            ]

            if equip_hist.empty:

                continue

            months_avail = sorted(

                equip_hist["측정일자"].str[:7].unique()

            )

            if not months_avail:

                continue

            # ---- "펌프,모터 진동측정 SHEET" 형식의 정보 박스 ----
            # (로고 + 제목 + 기기명/설치장소/점검일자/측정장비)

            sheet_logo_path = find_logo_file()

            if sheet_logo_path:

                try:

                    logo_p2 = doc.add_paragraph()

                    logo_p2.add_run().add_picture(

                        sheet_logo_path,

                        width=Inches(1.0)

                    )

                except Exception:

                    pass

            sheet_title_p = doc.add_paragraph()

            sheet_title_run = sheet_title_p.add_run(

                "펌프, 모터 진동측정 SHEET"

            )

            sheet_title_run.bold = True

            sheet_title_run.font.size = Pt(14)

            start_y, start_m = months_avail[0].split("-")

            end_y, end_m = months_avail[-1].split("-")

            period_text = (

                f"{start_y}년 {int(start_m)}월 ~ "
                f"{end_y}년 {int(end_m)}월"

            )

            _add_styled_table(

                doc,

                ["기기명", "설치장소", "점검일자", "측정장비"],

                [
                    [
                        pump["equip"],
                        pump["site"],
                        period_text,
                        "진동계"
                    ]
                ]

            )

            trend_rows = []

            for m in months_avail:

                m_df = equip_hist[

                    equip_hist["측정일자"].str.startswith(m)

                ]

                def _mv(pm, load, d):

                    r = m_df[

                        (m_df["펌프모터구분"] == pm)

                        &
                        (m_df["부하구분"] == load)

                        &
                        (m_df["측정방향"] == d)

                    ]

                    return (

                        round(r["측정값"].mean(), 1)

                        if not r.empty

                        else ""

                    )

                trend_rows.append(

                    [
                        m,
                        _mv("모터", "반부하", "수직"),
                        _mv("모터", "반부하", "수평"),
                        _mv("모터", "반부하", "축"),
                        _mv("모터", "부하", "수직"),
                        _mv("모터", "부하", "수평"),
                        _mv("펌프", "반부하", "수직"),
                        _mv("펌프", "반부하", "수평"),
                        _mv("펌프", "반부하", "축"),
                        _mv("펌프", "부하", "수직"),
                        _mv("펌프", "부하", "수평")
                    ]

                )

            # 그래프를 표보다 먼저 (이미지와 동일한 순서)

            month_display_labels = []

            _prev_year = ""

            for m in months_avail:

                disp, _prev_year = format_vib_month_label(

                    m,

                    _prev_year

                )

                month_display_labels.append(
                    disp
                )

            trend_fig = build_vibration_trend_chart_fig(

                pump["equip"],

                month_display_labels,

                trend_rows

            )

            if trend_fig is not None:

                _add_fig_to_doc(

                    doc,

                    trend_fig,

                    # 여백을 0.6인치씩 줄여둔 문서라 실제 표 폭은
                    # 8.5 - 0.6 - 0.6 = 7.3인치. 그래프도 이 폭에
                    # 맞춰서 표랑 나란히 꽉 차 보이게 한다.

                    width_inches=7.3

                )

            _add_styled_table(

                doc,

                [
                    "월별",
                    "모터 반부하V", "모터 반부하H", "모터 반부하축",
                    "모터 부하V", "모터 부하H",
                    "펌프 반부하V", "펌프 반부하H", "펌프 반부하축",
                    "펌프 부하V", "펌프 부하H"
                ],

                trend_rows

            )

            doc.add_paragraph(

                f"비고 : 판정기준 {VIB_JUDGE_GOOD}이하(양호), "
                f"{VIB_JUDGE_FAIR}이하(장시간 운전허용), "
                f"{VIB_JUDGE_BAD}이하(보수조치 필요)"

            )

            doc.add_page_break()

    else:

        doc.add_paragraph(

            "누적된 진동측정 이력이 없습니다. 매달 정기점검 "
            "실적목록을 업로드하면 이 자리에 추세표가 쌓입니다."

        )

    buffer = io.BytesIO()

    doc.save(buffer)

    return buffer.getvalue()


# ============================================================
# 11-1. 축정렬 · 진동분석 · 효율진단 보고서
#
# 아래 3개는 실제 회사 보고서 양식을 아직 받기 전이라, 업계
# 일반적인 항목(오프셋/각도, 스펙트럼 성분, 정격대비 효율)
# 기준으로 임시 구성했다. 실제 양식을 받으면 이 함수들의
# 표 구성만 바꾸면 되고, 데이터 저장 구조는 그대로 쓸 수
# 있다.
# ============================================================

def _set_table_col_widths(table, widths_inches):

    # python-docx는 헤더행에만 폭을 지정하면 무시되는 경우가
    # 있어서, 모든 행 × 모든 셀에 동일하게 폭을 강제 지정한다.

    table.autofit = False

    for row in table.rows:

        for i, cell in enumerate(row.cells):

            if i < len(widths_inches):

                cell.width = Inches(
                    widths_inches[i]
                )


def _add_table_row_helper(table, values, bold=False):

    cells = table.add_row().cells

    for i, v in enumerate(values):

        if i < len(cells):

            cells[i].text = str(v)

            if bold:

                for p in cells[i].paragraphs:

                    for r in p.runs:

                        r.bold = True


def _style_header_row(table):

    for cell in table.rows[0].cells:

        for p in cell.paragraphs:

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for r in p.runs:

                r.bold = True


def _add_landscape_section(doc):

    # 표 열이 많아 세로 페이지엔 비좁은 구간(붙임2 Data Sheet
    # 등)만 가로 페이지로 바꾼다. 이후 다시 세로로 돌아오려면
    # _add_portrait_section()을 호출하면 된다.

    new_section = doc.add_section(
        WD_SECTION.NEW_PAGE
    )

    new_section.orientation = WD_ORIENT.LANDSCAPE

    new_width, new_height = (

        new_section.page_height,
        new_section.page_width

    )

    new_section.page_width = new_width

    new_section.page_height = new_height

    new_section.left_margin = Inches(0.6)
    new_section.right_margin = Inches(0.6)
    new_section.top_margin = Inches(0.6)
    new_section.bottom_margin = Inches(0.6)

    return new_section


def _add_portrait_section(doc):

    new_section = doc.add_section(
        WD_SECTION.NEW_PAGE
    )

    new_section.orientation = WD_ORIENT.PORTRAIT

    new_width, new_height = (

        new_section.page_height,
        new_section.page_width

    )

    new_section.page_width = new_width

    new_section.page_height = new_height

    new_section.left_margin = Inches(0.7)
    new_section.right_margin = Inches(0.7)
    new_section.top_margin = Inches(0.6)
    new_section.bottom_margin = Inches(0.6)

    return new_section


def build_pump_motor_diagram():

    # 원본 보고서의 "펌프모터 진동측정" 지점 도식(1~4번 측정
    # 지점 표시)을 간단히 재현한다.

    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    from matplotlib.patches import Rectangle

    _font_path = None

    for _p in (
        "NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf"
    ):

        if os.path.exists(_p):

            _font_path = _p

            break

    if _font_path:

        fm.fontManager.addfont(_font_path)

        plt.rcParams["font.family"] = fm.FontProperties(
            fname=_font_path
        ).get_name()

    fig, ax = plt.subplots(figsize=(5.5, 2.2), dpi=110)

    ax.add_patch(Rectangle(
        (0.5, 0.5), 2.0, 1.4,
        facecolor="#D9D9D9", edgecolor="black"
    ))

    ax.text(1.5, 1.2, "Motor", ha="center", va="center", fontsize=11)

    ax.add_patch(Rectangle(
        (3.0, 0.3), 3.5, 1.8,
        facecolor="#D9D9D9", edgecolor="black"
    ))

    ax.text(4.75, 1.2, "Pump", ha="center", va="center", fontsize=11)

    ax.plot([2.5, 3.0], [1.2, 1.2], color="black", linewidth=2)

    ax.text(2.4, 1.85, "4", ha="center", fontsize=12, fontweight="bold")

    ax.text(2.4, 0.55, "3", ha="center", fontsize=12, fontweight="bold")

    ax.text(6.6, 1.85, "1", ha="center", fontsize=12, fontweight="bold")

    ax.text(6.6, 0.35, "2", ha="center", fontsize=12, fontweight="bold")

    ax.set_xlim(0, 7)

    ax.set_ylim(0, 2.3)

    ax.axis("off")

    fig.tight_layout()

    buf = io.BytesIO()

    fig.savefig(buf, format="png", transparent=True)

    plt.close(fig)

    return buf.getvalue()


def build_alignment_diagram():

    # 원본 붙임1의 "Pump-Motor Vertical Offset / Horizontal GAP"
    # 정렬 개념도를 간단히 재현한다.

    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import Rectangle
    import matplotlib.font_manager as fm

    _font_path = None

    for _p in (
        "NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf"
    ):

        if os.path.exists(_p):

            _font_path = _p

            break

    if _font_path:

        fm.fontManager.addfont(_font_path)

        plt.rcParams["font.family"] = fm.FontProperties(
            fname=_font_path
        ).get_name()

    fig, axes = plt.subplots(2, 1, figsize=(5.5, 3.2), dpi=110)

    for ax, offset_y, label in [

        (axes[0], 0.15, "Vertical  Offset"),
        (axes[1], 0.0, "Horizontal  Gap")

    ]:

        ax.add_patch(Rectangle(

            (0.3, 0.3), 2.0, 0.8,

            facecolor="#D9D9D9", edgecolor="black"

        ))

        ax.text(1.3, 0.7, "Pump", ha="center", va="center", fontsize=10)

        ax.add_patch(Rectangle(

            (3.2, 0.3 + offset_y), 2.0, 0.8,

            facecolor="#D9D9D9", edgecolor="black"

        ))

        ax.text(

            4.2, 0.7 + offset_y, "Motor",
            ha="center", va="center", fontsize=10

        )

        ax.annotate(

            "", xy=(3.2, 0.7 + offset_y), xytext=(2.3, 0.7),

            arrowprops=dict(arrowstyle="<->", color="#C00000")

        )

        ax.text(

            2.75, 0.85 + offset_y / 2, label,
            ha="center", fontsize=9, color="#C00000"

        )

        ax.set_xlim(0, 5.7)

        ax.set_ylim(-0.1, 1.4)

        ax.axis("off")

    fig.tight_layout()

    buf = io.BytesIO()

    fig.savefig(buf, format="png", transparent=True)

    plt.close(fig)

    return buf.getvalue()


def build_alignment_trend_chart(history_rows):

    # history_rows: [(측정일자, 각도or옵셋 수직값, 수평값), ...]
    # 실제 보고서의 "상태변화 및 추이분석" 그래프(수직/수평
    # 축정렬 추이 + 보통/양호 기준선)를 그대로 재현한다.

    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm

    _font_path = None

    for _p in (
        "NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf"
    ):

        if os.path.exists(_p):

            _font_path = _p

            break

    if _font_path:

        fm.fontManager.addfont(_font_path)

        plt.rcParams["font.family"] = fm.FontProperties(
            fname=_font_path
        ).get_name()

    dates = [r[0] for r in history_rows]

    v_vals = [r[1] for r in history_rows]

    h_vals = [r[2] for r in history_rows]

    fig, ax = plt.subplots(figsize=(6.5, 3.2), dpi=110)

    ax.plot(
        dates, v_vals,
        marker="o", color="#4472C4",
        label="수직 축정렬"
    )

    ax.plot(
        dates, h_vals,
        marker="o", color="#ED7D31",
        label="수평 축정렬"
    )

    ax.axhline(
        y=0.10, color="#A6A6A6",
        linewidth=1.5, label="보통(0.1)"
    )

    ax.axhline(
        y=0.05, color="#FFC000",
        linewidth=1.5, label="양호(0.05)"
    )

    ax.set_ylim(bottom=0)

    ax.legend(
        loc="upper center",
        bbox_to_anchor=(0.5, -0.15),
        ncol=4,
        fontsize=8,
        frameon=True
    )

    ax.grid(True, alpha=0.3)

    fig.tight_layout()

    buf = io.BytesIO()

    fig.savefig(buf, format="png")

    plt.close(fig)

    return buf.getvalue()


ALIGNMENT_CRITERIA_NOTE = (
    "환경부, 정수장 기술진단 매뉴얼, 2021"
)


def build_alignment_report_docx(

    pump,
    check_date_label,
    ai_data,
    image_bytes=None,
    history_rows=None

):

    # ai_data: call_claude_vision_analysis + extract_json_from_ai_response
    # 로 얻은 딕셔너리. 실패시에도 report는 만들어지도록
    # .get()으로 방어적으로 값을 꺼낸다.

    doc = Document()

    _set_korean_font(doc)

    for section in doc.sections:

        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    logo_path = find_logo_file()

    if logo_path:

        try:

            logo_p = doc.add_paragraph()

            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            logo_p.add_run().add_picture(
                logo_path,
                width=Inches(1.3)
            )

        except Exception:

            pass

    title = doc.add_heading(
        f"펌프모터 얼라인먼트 측정 보고서 ({check_date_label})",
        level=0
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 개요
    doc.add_heading("1. 개요", level=1)

    doc.add_paragraph(

        "모터–펌프 회전기계의 축 정렬 상태를 정밀 계측하여, "
        "운전 중 발생할 수 있는 베어링 과부하, 진동 증가, "
        "커플링 마모 및 에너지 손실을 사전에 예방하고, "
        "설비의 운전 안정성 확보 및 수명 연장을 목적으로 함."

    )

    info_table = doc.add_table(rows=1, cols=4)

    info_table.style = "Light Grid Accent 1"

    info_table.rows[0].cells[0].text = "설비명"
    info_table.rows[0].cells[1].text = pump["equip"]
    info_table.rows[0].cells[2].text = "사업장"
    info_table.rows[0].cells[3].text = pump["site"]

    _add_table_row_helper(

        info_table,

        [
            "점검일자",
            ai_data.get("점검일자", check_date_label),
            "측정장비",
            ai_data.get("측정장비", "")
        ]

    )

    doc.add_paragraph("")

    # 2. 측정데이터 요약 결과
    doc.add_heading("2. 측정데이터 요약 결과", level=1)

    summary_table = doc.add_table(rows=1, cols=6)

    summary_table.style = "Light Grid Accent 1"

    headers = [

        "방향", "①각도오차\n(mm/100mm)",
        "②옵셋오차\n(mm)", "③펌프전면변위\n(mm)",
        "④펌프후면변위\n(mm)", "판정"

    ]

    for i, h in enumerate(headers):

        summary_table.rows[0].cells[i].text = h

    _style_header_row(summary_table)

    _add_table_row_helper(

        summary_table,

        [
            "수평",
            ai_data.get("각도오차_수평", ""),
            ai_data.get("옵셋오차_수평", ""),
            ai_data.get("전면변위_수평", ""),
            ai_data.get("후면변위_수평", ""),
            ai_data.get("판정_수평", "")
        ]

    )

    _add_table_row_helper(

        summary_table,

        [
            "수직",
            ai_data.get("각도오차_수직", ""),
            ai_data.get("옵셋오차_수직", ""),
            ai_data.get("전면변위_수직", ""),
            ai_data.get("후면변위_수직", ""),
            ai_data.get("판정_수직", "")
        ]

    )

    _set_table_col_widths(

        summary_table,

        [0.7, 1.3, 1.15, 1.35, 1.35, 0.85]

    )

    doc.add_paragraph("")

    _note = doc.add_paragraph()

    _note.add_run(

        "*① 각도 오차 : 축이 얼마나 기울어져 있는지를 판단하는 "
        "수치(원인)  *② 옵셋 오차 : 커플링 위치에서의 축 중심 "
        "이탈 수치(원인)  *③ 펌프 전면 변위 : 각도+옵셋을 "
        "합해서 계산한 실제 축 변위(결과)  *④ 펌프 후면 변위 : "
        "펌프 뒤쪽 베어링 위치에서의 최종 변위(결과)"

    ).italic = True

    doc.add_paragraph("")

    # 3. 진단 의견
    doc.add_heading("3. 진단 의견", level=1)

    doc.add_paragraph(
        "☐ 종합 의견"
    ).runs[0].bold = True

    for _line in ai_data.get("종합의견", "").split("\n"):

        if _line.strip():

            doc.add_paragraph(
                "○ " + _line.strip()
            )

    doc.add_page_break()

    # 붙임1. 관리상태 판단기준
    doc.add_heading(
        "붙임1. 관리상태 판단기준",
        level=1
    )

    crit_table1 = doc.add_table(rows=1, cols=3)

    crit_table1.style = "Light Grid Accent 1"

    for i, h in enumerate(["구분", "GAP/직진도", "기준치(mm)"]):

        crit_table1.rows[0].cells[i].text = h

    _style_header_row(crit_table1)

    _add_table_row_helper(crit_table1, ["V(수직)", "센서 간격·직진도 확인", "0.05"])

    _add_table_row_helper(crit_table1, ["H(수평)", "센서 간격·직진도 확인", "0.05"])

    doc.add_paragraph("")

    crit_table2 = doc.add_table(rows=1, cols=5)

    crit_table2.style = "Light Grid Accent 1"

    for i, h in enumerate(

        ["회전수(RPM)", "Offset\nExcellent", "Offset\nAcceptable",
         "Angular\nExcellent", "Angular\nAcceptable"]

    ):

        crit_table2.rows[0].cells[i].text = h

    _style_header_row(crit_table2)

    for row in [

        ["0~1,000", "0.07", "0.13", "0.06", "0.10"],
        ["1,000~2,000", "0.05", "0.10", "0.05", "0.08"],
        ["2,000~3,000", "0.03", "0.07", "0.04", "0.07"],
        ["3,000~4,000", "0.02", "0.04", "0.03", "0.06"],
        ["4,000~5,000", "0.01", "0.03", "0.02", "0.05"],
        ["5,000~6,000", "<0.01", "<0.03", "0.01", "0.04"]

    ]:

        _add_table_row_helper(crit_table2, row)

    doc.add_paragraph("")

    try:

        _align_diagram_bytes = build_alignment_diagram()

        _diagram_p2 = doc.add_paragraph()

        _diagram_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _diagram_p2.add_run().add_picture(

            io.BytesIO(_align_diagram_bytes),

            width=Inches(4.0)

        )

    except Exception:

        pass

    _src = doc.add_paragraph()

    _src.add_run(
        f"출처 ○ {ALIGNMENT_CRITERIA_NOTE}"
    ).italic = True

    _add_landscape_section(doc)

    # 붙임2. 측정·시험 DATA
    doc.add_heading(
        "붙임2. 측정·시험 DATA",
        level=1
    )

    doc.add_paragraph(
        "펌프모터 얼라인먼트 Data Sheet"
    ).runs[0].bold = True

    ds_table = doc.add_table(rows=1, cols=4)

    ds_table.style = "Light Grid Accent 1"

    for i, h in enumerate(

        ["기기명", "점검일자", "설치장소", "측정장비"]

    ):

        ds_table.rows[0].cells[i].text = h

    _style_header_row(ds_table)

    _add_table_row_helper(

        ds_table,

        [
            pump["equip"],
            ai_data.get("점검일자", check_date_label),
            pump["site"],
            ai_data.get("측정장비", "")
        ]

    )

    doc.add_paragraph("")

    data_table = doc.add_table(rows=1, cols=8)

    data_table.style = "Light Grid Accent 1"

    for i, h in enumerate(

        ["구분", "방향", "각도오차\n(mm/100mm)", "옵셋오차\n(mm)",
         "기준거리\n(mm)", "전면변위\n(mm)", "후면변위\n(mm)", "판정"]

    ):

        data_table.rows[0].cells[i].text = h

    _style_header_row(data_table)

    for _label, _dist in [

        ("펌프모터 커플링 중심 ↔ 센서", ai_data.get("거리_커플링센서", "-")),
        ("센서 ↔ 센서", ai_data.get("거리_센서센서", "-")),
        ("모터 전(앞) Base 중심 → 센서(모터측)", ai_data.get("거리_모터전센서", "-")),
        ("모터 전(앞) Base 중심 ↔ 모터 후(뒤) Base 중심", ai_data.get("거리_모터전후", "-"))

    ]:

        _add_table_row_helper(

            data_table,

            ["", "", "", "", _dist, "", "", ""]

        )

        data_table.rows[-1].cells[0].text = _label

        data_table.rows[-1].cells[0].merge(

            data_table.rows[-1].cells[3]

        )

    _add_table_row_helper(

        data_table,

        [
            "모터-펌프", "수평",
            ai_data.get("각도오차_수평", ""),
            ai_data.get("옵셋오차_수평", ""),
            "-",
            ai_data.get("전면변위_수평", ""),
            ai_data.get("후면변위_수평", ""),
            ai_data.get("판정_수평", "")
        ]

    )

    _add_table_row_helper(

        data_table,

        [
            "모터-펌프", "수직",
            ai_data.get("각도오차_수직", ""),
            ai_data.get("옵셋오차_수직", ""),
            "-",
            ai_data.get("전면변위_수직", ""),
            ai_data.get("후면변위_수직", ""),
            ai_data.get("판정_수직", "")
        ]

    )

    _set_table_col_widths(

        data_table,

        [2.4, 0.7, 1.1, 1.0, 1.0, 1.0, 1.0, 0.8]

    )

    doc.add_paragraph("")

    # 추이 그래프 (이력이 쌓일수록 실제 추이가 나타남)
    if history_rows and len(history_rows) >= 1:

        doc.add_paragraph(
            "상태변화 및 추이분석"
        ).runs[0].bold = True

        try:

            chart_bytes = build_alignment_trend_chart(
                history_rows
            )

            chart_p = doc.add_paragraph()

            chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            chart_p.add_run().add_picture(

                io.BytesIO(chart_bytes),

                width=Inches(5.8)

            )

            if len(history_rows) < 2:

                doc.add_paragraph(

                    "※ 이번이 첫 측정이라 추이선이 점 하나로만 "
                    "표시됩니다. 다음 측정부터 자동으로 이력이 "
                    "누적되어 실제 추이 그래프가 됩니다."

                ).italic = True

        except Exception:

            pass

    if image_bytes:

        doc.add_paragraph("")

        doc.add_paragraph(
            "축정렬 장비 측정 자료"
        ).runs[0].bold = True

        try:

            img_p = doc.add_paragraph()

            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            img_p.add_run().add_picture(

                io.BytesIO(image_bytes),

                width=Inches(5.0)

            )

        except Exception:

            pass

    buffer = io.BytesIO()

    doc.save(buffer)

    return buffer.getvalue()


VIB_ISO10816_NOTE = "ISO10816-7 (펌프) / ISO10816-3 (모터) 기준 적용"

VA_GROUP_ORDER = [

    "펌프 반부하", "펌프 부하",
    "모터 부하", "모터 반부하"

]

VA_DEFECT_ITEMS = [

    "불평형", "축정렬불량", "전기적불평형",
    "베인이상", "베어링결함"

]


def build_vib_analysis_report_docx(

    pump,
    check_date_label,
    ai_data,
    images_dict=None

):

    doc = Document()

    _set_korean_font(doc)

    for section in doc.sections:

        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    logo_path = find_logo_file()

    if logo_path:

        try:

            logo_p = doc.add_paragraph()

            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            logo_p.add_run().add_picture(
                logo_path,
                width=Inches(1.3)
            )

        except Exception:

            pass

    title = doc.add_heading(
        f"{pump['site']} {pump['equip']} 진동분석 보고서 "
        f"({check_date_label})",
        level=0
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 개요
    doc.add_heading("1. 개요", level=1)

    doc.add_paragraph(

        "펌프모터의 설비상태 진단을 통한 불평형, 정렬불량, "
        "베어링 결함, 전기적 문제 등 결함 유형 식별과 고장이 "
        "발생하기 전에 스펙트럼·시간파형 분석을 통해 조기 경고 "
        "신호를 확인하고 정비 시기를 최적화하여 갑작스러운 "
        "안전사고를 예방함을 목적으로 함."

    )

    info_table = doc.add_table(rows=1, cols=4)

    info_table.style = "Light Grid Accent 1"

    info_table.rows[0].cells[0].text = "설비명"
    info_table.rows[0].cells[1].text = pump["equip"]
    info_table.rows[0].cells[2].text = "사업장"
    info_table.rows[0].cells[3].text = pump["site"]

    _add_table_row_helper(

        info_table,

        [
            "작업일자",
            check_date_label,
            "측정장비",
            ai_data.get("측정장비", "")
        ]

    )

    doc.add_paragraph("")

    # 2. 측정데이터 요약 결과
    doc.add_heading("2. 측정데이터 요약 결과", level=1)

    points = ai_data.get("points", {})

    sum_table = doc.add_table(rows=1, cols=7)

    sum_table.style = "Light Grid Accent 1"

    for i, h in enumerate(

        ["구분", "방향", "진동(rms,mm/s)", "판정",
         "pk-pk변위(µm)", "두드러진주파수(Hz)", "해석 키워드"]

    ):

        sum_table.rows[0].cells[i].text = h

    _style_header_row(sum_table)

    VA_POINT_LABELS = [

        ("모터 부하", "수평", "모터 부하 수평"),
        ("모터 부하", "수직", "모터 부하 수직"),
        ("모터 반부하", "수평", "모터 반부하 수평"),
        ("모터 반부하", "수직", "모터 반부하 수직"),
        ("모터 반부하", "축", "모터 반부하 축"),
        ("펌프 부하", "수평", "펌프 부하 수평"),
        ("펌프 부하", "수직", "펌프 부하 수직"),
        ("펌프 반부하", "수평", "펌프 반부하 수평"),
        ("펌프 반부하", "수직", "펌프 반부하 수직"),
        ("펌프 반부하", "축", "펌프 반부하 축")

    ]

    for group_label, direction, point_key in VA_POINT_LABELS:

        p_data = points.get(point_key, {})

        _add_table_row_helper(

            sum_table,

            [
                group_label,
                direction,
                p_data.get("overall_rms", ""),
                p_data.get("판정", ""),
                p_data.get("pk_pk", ""),
                p_data.get("주파수", ""),
                p_data.get("키워드", "")
            ]

        )

    _set_table_col_widths(

        sum_table,

        [0.85, 0.65, 0.85, 0.8, 0.75, 0.85, 1.5]

    )

    doc.add_paragraph("")

    doc.add_paragraph(
        "*연간점검으로 추이분석은 별도로 작성하지 않음"
    ).italic = True

    doc.add_paragraph("")

    # 3. 진동분석 의견
    doc.add_heading("3. 진동분석 의견", level=1)

    doc.add_paragraph(
        "☐ 종합 의견"
    ).runs[0].bold = True

    doc.add_paragraph(
        "○ " + ai_data.get("종합의견", "")
    )

    doc.add_paragraph(
        "☐ 측정 Point별 세부 의견"
    ).runs[0].bold = True

    groups = ai_data.get("groups", {})

    for group_name in VA_GROUP_ORDER:

        g = groups.get(group_name, {})

        doc.add_paragraph(
            f"○ {group_name}측"
        )

        for _line in g.get("세부의견", []):

            doc.add_paragraph(
                "- " + _line,
                style="List Bullet"
            )

    doc.add_page_break()

    # 붙임1. 관리상태 판단기준 (ISO 고정표)
    doc.add_heading(
        "붙임1. 관리상태 판단기준",
        level=1
    )

    doc.add_paragraph(
        "□ ISO10816-7 (적용대상: 1kW 이상 산업용 펌프, "
        "Category 2 일반펌프 기준)"
    )

    iso_pump = doc.add_table(rows=1, cols=2)

    iso_pump.style = "Light Grid Accent 1"

    for i, h in enumerate(["판정 Zone", "속도 기준(mm/s, rms)"]):

        iso_pump.rows[0].cells[i].text = h

    _style_header_row(iso_pump)

    for row in [

        ["A (양호)", "2.5 ~ 3.2 이하"],
        ["B (허용)", "3.2 ~ 5.0"],
        ["C (요주의)", "5.0 ~ 6.6"],
        ["D (불량)", "6.6 초과"]

    ]:

        _add_table_row_helper(iso_pump, row)

    doc.add_paragraph("")

    doc.add_paragraph(
        "□ ISO10816-3 (적용대상: 15kW 초과 산업용 모터, "
        "Group 2 중형모터·강성지지 기준)"
    )

    iso_motor = doc.add_table(rows=1, cols=2)

    iso_motor.style = "Light Grid Accent 1"

    for i, h in enumerate(["판정 Zone", "속도 기준(mm/s, rms)"]):

        iso_motor.rows[0].cells[i].text = h

    _style_header_row(iso_motor)

    for row in [

        ["A (양호)", "1.4 이하"],
        ["B (허용)", "1.4 ~ 2.3"],
        ["C (요주의)", "2.3 ~ 3.5"],
        ["D (불량)", "3.5 초과"]

    ]:

        _add_table_row_helper(iso_motor, row)

    doc.add_paragraph("")

    doc.add_paragraph(

        "※ 정수장 내 모터는 Group 2(강성지지대) 기준 적용. "
        f"{VIB_ISO10816_NOTE}"

    ).italic = True

    doc.add_page_break()

    # 붙임2. 측정·시험 DATA
    doc.add_heading(
        "붙임2. 측정·시험 DATA",
        level=1
    )

    doc.add_paragraph(
        "1. 총 진동값"
    ).runs[0].bold = True

    try:

        _diagram_bytes = build_pump_motor_diagram()

        _diagram_p = doc.add_paragraph()

        _diagram_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _diagram_p.add_run().add_picture(

            io.BytesIO(_diagram_bytes),

            width=Inches(4.5)

        )

    except Exception:

        pass

    doc.add_paragraph(
        "1-1. 진동 측정 값 (단위: mm/s rms)"
    )

    _pm_table = doc.add_table(rows=1, cols=4)

    _pm_table.style = "Light Grid Accent 1"

    for i, h in enumerate(["펌프측", "", "모터측", ""]):

        _pm_table.rows[0].cells[i].text = h

    _style_header_row(_pm_table)

    def _pt_lines(point_names):

        _lines = []

        for _dir, _pname in point_names:

            _p = points.get(_pname, {})

            _v = _p.get("overall_rms", "")

            if _v != "":

                _lines.append(f"{_dir}  {_v}")

        return "\n".join(_lines)

    _row1_left = _pt_lines([

        ("V", "펌프 반부하 수직"),
        ("H", "펌프 반부하 수평"),
        ("A", "펌프 반부하 축")

    ])

    _row1_right = _pt_lines([

        ("V", "모터 부하 수직"),
        ("H", "모터 부하 수평")

    ])

    _row2_left = _pt_lines([

        ("V", "펌프 부하 수직"),
        ("H", "펌프 부하 수평")

    ])

    _row2_right = _pt_lines([

        ("V", "모터 반부하 수직"),
        ("H", "모터 반부하 수평"),
        ("A", "모터 반부하 축")

    ])

    _add_table_row_helper(

        _pm_table,

        ["1  (반부하)", _row1_left, "3  (부하)", _row1_right]

    )

    _add_table_row_helper(

        _pm_table,

        ["2  (부하)", _row2_left, "4  (반부하)", _row2_right]

    )

    doc.add_paragraph("")

    doc.add_paragraph("")

    freq_table_data = ai_data.get("freq_table", {})

    # 설비정보는 원본엔 번호 없는 참고자료라, 있으면 조용히
    # 붙여둔다 (없어도 무방)

    for _comp in ["모터", "펌프"]:

        if images_dict and f"{_comp} 설비정보" in images_dict:

            try:

                doc.add_paragraph(
                    f"<{_comp} 설비정보>"
                )

                _info_p = doc.add_paragraph()

                _info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                _info_p.add_run().add_picture(

                    io.BytesIO(
                        images_dict[f"{_comp} 설비정보"]
                    ),

                    width=Inches(3.0)

                )

            except Exception:

                pass

    doc.add_paragraph("")

    doc.add_paragraph(
        "2. 주파수 분석"
    ).runs[0].bold = True

    for _comp, _label in [

        ("모터", "모터 베어링 주파수"),
        ("펌프", "펌프 베어링 주파수")

    ]:

        _key = f"{_comp} 주파수 분석"

        if images_dict and _key in images_dict:

            try:

                doc.add_paragraph(f"<{_label}>")

                img_p = doc.add_paragraph()

                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                img_p.add_run().add_picture(

                    io.BytesIO(images_dict[_key]),

                    width=Inches(3.5)

                )

            except Exception:

                pass

    doc.add_paragraph("")

    doc.add_paragraph(
        "2-1. 주파수 및 배수 확인"
    ).runs[0].bold = True

    for _comp in ["모터", "펌프"]:

        _freq_rows = freq_table_data.get(_comp, [])

        if _freq_rows:

            doc.add_paragraph(
                f"[{_comp}]"
            )

            freq_table = doc.add_table(rows=1, cols=3)

            freq_table.style = "Light Grid Accent 1"

            for i, h in enumerate(["구분", "주파수", "배수(nX)"]):

                freq_table.rows[0].cells[i].text = h

            _style_header_row(freq_table)

            for row in _freq_rows:

                _add_table_row_helper(

                    freq_table,

                    [
                        row.get("구분", ""),
                        row.get("주파수", ""),
                        row.get("배수", "")
                    ]

                )

            doc.add_paragraph("")

    doc.add_page_break()

    _GROUP_SECTION_TITLE = {

        "펌프 반부하": "진동 측정",
        "펌프 부하": "상태 확인",
        "모터 부하": "상태 확인",
        "모터 반부하": "상태 확인"

    }

    for _gi, group_name in enumerate(VA_GROUP_ORDER, start=3):

        g = groups.get(group_name, {})

        _section_word = _GROUP_SECTION_TITLE.get(
            group_name, "상태 확인"
        )

        doc.add_heading(
            f"{_gi}. {group_name}측 {_section_word}",
            level=2
        )

        # 해당 그룹의 원본 스펙트럼·시간파형 이미지들
        if images_dict:

            for point_key, img_bytes in images_dict.items():

                if point_key.startswith(group_name):

                    try:

                        doc.add_paragraph(
                            f"<{point_key}>"
                        )

                        img_p = doc.add_paragraph()

                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        img_p.add_run().add_picture(

                            io.BytesIO(img_bytes),

                            width=Inches(4.5)

                        )

                    except Exception:

                        continue

        doc.add_paragraph(
            f"{_gi}-1. {group_name}측 분석"
        ).runs[0].bold = True

        doc.add_paragraph(
            "1) 스펙트럼 그래프"
        )

        defect_table = doc.add_table(rows=1, cols=3)

        defect_table.style = "Light Grid Accent 1"

        for i, h in enumerate(["구분", "결함여부", "판단근거"]):

            defect_table.rows[0].cells[i].text = h

        _style_header_row(defect_table)

        _defects = {

            d.get("item"): d

            for d in g.get("defect_table", [])

        }

        for item in VA_DEFECT_ITEMS:

            d = _defects.get(item, {})

            _add_table_row_helper(

                defect_table,

                [
                    item,
                    d.get("result", "X"),
                    d.get("reason", "")
                ]

            )

        _set_table_col_widths(

            defect_table,

            [1.3, 0.9, 4.0]

        )

        doc.add_paragraph("")

        doc.add_paragraph(
            "2) 시간파형 그래프"
        )

        wave_table = doc.add_table(rows=1, cols=3)

        wave_table.style = "Light Grid Accent 1"

        for i, h in enumerate(["구분", "형상/특징", "시사점"]):

            wave_table.rows[0].cells[i].text = h

        _style_header_row(wave_table)

        for w in g.get("waveform_table", []):

            _add_table_row_helper(

                wave_table,

                [
                    w.get("방향", ""),
                    w.get("형상특징", ""),
                    w.get("시사점", "")
                ]

            )

        _set_table_col_widths(

            wave_table,

            [0.9, 2.9, 2.4]

        )

        doc.add_paragraph("")

        doc.add_paragraph(
            f"{_gi}-2. 해석"
        ).runs[0].bold = True

        for _i, _line in enumerate(g.get("해석", []), start=1):

            doc.add_paragraph(
                f"{_i}) {_line}"
            )

        doc.add_paragraph("")

    # 종합분석 / 향후대책
    doc.add_heading("7. 종합 분석", level=1)

    for _line in ai_data.get("종합분석", []):

        doc.add_paragraph(
            "- " + _line,
            style="List Bullet"
        )

    doc.add_heading("8. 향후 대책", level=1)

    for _i, _line in enumerate(

        ai_data.get("향후대책", []), start=1

    ):

        doc.add_paragraph(
            f"{_i}) {_line}"
        )

    buffer = io.BytesIO()

    doc.save(buffer)

    return buffer.getvalue()


def build_efficiency_report_docx(

    pump,
    month_label,
    ai_analysis_text,
    image_bytes=None

):

    doc = Document()

    _set_korean_font(doc)

    for section in doc.sections:

        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    logo_path = find_logo_file()

    if logo_path:

        try:

            logo_p = doc.add_paragraph()

            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            logo_p.add_run().add_picture(
                logo_path,
                width=Inches(1.3)
            )

        except Exception:

            pass

    title = doc.add_heading(
        "효율진단 보고서",
        level=0
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_p = doc.add_paragraph()

    info_p.add_run(

        f"설비명: {pump['equip']}  ·  사업장: {pump['site']}  ·  "
        f"대상기간: {month_label}"

    )

    doc.add_paragraph("")

    doc.add_heading(
        "AI 분석 소견",
        level=1
    )

    for _line in ai_analysis_text.split("\n"):

        if _line.strip():

            doc.add_paragraph(
                _line.strip()
            )

    if image_bytes:

        doc.add_paragraph("")

        doc.add_heading(
            "측정 원본 자료",
            level=1
        )

        try:

            img_p = doc.add_paragraph()

            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            img_p.add_run().add_picture(

                io.BytesIO(image_bytes),

                width=Inches(5.5)

            )

        except Exception:

            pass

    buffer = io.BytesIO()

    doc.save(buffer)

    return buffer.getvalue()


QR_UNIT_COST = {

    "QR 라벨 제작": 2500,

    "표면 세척/준비": 1000,

    "QR 부착 작업": 3000,

    "설비 등록 및 검수": 1500

}


def calculate_qr_cost(
    count
):

    unit = sum(
        QR_UNIT_COST.values()
    )

    return unit, unit * count



# ============================================================
# [pump_charts.py 내용 — 상태계산(pump_status) + 추세 차트]
# ============================================================
# ============================================================
# pump_charts.py
#
# 설비 상태 계산(pump_status) + 각종 시각화 차트 생성 함수 모음.
# pump_status는 df_history를 인자로 받는다 — 예전에는
# 전역변수를 몰래 참조해서 모듈 분리가 어려웠는데,
# 명시적으로 인자로 넘기게 고쳐서 어디서 호출하든
# 반드시 "어떤 이력 데이터 기준으로" 계산했는지 코드만 보고
# 알 수 있게 했다.
# ============================================================

import matplotlib.pyplot as plt
import pandas as pd



# ============================================================
# 11. 상태 계산
# ============================================================

def pump_status(
    pump,
    df_history
):

    name = pump["equip"]

    hours = pump["op_hours"]

    # ------------------------------------------------------
    # 1순위: 실제로 저장된 정밀진단 이력이 있으면 그 값을 사용한다.
    # (예전 버전은 설비명 문자열에 따라 값을 하드코딩해서
    #  정밀진단에서 아무리 값을 입력해도 다른 화면에 반영이
    #  안 되는 문제가 있었다. 이제는 최근 진단 결과를 읽어온다.)
    # ------------------------------------------------------

    latest = None

    if (

        not df_history.empty

        and
        "설비명" in df_history.columns

    ):

        pump_rows = df_history[

            df_history["설비명"] == name

        ]

        if not pump_rows.empty:

            latest = pump_rows.iloc[
                -1
            ]

    if (

        latest is not None

        and
        "효율측정값(%)" in latest.index

        and
        pd.notna(
            latest["효율측정값(%)"]
        )

    ):

        score = float(
            latest["종합점수"]
        )

        efficiency = float(
            latest["효율측정값(%)"]
        )

        vibration = float(
            latest["진동측정값(mm/s)"]
        ) if pd.notna(
            latest.get("진동측정값(mm/s)")
        ) else 2.2

        temperature = float(
            latest["온도측정값(°C)"]
        ) if pd.notna(
            latest.get("온도측정값(°C)")
        ) else round(
            43 + vibration * 2.3,
            1
        )

        current_amp = (

            float(latest["전류측정값(A)"])

            if (

                "전류측정값(A)" in latest.index

                and
                pd.notna(latest.get("전류측정값(A)"))

            )

            else None

        )

    else:

        # ------------------------------------------------------
        # 2순위: 아직 정밀진단 이력이 없는 설비는
        # 데모용 기본값(과거 방식)을 그대로 사용한다.
        # ------------------------------------------------------

        score = 85

        vibration = 2.2

        efficiency = 88

        if name == "가압펌프 #2":

            score = 68

            vibration = 5.8

            efficiency = 73

        elif name == "가압펌프 #3":

            score = 64

            vibration = 6.8

            efficiency = 68

        elif name == "가압펌프 #5":

            score = 74

            vibration = 4.2

            efficiency = 75

        elif hours > 11000:

            score = 77

            vibration = 3.8

            efficiency = 81

        elif hours > 10000:

            score = 82

            vibration = 2.8

            efficiency = 84

        temperature = round(
            43 + vibration * 2.3,
            1
        )

        current_amp = None

    grade = get_final_grade(
        score
    )

    if grade in ["A", "B"]:

        status = "정상"

    elif grade == "C":

        status = "관찰"

    else:

        status = "정비검토"

    next_due_hours, remaining_hours = estimate_next_overhaul(
        hours
    )

    return {

        "점수": score,

        "등급": grade,

        "상태": status,

        "진동": vibration,

        "효율": efficiency,

        "온도": temperature,

        "전류": current_amp,

        "실측이력있음": latest is not None,

        "다음오버홀예정시간": next_due_hours,

        "다음오버홀까지남은시간": remaining_hours

    }



# ============================================================
# 11-1. 진동 추세 차트 (AI 진단 공용)
#
# AI 이상징후 페이지와 QR 포털의 AI진단 탭이
# 동일한 그래프를 공유하므로 함수로 분리한다.
#
# 11-0. 실측 이력 조회 헬퍼
#
# 예전에는 그래프 5개가 전부 "현재값 하나를 기준으로 수식으로
# 만들어낸" 가짜 추세였다. 이제 실제 진단이력이 2건 이상
# 쌓여 있으면 그 실측값으로 그래프를 그리고, 부족하면
# 예시 데이터를 쓰되 화면에 "예시 데이터"라고 명확히 표시한다.
# ============================================================

def get_real_history_series(
    df_history,
    pump,
    column
):

    if (

        df_history is None

        or
        df_history.empty

        or
        "설비명" not in df_history.columns

        or
        column not in df_history.columns

    ):

        return [], []

    rows = df_history[

        df_history["설비명"] == pump["equip"]

    ].copy()

    if rows.empty:

        return [], []

    rows = rows[

        pd.notna(rows[column])

        &
        (rows[column].astype(str) != "")

    ]

    if rows.empty:

        return [], []

    rows = rows.sort_values(
        "점검일"
    )

    dates = rows["점검일"].astype(str).tolist()

    values = rows[column].astype(float).tolist()

    return dates, values


# ============================================================
# 11-0-1. 통계적 이상탐지 (이동평균 + 표준편차 기반)
#
# 예전에는 "진동 >= 7.1이면 위험"처럼 고정 임계값 if문뿐이라
# "AI 이상징후"라는 메뉴 이름이 무색했다. 실측 이력이
# 충분히 쌓이면(5건 이상), 최근값이 그 설비 자신의 평균적인
# 운전 패턴에서 통계적으로 벗어났는지(z-score)를 함께 보여준다.
# 데이터가 부족하면 "아직 활성화 안 됨"이라고 솔직하게 표시한다.
# ============================================================

ANOMALY_MIN_SAMPLES = 5

ANOMALY_Z_THRESHOLD = 2.0


def detect_statistical_anomaly(values):

    if len(values) < ANOMALY_MIN_SAMPLES:

        return None

    history = values[:-1]

    latest = values[-1]

    mean = sum(history) / len(history)

    variance = sum(

        (v - mean) ** 2 for v in history

    ) / len(history)

    stdev = variance ** 0.5

    if stdev == 0:

        z = 0.0

    else:

        z = (latest - mean) / stdev

    return {

        "mean": mean,

        "stdev": stdev,

        "latest": latest,

        "z": z,

        "is_anomaly": abs(z) >= ANOMALY_Z_THRESHOLD,

        "n": len(values)

    }


def build_vibration_trend_fig(
    pump,
    result,
    df_history=None,
    figsize=(9, 4)
):

    months, vibration, use_real = get_vibration_trend_data(
        pump,
        result,
        df_history
    )

    _th = get_alert_thresholds()

    fig, ax = plt.subplots(
        figsize=figsize
    )

    ax.plot(
        months,
        vibration,
        marker="o",
        linewidth=2
    )

    ax.axhline(
        _th["vib_watch"],
        linestyle="--",
        label="관찰 기준"
    )

    ax.axhline(
        _th["vib_danger"],
        linestyle=":",
        label="주의 기준"
    )

    ax.set_ylabel(
        "진동 (mm/s)"
    )

    title_suffix = (

        f"(실측 {len(vibration)}건)"

        if use_real

        else "(예시 데이터)"

    )

    ax.set_title(
        f"{pump['equip']} 진동 추세 {title_suffix}"
    )

    ax.grid(
        alpha=0.2
    )

    ax.legend()

    return fig


def build_trend_fig(

    pump,

    label,

    unit,

    months,

    values,

    watch_threshold=None,

    danger_threshold=None,

    figsize=(6.2, 3.4)

):

    # 진동 외 다른 지표(효율·온도·CBM Score 등)에도
    # 재사용하는 범용 추세 차트 함수.

    fig, ax = plt.subplots(
        figsize=figsize
    )

    ax.plot(
        months,
        values,
        marker="o",
        linewidth=2
    )

    if watch_threshold is not None:

        ax.axhline(
            watch_threshold,
            linestyle="--",
            label="관찰 기준"
        )

    if danger_threshold is not None:

        ax.axhline(
            danger_threshold,
            linestyle=":",
            label="주의 기준"
        )

    ax.set_ylabel(
        f"{label} ({unit})"
    )

    ax.set_title(
        f"{pump['equip']} {label} 추세"
    )

    ax.grid(
        alpha=0.2
    )

    if watch_threshold is not None or danger_threshold is not None:

        ax.legend()

    return fig


TREND_MONTHS = [

    "2025.01",
    "2025.06",
    "2025.12",
    "2026.03",
    "2026.08",
    "2026.12"

]


# ============================================================
# 11-0-2. 화면 그래프와 엑셀 그래프가 같은 숫자를 쓰도록
# 데이터 생성 부분만 뽑아낸 함수들.
#
# 엑셀로 내보낼 때도 화면에 보이는 것과 똑같은 추세 그래프가
# 나오게 하려면, 화면(matplotlib)과 엑셀(openpyxl)이 서로
# 다른 계산식을 쓰면 안 된다. 이 함수들이 "진짜 데이터"를
# 만들고, 화면 차트 함수와 엑셀 내보내기 양쪽에서 그대로
# 가져다 쓴다.
# ============================================================

def get_vibration_trend_data(
    pump,
    result,
    df_history=None
):

    real_dates, real_values = (

        get_real_history_series(
            df_history,
            pump,
            "진동측정값(mm/s)"
        )

        if df_history is not None

        else ([], [])

    )

    use_real = len(real_values) >= 2

    if use_real:

        return real_dates, real_values, True

    base = result["진동"]

    values = [

        max(1.2, base - 2.5),
        max(1.4, base - 2),
        max(1.6, base - 1.4),
        max(1.8, base - 0.8),
        base,
        base + 1.2

    ]

    return TREND_MONTHS, values, False


def get_efficiency_trend_data(
    pump,
    result,
    df_history=None
):

    real_dates, real_values = (

        get_real_history_series(
            df_history,
            pump,
            "효율측정값(%)"
        )

        if df_history is not None

        else ([], [])

    )

    use_real = len(real_values) >= 2

    if use_real:

        return real_dates, real_values, True

    base = result["효율"]

    values = [

        min(100, base + 8),
        min(100, base + 6),
        min(100, base + 4),
        min(100, base + 2),
        base,
        max(0, base - 3)

    ]

    return TREND_MONTHS, values, False


def get_temperature_trend_data(
    pump,
    result,
    df_history=None
):

    real_dates, real_values = (

        get_real_history_series(
            df_history,
            pump,
            "온도측정값(°C)"
        )

        if df_history is not None

        else ([], [])

    )

    use_real = len(real_values) >= 2

    if use_real:

        return real_dates, real_values, True

    base = result["온도"]

    values = [

        max(35, base - 6),
        max(37, base - 4.5),
        max(39, base - 3),
        max(40, base - 1.5),
        base,
        base + 2

    ]

    return TREND_MONTHS, values, False


def get_op_hours_trend_data(
    pump
):

    current_hours = pump["op_hours"]

    values = [

        max(0, current_hours - 5000),
        max(0, current_hours - 4000),
        max(0, current_hours - 3000),
        max(0, current_hours - 2000),
        max(0, current_hours - 1000),
        current_hours

    ]

    return TREND_MONTHS, values, False


def build_equipment_export_sheets(
    pump,
    result,
    df_history
):

    # 설비관리/QR포털/AI이상징후 페이지처럼 화면에 여러
    # 추세 그래프가 있는 페이지를 엑셀로 내보낼 때 쓰는
    # 공용 함수. 화면과 같은 데이터로 시트별 그래프를 만든다.

    summary_df = pd.DataFrame(

        [

            ["사업장", pump["site"]],
            ["설비명", pump["equip"]],
            ["제조사", pump["maker"]],
            ["모델명", pump["model"]],
            ["운전시간(h)", pump["op_hours"]],
            ["CBM Score", result["점수"]],
            ["등급", result["등급"]],
            ["상태", result["상태"]],
            ["효율(%)", result["효율"]],
            ["진동(mm/s)", result["진동"]],
            ["온도(°C)", result["온도"]]

        ],

        columns=["항목", "값"]

    )

    vib_dates, vib_values, vib_real = get_vibration_trend_data(
        pump, result, df_history
    )

    eff_dates, eff_values, eff_real = get_efficiency_trend_data(
        pump, result, df_history
    )

    temp_dates, temp_values, temp_real = get_temperature_trend_data(
        pump, result, df_history
    )

    hour_dates, hour_values, _ = get_op_hours_trend_data(
        pump
    )

    vib_df = pd.DataFrame(
        {
            "시점": vib_dates,
            "진동(mm/s)": vib_values
        }
    )

    eff_df = pd.DataFrame(
        {
            "시점": eff_dates,
            "효율(%)": eff_values
        }
    )

    temp_df = pd.DataFrame(
        {
            "시점": temp_dates,
            "온도(°C)": temp_values
        }
    )

    hour_df = pd.DataFrame(
        {
            "시점": hour_dates,
            "누적운전시간(h)": hour_values
        }
    )

    real_tag = lambda is_real: (

        "실측" if is_real else "예시데이터"

    )

    return [

        {
            "name": "요약정보",
            "df": summary_df,
            "chart": None
        },

        {
            "name": "진동추세",
            "df": vib_df,
            "chart": "line",
            "title": f"{pump['equip']} 진동 추세 ({real_tag(vib_real)})"
        },

        {
            "name": "효율추세",
            "df": eff_df,
            "chart": "bar",
            "title": f"{pump['equip']} 효율 추세 ({real_tag(eff_real)})"
        },

        {
            "name": "온도추세",
            "df": temp_df,
            "chart": "line",
            "title": f"{pump['equip']} 온도 추세 ({real_tag(temp_real)})"
        },

        {
            "name": "운전시간추세",
            "df": hour_df,
            "chart": "line",
            "title": f"{pump['equip']} 누적 운전시간 추세 (추정치)"
        }

    ]


def build_efficiency_trend_fig(
    pump,
    result,
    df_history=None,
    figsize=(6.2, 3.4)
):

    # 효율은 막대그래프로 표현

    months_e, values, use_real = get_efficiency_trend_data(
        pump,
        result,
        df_history
    )

    _th = get_alert_thresholds()

    colors = [

        "#e03131" if v < _th["eff_danger"]
        else "#f08c00" if v < _th["eff_watch"]
        else "#087ea4"

        for v in values

    ]

    fig, ax = plt.subplots(
        figsize=figsize
    )

    ax.bar(

        months_e,

        values,

        color=colors,

        width=0.55

    )

    ax.axhline(
        _th["eff_watch"],
        linestyle="--",
        color="#a16207",
        label="관찰 기준"
    )

    ax.axhline(
        _th["eff_danger"],
        linestyle=":",
        color="#c62828",
        label="주의 기준"
    )

    ax.set_ylabel(
        "효율 (%)"
    )

    title_suffix = (

        f"(실측 {len(values)}건)"

        if use_real

        else "(예시 데이터)"

    )

    ax.set_title(
        f"{pump['equip']} 효율 추세 {title_suffix}"
    )

    ax.set_ylim(
        0,
        105
    )

    ax.grid(
        alpha=0.2,
        axis="y"
    )

    ax.legend()

    return fig


def build_temperature_trend_fig(
    pump,
    result,
    df_history=None,
    figsize=(6.2, 3.4)
):

    # 온도는 영역(면적)그래프로 표현

    months_t, values, use_real = get_temperature_trend_data(
        pump,
        result,
        df_history
    )

    fig, ax = plt.subplots(
        figsize=figsize
    )

    ax.fill_between(

        months_t,

        values,

        color="#ff922b",

        alpha=0.35

    )

    ax.plot(

        months_t,

        values,

        marker="o",

        linewidth=2,

        color="#e8590c"

    )

    _th = get_alert_thresholds()

    ax.axhline(
        _th["temp_watch"],
        linestyle="--",
        color="#a16207",
        label="관찰 기준"
    )

    ax.axhline(
        _th["temp_danger"],
        linestyle=":",
        color="#c62828",
        label="주의 기준"
    )

    ax.set_ylabel(
        "온도 (°C)"
    )

    title_suffix = (

        f"(실측 {len(values)}건)"

        if use_real

        else "(예시 데이터)"

    )

    ax.set_title(
        f"{pump['equip']} 온도 추세 {title_suffix}"
    )

    ax.grid(
        alpha=0.2
    )

    ax.legend()

    return fig


@st.cache_data(show_spinner=False)
def build_status_donut_fig(
    normal,
    watch,
    repair,
    figsize=(1.8, 1.8)
):

    # 홈 화면 상단의 정상/관찰/정비검토 비율 도넛차트

    labels = []

    sizes = []

    colors = []

    for label, value, color in (

        (f"정상 {normal}대", normal, "#087f5b"),
        (f"관찰 {watch}대", watch, "#a16207"),
        (f"정비검토 {repair}대", repair, "#c62828")

    ):

        if value > 0:

            labels.append(label)

            sizes.append(value)

            colors.append(color)

    fig, ax = plt.subplots(
        figsize=figsize,
        dpi=100
    )

    if sizes:

        ax.pie(

            sizes,

            colors=colors,

            startangle=90,

            wedgeprops=dict(
                width=0.38,
                edgecolor="white"
            )

        )

        ax.text(

            0,

            0,

            f"{normal + watch + repair}대",

            ha="center",

            va="center",

            fontsize=11,

            fontweight="bold",

            color="#0f3552"

        )

        ax.legend(

            labels,

            loc="lower center",

            bbox_to_anchor=(0.5, -0.22),

            ncol=1,

            fontsize=7,

            frameon=False

        )

    else:

        ax.text(
            0.5,
            0.5,
            "데이터 없음",
            ha="center"
        )

    ax.set_aspect(
        "equal"
    )

    return fig


def build_svg_sparkline(
    values,
    color="#087ea4",
    width=90,
    height=26,
    tooltip=None
):

    # matplotlib 없이 순수 SVG로 만드는 미니 스파크라인.
    # 카드 10개마다 matplotlib을 새로 그리면 느려지므로,
    # 가벼운 인라인 SVG로 대체했다.
    #
    # <title> 요소를 넣어두면 마우스를 올렸을 때 브라우저가
    # 기본 툴팁으로 실제 수치를 보여준다 (별도 JS 없이 가능).

    if not values or len(values) < 2:

        return ""

    vmin = min(values)

    vmax = max(values)

    rng = (vmax - vmin) or 1

    step = width / (len(values) - 1)

    points = []

    for i, v in enumerate(values):

        x = i * step

        y = height - ((v - vmin) / rng) * (height - 4) - 2

        points.append(
            f"{x:.1f},{y:.1f}"
        )

    points_str = " ".join(
        points
    )

    title_elem = (

        f"<title>{tooltip}</title>"

        if tooltip

        else ""

    )

    return (

        f'<svg width="{width}" height="{height}" '
        f'viewBox="0 0 {width} {height}" '
        f'xmlns="http://www.w3.org/2000/svg">'
        f'<polyline points="{points_str}" fill="none" '
        f'stroke="{color}" stroke-width="2" '
        f'stroke-linecap="round" stroke-linejoin="round">'
        f'{title_elem}'
        f'</polyline>'
        f'</svg>'

    )


def grade_bar_color(
    score
):

    # 게이지바 색상을 등급(A~E) 5단계 기준과 통일한다.
    # (예전에는 80/60 기준 3단계뿐이라, 74점과 65점이
    #  똑같은 노란색으로 보여 등급 차이가 안 느껴졌다)

    if score >= 90:

        return "#087f5b"

    if score >= 80:

        return "#2f9e44"

    if score >= 70:

        return "#f08c00"

    if score >= 60:

        return "#e8590c"

    return "#c62828"


def build_score_gauge_fig(
    pump,
    result,
    figsize=(6.2, 2.3)
):

    # CBM Score는 게이지(구간별 색상 막대 + 현재값 표시)로 표현

    score = result["점수"]

    zones = [

        (0, 60, "#ffc9c9"),

        (60, 70, "#ffd8a8"),

        (70, 80, "#fff3bf"),

        (80, 90, "#d3f9d8"),

        (90, 100, "#b2f2bb")

    ]

    fig, ax = plt.subplots(
        figsize=figsize
    )

    for start, end, color in zones:

        ax.barh(

            0,

            end - start,

            left=start,

            color=color,

            height=0.6,

            edgecolor="white"

        )

    ax.barh(

        0,

        2,

        left=max(
            0,
            min(score, 98)
        ),

        color="#083b5c",

        height=0.9

    )

    ax.set_xlim(
        0,
        100
    )

    ax.set_yticks(
        []
    )

    ax.set_xlabel(
        "CBM Score (0~100)"
    )

    ax.set_title(
        f"{pump['equip']} 현재 CBM Score : {score}점 ({result['등급']}등급)"
    )

    return fig


def build_fleet_average_trend_fig(
    all_pumps,
    df_history,
    metric_key,
    label,
    unit,
    color,
    figsize=(9, 3.5)
):

    # 설비 하나씩 보는 추세 그래프만 있고, "전체 설비 평균이
    # 어떻게 변해왔나"를 보는 화면이 없었다. 각 설비의 현재값을
    # 기준으로 한 예시 추세(설비 개별 화면과 같은 방식)를 평균내서
    # 전사 차원의 큰 흐름을 보여준다.
    #
    # (개별 설비 화면처럼 실측/예시를 따로 구분하지 않고,
    #  여러 설비를 한 축(6개 시점)으로 맞추기 위해
    #  일관된 예시 추세식을 사용한다 — 개별 설비의 진짜 실측
    #  추세는 각 설비 화면에서 확인하는 것이 정확하다)

    series = []

    for pump in all_pumps:

        result = pump_status(
            pump,
            df_history
        )

        base = result[metric_key]

        if metric_key == "효율":

            values = [
                min(100, base + 8),
                min(100, base + 6),
                min(100, base + 4),
                min(100, base + 2),
                base,
                max(0, base - 3)
            ]

        elif metric_key == "진동":

            values = [
                max(1.2, base - 2.5),
                max(1.4, base - 2),
                max(1.6, base - 1.4),
                max(1.8, base - 0.8),
                base,
                base + 1.2
            ]

        elif metric_key == "온도":

            values = [
                max(35, base - 6),
                max(37, base - 4.5),
                max(39, base - 3),
                max(40, base - 1.5),
                base,
                base + 2
            ]

        else:

            values = [
                min(100, base + 12),
                min(100, base + 9),
                min(100, base + 6),
                min(100, base + 3),
                base,
                max(0, base - 4)
            ]

        series.append(
            values
        )

    avg_values = [

        sum(v[i] for v in series) / len(series)

        for i in range(6)

    ]

    fig, ax = plt.subplots(
        figsize=figsize
    )

    ax.plot(

        TREND_MONTHS,

        avg_values,

        marker="o",

        linewidth=2.4,

        color=color

    )

    ax.fill_between(

        TREND_MONTHS,

        avg_values,

        alpha=0.12,

        color=color

    )

    ax.set_ylabel(
        f"{label} ({unit})"
    )

    ax.set_title(

        f"전체 {len(all_pumps)}대 평균 {label} 추이 (참고용)"

    )

    ax.grid(
        alpha=0.2
    )

    return fig, avg_values


def build_fleet_compare_fig(
    pump,
    result,
    all_pumps,
    df_history,
    figsize=(9, 3)
):

    # 6번째 그래프 유형: 이 설비 vs 전체 설비 평균 비교(막대 3개짜리)
    # CBM 페이지의 비교뷰와 동일한 로직을 재사용한다.

    fleet_scores = [

        pump_status(p, df_history)

        for p in all_pumps

    ]

    fleet_avg_eff = sum(
        s["효율"] for s in fleet_scores
    ) / len(fleet_scores)

    fleet_avg_vib = sum(
        s["진동"] for s in fleet_scores
    ) / len(fleet_scores)

    fleet_avg_score = sum(
        s["점수"] for s in fleet_scores
    ) / len(fleet_scores)

    fig, (ax1, ax2, ax3) = plt.subplots(

        1,
        3,

        figsize=figsize

    )

    def _draw(
        ax,
        title,
        pump_value,
        fleet_value,
        unit
    ):

        bars = ax.bar(

            ["선택 설비", "전체 평균"],

            [pump_value, fleet_value],

            color=["#087ea4", "#adb5bd"]

        )

        ax.set_title(
            title
        )

        ax.set_ylabel(
            unit
        )

        for b in bars:

            ax.text(

                b.get_x() + b.get_width() / 2,

                b.get_height(),

                f"{b.get_height():.1f}",

                ha="center",

                va="bottom",

                fontsize=8

            )

    _draw(
        ax1,
        "효율",
        result["효율"],
        fleet_avg_eff,
        "%"
    )

    _draw(
        ax2,
        "진동",
        result["진동"],
        fleet_avg_vib,
        "mm/s"
    )

    _draw(
        ax3,
        "CBM Score",
        result["점수"],
        fleet_avg_score,
        "점"
    )

    fig.suptitle(
        f"{pump['equip']} vs 전체 {len(all_pumps)}대 평균"
    )

    fig.tight_layout()

    return fig


def build_op_hours_trend_fig(
    pump,
    figsize=(9, 3.2)
):

    # 누적 운전시간은 영역(면적)그래프로 표현

    _, values, _ = get_op_hours_trend_data(
        pump
    )

    fig, ax = plt.subplots(
        figsize=figsize
    )

    ax.fill_between(

        TREND_MONTHS,

        values,

        color="#087ea4",

        alpha=0.30

    )

    ax.plot(

        TREND_MONTHS,

        values,

        marker="o",

        linewidth=2,

        color="#063b63"

    )

    ax.set_ylabel(
        "누적 운전시간 (h)"
    )

    ax.set_title(
        f"{pump['equip']} 누적 운전시간 추세 (추정 — "
        f"과거 시점별 실측 기록은 별도로 저장하지 않습니다)"
    )

    ax.grid(
        alpha=0.2
    )

    return fig



# 배점 합계 검증 — 어긋나 있으면 화면에 크게 경고하고 멈춘다.
# (예전에는 이런 실수가 나도 조용히 잘못된 점수를 계속 뱉어냈다)

_weight_errors = validate_eval_weights()

if _weight_errors:

    st.error(
        "🚨 진단항목 배점 설정에 오류가 있습니다. "
        "관리자에게 문의하세요.\n\n"
        +
        "\n".join(
            f"- {e}" for e in _weight_errors
        )
    )

    st.stop()


# df_history / ALL_PUMPS는 매 요청(rerun)마다 최신 데이터를
# 반영해야 하므로, 모듈 import 시점이 아니라 여기서 매번 새로 읽는다.

df_history = read_excel(
    DB_FILE_PATH,
    "진단이력"
)

ALL_PUMPS = get_all_pumps()

if read_errors:

    st.warning(
        "⚠️ 일부 데이터 파일을 읽는 중 문제가 발생했습니다 "
        "(파일이 없는 것과는 다른 문제입니다):\n\n"
        +
        "\n".join(
            f"- {e}" for e in read_errors
        )
    )


# ============================================================
# 12-1. 로그인 게이트 · 보기 전용 모드
#
# URL만 알면 누구나 데이터를 수정할 수 있던 문제에 대한
# 최소한의 보완. 정식 계정 시스템은 아니고
# 단순 PIN 게이트 수준이라는 점을 감안해서 사용할 것.
# ============================================================

if "authenticated" not in st.session_state:

    st.session_state.authenticated = False

if "read_only" not in st.session_state:

    st.session_state.read_only = False

if "entered_as_viewer" not in st.session_state:

    st.session_state.entered_as_viewer = False

if "user_name" not in st.session_state:

    # 예전에는 모든 저장 기록에 "최진욱"이 하드코딩되어 있어서
    # 여러 사람이 쓰기 시작하면 누가 입력했는지 구분할 수 없었다.
    # 이제 접속할 때 이름을 입력받아 세션에 저장해두고,
    # 저장 시점마다 이 이름을 사용한다.

    st.session_state.user_name = "최진욱"


def is_read_only():

    return st.session_state.read_only


SESSION_TIMEOUT_MINUTES = 60

if (

    st.session_state.authenticated

    and "_login_time" in st.session_state

):

    _elapsed_minutes = (

        datetime.now()

        -
        datetime.fromisoformat(
            st.session_state["_login_time"]
        )

    ).total_seconds() / 60

    if _elapsed_minutes > SESSION_TIMEOUT_MINUTES:

        st.session_state.authenticated = False

        log_audit(

            "세션 자동 로그아웃",

            st.session_state.user_name,

            f"{SESSION_TIMEOUT_MINUTES}분 비활동"

        )

        st.info(

            f"🔒 {SESSION_TIMEOUT_MINUTES}분 이상 활동이 "
            "없어 자동으로 로그아웃되었습니다. 다시 "
            "로그인해주세요."

        )

if st.session_state.authenticated:

    # 활동(페이지 조작)이 있을 때마다 기준 시각을 지금으로
    # 갱신한다. 다음 요청이 SESSION_TIMEOUT_MINUTES보다 늦게
    # 들어오면 위 검사에서 자동 로그아웃된다.

    st.session_state["_login_time"] = (

        datetime.now().isoformat()

    )


if LOGIN_GATE_ENABLED and not st.session_state.authenticated:

    _login_logo_html = get_logo_base64_html(
        max_height_px=52
    )

    _login_mascot_html = get_mascot_base64_html(
        max_height_px=110,
        name="mascot_safety"
    )

    st.markdown(

        f"""
        <div style="
        background: linear-gradient(
            135deg,
            #052c4a 0%, #063b63 35%, #087ea4 75%, #11a6c9 100%
        );
        border-radius: 20px;
        padding: 46px 36px 40px 36px;
        margin-bottom: 22px;
        text-align: center;
        box-shadow: 0 14px 34px rgba(3, 65, 100, 0.22);
        position: relative;">

            {f'<div style="position:absolute; top:14px; right:22px;">{_login_mascot_html}</div>' if _login_mascot_html else ''}

            <div style="
            background:white; border-radius:14px;
            display:inline-block; padding:12px 22px;
            margin-bottom:18px;">
            {_login_logo_html if _login_logo_html else '<span style="font-size:2.2rem;">💧</span>'}
            </div>

            <div style="
            color:white; font-size:1.6rem; font-weight:900;
            letter-spacing:-0.5px;">
            설비관리 통합 플랫폼
            </div>

            <div style="
            color:rgba(255,255,255,0.85); font-size:0.95rem;
            margin-top:8px;">
            접속하려면 이름과 PIN 번호를 입력하세요.
            </div>

        </div>
        """,

        unsafe_allow_html=True

    )

    login_col1, login_col2, login_col3 = st.columns(
        [1, 2, 1]
    )

    with login_col2:

        name_input = st.text_input(

            "이름",

            value="최진욱",

            key="name_input"

        )

        pin_input = st.text_input(

            "PIN 번호",

            type="password",

            key="pin_input"

        )

        LOGIN_LOCKOUT_THRESHOLD = 5

        LOGIN_LOCKOUT_MINUTES = 5

        _locked_until = st.session_state.get(
            "_login_locked_until"
        )

        _is_locked = False

        if _locked_until:

            _remaining = (

                datetime.fromisoformat(_locked_until)

                -
                datetime.now()

            ).total_seconds()

            if _remaining > 0:

                _is_locked = True

                st.error(

                    f"🔒 PIN을 너무 많이 틀려서 "
                    f"{int(_remaining // 60) + 1}분간 "
                    "접속이 잠겼습니다. 잠시 후 다시 "
                    "시도해주세요."

                )

            else:

                st.session_state["_login_locked_until"] = None

                st.session_state["_login_fail_count"] = 0

        if st.button(

            "🔓 접속",

            type="primary",

            use_container_width=True,

            disabled=_is_locked

        ):

            if pin_input == AUTH_PIN:

                st.session_state.authenticated = True

                st.session_state.entered_as_viewer = False

                st.session_state["_login_fail_count"] = 0

                st.session_state["_login_locked_until"] = None

                st.session_state.user_name = (

                    name_input.strip()

                    or
                    "최진욱"

                )

                st.session_state["_login_time"] = (

                    datetime.now().isoformat()

                )

                log_audit(

                    "로그인 성공",

                    st.session_state.user_name

                )

                st.rerun()

            else:

                log_audit(

                    "로그인 실패",

                    name_input.strip() or "(이름 미입력)"

                )

                _fail_count = st.session_state.get(

                    "_login_fail_count",
                    0

                ) + 1

                st.session_state["_login_fail_count"] = _fail_count

                if _fail_count >= LOGIN_LOCKOUT_THRESHOLD:

                    st.session_state["_login_locked_until"] = (

                        (
                            datetime.now()
                            +
                            timedelta(
                                minutes=LOGIN_LOCKOUT_MINUTES
                            )
                        ).isoformat()

                    )

                    log_audit(

                        "로그인 잠금",

                        name_input.strip() or "(이름 미입력)",

                        f"{LOGIN_LOCKOUT_THRESHOLD}회 연속 실패로 "
                        f"{LOGIN_LOCKOUT_MINUTES}분 잠금"

                    )

                    st.rerun()

                else:

                    st.error(

                        f"PIN 번호가 올바르지 않습니다. "
                        f"({_fail_count}/{LOGIN_LOCKOUT_THRESHOLD}회, "
                        f"{LOGIN_LOCKOUT_THRESHOLD}회 실패 시 "
                        f"{LOGIN_LOCKOUT_MINUTES}분간 잠깁니다)"

                    )

    st.stop()


# ============================================================
# 13. 사이드바
# ============================================================

if "page" not in st.session_state:

    # QR 스캔으로 들어온 경우 URL의 ?page=QR&equip=... 를
    # 그대로 반영해서 해당 설비 화면으로 바로 이동시킨다.

    st.session_state.page = st.query_params.get(
        "page",
        "홈"
    )


def go_to_page(
    page_key
):

    # 버튼의 on_click 콜백으로 사용.
    # 상태 변경 직후 Streamlit이 자동으로 rerun을 실행하므로
    # 별도의 st.rerun() 호출이 필요 없고,
    # 모바일 인앱 브라우저(카카오톡 등)에서도
    # 클릭 반응이 더 안정적이다.

    st.session_state.page = page_key

    st.query_params["page"] = page_key

    if "equip" in st.query_params:

        del st.query_params["equip"]


PAGE_GROUP_MAP = {

    "통합검색": "📍 현장 업무",
    "QR": "📍 현장 업무",

    "진단": "🔍 진단 · 판단",
    "CBM": "🔍 진단 · 판단",
    "오버홀": "🔍 진단 · 판단",

    "AI": "📊 분석 · 리포트",
    "전사트렌드": "📊 분석 · 리포트",
    "ROI": "📊 분석 · 리포트",
    "KPI": "📊 분석 · 리포트",
    "보고서": "📊 분석 · 리포트",

    "설비": "⚙️ 관리",
    "노하우": "⚙️ 관리",
    "설계적산": "⚙️ 관리",
    "백업": "⚙️ 관리",
    "업데이트내역": "⚙️ 관리",

    "토목": "🏗️ 부서별 관리",
    "전기": "🏗️ 부서별 관리",
    "전자": "🏗️ 부서별 관리",
    "행정": "🏗️ 부서별 관리",
    "안전": "🏗️ 부서별 관리",

}


def render_back_to_home_button(key_suffix):

    # 브라우저 뒤로가기가 Streamlit 특성상 두 번 눌러야 하는
    # 경우가 있어서, 확실하게 한 번에 홈으로 돌아갈 수 있는
    # 버튼을 각 페이지 상단에도 심어둔다.

    _current_page = st.session_state.page

    _group = PAGE_GROUP_MAP.get(_current_page)

    try:

        _current_label = MENU_LABEL_BY_KEY.get(

            _current_page,

            _current_page

        )

    except NameError:

        _current_label = _current_page

    _breadcrumb_parts = ["🏠 홈"]

    if _group:

        _breadcrumb_parts.append(_group)

    _breadcrumb_parts.append(_current_label)

    st.caption(
        " › ".join(_breadcrumb_parts)
    )

    if st.button(

        "← 홈으로",

        key=f"back_to_home_{key_suffix}"

    ):

        st.session_state.page = "홈"

        st.query_params["page"] = "홈"

        if "equip" in st.query_params:

            del st.query_params["equip"]

        st.rerun()


with st.sidebar:


    _sidebar_logo_html = get_logo_base64_html(
        max_height_px=40
    )

    st.markdown(
        f"""
        <div style="
        text-align:center;
        padding:8px 0 18px 0;
        ">
            <div style="
            background:white; border-radius:10px;
            display:inline-block; padding:8px 14px;
            ">
            {_sidebar_logo_html if _sidebar_logo_html else '<span style="font-size:2rem;">💧</span>'}
            </div>
            <div style="
            font-size:0.72rem;
            opacity:0.7;
            margin-top:6px;
            color:white;
            ">
            설비관리 통합 플랫폼
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    home_menu_item = ("홈", "🏠 설비관리 홈")

    st.button(

        home_menu_item[1],

        key=f"menu_{home_menu_item[0]}",

        use_container_width=True,

        on_click=go_to_page,

        args=(home_menu_item[0],),

        type=(
            "primary"
            if st.session_state.page == home_menu_item[0]
            else "secondary"
        )

    )

    st.markdown(
        "<div class='menu-caption'>📍 현장 업무</div>",
        unsafe_allow_html=True
    )

    main_menus = [

        ("통합검색", "🔎 통합검색"),

        ("QR", "📱 QR 설비 포털")

    ]

    for key, label in main_menus:

        st.button(
            label,
            key=f"menu_{key}",
            use_container_width=True,
            on_click=go_to_page,
            args=(key,),
            type=(
                "primary"
                if st.session_state.page == key
                else "secondary"
            )
        )

    st.markdown(
        "<div class='menu-caption'>🔍 진단 · 판단</div>",
        unsafe_allow_html=True
    )

    analysis_menus = [

        ("진단", "🔍 정밀 진단"),

        ("CBM", "🎯 CBM 정비판단"),

        ("오버홀", "🛠️ 오버홀 관리")

    ]

    for key, label in analysis_menus:

        st.button(
            label,
            key=f"menu_{key}",
            use_container_width=True,
            on_click=go_to_page,
            args=(key,),
            type=(
                "primary"
                if st.session_state.page == key
                else "secondary"
            )
        )

    st.markdown(
        "<div class='menu-caption'>📊 분석 · 리포트</div>",
        unsafe_allow_html=True
    )

    knowledge_menus = [

        ("AI", "📈 AI 이상징후"),

        ("전사트렌드", "🌐 전사 트렌드"),

        ("ROI", "💰 정비효과·ROI"),

        ("KPI", "📊 성과관리"),

        ("보고서", "📄 진단 보고서")

    ]

    for key, label in knowledge_menus:

        st.button(
            label,
            key=f"menu_{key}",
            use_container_width=True,
            on_click=go_to_page,
            args=(key,),
            type=(
                "primary"
                if st.session_state.page == key
                else "secondary"
            )
        )

    st.markdown(
        "<div class='menu-caption'>⚙️ 관리</div>",
        unsafe_allow_html=True
    )

    manage_menus = [

        ("설비", "🏭 설비 관리"),

        ("노하우", "💡 기술 노하우"),

        ("설계적산", "📐 설계적산"),

        ("백업", "💾 데이터 관리"),

        ("업데이트내역", "📋 업데이트 내역")

    ]

    for key, label in manage_menus:

        st.button(
            label,
            key=f"menu_{key}",
            use_container_width=True,
            on_click=go_to_page,
            args=(key,),
            type=(
                "primary"
                if st.session_state.page == key
                else "secondary"
            )
        )

    with st.expander(

        "🏗️ 부서별 관리 (개발예정)",

        expanded=False

    ):

        dept_menus = [

            ("토목", "🏗️ 토목"),

            ("전기", "⚡ 전기"),

            ("전자", "💻 전자"),

            ("행정", "📋 행정"),

            ("안전", "🦺 안전")

        ]

        for key, label in dept_menus:

            st.button(
                label,
                key=f"menu_{key}",
                use_container_width=True,
                on_click=go_to_page,
                args=(key,),
                type=(
                    "primary"
                    if st.session_state.page == key
                    else "secondary"
                )
            )

    st.markdown("---")

    st.caption(

        f"{st.session_state.user_name} · "
        +
        (
            "보기 전용 방문자"

            if is_read_only()

            else "정밀진단원 / 관리자"
        )

    )

    st.caption(
        "밀양정수장"
    )

    _manual_pdf_path = find_manual_pdf()

    if _manual_pdf_path:

        try:

            with open(

                _manual_pdf_path,

                "rb"

            ) as f:

                _manual_pdf_bytes = f.read()

            st.download_button(

                "📘 사용자 매뉴얼 다운로드",

                data=_manual_pdf_bytes,

                file_name="설비관리_플랫폼_사용자매뉴얼.pdf",

                mime="application/pdf",

                use_container_width=True,

                key="manual_pdf_download_btn"

            )

        except Exception:

            pass

    st.toggle(

        "🌙 다크모드",

        key="dark_mode",

        value=False

    )

    if st.session_state.get("dark_mode"):

        st.markdown(

            """
            <style>

            [data-testid="stAppViewContainer"],
            [data-testid="stMain"],
            [data-testid="stHeader"] {
                background: #0f172a !important;
            }

            [data-testid="stMain"] * {
                color: #e2e8f0;
            }

            .section-title { color: #f1f5f9 !important; }

            .section-caption { color: #94a3b8 !important; }

            .kpi-card,
            .equip-card,
            .feature-card,
            .platform-card,
            .story-card {
                background: #1e293b !important;
                box-shadow: 0 3px 12px rgba(0,0,0,0.5) !important;
            }

            .search-result-wrap {
                border-color: #334155 !important;
            }

            .search-result-wrap th {
                background: #1e293b !important;
                border-bottom-color: #334155 !important;
            }

            .search-result-wrap td {
                border-bottom-color: #334155 !important;
            }

            .search-result-wrap mark {
                background: #a16207 !important;
                color: #fff8e1 !important;
            }

            .cbm-bar-track { background: #334155 !important; }

            [data-testid="stMain"] [data-testid="stDataFrame"],
            [data-testid="stMain"] .stDataFrame {
                filter: invert(0.92) hue-rotate(180deg);
            }

            [data-testid="stMain"] input,
            [data-testid="stMain"] textarea,
            [data-testid="stMain"] select {
                background: #1e293b !important;
                color: #e2e8f0 !important;
            }

            [data-testid="stMain"] .stButton > button,
            [data-testid="stMain"] .stDownloadButton > button {
                background: #1e293b !important;
                color: #e2e8f0 !important;
                border: 1px solid #334155 !important;
            }

            </style>
            """,

            unsafe_allow_html=True

        )

    if st.session_state.entered_as_viewer:

        # 보기전용으로 들어온 사람은 토글 자체를 못 보게 해서
        # 스스로 수정권한을 켜는 것을 막는다.
        # 관리자 권한이 필요하면 로그아웃 후 PIN으로
        # 다시 들어와야 한다.

        st.caption(
            "🔒 보기 전용 계정입니다 (수정 불가)"
        )

        if st.button(
            "🔓 관리자로 다시 로그인",
            use_container_width=True
        ):

            st.session_state.authenticated = False

            st.session_state.read_only = False

            st.session_state.entered_as_viewer = False

            st.rerun()

    else:

        st.toggle(

            "🔒 보기 전용 모드",

            key="read_only",

            help="켜면 모든 저장·삭제 버튼이 비활성화됩니다. "
                 "외부인에게 시연할 때 실수로 데이터가 "
                 "바뀌는 것을 막아줍니다."

        )

    _risk_count = sum(

        1

        for p in ALL_PUMPS

        if pump_status(p, df_history)["상태"] == "정비검토"

    )

    if _risk_count > 0:

        st.error(
            f"🚨 정비검토 필요 설비 {_risk_count}대"
        )


# ============================================================
# 14. 공통 헤더 및 상단 이동 메뉴
#
# - 큰 배너(K-water tech ... 소개문구)는 "홈" 화면에서만 보여준다.
#   다른 화면에서는 각 화면 자체의 section-title/caption이
#   맨 위에 바로 보이도록 한다. (요청 6)
#
# - 사이드바(왼쪽 서랍 메뉴)는 모바일 브라우저에서 버튼을 눌러도
#   서랍이 자동으로 닫히지 않는 Streamlit 자체의 한계가 있어,
#   대신 본문 맨 위에 항상 떠 있는 "메뉴 이동" 드롭다운을 추가한다.
#   드롭다운은 선택하면 자동으로 닫히는 표준 UI라 이 문제가 없다.
#   (요청 3)
# ============================================================

ALL_MENUS = (
    [home_menu_item]
    +
    main_menus
    +
    analysis_menus
    +
    knowledge_menus
    +
    manage_menus
    +
    dept_menus
)

MENU_LABEL_BY_KEY = dict(
    ALL_MENUS
)

MENU_KEY_BY_LABEL = {

    label: key

    for key, label in ALL_MENUS

}

_current_page = st.session_state.page

_current_label = MENU_LABEL_BY_KEY.get(
    _current_page,
    "🏠 설비관리 홈"
)

# Streamlit 자체 query_params 동기화 메커니즘을 매 렌더링마다
# 그대로 이용한다. 이건 내가 만든 JS 타이밍과 무관하게
# Streamlit이 직접 관리하는 안정적인 경로라서, 주소창의
# page 값이 항상 지금 보고 있는 페이지와 일치하도록
# 보장해준다 (뒤로가기 시 이 값을 기준으로 페이지를 되돌림).

st.query_params["page"] = _current_page

# ------------------------------------------------------------
# 메뉴를 눌러서 페이지가 실제로 바뀐 경우에만, 스크롤을
# 맨 위로 되돌린다. (같은 페이지 안에서 체크박스·입력창 같은
# 걸 조작할 때는 스크롤 위치를 그대로 유지해야 하므로,
# "페이지 자체가 바뀌었을 때"만 동작하도록 구분했다.)
# ------------------------------------------------------------

if st.session_state.get("_last_rendered_page") != _current_page:

    st.session_state["_last_rendered_page"] = _current_page

    components.html(

        f"""
        <!-- page-nonce: {_current_page} -->
        <script>
        function scrollAppToTop() {{
            try {{
                window.parent.scrollTo({{top: 0, left: 0}});
            }} catch (e) {{
                console.log("DEBUG scrollAppToTop window.parent.scrollTo 실패:", e);
            }}
            try {{
                var stMain = window.parent.document.querySelector(
                    'section[data-testid="stMain"]'
                );
                if (stMain) {{
                    stMain.scrollTop = 0;
                    console.log("DEBUG stMain 찾음, scrollTop 설정 후:", stMain.scrollTop);
                }} else {{
                    console.log("DEBUG stMain 못찾음");
                }}
            }} catch (e) {{
                console.log("DEBUG stMain 접근 실패:", e);
            }}
        }}

        console.log("DEBUG scrollAppToTop 스크립트 로드됨 ({_current_page})");

        scrollAppToTop();

        // 예전에는 100ms마다 3초 동안(총 30번) 계속 스크롤을
        // 강제로 위로 되돌렸는데, 그 3초 사이에 사용자가
        // 손으로 아래로 내리면 다음 반복에서 또 위로
        // 끌어올려져서 스크롤을 아예 못 하는 문제가 있었다.
        // 이제는 초기 렌더링 타이밍만 짧게 커버하도록
        // 아주 잠깐(최대 0.3초) 동안 몇 번만 시도하고 멈춘다.

        setTimeout(scrollAppToTop, 60);

        setTimeout(scrollAppToTop, 150);

        setTimeout(scrollAppToTop, 300);

        // Streamlit은 st.query_params가 바뀌면 자체적으로
        // 이미 브라우저 히스토리에 기록을 남긴다(확인 완료).
        // 그래서 여기서 따로 pushState를 또 하지 않는다.
        //
        // 뒤로가기를 누르면 브라우저가 이전 URL로 이동하는데,
        // Streamlit은 그 URL 변화를 스스로 감지해 다시 그리지
        // 않으므로, 우리가 강제로 새로고침해서 그 시점의
        // query_params를 다시 읽게 만든다.
        //
        // (페이지 전환마다 리스너를 다시 등록한다. 중복 등록돼도
        //  reload를 여러 번 트리거하는 것뿐이라 무해하고,
        //  최초 1회 등록이 어떤 이유로 실패해도 다음 페이지
        //  전환 때 다시 시도되어 결국 붙게 된다.)

        window.parent.addEventListener("popstate", function (event) {{

            console.log("DEBUG popstate 감지, 새로고침");

            window.parent.location.reload();

        }});
        </script>
        """,

        height=0

    )

_nav_labels = [
    label
    for key, label in ALL_MENUS
]


def _on_top_nav_change():

    chosen_label = st.session_state[
        f"top_nav_{_current_page}"
    ]

    new_page = MENU_KEY_BY_LABEL[
        chosen_label
    ]

    st.session_state.page = new_page

    st.query_params["page"] = new_page

    if "equip" in st.query_params:

        del st.query_params["equip"]


st.selectbox(

    "메뉴 이동",

    _nav_labels,

    index=_nav_labels.index(
        _current_label
    ),

    key=f"top_nav_{_current_page}",

    on_change=_on_top_nav_change,

    label_visibility="collapsed"

)

# 예전에는 홈 화면 맨 위에 "K-water tech 설비관리 통합
# 플랫폼"이라는 일반 배너가 있었는데, 이제 홈 화면 자체에
# 더 화려한 히어로 배너(아치 로고 + 색상 강조 타이틀)가
# 항상 나오므로, 같은 내용을 두 번 겹쳐 보여주지 않도록
# 이 일반 배너는 제거했다.


# ============================================================
# 15. 홈
# ============================================================

if st.session_state.page == "홈":

    _hour = datetime.now().hour

    if _hour < 12:

        _greeting = "좋은 아침입니다"

    elif _hour < 18:

        _greeting = "안녕하세요"

    else:

        _greeting = "수고 많으셨습니다"

    # 실제 회사 홈페이지들이 방문할 때마다 똑같은 히어로
    # 배너를 보여주듯, 여기서도 홈에 올 때마다 항상 큰
    # 히어로 배너를 보여준다. (예전에는 세션 첫 방문에만
    # 크게 보여주고 그 다음부터 작게 줄였는데, 그러면
    # "히어로 배너가 안 보인다"고 느끼기 쉬웠다)

    _hero_logo_html = get_logo_base64_html(
        max_height_px=44
    )

    _hero_mascot_html = get_mascot_base64_html(
        max_height_px=96
    )

    st.markdown(

        f"""
        <div class="hero-banner">
        <div class="hero-inner">

            <div class="hero-logo-wrap" style="
            background:white; border-radius:12px;
            padding:8px 14px; display:inline-block;">
            {_hero_logo_html if _hero_logo_html else '<span style="font-size:2rem;">💧</span>'}
            </div>

            <div class="hero-text-wrap">
                <div class="hero-title">
                설비를 <span class="hl-cyan">살아있게</span>
                관리하는 <span class="hl-amber">기술</span>
                </div>
                <div class="hero-sub">
                K-water tech 설비관리 통합 플랫폼 ·
                Smart Pump Management
                </div>
                <div class="hero-greeting">
                {_greeting}, {st.session_state.get('user_name', '')}님 👋
                오늘도 설비 현황을 한눈에 확인해보세요.
                </div>
            </div>

            {f'<div class="hero-mascot-wrap">{_hero_mascot_html}</div>' if _hero_mascot_html else ''}

        </div>
        </div>
        """,

        unsafe_allow_html=True

    )

    st.markdown(
        """
        <div class="section-title">
        설비관리 현황
        </div>

        <div class="section-caption">
        현장의 설비 상태를 한 화면에서 확인하고
        정비 우선순위를 판단합니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    # ---------------- 백업 알림 배너 ----------------
    #
    # "마지막으로 언제 백업했는지"가 어디서도 안 보이면
    # 깜빡하고 안 하게 되므로, 홈 화면에서 눈에 띄게 알려준다.

    _days_since_backup = get_days_since_last_backup()

    if _days_since_backup is None:

        st.warning(

            "⚠️ 아직 한 번도 백업하지 않았습니다. "
            "'데이터 관리 > 💾 백업·리포트'에서 백업을 "
            "권장합니다."

        )

    elif _days_since_backup >= 7:

        st.warning(

            f"⚠️ 마지막 백업: {_days_since_backup}일 전. "
            "정기적인 백업을 권장합니다."

        )

    # ---------------- 사용 설명서 ----------------
    #
    # 매뉴얼 파일을 따로 안 열어봐도, 앱 안에서 바로 넘겨볼
    # 수 있게 한다(신규 직원 온보딩에 특히 유용).

    _manual_pdf_path_home = find_manual_pdf()

    if _manual_pdf_path_home:

        with st.expander(

            "📖 사용 설명서 (앱 안에서 바로 보기)"

        ):

            _manual_html = render_manual_pdf_viewer_html(

                _manual_pdf_path_home

            )

            if _manual_html:

                st.markdown(

                    _manual_html,

                    unsafe_allow_html=True

                )

            else:

                st.info(

                    "매뉴얼을 불러오지 못했습니다. "
                    "사이드바의 다운로드 버튼을 이용해주세요."

                )

    # ---------------- 빠른 이동 ----------------
    #
    # 사이드바를 스크롤하며 찾는 대신, 메뉴 이름을 타이핑해서
    # 바로 이동할 수 있게 한다.

    _quick_jump_options = ["🔎 어디로 갈까요? (메뉴 이름 입력)"] + [

        label for _, label in ALL_MENUS

        if label != "🏠 설비관리 홈"

    ]

    _quick_jump_selected = st.selectbox(

        "빠른 이동",

        _quick_jump_options,

        key="home_quick_jump",

        label_visibility="collapsed"

    )

    if _quick_jump_selected != _quick_jump_options[0]:

        _target_key = MENU_KEY_BY_LABEL.get(
            _quick_jump_selected
        )

        if _target_key:

            st.session_state.page = _target_key

            st.query_params["page"] = _target_key

            st.rerun()

    # ---------------- 데이터 최신성 표시 ----------------
    #
    # "이 화면이 언제 기준 데이터인지"를 보여줘서
    # 실시간으로 갱신되는 데이터라는 신뢰를 준다.

    _mtimes = []

    for _path in (DB_FILE_PATH, OVERHAUL_DB_PATH):

        if os.path.exists(_path):

            _mtimes.append(
                os.path.getmtime(_path)
            )

    if _mtimes:

        _last_update = datetime.fromtimestamp(

            max(_mtimes)

        ).strftime("%Y-%m-%d %H:%M")

        st.caption(

            f"🕒 데이터 최종 갱신 : {_last_update}"

        )

    # ---------------- 바로가기 피처 카드 ----------------
    #
    # K-water Tech 사이트의 "부서안내" 카드처럼, 그라데이션
    # 배경에 굵은 흰 글자로 된 카드로 자주 쓰는 기능 3개를
    # 바로가기로 보여준다. (실제 현장 사진 대신 색으로 구분)

    qc1, qc2, qc3 = st.columns(3)

    _qc_mascot_diag = get_mascot_base64_html(
        max_height_px=56,
        name="mascot_control"
    )

    _qc_mascot_report = get_mascot_base64_html(
        max_height_px=56,
        name="mascot"
    )

    _qc_mascot_qr = get_mascot_base64_html(
        max_height_px=56,
        name="mascot_patrol"
    )

    with qc1:

        st.markdown(

            f"""
            <div class="feature-card" style="
            background: linear-gradient(135deg, #063b63, #0891b2);
            position: relative;">
                {f'<div style="position:absolute; top:8px; right:10px; opacity:0.9;">{_qc_mascot_diag}</div>' if _qc_mascot_diag else ''}
                <div class="feature-card-icon">🔍</div>
                <div class="feature-card-title">정밀진단</div>
                <div class="feature-card-desc">
                17개 항목으로 설비 상태를 점검하고
                등급을 산출합니다.
                </div>
                <div class="feature-card-link">시작하기 →</div>
            </div>
            """,

            unsafe_allow_html=True

        )

        st.button(

            "🔍 정밀진단 시작하기",

            key="home_quick_diag",

            use_container_width=True,

            on_click=go_to_page,

            args=("진단",)

        )

    with qc2:

        st.markdown(

            f"""
            <div class="feature-card" style="
            background: linear-gradient(135deg, #0f5132, #2f9e44);
            position: relative;">
                {f'<div style="position:absolute; top:8px; right:10px; opacity:0.9;">{_qc_mascot_report}</div>' if _qc_mascot_report else ''}
                <div class="feature-card-icon">📝</div>
                <div class="feature-card-title">월간 보고서</div>
                <div class="feature-card-desc">
                설비별 점검·오버홀 이력을 담은
                Word 보고서를 만듭니다.
                </div>
                <div class="feature-card-link">만들기 →</div>
            </div>
            """,

            unsafe_allow_html=True

        )

        st.button(

            "📝 월간 보고서 만들기",

            key="home_quick_report",

            use_container_width=True,

            on_click=go_to_page,

            args=("보고서",)

        )

    with qc3:

        st.markdown(

            f"""
            <div class="feature-card" style="
            background: linear-gradient(135deg, #9a3412, #f08c00);
            position: relative;">
                {f'<div style="position:absolute; top:8px; right:10px; opacity:0.9;">{_qc_mascot_qr}</div>' if _qc_mascot_qr else ''}
                <div class="feature-card-icon">📱</div>
                <div class="feature-card-title">QR 설비 포털</div>
                <div class="feature-card-desc">
                QR을 스캔하면 바로 뜨는
                설비별 실시간 화면입니다.
                </div>
                <div class="feature-card-link">열기 →</div>
            </div>
            """,

            unsafe_allow_html=True

        )

        st.button(

            "📱 QR 포털 열기",

            key="home_quick_qr",

            use_container_width=True,

            on_click=go_to_page,

            args=("QR",)

        )

    # ---------------- 사업장 필터 ----------------
    #
    # 지금은 밀양정수장 설비만 있지만, 다른 사업장 설비가
    # 추가되면(설비 추가 화면에서 사업장을 다르게 입력) 홈에서
    # 사업장별로 나눠 볼 수 있게 한다.

    _site_list = sorted(

        set(
            p["site"] for p in ALL_PUMPS
        )

    )

    if len(_site_list) > 1:

        _site_filter = st.selectbox(

            "🏢 사업장 필터",

            ["전체"] + _site_list,

            key="home_site_filter"

        )

    else:

        _site_filter = "전체"

    if _site_filter == "전체":

        _filtered_pumps = ALL_PUMPS

    else:

        _filtered_pumps = [

            p for p in ALL_PUMPS

            if p["site"] == _site_filter

        ]

    total = len(
        _filtered_pumps
    )

    normal = 0
    watch = 0
    repair = 0

    _all_results = []

    for pump in _filtered_pumps:

        result = pump_status(
            pump,
            df_history
        )

        _all_results.append(
            (pump, result)
        )

        if result["상태"] == "정상":

            normal += 1

        elif result["상태"] == "관찰":

            watch += 1

        else:

            repair += 1

    if repair > 0:

        st.error(

            f"🚨 정비검토 필요 설비가 {repair}대 있습니다."

        )

        if st.button(

            "→ CBM 정비판단에서 확인하기",

            key="home_repair_alert_goto_cbm",

            use_container_width=True

        ):

            st.session_state["_cbm_status_filter"] = "정비검토"

            st.session_state.page = "CBM"

            st.query_params["page"] = "CBM"

            st.rerun()

    with st.expander(

        "📤 오늘 현황 공유하기"

    ):

        _repair_names = [

            p["equip"]

            for p, r in _all_results

            if r["상태"] == "정비검토"

        ]

        _share_text = (

            f"[K-water tech 설비관리 현황] "
            f"{datetime.now().strftime('%Y-%m-%d')}\n"
            f"정상 {normal}대 · 관찰 {watch}대 · "
            f"정비검토 {repair}대\n"

        )

        if _repair_names:

            _share_text += (

                "정비검토 필요 설비: "
                +
                ", ".join(_repair_names)
                +
                "\n"

            )

        st.text_area(

            "아래 내용을 복사해서 문자·카톡 등으로 보내세요",

            value=_share_text,

            height=120,

            key="home_share_text_area"

        )

        st.link_button(

            "💬 문자로 바로 보내기",

            "sms:?body="
            +
            urllib.parse.quote(_share_text),

            use_container_width=True

        )

    _recent_viewed = [

        e for e in get_recent_view_list(limit=10)

        if e in {p["equip"] for p in ALL_PUMPS}

    ]

    if _recent_viewed:

        st.markdown(
            "##### 🕘 최근 본 설비"
        )

        recent_cols = st.columns(
            min(len(_recent_viewed), 3)
        )

        for rcol, equip_name in zip(

            recent_cols,

            _recent_viewed[:3]

        ):

            with rcol:

                if st.button(

                    equip_name,

                    key=f"recent_goto_{equip_name}",

                    use_container_width=True

                ):

                    st.session_state.page = "QR"

                    st.query_params["page"] = "QR"

                    st.query_params["equip"] = equip_name

                    st.rerun()

    kpi_defs = [

        ("관리 설비", f"{total}대", "등록 설비", "설비", None),
        ("정상", f"{normal}대", "정상 운전", "CBM", "정상"),
        ("관찰", f"{watch}대", "추이관리", "CBM", "관찰"),
        ("정비 검토", f"{repair}대", "CBM 우선관리", "CBM", "정비검토"),
        ("QR 관리", f"{total}개", "설비별 1개 기준", "QR", None)

    ]

    kpi_cols = st.columns(5)

    for kcol, (label, value, sub, target_page, status_filter) in zip(

        kpi_cols,
        kpi_defs

    ):

        with kcol:

            st.markdown(

                f"""
                <div class="kpi-card">
                    <div class="kpi-label">{label}</div>
                    <div class="kpi-value">{value}</div>
                    <div class="kpi-sub">{sub}</div>
                </div>
                """,

                unsafe_allow_html=True

            )

            if st.button(

                "바로가기",

                key=f"kpi_goto_{label}",

                use_container_width=True

            ):

                st.session_state["_cbm_status_filter"] = status_filter

                st.session_state.page = target_page

                st.query_params["page"] = target_page

                st.rerun()

    # ---------------- 상태 비율 도넛 + 이달의 하이라이트 ----------------

    donut_col, highlight_col = st.columns(
        [1, 1.3]
    )

    with donut_col:

        st.markdown(
            """
            <div class="platform-card">
            <div class="card-title">📊 설비 상태 비율</div>
            """,
            unsafe_allow_html=True
        )

        st.pyplot(

            build_status_donut_fig(
                normal,
                watch,
                repair
            ),

            use_container_width=False

        )

        st.markdown(
            "</div>",
            unsafe_allow_html=True
        )

    with highlight_col:

        st.markdown(
            """
            <div class="platform-card">
            <div class="card-title">📍 현재 하이라이트</div>
            """,
            unsafe_allow_html=True
        )

        if not _all_results:

            st.info(
                "등록된 설비가 없습니다."
            )

        else:

            best_pump, best_result = max(

                _all_results,

                key=lambda pr: pr[1]["점수"]

            )

            worst_pump, worst_result = min(

                _all_results,

                key=lambda pr: pr[1]["점수"]

            )

            st.markdown(

                f"""
                <span class="highlight-pill" style="
                background:#e7f8f1; color:#087f5b;">
                🏆 최우수 · {best_pump['equip']} ({best_result['점수']}점)
                </span>

                <span class="highlight-pill" style="
                background:#fff0f0; color:#c62828;">
                ⚠️ 최우선 관리 · {worst_pump['equip']} ({worst_result['점수']}점)
                </span>
                """,

                unsafe_allow_html=True

            )

            st.caption(

                "※ 현재 시점 CBM Score 기준입니다 "
                "(이번 달 진단 여부와 무관)"

            )

            st.write("")

            _next_due = [

                (p, r["다음오버홀까지남은시간"])

                for p, r in _all_results

                if r.get("다음오버홀까지남은시간") is not None

            ]

            if _next_due:

                _soonest_pump, _soonest_hours = min(

                    _next_due,

                    key=lambda x: x[1]

                )

                st.caption(

                    f"🔧 오버홀이 가장 임박한 설비 : "
                    f"{_soonest_pump['equip']} "
                    f"(약 {_soonest_hours:,}시간 남음)"

                )

        st.markdown(
            "</div>",
            unsafe_allow_html=True
        )

    col1, col2 = st.columns(
        [1.45, 1]
    )

    with col1:

        st.markdown(
            """
            <div class="platform-card">

            <div class="card-title">
            🏭 설비 건전성 현황
            </div>

            """
            ,
            unsafe_allow_html=True
        )

        home_search = st.text_input(

            "🔎 설비명 검색",

            key="home_equip_search",

            placeholder="예: 가압펌프 #3"

        )

        # ---------------- 즐겨찾기 ----------------
        #
        # 관리하는 설비가 늘어나면 자주 보는 몇 대만 따로
        # 모아보고 싶어질 수 있다. 예전엔 세션 동안만 유지되는
        # 임시 목록이라 재접속하면 사라졌는데, 이제 파일로
        # 저장해서 다음에 다시 들어와도 그대로 남는다.

        _favorite_equips = set(
            get_favorite_equipment_list()
        )

        fav_c1, fav_c2 = st.columns(
            [2, 1]
        )

        with fav_c1:

            selected_favorites = st.multiselect(

                "⭐ 즐겨찾기 설비 선택",

                [p["equip"] for p in _filtered_pumps],

                default=[

                    e for e in _favorite_equips

                    if e in [p["equip"] for p in _filtered_pumps]

                ],

                key="home_favorite_select"

            )

            _new_favorite_set = set(
                selected_favorites
            )

            if _new_favorite_set != _favorite_equips:

                for _added in _new_favorite_set - _favorite_equips:

                    toggle_favorite_equipment(_added)

                for _removed in _favorite_equips - _new_favorite_set:

                    toggle_favorite_equipment(_removed)

                _favorite_equips = _new_favorite_set

        with fav_c2:

            show_favorites_only = st.toggle(

                "즐겨찾기만 보기",

                key="home_favorites_only_toggle",

                disabled=len(_favorite_equips) == 0

            )

        rows = []

        rendered_card_count = 0

        status_color_map = {

            "정상": ("#087f5b", "#e7f8f1"),

            "관찰": ("#a16207", "#fff7df"),

            "정비검토": ("#c62828", "#fff0f0")

        }

        # 8. 카드 정렬 순서: 급한 것부터 보이도록
        # 정비검토 -> 관찰 -> 정상 순, 같은 상태 안에서는
        # 점수가 낮은(더 안 좋은) 순서로. 오른쪽 CBM 순위표와
        # 같은 "급한 순" 기준으로 통일했다.

        status_priority = {

            "정비검토": 0,

            "관찰": 1,

            "정상": 2

        }

        sorted_results = sorted(

            _all_results,

            key=lambda pr: (

                status_priority.get(pr[1]["상태"], 9),

                pr[1]["점수"]

            )

        )

        # 검색어로 걸러진 설비명 목록 (CBM 순위표에도 동일하게 적용)

        home_search_matched_names = set()

        # 카드가 세로로 한 줄씩 쭉 나열되면 폭이 너무 넓어서
        # 오히려 보기 불편했다. 2열로 나눠서 옆으로 배치한다.

        home_cols = st.columns(2)

        for _card_idx, (pump, result) in enumerate(sorted_results):

            rows.append(

                {
                    "설비": pump["equip"],
                    "운전시간": f'{pump["op_hours"]:,} h',
                    "효율": f'{result["효율"]:.1f}%',
                    "진동": f'{result["진동"]:.1f} mm/s',
                    "CBM": result["점수"],
                    "등급": result["등급"],
                    "판정": result["상태"]
                }

            )

            if (

                home_search

                and
                home_search.strip().lower() not in pump["equip"].lower()

            ):

                continue

            if (

                show_favorites_only

                and
                pump["equip"] not in _favorite_equips

            ):

                continue

            home_search_matched_names.add(
                pump["equip"]
            )

            _col_idx = rendered_card_count % 2

            with home_cols[_col_idx]:

                text_color, bg_color = status_color_map.get(

                    result["상태"],
                    ("#64748b", "#f1f5f9")

                )

                # 6. 게이지바 색상을 등급 5단계 기준과 통일

                gauge_color = grade_bar_color(
                    result["점수"]
                )

                _, vib_values, _ = get_vibration_trend_data(

                    pump,

                    result,

                    df_history

                )

                # 7. 스파크라인에 마우스를 올리면 실제 수치가
                # 툴팁으로 보이도록 <title>을 넣는다.

                spark_tooltip = (

                    "진동 추세(mm/s): "
                    +
                    " → ".join(
                        f"{v:.1f}" for v in vib_values
                    )

                )

                spark_svg = build_svg_sparkline(

                    vib_values,

                    color=gauge_color,

                    tooltip=spark_tooltip

                )

                st.markdown(

                    f"""
                    <div class="equip-card" style="
                    border-left:4px solid {text_color};">

                        <div class="equip-card-title">
                        {'⭐ ' if pump['equip'] in _favorite_equips else ''}{pump['equip']}
                        </div>

                        <span style="
                        background:{bg_color}; color:{text_color};
                        border-radius:999px; padding:2px 9px;
                        font-size:0.68rem; font-weight:800;">
                        {status_icon(result['상태'])} {result['상태']} · {result['등급']}등급
                        </span>

                        <div class="cbm-bar-track">
                            <div class="cbm-bar-fill" style="
                            width:{result['점수']}%;
                            background:{gauge_color};"></div>
                        </div>

                        <div class="equip-card-meta">
                            <span>CBM {result['점수']}점</span>
                            <span>효율 {result['효율']:.1f}%</span>
                        </div>

                        <div style="margin-top:6px;">
                        {spark_svg}
                        <span style="
                        font-size:0.68rem; color:#94a3b8;
                        margin-left:4px;">진동 추세</span>
                        </div>

                    </div>
                    """,

                    unsafe_allow_html=True

                )

                # 카드 자체는 그림(HTML)이라 눌러도 반응이 없었다.
                # 카드 바로 밑에 진짜 버튼을 붙여서, 누르면 그
                # 설비를 선택한 채로 QR 포털로 이동한다.

                if st.button(

                    f"🔍 상세보기 →",

                    key=f"card_goto_{_card_idx}_{pump['equip']}",

                    use_container_width=True

                ):

                    st.session_state.page = "QR"

                    st.query_params["page"] = "QR"

                    st.query_params["equip"] = pump["equip"]

                    st.rerun()

            rendered_card_count += 1

        if rendered_card_count == 0:

            st.info(
                "검색 결과가 없습니다."
            )

    with col2:

        st.markdown(
            """
            <div class="platform-card">

            <div class="card-title">
            🎯 CBM 정비 우선순위
            </div>

            """,
            unsafe_allow_html=True
        )

        ranking = []

        for pump, result in _all_results:

            if pump["equip"] not in home_search_matched_names:

                continue

            ranking.append(

                [
                    pump["equip"],
                    result["점수"],
                    result["등급"],
                    pump["op_hours"]
                ]

            )

        ranking = sorted(
            ranking,
            key=lambda x: x[1]
        )

        if ranking:

            ranking_df = pd.DataFrame(

                ranking[:10],

                columns=[
                    "설비",
                    "CBM Score",
                    "등급",
                    "운전시간"
                ]

            )

            st.dataframe(
                ranking_df,
                use_container_width=True,
                hide_index=True,
                height=350
            )

            if st.button(

                "→ CBM 정비판단에서 전체 보기",

                key="home_cbm_ranking_more_btn",

                use_container_width=True

            ):

                st.session_state.page = "CBM"

                st.query_params["page"] = "CBM"

                st.rerun()

        else:

            st.info(
                "검색 결과가 없습니다."
            )

        st.markdown(
            "</div>",
            unsafe_allow_html=True
        )

    # ---------------- 최근 활동 피드 + 임박한 오버홀 ----------------

    col3, col4 = st.columns(
        [1.2, 1]
    )

    with col3:

        st.markdown(
            """
            <div class="platform-card">

            <div class="card-title">
            🕒 최근 활동
            </div>

            """,
            unsafe_allow_html=True
        )

        activities = []

        # 진단·오버홀 이력이 계속 쌓여도 홈 화면이 느려지지
        # 않도록, 정렬 전에 미리 최근 30건만 잘라서 순회한다.
        # (점검일/작업일자는 날짜순으로 계속 append되므로
        #  tail이 곧 최근 기록이다)

        if (

            not df_history.empty

            and
            "설비명" in df_history.columns

        ):

            for _, r in df_history.tail(30).iterrows():

                activities.append(

                    {
                        "일자": str(r.get("점검일", "")),
                        "내용": (

                            f"{r.get('설비명', '')} 정밀진단 저장 · "
                            f"{_clean_cell_value(r.get('최종등급'), '-')}등급"

                        )
                    }

                )

        df_overhaul_home = read_excel(

            OVERHAUL_DB_PATH,

            "오버홀이력"

        )

        if (

            not df_overhaul_home.empty

            and
            "설비명" in df_overhaul_home.columns

        ):

            for _, r in df_overhaul_home.tail(30).iterrows():

                activities.append(

                    {
                        "일자": str(r.get("작업일자", "")),
                        "내용": (

                            f"{r.get('설비명', '')} "
                            f"{_clean_cell_value(r.get('공정단계'), '작업')} 기록됨"

                        )
                    }

                )

        activities = sorted(

            activities,

            key=lambda x: x["일자"],

            reverse=True

        )[:8]

        if activities:

            for act in activities:

                st.markdown(

                    f"<div style='font-size:0.85rem; padding:4px 0; "
                    f"border-bottom:1px solid #eef2f5;'>"
                    f"<b>{act['일자']}</b> · {act['내용']}</div>",

                    unsafe_allow_html=True

                )

        else:

            st.info(
                "아직 기록된 활동이 없습니다."
            )

        st.markdown(
            "</div>",
            unsafe_allow_html=True
        )

    with col4:

        st.markdown(
            """
            <div class="platform-card">

            <div class="card-title">
            🔧 오버홀 임박 설비
            </div>

            """,
            unsafe_allow_html=True
        )

        overhaul_soon = []

        for pump, result in _all_results:

            remaining = result.get(
                "다음오버홀까지남은시간"
            )

            if remaining is not None:

                overhaul_soon.append(

                    {
                        "설비": pump["equip"],
                        "남은시간(h)": remaining
                    }

                )

        overhaul_soon = sorted(

            overhaul_soon,

            key=lambda x: x["남은시간(h)"]

        )[:5]

        if overhaul_soon:

            st.dataframe(

                pd.DataFrame(
                    overhaul_soon
                ),

                use_container_width=True,

                hide_index=True

            )

        else:

            st.info(
                "표시할 설비가 없습니다."
            )

        st.markdown(
            "</div>",
            unsafe_allow_html=True
        )

    st.markdown(
        """
        <div class="platform-card">

        <div class="card-title">
        🔄 플랫폼 운영 흐름
        </div>

        <div style="
        display:flex;
        flex-wrap:wrap;
        gap:8px;
        align-items:center;
        font-size:0.82rem;
        color:#164e63;
        font-weight:700;
        ">

        <span>📱 QR 접속</span>
        →
        <span>🏭 설비정보</span>
        →
        <span>🔍 상태진단</span>
        →
        <span>🎯 CBM Score</span>
        →
        <span>📌 정비 Ranking</span>
        →
        <span>🛠️ 오버홀</span>
        →
        <span>📈 효과검증</span>
        →
        <span>📚 이력축적</span>

        </div>

        </div>
        """,
        unsafe_allow_html=True
    )

    render_excel_export_section(

        "home",

        "설비관리_현황.xlsx",

        lambda: pd.DataFrame(rows)

    )


# ============================================================
# 16. 설비 관리
# ============================================================

elif st.session_state.page == "설비":

    st.markdown(
        """
        <div class="section-title">
        🏭 설비 관리
        </div>

        <div class="section-caption">
        관리 중인 펌프의 기본 제원과 운전상태를 한눈에 확인합니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    if not ALL_PUMPS:

        st.info(

            "등록된 설비가 없습니다. '데이터 관리' 메뉴에서 "
            "설비를 추가하거나 엑셀을 업로드해주세요."

        )

        st.stop()

    selected = st.selectbox(

        "설비 선택",

        [
            p["equip"]
            for p in ALL_PUMPS
        ]

    )

    pump = next(
        p for p in ALL_PUMPS
        if p["equip"] == selected
    )

    log_recent_view(
        pump["equip"]
    )

    result = pump_status(
        pump,
        df_history
    )

    status_class = {

        "정상": "status-normal",

        "관찰": "status-watch",

        "정비검토": "status-danger"

    }.get(
        result["상태"],
        "status-watch"
    )

    st.markdown(
        f"""
        <div class="platform-card">

        <div class="card-title">
        {pump['equip']} 종합 현황
        </div>

        <span class="{status_class}">
        {status_icon(result['상태'])} {result['상태']}
        </span>

        </div>
        """,
        unsafe_allow_html=True
    )

    a, b, c, d = st.columns(4)

    a.metric(
        "CBM Score",
        f'{result["점수"]}점'
    )

    b.metric(
        "현재 효율",
        f'{result["효율"]:.1f}%'
    )

    c.metric(
        "진동",
        f'{result["진동"]:.1f} mm/s'
    )

    d.metric(
        "온도",
        f'{result["온도"]:.1f}°C'
    )

    if result.get("전류"):

        st.caption(
            f"⚡ 측정 전류 : {result['전류']:.1f} A"
        )

    if result.get("다음오버홀까지남은시간") is not None:

        remaining = result["다음오버홀까지남은시간"]

        if remaining <= 500:

            st.error(
                f"🔧 다음 오버홀 예정까지 {remaining:,}시간 남음 "
                f"(기준 {OVERHAUL_INTERVAL_HOURS:,}h 주기) — 임박"
            )

        elif remaining <= 2000:

            st.warning(
                f"🔧 다음 오버홀 예정까지 {remaining:,}시간 남음 "
                f"(기준 {OVERHAUL_INTERVAL_HOURS:,}h 주기)"
            )

        else:

            st.caption(
                f"🔧 다음 오버홀 예정까지 {remaining:,}시간 남음 "
                f"(기준 {OVERHAUL_INTERVAL_HOURS:,}h 주기)"
            )

    st.markdown(
        "##### 📊 설비 상태 종합 그래프"
    )

    gc1, gc2 = st.columns(2)

    with gc1:

        st.pyplot(

            build_score_gauge_fig(

                pump,

                result

            )

        )

    with gc2:

        st.pyplot(

            build_vibration_trend_fig(

                pump,

                result,

                df_history=df_history,

                figsize=(6.2, 2.3)

            )

        )

    gc3, gc4 = st.columns(2)

    with gc3:

        st.pyplot(

            build_efficiency_trend_fig(

                pump,

                result,

                df_history=df_history,

                figsize=(6.2, 3.0)

            )

        )

    with gc4:

        st.pyplot(

            build_temperature_trend_fig(

                pump,

                result,

                df_history=df_history,

                figsize=(6.2, 3.0)

            )

        )

    gc5, gc6 = st.columns(2)

    with gc5:

        st.pyplot(

            build_op_hours_trend_fig(

                pump,

                figsize=(6.2, 3.0)

            )

        )

    with gc6:

        st.pyplot(

            build_fleet_compare_fig(

                pump,

                result,

                ALL_PUMPS,

                df_history,

                figsize=(6.2, 3.0)

            )

        )

    st.write("")

    tab1, tab2, tab3 = st.tabs(

        [
            "기본정보",
            "운전상태",
            "정비이력"
        ]

    )

    with tab1:

        st.dataframe(

            pd.DataFrame(

                [

                    ["사업장", pump["site"]],

                    ["설비명", pump["equip"]],

                    ["제조사", pump["maker"]],

                    ["모델명", pump["model"]],

                    ["정격출력(HP)", f'{pump["hp"]}'],

                    ["정격양정(m)", f'{pump["head"]}'],

                    ["정격유량(m³/h)", f'{pump["flow"]}'],

                    ["회전수(RPM)", f'{pump["rpm"]}'],

                    ["준공일", pump["build_date"]]

                ],

                columns=[
                    "항목",
                    "값"
                ]

            ),

            use_container_width=True,

            hide_index=True

        )

    with tab2:

        st.dataframe(

            pd.DataFrame(
                [
                    [
                        "누적 운전시간",
                        f'{pump["op_hours"]:,} h'
                    ],
                    [
                        "현재 효율",
                        f'{result["효율"]:.1f}%'
                    ],
                    [
                        "현재 진동",
                        f'{result["진동"]:.1f} mm/s'
                    ],
                    [
                        "현재 온도",
                        f'{result["온도"]:.1f}°C'
                    ],
                    [
                        "CBM 상태",
                        result["상태"]
                    ]
                ],
                columns=[
                    "항목",
                    "현재값"
                ]
            ),

            use_container_width=True,

            hide_index=True

        )

    with tab3:

        df_overhaul = read_excel(
            OVERHAUL_DB_PATH,
            "오버홀이력"
        )

        pump_overhaul = df_overhaul[
            df_overhaul["설비명"] == pump["equip"]
        ] if not df_overhaul.empty else pd.DataFrame()

        if not pump_overhaul.empty:

            st.dataframe(

                pump_overhaul[
                    [
                        "작업일자",
                        "공정단계",
                        "작업내용"
                    ]
                ].tail(5),

                use_container_width=True,

                hide_index=True

            )

        else:

            st.info(
                "저장된 오버홀·작업 이력이 없습니다. "
                "QR 포털의 '이력' 탭에서 현장 작업기록을 추가할 수 있습니다."
            )

    render_excel_export_section(

        f"equip_{pump['equip']}",

        f"{pump['equip']}_설비정보.xlsx",

        lambda: build_equipment_export_sheets(

            pump,

            result,

            df_history

        )

    )


# ============================================================
# 17. QR 설비 포털
# ============================================================

elif st.session_state.page == "QR":

    if st.button(

        "← 홈으로",

        key="qr_back_to_home_btn"

    ):

        st.session_state.page = "홈"

        st.query_params["page"] = "홈"

        if "equip" in st.query_params:

            del st.query_params["equip"]

        st.rerun()

    _qr_mascot_html = get_mascot_base64_html(
        max_height_px=100,
        name="mascot_patrol"
    )

    st.markdown(

        f"""
        <div class="page-header-mascot">

            <div class="page-header-text">
                <div class="section-title">
                📱 QR 설비 디지털 포털
                </div>

                <div class="section-caption">
                펌프 옆 QR코드를 스마트폰으로 스캔하면
                해당 설비의 제원·도면·이력·AI진단을
                한 화면에서 바로 확인합니다.
                </div>
            </div>

            {f'<div class="page-header-mascot-img">{_qr_mascot_html}</div>' if _qr_mascot_html else ''}

        </div>
        """,

        unsafe_allow_html=True

    )

    if not ALL_PUMPS:

        st.info(

            "등록된 설비가 없습니다. '데이터 관리' 메뉴에서 "
            "설비를 추가하거나 엑셀을 업로드해주세요."

        )

        st.stop()

    _pump_names = [
        p["equip"]
        for p in ALL_PUMPS
    ]

    _query_equip = st.query_params.get(
        "equip"
    )

    _default_index = (

        _pump_names.index(_query_equip)

        if _query_equip in _pump_names

        else 0

    )

    selected = st.selectbox(

        "설비 선택",

        _pump_names,

        index=_default_index

    )

    pump = next(
        p for p in ALL_PUMPS
        if p["equip"] == selected
    )

    log_recent_view(
        pump["equip"]
    )

    result = pump_status(
        pump,
        df_history
    )

    equip_no = ALL_PUMPS.index(
        pump
    ) + 1

    qr_id = f"PUMP-MLY-{equip_no:03d}"

    # 실제로 스캔 가능한 QR 이미지 생성.
    # 배포 주소는 코드에 박아두지 않고 설정 파일에서 읽어온다.
    # (데이터관리 페이지에서 관리자가 바꿀 수 있음)

    APP_BASE_URL = get_app_base_url()

    if APP_BASE_URL == DEFAULT_APP_BASE_URL:

        st.caption(
            "⚠️ 배포 주소가 기본값(예시)으로 설정되어 있습니다. "
            "실제 배포 주소와 다르면 QR을 스캔해도 연결되지 않습니다. "
            "데이터관리 페이지에서 실제 주소로 바꿔주세요."
        )

    qr_target_url = (

        f"{APP_BASE_URL}/?page=QR&equip="

        +
        urllib.parse.quote(
            selected
        )

    )

    qr_img = qrcode.make(
        qr_target_url
    )

    qr_buffer = io.BytesIO()

    qr_img.save(
        qr_buffer,
        format="PNG"
    )

    c1, c2 = st.columns(
        [1, 2]
    )

    with c1:

        st.image(

            qr_buffer.getvalue(),

            use_container_width=True,

            caption=qr_id

        )

        st.caption(
            "실제 스캔 가능한 QR입니다. "
            "스캔하면 이 설비 페이지로 바로 연결됩니다."
        )

    with c2:

        st.markdown(
            f"""
            <div class="platform-card">

            <div class="card-title">
            {pump["equip"]} 디지털 이력카드
            </div>

            <b>사업장</b> : {pump["site"]}<br>
            <b>제조사</b> : {pump["maker"]}<br>
            <b>모델</b> : {pump["model"]}<br>
            <b>출력</b> : {pump["hp"]} HP<br>
            <b>운전시간</b> : {pump["op_hours"]:,} h<br>
            <b>현재 효율</b> : {result["효율"]:.1f}%<br>
            <b>현재 진동</b> : {result["진동"]:.1f} mm/s<br>
            <b>현재 온도</b> : {result["온도"]:.1f}°C<br>
            <b>CBM Score</b> : {result["점수"]}점<br>
            <b>상태</b> : {result["상태"]}


            </div>
            """,
            unsafe_allow_html=True
        )

    st.write("")

    t1, t2, t3, t4, t5 = st.tabs(

        [
            "제원",
            "도면",
            "이력",
            "AI진단",
            "📖 설비 스토리"
        ]

    )

    with t1:

        st.dataframe(

            pd.DataFrame(

                [

                    ["사업장", pump["site"]],

                    ["설비명", pump["equip"]],

                    ["제조사", pump["maker"]],

                    ["모델명", pump["model"]],

                    ["정격출력(HP)", f'{pump["hp"]}'],

                    ["정격양정(m)", f'{pump["head"]}'],

                    ["정격유량(m³/h)", f'{pump["flow"]}'],

                    ["회전수(RPM)", f'{pump["rpm"]}'],

                    ["준공일", pump["build_date"]],

                    [
                        "누적 운전시간",
                        f'{pump["op_hours"]:,} h'
                    ]

                ],

                columns=[
                    "항목",
                    "값"
                ]

            ),

            use_container_width=True,

            hide_index=True

        )

    with t2:

        drawing_file = st.file_uploader(

            "설비 도면 업로드 (배관도·조립도 등)",

            type=["png", "jpg", "jpeg"],

            key=f"drawing_upload_{pump['equip']}",

            disabled=is_read_only()

        )

        if drawing_file is not None and not is_read_only():

            save_drawing(

                pump["equip"],

                drawing_file.getvalue(),

                drawing_file.name

            )

            st.success(
                "도면이 저장되었습니다."
            )

        df_drawings = read_excel(

            DRAWING_DB_PATH,

            "도면이력"

        )

        pump_drawings = pd.DataFrame()

        if (

            not df_drawings.empty

            and
            "설비명" in df_drawings.columns

        ):

            pump_drawings = df_drawings[

                df_drawings["설비명"] == pump["equip"]

            ].sort_values("등록일자")

        if not pump_drawings.empty:

            latest = pump_drawings.iloc[-1]

            latest_path = os.path.join(

                DRAWING_DIR,

                str(latest["파일명"])

            )

            if os.path.exists(latest_path):

                st.image(

                    latest_path,

                    use_container_width=True,

                    caption=(

                        f"{pump['equip']} 최신 도면 "
                        f"({latest['등록일자']})"

                    )

                )

                st.caption(
                    "스마트폰에서는 두 손가락으로 확대(핀치 줌)해서 볼 수 있습니다."
                )

            if len(pump_drawings) > 1:

                with st.expander(

                    f"📜 이전 도면 이력 보기 (총 {len(pump_drawings)}건)"

                ):

                    for _, drow in pump_drawings.iloc[:-1][::-1].iterrows():

                        old_path = os.path.join(

                            DRAWING_DIR,

                            str(drow["파일명"])

                        )

                        if os.path.exists(old_path):

                            st.image(

                                old_path,

                                use_container_width=True,

                                caption=f"등록일 : {drow['등록일자']}"

                            )

        else:

            st.info(
                "등록된 도면이 없습니다. "
                "배관도·조립도 이미지를 업로드하면 "
                "현장에서 바로 확대해 볼 수 있습니다."
            )

    with t3:

        st.markdown(
            "##### 🕘 통합 타임라인"
        )

        st.caption(

            "정밀진단·오버홀·진동측정 이력을 시간순으로 "
            "한 화면에서 봅니다."

        )

        _timeline_events = build_equipment_timeline(

            pump,

            df_history,

            read_excel(OVERHAUL_DB_PATH, "오버홀이력"),

            read_excel(VIBRATION_DB_PATH, "진동측정이력")

        )

        if _timeline_events:

            _timeline_color = {

                "정밀진단": "#087ea4",
                "오버홀": "#e8590c",
                "진동측정": "#2f9e44"

            }

            _timeline_html_parts = []

            for ev in _timeline_events[:20]:

                _c = _timeline_color.get(

                    ev["type"],
                    "#64748b"

                )

                _timeline_html_parts.append(

                    f"""
                    <div style="
                    display:flex; gap:10px;
                    padding:6px 0;
                    border-left:3px solid {_c};
                    padding-left:12px; margin-bottom:6px;">
                        <div style="font-size:1.1rem;">
                        {ev['icon']}
                        </div>
                        <div>
                            <div style="
                            font-size:0.78rem; color:#64748b;">
                            {ev['date']} · {ev['type']}
                            </div>
                            <div style="font-size:0.9rem;">
                            {ev['desc']}
                            </div>
                        </div>
                    </div>
                    """

                )

            st.markdown(

                "".join(_timeline_html_parts),

                unsafe_allow_html=True

            )

            if len(_timeline_events) > 20:

                st.caption(

                    f"최근 20건만 표시됩니다 "
                    f"(전체 {len(_timeline_events)}건)."

                )

        else:

            st.info(
                "아직 기록된 이력이 없습니다."
            )

        st.markdown(
            "##### 최근 정밀진단 이력"
        )

        pump_history = df_history[
            df_history["설비명"] == pump["equip"]
        ] if not df_history.empty else pd.DataFrame()

        if not pump_history.empty:

            st.dataframe(

                pump_history[
                    [
                        "점검일",
                        "종합점수",
                        "최종등급"
                    ]
                ].tail(5),

                use_container_width=True,

                hide_index=True

            )

            if len(pump_history) > 5:

                with st.expander(
                    f"전체 진단이력 펼쳐보기 (총 {len(pump_history)}건)"
                ):

                    st.dataframe(

                        pump_history[
                            [
                                "점검일",
                                "종합점수",
                                "최종등급"
                            ]
                        ],

                        use_container_width=True,

                        hide_index=True

                    )

        else:

            st.info(
                "저장된 정밀진단 이력이 없습니다."
            )

        st.markdown(
            "##### 오버홀·현장 작업 이력"
        )

        df_overhaul = read_excel(
            OVERHAUL_DB_PATH,
            "오버홀이력"
        )

        pump_overhaul = df_overhaul[
            df_overhaul["설비명"] == pump["equip"]
        ] if not df_overhaul.empty else pd.DataFrame()

        if not pump_overhaul.empty:

            st.dataframe(

                pump_overhaul[
                    [
                        "작업일자",
                        "공정단계",
                        "작업내용"
                    ]
                ].tail(5),

                use_container_width=True,

                hide_index=True

            )

            if len(pump_overhaul) > 5:

                with st.expander(
                    f"전체 작업이력 펼쳐보기 (총 {len(pump_overhaul)}건)"
                ):

                    st.dataframe(

                        pump_overhaul[
                            [
                                "작업일자",
                                "공정단계",
                                "작업내용"
                            ]
                        ],

                        use_container_width=True,

                        hide_index=True

                    )

            photo_rows = pump_overhaul[

                pump_overhaul["사진파일명"].astype(str) != ""

            ] if "사진파일명" in pump_overhaul.columns else pd.DataFrame()

            for _, prow in photo_rows.tail(3).iterrows():

                photo_path = os.path.join(
                    PHOTO_DIR,
                    str(prow["사진파일명"])
                )

                if os.path.exists(photo_path):

                    st.image(

                        photo_path,

                        caption=f'{prow["작업일자"]} 작업사진',

                        use_container_width=True

                    )

        else:

            st.info(
                "저장된 오버홀·작업 이력이 없습니다."
            )

        st.write("")

        work = st.text_area(

            "현장 작업 메모 추가",

            placeholder="점검 및 작업 내용을 입력하세요.",

            key=f"qr_note_{pump['equip']}",

            disabled=is_read_only()

        )

        if is_read_only():

            st.info(
                "🔒 보기 전용 모드에서는 저장할 수 없습니다."
            )

        elif st.button(

            "작업기록 저장",

            type="primary",

            key=f"qr_save_{pump['equip']}"

        ):

            safe_append_row(

                OVERHAUL_DB_PATH,

                "오버홀이력",

                [
                    datetime.now().strftime(
                        "%Y-%m-%d"
                    ),

                    pump["site"],

                    pump["equip"],

                    "현장메모",

                    st.session_state.user_name,

                    work,

                    "",

                    "",

                    ""

                ]

            )

            st.success(
                "현장 작업기록이 저장되었습니다."
            )

    with t4:

        status_class = {

            "정상": "status-normal",

            "관찰": "status-watch",

            "정비검토": "status-danger"

        }.get(
            result["상태"],
            "status-watch"
        )

        st.markdown(
            f"""
            <span class="{status_class}">
            {status_icon(result['상태'])} {result['상태']}
            </span>
            """,
            unsafe_allow_html=True
        )

        a, b, c = st.columns(3)

        a.metric(
            "CBM Score",
            f'{result["점수"]}점'
        )

        b.metric(
            "진동(RMS)",
            f'{result["진동"]:.1f} mm/s'
        )

        c.metric(
            "온도",
            f'{result["온도"]:.1f}°C'
        )

        fig = build_vibration_trend_fig(

            pump,

            result,

            df_history=df_history,

            figsize=(7, 3)

        )

        st.pyplot(
            fig
        )

        _th_qr = get_alert_thresholds()

        if result["진동"] >= _th_qr["vib_danger"]:

            st.error(
                "CBM 예측 : 고위험 상태 · 정비검토 필요"
            )

        elif result["진동"] >= _th_qr["vib_watch"]:

            st.warning(
                "CBM 예측 : 주의 상태 · 추이관찰 및 정밀진단 권고"
            )

        else:

            st.success(
                "CBM 예측 : 현재 상태 양호"
            )

        if result.get("다음오버홀까지남은시간") is not None:

            st.caption(

                f"🔧 다음 오버홀 예정까지 "
                f"{result['다음오버홀까지남은시간']:,}시간 남음"

            )

    with t5:

        st.caption(

            "설비 하나의 일생을 이야기 형식으로 정리합니다. "
            "경진대회나 보고 자리에서 설비 하나만 보여줘도 "
            "전체 시스템이 어떻게 돌아가는지 보여주기 좋습니다."

        )

        df_overhaul_story = read_excel(

            OVERHAUL_DB_PATH,

            "오버홀이력"

        )

        _story_overhaul = (

            df_overhaul_story[

                df_overhaul_story["설비명"] == pump["equip"]

            ]

            if not df_overhaul_story.empty
            and "설비명" in df_overhaul_story.columns

            else pd.DataFrame()

        )

        _story_diag_count = (

            len(

                df_history[

                    df_history["설비명"] == pump["equip"]

                ]

            )

            if not df_history.empty
            and "설비명" in df_history.columns

            else 0

        )

        _tco, _overhaul_count = calculate_equipment_tco(

            pump["equip"],

            df_overhaul_story

        )

        try:

            _build_date = datetime.strptime(

                str(pump.get("build_date", ""))[:10],

                "%Y-%m-%d"

            ).date()

            _age_years = round(

                (datetime.now().date() - _build_date).days

                / 365,

                1

            )

            _age_text = f"{_build_date} 준공, 약 {_age_years}년째 가동 중"

        except Exception:

            _age_text = "준공일 정보 없음"

        _df_vib_story = read_excel(

            VIBRATION_DB_PATH,

            "진동측정이력"

        )

        _story_pred = predict_failure_date_regression(

            pump,

            _df_vib_story

        )

        if _story_pred and _story_pred.get("status") == "predicted":

            _pred_text = (

                f"AI 회귀분석 결과, 약 {_story_pred['predicted_date']}"
                f"경 위험기준에 도달할 것으로 예측됩니다."

            )

        elif _story_pred and _story_pred.get("status") == "stable":

            _pred_text = "AI 분석 결과 현재 진동 추세는 안정적입니다."

        else:

            _pred_text = "AI 예측을 위한 진동 데이터가 아직 부족합니다."

        st.markdown(

            f"""
            <div class="story-card">

            <div style="font-size:1.1rem; font-weight:800;
            color:#0f3552; margin-bottom:10px;">
            📖 {pump['equip']} ({pump['site']})의 이야기
            </div>

            <p style="line-height:1.9; font-size:0.95rem;
            color:#334155;">
            이 설비는 <b>{_age_text}</b>입니다.<br>
            지금까지 <b>정밀진단 {_story_diag_count}회</b>,
            <b>오버홀 {_overhaul_count}회</b>를 거쳤고,
            누적된 오버홀 비용은
            <b>{_tco:,.0f}원</b>입니다.<br>
            현재 CBM Score는 <b>{result['점수']}점
            ({result['등급']}등급)</b>이며,
            {_pred_text}
            </p>

            </div>
            """,

            unsafe_allow_html=True

        )

        if not _story_overhaul.empty:

            st.markdown(
                "##### 오버홀 이력"
            )

            st.dataframe(

                _story_overhaul[

                    [
                        "작업일자",
                        "작업내용",
                        "작업자",
                        "비용"
                    ]

                ],

                use_container_width=True,

                hide_index=True

            )

    def _build_qr_export_sheets():

        sheets = build_equipment_export_sheets(

            pump,

            result,

            df_history

        )

        qr_id_row = pd.DataFrame(

            [["QR ID", qr_id]],

            columns=["항목", "값"]

        )

        sheets[0]["df"] = pd.concat(

            [qr_id_row, sheets[0]["df"]],

            ignore_index=True

        )

        return sheets

    render_excel_export_section(

        f"qr_{pump['equip']}",

        f"{pump['equip']}_QR정보.xlsx",

        _build_qr_export_sheets

    )


# ============================================================
# 18. 정밀진단
# ============================================================

elif st.session_state.page == "진단":

    render_back_to_home_button("diag")

    _diag_mascot_html = get_mascot_base64_html(
        max_height_px=100,
        name="mascot_control"
    )

    st.markdown(

        f"""
        <div class="page-header-mascot">

            <div class="page-header-text">
                <div class="section-title">
                🔍 펌프 정밀 상태진단
                </div>

                <div class="section-caption">
                17개 핵심 진단항목을 기반으로 설비 상태를
                점수화하고 CBM 정비판단에 활용합니다.
                </div>
            </div>

            {f'<div class="page-header-mascot-img">{_diag_mascot_html}</div>' if _diag_mascot_html else ''}

        </div>
        """,

        unsafe_allow_html=True

    )

    with st.expander(

        "📋 점검 전 체크리스트 템플릿 보기"

    ):

        _checklist_names = get_checklist_template_names()

        if not _checklist_names:

            st.caption(

                "등록된 체크리스트 템플릿이 없습니다."

            )

        else:

            _selected_checklist = st.selectbox(

                "설비 종류에 맞는 템플릿 선택",

                _checklist_names,

                key="diag_checklist_select"

            )

            _checklist_items = get_checklist_template_items(

                _selected_checklist

            )

            for _item in _checklist_items:

                st.checkbox(
                    _item,
                    key=f"diag_checklist_{_selected_checklist}_{_item}"
                )

    if not ALL_PUMPS:

        st.info(

            "등록된 설비가 없습니다. '데이터 관리' 메뉴에서 "
            "설비를 추가하거나 엑셀을 업로드해주세요."

        )

        st.stop()

    selected = st.selectbox(

        "진단 대상",

        [
            p["equip"]
            for p in ALL_PUMPS
        ]

    )

    pump = next(
        p for p in ALL_PUMPS
        if p["equip"] == selected
    )

    log_recent_view(
        pump["equip"]
    )

    # ------------------------------------------------------
    # 임시저장(초안) 공간.
    #
    # 예전에는 number_input/selectbox의 key만 믿고 있었는데,
    # 실제로 다른 메뉴로 이동했다가 돌아오면 Streamlit이
    # 렌더링되지 않은 위젯의 값을 초기화해버려서
    # 17개 항목을 입력하다가 실수로 다른 메뉴를 누르면
    # 처음부터 다시 입력해야 하는 문제가 있었다.
    # (실제로 재현 테스트해서 확인함)
    #
    # 이제 위젯 key와 별개로 diag_draft 딕셔너리에 값을
    # on_change 콜백으로 복사해두고, 위젯을 다시 그릴 때
    # 그 딕셔너리 값을 초기값으로 사용한다.
    # ------------------------------------------------------

    if "diag_draft" not in st.session_state:

        st.session_state.diag_draft = {}

    if pump["equip"] not in st.session_state.diag_draft:

        # 세션에 없으면(새로고침 등으로 세션이 새로 시작된 경우
        # 포함) 디스크에 저장해둔 초안이 있는지 먼저 확인한다.

        st.session_state.diag_draft[pump["equip"]] = load_draft_from_disk(
            pump["equip"]
        )

    draft = st.session_state.diag_draft[
        pump["equip"]
    ]

    if draft:

        st.caption(
            "📝 이전에 입력하던 내용이 있어 자동으로 불러왔습니다 "
            "(새로고침해도 유지됩니다)."
        )

    st.info(
        f'{pump["site"]} / '
        f'{pump["equip"]} / '
        f'{pump["maker"]} {pump["model"]}'
    )

    if pump.get("기준진동") or pump.get("기준효율"):

        st.caption(
            f"⚙️ 이 설비는 맞춤 기준 적용 중 · "
            f"기준진동 {pump.get('기준진동') or '공통기준'} mm/s · "
            f"기준효율 {pump.get('기준효율') or '공통기준'} %"
        )

    total_score = 0

    category_score = {
        "성능": 0,
        "내부상태": 0,
        "기계상태": 0,
        "정비이력": 0
    }

    details = []

    auto_raw_values = {}

    def render_eval_item(
        idx,
        item
    ):

        category = item[0]

        name = item[1]

        weight = item[2]

        standard = item[3]

        options = item[4]

        score_map = item[5]

        auto_fn = item[6]

        with st.expander(
            f"{idx + 1}. {name} · {weight}점",
            expanded=False
        ):

            st.caption(
                f"판정기준 : {standard}"
            )

            glossary_text = term_help(
                name
            )

            if glossary_text:

                st.caption(
                    f"ℹ️ {glossary_text}"
                )

            if auto_fn is not None:

                min_v, max_v, default_v, step_v, unit = AUTO_INPUT_CONFIG[
                    name
                ]

                def _save_val_draft(
                    idx=idx
                ):

                    draft[
                        f"val_{idx}"
                    ] = st.session_state[
                        f"val_{idx}"
                    ]

                    save_draft_to_disk(
                        pump["equip"],
                        draft
                    )

                raw_value = st.number_input(

                    f"측정값 입력 ({unit})",

                    min_value=float(min_v),

                    max_value=float(max_v),

                    value=float(
                        draft.get(
                            f"val_{idx}",
                            default_v
                        )
                    ),

                    step=float(step_v),

                    key=f"val_{idx}",

                    on_change=_save_val_draft,

                    disabled=is_read_only()

                )

                auto_raw_values[
                    name
                ] = raw_value

                effective_fn = get_effective_auto_fn(
                    name,
                    auto_fn,
                    pump
                )

                grade = effective_fn(
                    raw_value
                )

                st.metric(
                    "자동판정",
                    grade
                )

                if grade == "E":

                    st.caption(

                        "⚠️ 측정값이 정상 범위를 크게 벗어난 "
                        "최하 등급입니다. 오타나 측정 오류가 "
                        "아닌지 한 번 더 확인해주세요."

                    )

            elif name == "NPSH 여유율/캐비테이션":

                use_calc = st.toggle(

                    "📐 NPSH 계산기 사용",

                    key=f"npsh_calc_toggle_{idx}",

                    disabled=is_read_only()

                )

                if use_calc:

                    nc1, nc2 = st.columns(2)

                    atm = nc1.number_input(
                        "대기압수두(m) — 보통 10.33",
                        value=10.33,
                        key=f"npsh_atm_{idx}",
                        disabled=is_read_only()
                    )

                    vapor = nc2.number_input(
                        "포화증기압수두(m) — 상온 물 약 0.24",
                        value=0.24,
                        key=f"npsh_vapor_{idx}",
                        disabled=is_read_only()
                    )

                    lift = nc1.number_input(
                        "흡입양정(m)",
                        value=2.0,
                        key=f"npsh_lift_{idx}",
                        disabled=is_read_only()
                    )

                    friction = nc2.number_input(
                        "배관 마찰손실(m)",
                        value=0.5,
                        key=f"npsh_friction_{idx}",
                        disabled=is_read_only()
                    )

                    npshr = st.number_input(
                        "NPSHr (펌프 제조사 요구값, m)",
                        value=3.0,
                        key=f"npsh_r_{idx}",
                        disabled=is_read_only()
                    )

                    npsha = calc_npsha(
                        atm,
                        vapor,
                        lift,
                        friction
                    )

                    margin_ratio = (

                        npsha / npshr * 100

                        if npshr > 0

                        else 0

                    )

                    grade = calc_npsh_margin_grade(
                        margin_ratio
                    )

                    st.metric(

                        "계산된 NPSHa / 여유율",

                        f"{npsha:.2f} m / {margin_ratio:.0f}%"

                    )

                    st.metric(
                        "자동판정",
                        grade
                    )

                else:

                    def _save_grade_draft(
                        idx=idx
                    ):

                        draft[
                            f"grade_{idx}"
                        ] = st.session_state[
                            f"grade_{idx}"
                        ]

                        save_draft_to_disk(
                            pump["equip"],
                            draft
                        )

                    grade = st.selectbox(

                        "판정",

                        options,

                        index=options.index(
                            draft.get(
                                f"grade_{idx}",
                                options[0]
                            )
                        ),

                        key=f"grade_{idx}",

                        on_change=_save_grade_draft,

                        disabled=is_read_only()

                    )

            else:

                def _save_grade_draft(
                    idx=idx
                ):

                    draft[
                        f"grade_{idx}"
                    ] = st.session_state[
                        f"grade_{idx}"
                    ]

                    save_draft_to_disk(
                        pump["equip"],
                        draft
                    )

                grade = st.selectbox(

                    "판정",

                    options,

                    index=options.index(
                        draft.get(
                            f"grade_{idx}",
                            options[0]
                        )
                    ),

                    key=f"grade_{idx}",

                    on_change=_save_grade_draft,

                    disabled=is_read_only()

                )

        score = score_map[
            grade
        ]

        return category, name, grade, score

    category_order = [
        "성능",
        "내부상태",
        "기계상태",
        "정비이력"
    ]

    category_tabs = st.tabs(

        [
            f"{cat} · {CATEGORIES[cat]}점"
            for cat in category_order
        ]

    )

    for tab, cat in zip(
        category_tabs,
        category_order
    ):

        with tab:

            cat_items = [

                (idx, item)

                for idx, item in enumerate(EVAL_ITEMS)

                if item[0] == cat

            ]

            for i in range(
                0,
                len(cat_items),
                2
            ):

                pair = cat_items[i:i + 2]

                cols = st.columns(
                    len(pair)
                )

                for col, (idx, item) in zip(
                    cols,
                    pair
                ):

                    with col:

                        category, name, grade, score = render_eval_item(
                            idx,
                            item
                        )

                        category_score[
                            category
                        ] += score

                        total_score += score

                        details.append(
                            {
                                "항목": name,
                                "등급": grade,
                                "점수": score
                            }
                        )

    total_score = round(
        total_score,
        2
    )

    final_grade = get_final_grade(
        total_score
    )

    st.write("---")

    c1, c2, c3 = st.columns(3)

    c1.metric(
        "종합점수",
        f"{total_score}점"
    )

    c2.metric(
        "최종등급",
        final_grade
    )

    c3.metric(
        "판정",
        get_grade_text(
            final_grade
        )
    )

    st.markdown(
        "### 분야별 상태"
    )

    cat_df = pd.DataFrame(

        [
            [
                key,
                round(
                    value,
                    1
                ),
                CATEGORIES[key]
            ]

            for key, value
            in category_score.items()
        ],

        columns=[
            "분야",
            "획득점수",
            "배점"
        ]

    )

    st.dataframe(
        cat_df,
        use_container_width=True,
        hide_index=True
    )

    def _save_temp_draft():

        draft["diag_temp"] = st.session_state.diag_temp

        save_draft_to_disk(
            pump["equip"],
            draft
        )

    temp_measured = st.number_input(

        "측정 온도 (°C) — EVAL_ITEMS에는 없지만 "
        "AI 이상징후 추세에 함께 반영됩니다",

        min_value=0.0,

        max_value=120.0,

        value=float(
            draft.get("diag_temp", 45.0)
        ),

        step=0.5,

        key="diag_temp",

        on_change=_save_temp_draft,

        disabled=is_read_only()

    )

    def _save_current_draft():

        draft["diag_current"] = st.session_state.diag_current

        save_draft_to_disk(
            pump["equip"],
            draft
        )

    current_measured = st.number_input(

        "측정 전류 (A) — 부하 이상은 진동보다 "
        "전류에서 먼저 나타나는 경우가 많습니다",

        min_value=0.0,

        max_value=1000.0,

        value=float(
            draft.get("diag_current", 0.0)
        ),

        step=0.5,

        key="diag_current",

        on_change=_save_current_draft,

        disabled=is_read_only()

    )

    if is_read_only():

        st.info(
            "🔒 보기 전용 모드입니다. "
            "저장하려면 사이드바에서 보기 전용 모드를 해제하세요."
        )

    else:

        if st.button(
            "💾 진단결과 저장",
            type="primary",
            use_container_width=True
        ):

            save_time = datetime.now().strftime(
                "%Y-%m-%d"
            )

            row = [

                save_time,

                pump["site"],

                pump["equip"],

                pump["maker"],

                pump["model"],

                pump["hp"],

                pump["head"],

                pump["build_date"],

                st.session_state.user_name,

                total_score,

                final_grade

            ] + [

                x["등급"]
                for x in details

            ] + [

                auto_raw_values.get(
                    "펌프 효율 유지율 (%)",
                    ""
                ),

                auto_raw_values.get(
                    "Overall 진동 (mm/s)",
                    ""
                ),

                temp_measured,

                current_measured

            ]

            safe_append_row(

                DB_FILE_PATH,

                "진단이력",

                row

            )

            log_audit(

                "정밀진단 저장",

                pump["equip"],

                f"종합점수={total_score}, 등급={final_grade}"

            )

            st.success(
                f"{pump['equip']} "
                f"진단결과가 저장되었습니다."
            )

            # 위험 등급(D/E)이면 설정된 웹훅으로 알림을 보낸다.
            # 웹훅이 설정 안 돼 있으면 아무 일도 안 일어난다.

            if final_grade in ("D", "E"):

                send_webhook_notification(

                    f"🚨 [K-water tech] {pump['equip']} 정밀진단 결과 "
                    f"{final_grade}등급({total_score}점) - 정비검토가 필요합니다."

                )

            # 저장이 끝났으니 이 설비의 임시저장(초안)은
            # 세션과 디스크 양쪽에서 모두 비운다.

            clear_draft(
                pump["equip"]
            )

            # 저장 직후 방금 저장한 값이 실제로
            # 반영됐는지 바로 눈으로 확인할 수 있도록
            # 파일을 다시 읽어서 보여준다.

            df_just_saved = read_excel(
                DB_FILE_PATH,
                "진단이력"
            )

            st.markdown(
                "###### ✅ 방금 저장된 기록"
            )

            st.dataframe(

                df_just_saved[
                    df_just_saved["설비명"] == pump["equip"]
                ].tail(1),

                use_container_width=True,

                hide_index=True

            )

    render_excel_export_section(

        f"diag_{pump['equip']}",

        f"{pump['equip']}_정밀진단결과.xlsx",

        lambda: pd.DataFrame(details)

    )


# ============================================================
# 19. CBM
# ============================================================

elif st.session_state.page == "CBM":

    render_back_to_home_button("cbm")

    _cbm_mascot_html = get_mascot_base64_html(
        max_height_px=100,
        name="mascot_water"
    )

    st.markdown(

        f"""
        <div class="page-header-mascot">

            <div class="page-header-text">
                <div class="section-title">
                🎯 CBM 정비판단
                </div>

                <div class="section-caption">
                운전시간만으로 오버홀을 결정하지 않고
                성능·진동·내부상태·정비이력을 종합하여
                정비 우선순위를 판단합니다.
                </div>
            </div>

            {f'<div class="page-header-mascot-img">{_cbm_mascot_html}</div>' if _cbm_mascot_html else ''}

        </div>
        """,

        unsafe_allow_html=True

    )

    if not ALL_PUMPS:

        st.info(

            "등록된 설비가 없습니다. '데이터 관리' 메뉴에서 "
            "설비를 추가하거나 엑셀을 업로드해주세요."

        )

        st.stop()

    ranking = []

    df_vib_for_cbm = read_excel(

        VIBRATION_DB_PATH,

        "진동측정이력"

    )

    for pump in ALL_PUMPS:

        result = pump_status(
            pump,
            df_history
        )

        risk = round(
            100 - result["점수"],
            1
        )

        pred = predict_failure_date_regression(

            pump,
            df_vib_for_cbm

        )

        if pred and pred.get("status") == "predicted":

            ai_pred_text = str(
                pred["predicted_date"]
            )

            _days_until = (

                pred["predicted_date"]

                -
                datetime.now().date()

            ).days

            if _days_until <= 90:

                ai_urgency = "🔴 임박"

            elif _days_until <= 180:

                ai_urgency = "🟡 관찰"

            else:

                ai_urgency = "🟢 여유"

            _sort_key = _days_until

        elif pred and pred.get("status") == "stable":

            ai_pred_text = "안정적"

            ai_urgency = "🟢 여유"

            _sort_key = 99999

        else:

            ai_pred_text = "예측불가"

            ai_urgency = "-"

            _sort_key = 100000

        ranking.append(

            {
                "우선순위": 0,
                "설비": pump["equip"],
                "CBM Score": result["점수"],
                "Risk": risk,
                "운전시간": pump["op_hours"],
                "효율": result["효율"],
                "진동": result["진동"],
                "등급": result["등급"],
                "정비판단": result["상태"],
                "AI긴급도": ai_urgency,
                "AI예측 위험도달일": ai_pred_text,
                "_ai_sort_key": _sort_key
            }

        )

    ranking = sorted(
        ranking,
        key=lambda x: x["Risk"],
        reverse=True
    )

    for i, row in enumerate(
        ranking
    ):

        row["우선순위"] = i + 1

    # ------------------------------------------------------
    # 설비 간 비교(벤치마킹): "이 설비가 나쁘다"를 절대기준이
    # 아니라 "우리 설비군 평균 대비 얼마나 벗어나 있는지"로도
    # 보여준다. 예전에는 설비 하나만 보게 되어 있어서
    # "가압펌프 #3이 왜 유독 나쁘지?"를 판단할 비교 대상이 없었다.
    # ------------------------------------------------------

    fleet_avg_eff = sum(

        r["효율"] for r in ranking

    ) / len(ranking)

    fleet_avg_vib = sum(

        r["진동"] for r in ranking

    ) / len(ranking)

    fleet_avg_score = sum(

        r["CBM Score"] for r in ranking

    ) / len(ranking)

    for row in ranking:

        row["효율(전체평균대비)"] = round(

            row["효율"] - fleet_avg_eff,
            1

        )

        row["진동(전체평균대비)"] = round(

            row["진동"] - fleet_avg_vib,
            1

        )

    df_rank = pd.DataFrame(
        ranking
    )

    _status_options = ["전체", "정상", "관찰", "정비검토"]

    _preset_status = st.session_state.get(
        "_cbm_status_filter"
    )

    _status_default_index = (

        _status_options.index(_preset_status)

        if _preset_status in _status_options

        else 0

    )

    cbm_status_filter = st.selectbox(

        "정비판단 필터",

        _status_options,

        index=_status_default_index,

        key="cbm_status_filter_select"

    )

    st.session_state["_cbm_status_filter"] = None

    if cbm_status_filter != "전체":

        df_rank = df_rank[

            df_rank["정비판단"] == cbm_status_filter

        ]

    sort_by_ai = st.checkbox(

        "🤖 AI예측 임박한 순으로 정렬",

        key="sort_by_ai_prediction"

    )

    if sort_by_ai and "_ai_sort_key" in df_rank.columns:

        df_rank = df_rank.sort_values(
            "_ai_sort_key"
        )

    df_rank_display = df_rank.drop(

        columns=["_ai_sort_key"],

        errors="ignore"

    )

    st.dataframe(
        df_rank_display,
        use_container_width=True,
        hide_index=True
    )

    st.caption(

        f"전체 {len(ranking)}대 평균 · "

        f"효율 {fleet_avg_eff:.1f}% · "

        f"진동 {fleet_avg_vib:.1f} mm/s · "

        f"CBM Score {fleet_avg_score:.1f}점 "

        "— '전체평균대비' 컬럼이 음수면 평균보다 나쁜 상태입니다."

    )

    st.write("")

    st.markdown(
        "### 📊 설비 간 비교 (벤치마킹)"
    )

    compare_target = st.selectbox(

        "비교할 설비 선택",

        [
            r["설비"]
            for r in ranking
        ],

        key="cbm_compare_select"

    )

    compare_row = next(

        r for r in ranking

        if r["설비"] == compare_target

    )

    fig, (ax1, ax2, ax3) = plt.subplots(

        1,
        3,

        figsize=(10, 3.2)

    )

    def _draw_compare_bar(
        ax,
        title,
        pump_value,
        fleet_value,
        unit
    ):

        bars = ax.bar(

            ["선택 설비", "전체 평균"],

            [pump_value, fleet_value],

            color=["#087ea4", "#adb5bd"]

        )

        ax.set_title(
            title
        )

        ax.set_ylabel(
            unit
        )

        for b in bars:

            ax.text(

                b.get_x() + b.get_width() / 2,

                b.get_height(),

                f"{b.get_height():.1f}",

                ha="center",

                va="bottom",

                fontsize=8

            )

    _draw_compare_bar(
        ax1,
        "효율",
        compare_row["효율"],
        fleet_avg_eff,
        "%"
    )

    _draw_compare_bar(
        ax2,
        "진동",
        compare_row["진동"],
        fleet_avg_vib,
        "mm/s"
    )

    _draw_compare_bar(
        ax3,
        "CBM Score",
        compare_row["CBM Score"],
        fleet_avg_score,
        "점"
    )

    fig.tight_layout()

    st.pyplot(
        fig
    )

    st.write("")

    top = ranking[0]

    st.warning(
        f"최우선 정비 대상 : "
        f"{top['설비']} / "
        f"CBM Score {top['CBM Score']}점 / "
        f"Risk {top['Risk']}%"
    )

    st.markdown(
        """
        <div class="platform-card">

        <div class="card-title">
        CBM 판단 로직
        </div>

        <b>① 성능상태</b>
        → 효율·양정·유량·BEP

        <br>

        <b>② 내부상태</b>
        → 링간극·축슬리브·임펠러·코팅

        <br>

        <b>③ 기계상태</b>
        → 진동·베어링·센터링·배관응력

        <br>

        <b>④ 정비이력</b>
        → 운전시간·소모품·오버홀 이력

        <br><br>

        <b>⑤ 종합 CBM Score</b>
        → 정비 우선순위 Ranking

        </div>
        """,
        unsafe_allow_html=True
    )

    render_excel_export_section(

        "cbm",

        "CBM_정비우선순위.xlsx",

        lambda: [

            {
                "name": "정비우선순위",
                "df": df_rank,
                "chart": "auto"
            },

            {
                "name": "설비별비교",
                "df": pd.DataFrame(

                    {
                        "지표": ["효율(%)", "진동(mm/s)", "CBM Score"],
                        compare_target: [
                            compare_row["효율"],
                            compare_row["진동"],
                            compare_row["CBM Score"]
                        ],
                        "전체평균": [
                            round(fleet_avg_eff, 1),
                            round(fleet_avg_vib, 1),
                            round(fleet_avg_score, 1)
                        ]
                    }

                ),
                "chart": "auto",
                "title": f"{compare_target} vs 전체 {len(ranking)}대 평균"
            }

        ]

    )


# ============================================================
# 20. 오버홀 관리
# ============================================================

elif st.session_state.page == "오버홀":

    render_back_to_home_button("overhaul")

    _overhaul_mascot_html = get_mascot_base64_html(
        max_height_px=100,
        name="mascot_pump"
    )

    st.markdown(

        f"""
        <div class="page-header-mascot">

            <div class="page-header-text">
                <div class="section-title">
                🛠️ 오버홀 관리
                </div>

                <div class="section-caption">
                진단 결과를 실제 정비활동과 연결하고
                오버홀 전·후 상태변화를 데이터로 축적합니다.
                </div>
            </div>

            {f'<div class="page-header-mascot-img">{_overhaul_mascot_html}</div>' if _overhaul_mascot_html else ''}

        </div>
        """,

        unsafe_allow_html=True

    )

    if not ALL_PUMPS:

        st.info(

            "등록된 설비가 없습니다. '데이터 관리' 메뉴에서 "
            "설비를 추가하거나 엑셀을 업로드해주세요."

        )

        st.stop()

    with st.expander(

        "🗓️ 전체 설비 오버홀 예정 일정 (진동추세 반영)",

        expanded=False

    ):

        st.caption(

            "누적운전시간·준공일 기반 예상일에 실제 오버홀 "
            "이력과 최근 진동 추세까지 반영해서 날짜를 다시 "
            "계산합니다. 진동이 최근 뚜렷하게 오르는 추세면 "
            "예정일을 앞당깁니다."

        )

        df_overhaul_all = read_excel(

            OVERHAUL_DB_PATH,
            "오버홀이력"

        )

        df_vib_all = read_excel(

            VIBRATION_DB_PATH,
            "진동측정이력"

        )

        schedule_rows = []

        for p in ALL_PUMPS:

            est_date, basis, vib_adj = estimate_next_overhaul_advanced(

                p,

                p.get("op_hours"),

                df_overhaul_all,

                df_vib_all

            )

            if est_date is not None:

                schedule_rows.append(

                    (p["equip"], p["site"], est_date, basis, vib_adj)

                )

        schedule_rows.sort(
            key=lambda r: r[2]
        )

        if schedule_rows:

            _cal_view = st.toggle(

                "🗓️ 캘린더로 보기",

                key="overhaul_calendar_toggle"

            )

            if _cal_view:

                _today = datetime.now().date()

                _cal_col1, _cal_col2 = st.columns(2)

                _cal_year = _cal_col1.selectbox(

                    "연도",

                    list(range(_today.year - 1, _today.year + 2)),

                    index=1,

                    key="overhaul_cal_year"

                )

                _cal_month = _cal_col2.selectbox(

                    "월",

                    list(range(1, 13)),

                    index=_today.month - 1,

                    key="overhaul_cal_month"

                )

                _events_by_day = {}

                for equip, site, est_date, basis, vib_adj in (

                    schedule_rows

                ):

                    if (

                        est_date.year == _cal_year

                        and est_date.month == _cal_month

                    ):

                        _events_by_day.setdefault(

                            est_date.day,
                            []

                        ).append(equip)

                st.markdown(

                    build_month_calendar_html(

                        _cal_year,
                        _cal_month,
                        _events_by_day

                    ),

                    unsafe_allow_html=True

                )

                st.caption(

                    "🔧 표시는 그 날짜에 예정된 오버홀 설비입니다."

                )

            today_d = datetime.now().date()

            table_rows = []

            for equip, site, est_date, basis, vib_adj in schedule_rows:

                d_left = (est_date - today_d).days

                if d_left <= 90:

                    urgency = "🔴 임박"

                elif d_left <= 180:

                    urgency = "🟡 관찰"

                else:

                    urgency = "🟢 여유"

                _p_for_compare = next(

                    p for p in ALL_PUMPS

                    if p["equip"] == equip

                )

                _ai_pred = predict_failure_date_regression(

                    _p_for_compare,

                    df_vib_all

                )

                if _ai_pred and _ai_pred.get("status") == "predicted":

                    _ai_date_text = str(

                        _ai_pred["predicted_date"]

                    )

                    _gap_days = (

                        est_date - _ai_pred["predicted_date"]

                    ).days

                    if _gap_days > 30:

                        _ai_date_text += " (오버홀 예정보다 이름)"

                elif _ai_pred and _ai_pred.get("status") == "stable":

                    _ai_date_text = "안정적"

                else:

                    _ai_date_text = "예측불가"

                table_rows.append(

                    {
                        "긴급도": urgency,
                        "사업장": site,
                        "설비명": equip,
                        "예상일(가동률기반)": est_date.strftime(
                            "%Y-%m-%d"
                        ),
                        "AI예측(진동추세기반)": _ai_date_text,
                        "남은일수": d_left,
                        "진동추세 반영": "✓" if vib_adj else "",
                        "산정 근거": basis
                    }

                )

            st.dataframe(

                pd.DataFrame(table_rows),

                use_container_width=True,

                hide_index=True

            )

            ics_bytes = build_overhaul_ics_bytes(

                [

                    (equip, site, est_date, basis)

                    for equip, site, est_date, basis, vib_adj

                    in schedule_rows

                ]

            )

            st.download_button(

                "📅 캘린더 파일(.ics)로 내보내기",

                data=ics_bytes,

                file_name="오버홀_예정일정.ics",

                mime="text/calendar",

                use_container_width=True,

                key="overhaul_ics_download_btn"

            )

            # 임박한 오버홀이 있는데 소모품 재고가 부족하면
            # 미리 발주할 수 있게 교차로 경고해준다.

            _imminent_count = sum(

                1 for _, _, _est, _, _ in schedule_rows

                if (_est - datetime.now().date()).days <= 90

            )

            if _imminent_count > 0:

                _low_stock_oh = get_low_stock_consumables()

                if not _low_stock_oh.empty:

                    st.error(

                        f"🔴 90일 이내 오버홀 예정이 "
                        f"{_imminent_count}건 있는데, 재고가 "
                        "부족한 소모품이 있습니다 — 미리 "
                        "발주를 검토하세요: "
                        +
                        ", ".join(

                            f"{r['소모품명']}(현재고 {r['현재고']}/"
                            f"안전재고 {r['안전재고']})"

                            for _, r in _low_stock_oh.iterrows()

                        )

                    )

        else:

            st.caption(

                "예정일을 계산할 수 있는 설비가 없습니다 "
                "(준공일·운전시간 정보 확인 필요)."

            )

    selected = st.selectbox(

        "작업 대상",

        [
            p["equip"]
            for p in ALL_PUMPS
        ]

    )

    pump = next(
        p for p in ALL_PUMPS
        if p["equip"] == selected
    )

    log_recent_view(
        pump["equip"]
    )

    tab1, tab2, tab3 = st.tabs(

        [
            "작업계획",
            "작업기록",
            "전후 효과"
        ]

    )

    with tab1:

        st.write(
            f"대상설비 : **{pump['equip']}**"
        )

        _est_date, _basis, _vib_adj = estimate_next_overhaul_advanced(

            pump,

            pump.get("op_hours"),

            read_excel(OVERHAUL_DB_PATH, "오버홀이력"),

            read_excel(VIBRATION_DB_PATH, "진동측정이력")

        )

        if _est_date is not None:

            _d_left = (
                _est_date - datetime.now().date()
            ).days

            st.info(

                f"🗓️ 예상 오버홀 시점 : **{_est_date}** "
                f"(약 {_d_left}일 후)"
                +
                (" · ⚠️ 진동추세 반영으로 앞당겨짐" if _vib_adj else "")

            )

            st.caption(
                f"산정 근거: {_basis}"
            )

        steps = [

            "분해 및 상태확인",

            "부품 측정 및 마모판정",

            "가공·교체",

            "조립 및 축정렬",

            "시운전",

            "최종 성능확인"

        ]

        for i, step in enumerate(
            steps
        ):

            st.checkbox(
                f"{i + 1}. {step}",
                key=f"oh_{i}"
            )

    with tab2:

        work = st.text_area(
            "작업내용",
            height=150,
            placeholder=
            "분해점검, 마모측정, "
            "부품교체, 센터링 등",
            disabled=is_read_only()
        )

        work_photo = st.file_uploader(

            "작업 전후 사진 첨부 (선택)",

            type=["png", "jpg", "jpeg"],

            key="overhaul_photo_upload",

            disabled=is_read_only()

        )

        work_cost = st.number_input(

            "이번 작업 비용(원)",

            min_value=0,

            value=0,

            step=100000,

            disabled=is_read_only(),

            help="부품비+인건비+외주비 등 이번 오버홀에 실제로 "
            "든 총비용. 나중에 설비별 총소유비용(TCO) 계산에 "
            "쓰입니다."

        )

        if is_read_only():

            st.info(
                "🔒 보기 전용 모드에서는 저장할 수 없습니다."
            )

        elif st.button(
            "오버홀 작업기록 저장",
            type="primary"
        ):

            photo_filename = ""

            if work_photo is not None:

                os.makedirs(
                    PHOTO_DIR,
                    exist_ok=True
                )

                photo_filename = (

                    f"{pump['equip']}_"

                    +
                    datetime.now().strftime(
                        "%Y%m%d_%H%M%S"
                    )

                    +
                    "_"
                    +
                    work_photo.name

                )

                with open(

                    os.path.join(
                        PHOTO_DIR,
                        photo_filename
                    ),

                    "wb"

                ) as f:

                    f.write(
                        work_photo.getbuffer()
                    )

            safe_append_row(

                OVERHAUL_DB_PATH,

                "오버홀이력",

                [
                    datetime.now().strftime(
                        "%Y-%m-%d"
                    ),

                    pump["site"],

                    pump["equip"],

                    "오버홀",

                    st.session_state.user_name,

                    work,

                    photo_filename,

                    "",

                    "",

                    work_cost

                ]

            )

            st.success(
                "오버홀 작업기록이 저장되었습니다."
            )

    with tab3:

        c1, c2 = st.columns(2)

        before_eff = c1.number_input(
            "정비 전 효율 (%)",
            min_value=0.0,
            max_value=100.0,
            value=72.3
        )

        after_eff = c2.number_input(
            "정비 후 효율 (%)",
            min_value=0.0,
            max_value=100.0,
            value=77.7
        )

        before_vib = c1.number_input(
            "정비 전 진동",
            min_value=0.0,
            value=5.8
        )

        after_vib = c2.number_input(
            "정비 후 진동",
            min_value=0.0,
            value=3.7
        )

        st.metric(
            "효율 개선",
            f"{after_eff-before_eff:+.1f}%p"
        )

        st.metric(
            "진동 감소",
            f"{before_vib-after_vib:+.1f} mm/s"
        )

    render_excel_export_section(

        f"overhaul_{pump['equip']}",

        f"{pump['equip']}_오버홀_전후효과.xlsx",

        lambda: pd.DataFrame(

            [

                ["설비명", pump["equip"]],
                ["정비 전 효율(%)", before_eff],
                ["정비 후 효율(%)", after_eff],
                ["효율 개선(%p)", round(after_eff - before_eff, 2)],
                ["정비 전 진동(mm/s)", before_vib],
                ["정비 후 진동(mm/s)", after_vib],
                ["진동 감소(mm/s)", round(before_vib - after_vib, 2)]

            ],

            columns=["항목", "값"]

        )

    )


# ============================================================
# 21. AI 이상징후
# ============================================================

elif st.session_state.page == "AI":

    render_back_to_home_button("ai")

    st.markdown(
        """
        <div class="section-title">
        📈 AI 이상징후 및 추세분석
        </div>

        <div class="section-caption">
        축적된 진동·효율·운전시간·정비이력을 기반으로
        설비 상태변화를 시각화하고 향후 정비시점을 판단합니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    if not ALL_PUMPS:

        st.info(

            "등록된 설비가 없습니다. '데이터 관리' 메뉴에서 "
            "설비를 추가하거나 엑셀을 업로드해주세요."

        )

        st.stop()

    selected = st.selectbox(

        "분석 대상",

        [
            p["equip"]
            for p in ALL_PUMPS
        ]

    )

    pump = next(
        p for p in ALL_PUMPS
        if p["equip"] == selected
    )

    log_recent_view(
        pump["equip"]
    )

    result = pump_status(
        pump,
        df_history
    )

    st.write("---")

    _ai_tab1, _ai_tab2, _ai_tab3 = st.tabs(
        [
            "🔮 고장예측",
            "🧠 AI 종합의견",
            "📊 통계 이상탐지"
        ]
    )

    with _ai_tab1:

        st.markdown(
            "### 🤖 AI 고장예측 (회귀분석)"
        )

        st.caption(

            "규칙(임계값 비교)이 아니라, 실제 진동 이력에 선형회귀를 "
            "적합시켜 위험 기준(8.5mm/s) 도달 시점을 통계적으로 "
            "예측합니다."

        )

        df_vib_ai = read_excel(

            VIBRATION_DB_PATH,

            "진동측정이력"

        )

        prediction = predict_failure_date_regression(

            pump,
            df_vib_ai

        )

        if prediction is None:

            st.info(
                "이 설비의 진동측정 이력이 없습니다."
            )

        elif prediction["status"] == "insufficient_data":

            st.info(

                f"예측하려면 최소 3개월치 데이터가 필요합니다 "
                f"(현재 {prediction['months_count']}개월)."

            )

        else:

            pred_fig = build_failure_prediction_chart_fig(

                pump,
                prediction

            )

            st.pyplot(
                pred_fig
            )

            if prediction["status"] == "predicted":

                st.warning(

                    f"⚠️ 현재 추세가 이어지면 약 "
                    f"**{prediction['predicted_date']}**경 "
                    f"위험 기준(8.5mm/s)에 도달할 것으로 "
                    f"예측됩니다 (결정계수 R²="
                    f"{prediction['r_squared']:.2f}, 1에 "
                    "가까울수록 추세가 뚜렷함)."
                    +
                    (
                        " 6개월 이상 데이터가 쌓여 월별 계절편차"
                        "(예: 여름철 부하 상승분)까지 반영한 "
                        "예측입니다."

                        if prediction.get("seasonal_applied")

                        else " (데이터가 6개월 미만이라 계절성은 "
                        "반영하지 못했습니다.)"
                    )

                )

            else:

                st.success(

                    f"✅ 진동값이 증가 추세가 아니거나 안정적입니다 "
                    f"(R²={prediction['r_squared']:.2f})."

                )

            log_prediction(
                pump,
                prediction
            )

        with st.expander(

            "📊 예측 검증 (지난 예측이 실제로 맞았는지)"

        ):

            st.caption(

                "예측일이 이미 지난 과거 예측들을 모아서, 그 시점에 "
                "실제로 위험기준에 도달했었는지 확인합니다. 예측이 "
                "쌓일수록 이 표도 같이 쌓입니다."

            )

            verify_df = verify_past_predictions(

                ALL_PUMPS,

                df_vib_ai

            )

            if verify_df.empty:

                st.info(

                    "아직 검증할 만큼 시간이 지난 예측이 없습니다."

                )

            else:

                st.dataframe(

                    verify_df,

                    use_container_width=True,

                    hide_index=True

                )

        st.write("---")



    with _ai_tab2:

        st.markdown(
            "### 🤖 AI 종합의견 생성 (LLM)"
        )

        st.caption(

            "정해진 문장 틀이 아니라, 실제 Claude API를 호출해서 "
            "이 설비의 데이터를 보고 종합의견을 새로 작성합니다. "
            "(Streamlit Secrets에 ANTHROPIC_API_KEY 등록 필요)"

        )

        if st.button(

            "🤖 AI에게 종합의견 작성 요청",

            type="primary",

            key="generate_ai_commentary_btn"

        ):

            with st.status(

                "AI 종합의견 생성 중...",

                expanded=True

            ) as status_box:

                st.write("1/3 · 설비 데이터 정리 중...")

                st.write(

                    f"CBM {result['점수']}점, "
                    f"진동 {result['진동']:.1f}mm/s 확인 완료"

                )

                st.write("2/3 · Claude API 호출 중... "
                         "(몇 초 정도 걸릴 수 있습니다)")

                commentary, error = generate_ai_commentary_for_equipment(

                    pump,
                    result,
                    prediction

                )

                if error:

                    status_box.update(

                        label="AI 종합의견 생성 실패",

                        state="error"

                    )

                else:

                    st.write("3/3 · 완료")

                    status_box.update(

                        label="AI 종합의견 생성 완료",

                        state="complete"

                    )

            if error:

                st.error(
                    error
                )

            else:

                st.session_state["_ai_commentary"] = commentary

        if "_ai_commentary" in st.session_state:

            st.info(

                st.session_state["_ai_commentary"]

            )

        g1, g2 = st.columns(2)

        with g1:

            st.pyplot(

                build_vibration_trend_fig(

                    pump,

                    result,

                    df_history=df_history,

                    figsize=(6.2, 3.4)

                )

            )

        with g2:

            st.pyplot(

                build_efficiency_trend_fig(

                    pump,

                    result,

                    df_history=df_history

                )

            )

        g3, g4 = st.columns(2)

        with g3:

            st.pyplot(

                build_temperature_trend_fig(

                    pump,

                    result,

                    df_history=df_history

                )

            )

        with g4:

            st.pyplot(

                build_score_gauge_fig(

                    pump,

                    result

                )

            )

        st.pyplot(

            build_op_hours_trend_fig(

                pump

            )

        )

        st.caption(
            "📊 진동·효율·온도 그래프는 저장된 진단이력이 "
            "2건 이상 쌓이면 실측값으로, 그전까지는 이해를 돕기 "
            "위한 예시 데이터로 표시됩니다. 운전시간 그래프는 "
            "과거 시점별 기록을 저장하지 않아 항상 추정치입니다."
        )



    with _ai_tab3:

        st.markdown(
            "### 🤖 AI 통계 이상탐지"
        )

        _vib_dates, _vib_values = get_real_history_series(
            df_history,
            pump,
            "진동측정값(mm/s)"
        )

        anomaly = detect_statistical_anomaly(
            _vib_values
        )

        if anomaly is None:

            st.info(

                f"진동 실측 이력이 {len(_vib_values)}건입니다. "
                f"최소 {ANOMALY_MIN_SAMPLES}건 이상 쌓이면 "
                f"이 설비 고유의 평균 운전패턴 대비 "
                f"통계적 이상 여부(z-score)를 자동으로 계산합니다."

            )

        elif anomaly["is_anomaly"]:

            st.error(

                f"🚨 통계적 이상치 감지 · z-score {anomaly['z']:.2f} "
                f"(과거 평균 {anomaly['mean']:.2f} mm/s, "
                f"표준편차 {anomaly['stdev']:.2f} 대비 "
                f"현재 {anomaly['latest']:.2f} mm/s는 "
                f"{ANOMALY_Z_THRESHOLD}표준편차 이상 벗어남, "
                f"실측 {anomaly['n']}건 기준)"

            )

        else:

            st.success(

                f"✅ 통계적으로 정상 범위 · z-score {anomaly['z']:.2f} "
                f"(실측 {anomaly['n']}건 기준, "
                f"과거 평균 {anomaly['mean']:.2f} mm/s 대비 "
                f"{ANOMALY_Z_THRESHOLD}표준편차 이내)"

            )

        st.write("")

        st.write("")

        _th_ai = get_alert_thresholds()

        if (

            result["진동"] >= _th_ai["vib_danger"]

            or
            result["효율"] <= _th_ai["eff_danger"]

            or
            result["온도"] >= _th_ai["temp_danger"]

            or
            result["점수"] < 60

        ):

            st.error(
                "고위험 상태 · 정비검토 필요"
            )

        elif (

            result["진동"] >= _th_ai["vib_watch"]

            or
            result["효율"] <= _th_ai["eff_watch"]

            or
            result["온도"] >= _th_ai["temp_watch"]

            or
            result["점수"] < 80

        ):

            st.warning(
                "주의 상태 · 추이관찰 및 정밀진단 권고"
            )

        else:

            st.success(
                "현재 상태 양호 · 모든 지표 정상범위"
            )

        render_excel_export_section(

            f"ai_{pump['equip']}",

            f"{pump['equip']}_AI이상징후.xlsx",

            lambda: build_equipment_export_sheets(

                pump,

                result,

                df_history

            )

        )


    # ============================================================
    # 21-1. 전사 트렌드
    # ============================================================


elif st.session_state.page == "전사트렌드":

    st.markdown(
        """
        <div class="section-title">
        🌐 전사 트렌드
        </div>

        <div class="section-caption">
        설비 하나하나가 아니라, 전체 설비의 평균적인 추이를
        큰 그림으로 봅니다. 경영진 보고나 전체 현황 파악에
        활용하세요.
        </div>
        """,
        unsafe_allow_html=True
    )

    if not ALL_PUMPS:

        st.info(

            "등록된 설비가 없습니다. '데이터 관리' 메뉴에서 "
            "설비를 추가하거나 엑셀을 업로드해주세요."

        )

        st.stop()

    st.caption(

        "※ 개별 설비 화면과 달리, 실측/예시 데이터를 구분하지 "
        "않고 현재값 기준 참고용 추이를 평균낸 것입니다. "
        "정확한 실측 추세는 각 설비의 AI 이상징후 화면에서 "
        "확인하세요."

    )

    fg1, fg2 = st.columns(2)

    with fg1:

        fig_eff, avg_eff = build_fleet_average_trend_fig(

            ALL_PUMPS,

            df_history,

            "효율",

            "평균 효율",

            "%",

            "#087ea4"

        )

        st.pyplot(
            fig_eff
        )

    with fg2:

        fig_vib, avg_vib = build_fleet_average_trend_fig(

            ALL_PUMPS,

            df_history,

            "진동",

            "평균 진동",

            "mm/s",

            "#e8590c"

        )

        st.pyplot(
            fig_vib
        )

    fg3, fg4 = st.columns(2)

    with fg3:

        fig_temp, avg_temp = build_fleet_average_trend_fig(

            ALL_PUMPS,

            df_history,

            "온도",

            "평균 온도",

            "°C",

            "#c62828"

        )

        st.pyplot(
            fig_temp
        )

    with fg4:

        fig_score, avg_score = build_fleet_average_trend_fig(

            ALL_PUMPS,

            df_history,

            "점수",

            "평균 CBM Score",

            "점",

            "#2f9e44"

        )

        st.pyplot(
            fig_score
        )

    render_excel_export_section(

        "fleet_trend",

        "전사트렌드.xlsx",

        lambda: pd.DataFrame(

            {
                "시점": TREND_MONTHS,
                "평균효율(%)": [round(v, 1) for v in avg_eff],
                "평균진동(mm/s)": [round(v, 2) for v in avg_vib],
                "평균온도(°C)": [round(v, 1) for v in avg_temp],
                "평균CBMScore": [round(v, 1) for v in avg_score]
            }

        )

    )


# ============================================================
# 22. ROI
# ============================================================

elif st.session_state.page == "ROI":

    st.markdown(
        """
        <div class="section-title">
        💰 정비효과 및 ROI 분석
        </div>

        <div class="section-caption">
        오버홀 전·후 효율 변화와 실제 부하율을 반영하여
        에너지 절감 및 투자회수 효과를 계산합니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    st.markdown(
        "### 🏆 종합 성과 요약"
    )

    st.caption(

        "이 플랫폼을 도입한 뒤 지금까지 누적된 실제 활동 "
        "지표입니다. 경진대회·보고 자료에 그대로 쓰셔도 됩니다."

    )

    _df_diag_all = read_excel(

        DB_FILE_PATH,
        "진단이력"

    )

    _df_overhaul_all_roi = read_excel(

        OVERHAUL_DB_PATH,
        "오버홀이력"

    )

    _total_tco = 0

    if (

        not _df_overhaul_all_roi.empty

        and "비용" in _df_overhaul_all_roi.columns

    ):

        _total_tco = pd.to_numeric(

            _df_overhaul_all_roi["비용"],

            errors="coerce"

        ).fillna(0).sum()

    _df_vib_all_roi = read_excel(

        VIBRATION_DB_PATH,
        "진동측정이력"

    )

    _ai_predicted_count = 0

    for _p_roi in ALL_PUMPS:

        _pred_roi = predict_failure_date_regression(

            _p_roi,

            _df_vib_all_roi

        )

        if _pred_roi and _pred_roi.get("status") == "predicted":

            _ai_predicted_count += 1

    s1, s2 = st.columns(2)

    s1.metric(

        "관리 설비 수",

        f"{len(ALL_PUMPS)}대"

    )

    s2.metric(

        "누적 정밀진단",

        f"{len(_df_diag_all)}건"

    )

    s3, s4 = st.columns(2)

    s3.metric(

        "AI 조기예측 설비",

        f"{_ai_predicted_count}대"

    )

    s4.metric(

        "누적 오버홀 비용",

        f"{_total_tco/10000:,.0f}만원"

    )

    st.write("---")

    # 결과 카드를 입력값보다 먼저 화면에 보여주기 위해
    # 자리(placeholder)를 미리 만들어 둔다.
    # 실제 값은 아래에서 입력을 받은 뒤 이 자리에 채워 넣는다.

    result_slot = st.container()

    with st.expander(
        "🔧 계산 조건 입력",
        expanded=True
    ):

        c1, c2 = st.columns(2)

        rated_kw = c1.number_input(
            "모터 정격출력(kW)",
            min_value=0.0,
            value=110.0
        )

        load = c2.number_input(
            "평균 부하율(%)",
            min_value=0.0,
            max_value=100.0,
            value=80.0
        )

        c1, c2 = st.columns(2)

        before = c1.number_input(
            "정비 전 효율(%)",
            min_value=0.0,
            max_value=100.0,
            value=73.0
        )

        after = c2.number_input(
            "정비 후 효율(%)",
            min_value=0.0,
            max_value=100.0,
            value=82.0
        )

        c1, c2 = st.columns(2)

        hours = c1.number_input(
            "연간 운전시간",
            min_value=0,
            value=6000
        )

        price = c2.number_input(
            "전력단가(원/kWh)",
            min_value=0,
            value=140
        )

        repair = st.number_input(
            "오버홀 비용(원)",
            min_value=0,
            value=35_000_000
        )

    effective_kw = (

        rated_kw
        *
        load
        /
        100

    )

    saved_kwh = (

        effective_kw
        *
        (
            1 / (before / 100)
            -
            1 / (after / 100)
        )
        *
        hours

    )

    saved_money = (
        saved_kwh
        *
        price
    )

    payback = (

        repair / saved_money

        if saved_money > 0

        else 0

    )

    with result_slot:

        st.markdown(
            f"""
            <div class="kpi-grid" style="
            grid-template-columns:
            repeat(3, 1fr);
            ">

                <div class="kpi-card">
                    <div class="kpi-label">연간 절감전력</div>
                    <div class="kpi-value">{saved_kwh:,.0f}</div>
                    <div class="kpi-sub">kWh / 년</div>
                </div>

                <div class="kpi-card">
                    <div class="kpi-label">연간 절감액</div>
                    <div class="kpi-value">{saved_money:,.0f}</div>
                    <div class="kpi-sub">원 / 년</div>
                </div>

                <div class="kpi-card">
                    <div class="kpi-label">투자회수 기간</div>
                    <div class="kpi-value">{payback:.1f}년</div>
                    <div class="kpi-sub">오버홀 비용 대비</div>
                </div>

            </div>
            """,
            unsafe_allow_html=True
        )

    render_excel_export_section(

        "roi",

        "정비효과_ROI분석.xlsx",

        lambda: pd.DataFrame(

            [

                ["연간 절감전력(kWh)", round(saved_kwh, 0)],
                ["연간 절감액(원)", round(saved_money, 0)],
                ["투자회수 기간(년)", round(payback, 2)],
                ["오버홀 비용(원)", repair]

            ],

            columns=["항목", "값"]

        )

    )


# ============================================================
# 23. KPI
# ============================================================

elif st.session_state.page == "KPI":

    st.markdown(
        """
        <div class="section-title">
        📊 플랫폼 성과관리
        </div>

        <div class="section-caption">
        PoC와 사업화 과정에서 시스템의 실제 성과를 관리합니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    kpis = [

        [
            "설비정보 조회시간 단축률",
            50,
            "%"
        ],

        [
            "QR 스캔 성공률",
            95,
            "%"
        ],

        [
            "설비정보 정확성",
            98,
            "%"
        ],

        [
            "현장 작업자 수용성",
            70,
            "%"
        ]

    ]

    kpi_values = []

    for name, target, unit in kpis:

        value = st.number_input(
            name,
            min_value=0.0,
            value=0.0,
            key=f"kpi_{name}"
        )

        kpi_values.append(

            {
                "지표": name,
                "실적": value,
                "목표": target,
                "단위": unit
            }

        )

        st.progress(

            min(
                value / target
                if target
                else 0,
                1
            )

        )

        st.caption(
            f"목표 : {target} {unit}"
        )

    if is_read_only():

        st.info(
            "🔒 보기 전용 모드에서는 저장할 수 없습니다."
        )

    elif st.button(

        "💾 KPI 실적 저장",

        type="primary",

        use_container_width=True,

        key="kpi_save_btn"

    ):

        save_time = datetime.now().strftime(
            "%Y-%m-%d"
        )

        for row in kpi_values:

            safe_append_row(

                KPI_DB_PATH,

                "KPI실적",

                [
                    save_time,
                    row["지표"],
                    row["목표"],
                    row["실적"],
                    row["단위"],
                    ""
                ]

            )

        st.success(
            "KPI 실적이 저장되었습니다."
        )

    df_kpi_history = read_excel(
        KPI_DB_PATH,
        "KPI실적"
    )

    if not df_kpi_history.empty:

        with st.expander(
            f"📈 저장된 KPI 이력 (총 {len(df_kpi_history)}건)"
        ):

            st.dataframe(

                df_kpi_history.tail(20),

                use_container_width=True,

                hide_index=True

            )

    else:

        st.caption(
            "아직 저장된 KPI 실적 이력이 없습니다. "
            "위 '실적 저장' 버튼을 누르면 여기에 쌓입니다."
        )

    render_excel_export_section(

        "kpi",

        "성과관리_KPI.xlsx",

        lambda: (

            df_kpi_history

            if not df_kpi_history.empty

            else pd.DataFrame(kpi_values)

        )

    )


# ============================================================
# 24. 기술 노하우
# ============================================================

elif st.session_state.page == "노하우":

    st.markdown(
        """
        <div class="section-title">
        💡 기술 노하우 DB
        </div>

        <div class="section-caption">
        현장에서 축적된 정비경험과 결함원인·조치방법을
        조직의 기술자산으로 관리합니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    with st.expander(

        "📋 정비 체크리스트 템플릿 관리"

    ):

        st.caption(

            "설비 종류별로 표준화된 점검항목을 만들어두면, "
            "정밀진단 페이지에서 누구나 빠짐없이 점검할 수 "
            "있습니다."

        )

        _template_names = get_checklist_template_names()

        if _template_names:

            _manage_template = st.selectbox(

                "수정할 템플릿 선택",

                ["(새 템플릿 만들기)"] + _template_names,

                key="knowhow_template_manage_select"

            )

        else:

            _manage_template = "(새 템플릿 만들기)"

        if _manage_template == "(새 템플릿 만들기)":

            _template_name_input = st.text_input(

                "새 템플릿 이름",

                placeholder="예: 원심펌프용",

                key="new_template_name_input"

            )

            _existing_items = []

        else:

            _template_name_input = _manage_template

            _existing_items = get_checklist_template_items(

                _manage_template

            )

        _items_text = st.text_area(

            "점검항목 (한 줄에 하나씩)",

            value="\n".join(_existing_items),

            height=180,

            key="template_items_textarea"

        )

        tcol1, tcol2 = st.columns(2)

        if is_read_only():

            st.info(
                "🔒 보기 전용 모드에서는 저장할 수 없습니다."
            )

        else:

            if tcol1.button(

                "💾 템플릿 저장",

                type="primary",

                use_container_width=True,

                key="save_template_btn"

            ):

                if not _template_name_input.strip():

                    st.error(
                        "템플릿 이름을 입력해주세요."
                    )

                else:

                    _new_items = [

                        line.strip()

                        for line in _items_text.split("\n")

                        if line.strip()

                    ]

                    save_checklist_template(

                        _template_name_input.strip(),

                        _new_items

                    )

                    st.success(

                        f"'{_template_name_input}' 템플릿을 "
                        f"저장했습니다 ({len(_new_items)}개 항목)."

                    )

                    st.rerun()

            if (

                _manage_template != "(새 템플릿 만들기)"

                and tcol2.button(

                    "🗑️ 이 템플릿 삭제",

                    use_container_width=True,

                    key="delete_template_btn"

                )

            ):

                delete_checklist_template(
                    _manage_template
                )

                st.success(
                    f"'{_manage_template}' 템플릿을 삭제했습니다."
                )

                st.rerun()


    df_knowhow = read_excel(
        KNOWHOW_DB_PATH,
        "노하우DB"
    )

    if not df_knowhow.empty:

        fc1, fc2 = st.columns(
            [2, 1]
        )

        search_word = fc1.text_input(

            "🔍 키워드 검색 (현상·원인·해결노하우·모델명)",

            key="knowhow_search"

        )

        category_options = ["전체"] + sorted(

            df_knowhow["분류"].dropna().unique().tolist()

        )

        category_filter = fc2.selectbox(

            "분류 필터",

            category_options,

            key="knowhow_category_filter"

        )

        filtered = df_knowhow.copy()

        if category_filter != "전체":

            filtered = filtered[

                filtered["분류"] == category_filter

            ]

        if search_word:

            mask = (

                filtered["관련모델"].astype(str).str.contains(search_word, case=False, na=False)

                |
                filtered["현상및원인"].astype(str).str.contains(search_word, case=False, na=False)

                |
                filtered["해결노하우"].astype(str).str.contains(search_word, case=False, na=False)

            )

            filtered = filtered[mask]

        st.caption(
            f"총 {len(df_knowhow)}건 중 {len(filtered)}건 표시"
        )

        st.dataframe(
            filtered,
            use_container_width=True,
            hide_index=True
        )

        if not is_read_only() and not filtered.empty:

            with st.expander(
                "🗑️ 노하우 삭제"
            ):

                delete_options = [

                    f'{i} · {row["등록일자"]} · {row["관련모델"]} · {str(row["현상및원인"])[:20]}'

                    for i, row in filtered.iterrows()

                ]

                delete_choice = st.selectbox(

                    "삭제할 항목 선택",

                    delete_options,

                    key="knowhow_delete_select"

                )

                if st.button(
                    "선택 항목 삭제",
                    key="knowhow_delete_btn"
                ):

                    del_idx = int(
                        delete_choice.split(" · ")[0]
                    )

                    with get_lock(KNOWHOW_DB_PATH):

                        wb = load_workbook(
                            KNOWHOW_DB_PATH
                        )

                        ws = wb[
                            "노하우DB"
                        ]

                        headers = [
                            c.value for c in ws[1]
                        ]

                        # 엑셀 행 번호 = 데이터프레임 인덱스 + 2
                        # (1행은 헤더이고, iterrows 인덱스는 0부터 시작)

                        _row_cells = ws[del_idx + 2]

                        _row_dict = {

                            headers[i]: cell.value

                            for i, cell in enumerate(_row_cells)

                            if i < len(headers)

                        }

                        move_to_trash(

                            "노하우",

                            delete_choice,

                            "노하우DB",

                            _row_dict

                        )

                        ws.delete_rows(
                            del_idx + 2
                        )

                        wb.save(
                            KNOWHOW_DB_PATH
                        )

                        wb.close()

                    _read_excel_cached.clear()

                    log_audit(

                        "노하우 삭제",

                        delete_choice

                    )

                    st.success(
                        "삭제되었습니다. 휴지통에서 30일간 복원할 수 있습니다."
                    )

    else:

        st.info(
            "등록된 노하우가 아직 없습니다. "
            "아래 '새로운 기술 노하우 등록'에서 첫 사례를 "
            "남겨보세요."
        )

    st.write("")

    with st.expander(
        "➕ 새로운 기술 노하우 등록"
    ):

        category = st.selectbox(
            "분류",
            [
                "펌프",
                "모터",
                "진동",
                "축정렬",
                "오버홀",
                "누수",
                "기타"
            ],
            disabled=is_read_only()
        )

        model = st.text_input(
            "관련 모델",
            disabled=is_read_only()
        )

        phenomenon = st.text_area(
            "현상 및 원인",
            disabled=is_read_only()
        )

        solution = st.text_area(
            "해결 노하우",
            disabled=is_read_only()
        )

        if is_read_only():

            st.info(
                "🔒 보기 전용 모드에서는 등록할 수 없습니다."
            )

        elif st.button(
            "노하우 저장",
            type="primary"
        ):

            safe_append_row(

                KNOWHOW_DB_PATH,

                "노하우DB",

                [
                    datetime.now().strftime(
                        "%Y-%m-%d"
                    ),

                    category,

                    model,

                    phenomenon,

                    solution,

                    st.session_state.user_name

                ]

            )

            st.success(
                "기술 노하우가 등록되었습니다."
            )

    render_excel_export_section(

        "knowhow",

        "기술노하우DB.xlsx",

        lambda: (

            filtered

            if "filtered" in dir()

            else df_knowhow

        )

    )


# ============================================================
# 25. 보고서
# ============================================================

elif st.session_state.page == "보고서":

    st.markdown(
        """
        <div class="section-title">
        📄 보고서
        </div>

        <div class="section-caption">
        용도에 맞는 보고서 유형을 선택하세요.
        </div>
        """,
        unsafe_allow_html=True
    )

    report_tab1, report_tab2, report_tab3, report_tab4, report_tab5 = st.tabs(
        [
            "📄 설비별 CBM 월간보고서",
            "📈 월간 진동측정 보고서",
            "📐 연간 축정렬 보고서",
            "🔬 연간 진동분석 보고서",
            "⚡ 효율진단 보고서"
        ]
    )

    with report_tab1:

        st.markdown(
            """
            <div class="section-title">
            📄 설비별 월간 보고서
            </div>

            <div class="section-caption">
            설비 하나를 골라 그 달의 점검·오버홀 이력, 점수, 추세 그래프까지
            담은 Word 월간 보고서를 만듭니다.
            </div>
            """,
            unsafe_allow_html=True
        )

        if not ALL_PUMPS:

            st.info(

                "등록된 설비가 없습니다. '데이터 관리' 메뉴에서 "
                "설비를 추가하거나 엑셀을 업로드해주세요."

            )

            st.stop()

        report_pump_name = st.selectbox(

            "설비 선택",

            [
                p["equip"]
                for p in ALL_PUMPS
            ],

            key="report_pump_select"

        )

        report_pump = next(

            p for p in ALL_PUMPS

            if p["equip"] == report_pump_name

        )

        today = datetime.now()

        rc1, rc2 = st.columns(2)

        report_year = rc1.selectbox(

            "보고 연도",

            list(range(today.year - 2, today.year + 1)),

            index=2,

            key="report_year_select"

        )

        report_month = rc2.selectbox(

            "보고 월",

            list(range(1, 13)),

            index=today.month - 1,

            key="report_month_select"

        )

        month_label = f"{report_year}-{report_month:02d}"

        report_result = pump_status(
            report_pump,
            df_history
        )

        st.markdown(

            f"""
            <div class="platform-card">

            <div class="card-title">
            보고서 미리보기 — {report_pump['equip']} · {month_label}
            </div>

            <b>사업장</b> : {report_pump['site']}<br>
            <b>CBM Score</b> : {report_result['점수']}점 ({report_result['등급']}등급)<br>
            <b>상태</b> : {report_result['상태']}<br>
            <b>구성</b> : 표지(문서번호) · 종합소견 · 설비스펙 · 펌프점수(전월비교) ·
            점검이력 · 오버홀내역(작업사진) · 정비권고사항 · 등급판정기준 ·
            추세그래프(진동·효율·온도·운전시간·설비비교) · 용어부록

            </div>
            """,

            unsafe_allow_html=True

        )

        use_ai_commentary_checkbox = st.checkbox(

            "🤖 AI(Claude)가 종합소견을 직접 작성하도록 하기",

            key="use_ai_commentary_checkbox"

        )

        st.caption(

            "체크하면 미리 정해둔 문장 틀 대신, 실제 Claude API가 "
            "이 설비 데이터를 보고 종합소견을 새로 씁니다 "
            "(Streamlit Secrets에 ANTHROPIC_API_KEY 필요, 없으면 "
            "기본 요약으로 자동 대체됩니다)."

        )

        if st.button(

            "📝 Word 월간 보고서 생성",

            type="primary",

            use_container_width=True,

            key="generate_word_report_btn"

        ):

            with st.spinner(
                "보고서를 만드는 중입니다..."
            ):

                docx_bytes = build_pump_monthly_report_docx(

                    report_pump,

                    report_result,

                    month_label,

                    df_history,

                    ALL_PUMPS,

                    use_ai_commentary=use_ai_commentary_checkbox

                )

            st.session_state[
                "_word_report_bytes"
            ] = docx_bytes

            st.session_state[
                "_word_report_filename"
            ] = f"{report_pump['equip']}_{month_label}_월간보고서.docx"

        if "_word_report_bytes" in st.session_state:

            st.download_button(

                "⬇️ Word 보고서 다운로드",

                data=st.session_state["_word_report_bytes"],

                file_name=st.session_state["_word_report_filename"],

                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                ),

                type="primary",

                use_container_width=True,

                key="download_word_report_btn"

            )

            if st.button(

                "🖨️ PDF로 변환해서 받기",

                use_container_width=True,

                key="convert_word_report_pdf_btn"

            ):

                with st.spinner(
                    "PDF로 변환하는 중입니다..."
                ):

                    _pdf_bytes = convert_docx_bytes_to_pdf_bytes(

                        st.session_state["_word_report_bytes"]

                    )

                if _pdf_bytes:

                    st.session_state["_word_report_pdf_bytes"] = _pdf_bytes

                else:

                    st.error(

                        "PDF 변환에 실패했습니다. 서버에 "
                        "LibreOffice가 설치돼 있는지 확인해주세요 "
                        "(packages.txt에 libreoffice-writer 필요)."

                    )

            if "_word_report_pdf_bytes" in st.session_state:

                st.download_button(

                    "⬇️ PDF 보고서 다운로드",

                    data=st.session_state["_word_report_pdf_bytes"],

                    file_name=st.session_state[

                        "_word_report_filename"

                    ].replace(".docx", ".pdf"),

                    mime="application/pdf",

                    use_container_width=True,

                    key="download_word_report_pdf_btn"

                )

        st.write("---")

        st.markdown(
            "##### 📦 전체 설비 일괄 생성"
        )

        st.caption(

            f"{month_label} 기준으로 등록된 설비 {len(ALL_PUMPS)}대의 "
            "월간 보고서를 한 번에 만들어 zip으로 받습니다. "
            "설비 수가 많으면 시간이 다소 걸릴 수 있습니다."

        )

        if st.button(

            "📦 전체 설비 일괄 생성 (zip)",

            use_container_width=True,

            key="generate_all_reports_btn"

        ):

            with st.spinner(

                f"{len(ALL_PUMPS)}대 설비 보고서를 만드는 중입니다..."

            ):

                zip_buffer = io.BytesIO()

                with zipfile.ZipFile(

                    zip_buffer,

                    "w",

                    zipfile.ZIP_DEFLATED

                ) as zf:

                    progress = st.progress(
                        0.0
                    )

                    for i, one_pump in enumerate(
                        ALL_PUMPS
                    ):

                        one_result = pump_status(

                            one_pump,

                            df_history

                        )

                        one_bytes = build_pump_monthly_report_docx(

                            one_pump,

                            one_result,

                            month_label,

                            df_history,

                            ALL_PUMPS

                        )

                        zf.writestr(

                            f"{one_pump['equip']}_{month_label}_월간보고서.docx",

                            one_bytes

                        )

                        progress.progress(

                            (i + 1) / len(ALL_PUMPS)

                        )

            st.session_state[
                "_all_reports_zip_bytes"
            ] = zip_buffer.getvalue()

            st.session_state[
                "_all_reports_zip_filename"
            ] = f"전체설비_{month_label}_월간보고서.zip"

        if "_all_reports_zip_bytes" in st.session_state:

            st.download_button(

                "⬇️ 전체 보고서 zip 다운로드",

                data=st.session_state["_all_reports_zip_bytes"],

                file_name=st.session_state["_all_reports_zip_filename"],

                mime="application/zip",

                type="primary",

                use_container_width=True,

                key="download_all_reports_zip_btn"

            )

        render_excel_export_section(

            f"report_{report_pump['equip']}_{month_label}",

            f"{report_pump['equip']}_{month_label}_보고서데이터.xlsx",

            lambda: build_equipment_export_sheets(

                report_pump,

                report_result,

                df_history

            )

        )

    with report_tab2:

        st.markdown(
            """
            <div class="section-title">
            📈 월간 진동측정 보고서
            </div>

            <div class="section-caption">
            매달 받는 "정기점검 실적목록" 엑셀을 그대로 올리면
            누적 DB에 쌓이고, 그걸로 실제 사업소 양식 그대로
            월간 진동측정 분석 보고서(Word)를 만듭니다.
            </div>
            """,
            unsafe_allow_html=True
        )

        st.markdown(
            "#### 1. 이번 달 정기점검 실적목록 업로드"
        )

        vib_upload_file = st.file_uploader(

            "정기점검 실적목록 엑셀",

            type=["xlsx"],

            key="vib_measure_uploader"

        )

        if vib_upload_file is not None:

            try:

                _vib_rows, _vib_skipped, _vib_unmatched = parse_vibration_measure_upload(

                    vib_upload_file.getvalue(),

                    ALL_PUMPS

                )

                st.success(

                    f"{len(_vib_rows)}건의 측정점을 확인했습니다."
                    +
                    (

                        f" ({_vib_skipped}건은 매칭 실패/누락으로 제외)"

                        if _vib_skipped

                        else ""
                    )

                )

                if _vib_unmatched:

                    st.warning(

                        "설비 매칭에 실패한 사업장/호기: "
                        +
                        ", ".join(sorted(_vib_unmatched))

                    )

                with st.expander(

                    f"업로드 데이터 미리보기 ({len(_vib_rows)}건)"

                ):

                    st.dataframe(

                        pd.DataFrame(

                            _vib_rows,

                            columns=[
                                "측정일자", "사업장", "설비명",
                                "펌프모터구분", "부하구분", "측정방향",
                                "측정값", "평가코드내역", "측정인"
                            ]

                        ),

                        use_container_width=True,

                        hide_index=True

                    )

                if is_read_only():

                    st.info(
                        "🔒 보기 전용 모드에서는 반영할 수 없습니다."
                    )

                elif st.button(

                    "📥 진동측정 데이터 반영하기",

                    type="primary",

                    use_container_width=True,

                    key="vib_upload_apply_btn"

                ):

                    apply_vibration_measure_upload(
                        _vib_rows
                    )

                    st.success(

                        f"{len(_vib_rows)}건이 누적 DB에 저장되었습니다."

                    )

            except Exception as e:

                st.error(
                    f"엑셀을 읽는 중 문제가 발생했습니다: {e}"
                )

        st.markdown(
            "#### 1-2. 과거 이력 일괄 업로드 (진동 그래프 엑셀, 선택)"
        )

        st.caption(

            "설비마다 시트 하나씩, 월별로 이미 집계된 값이 담긴 "
            "'진동 그래프' 엑셀이 있으면 여기 올리세요. 정기점검 "
            "실적목록을 매달 올리는 대신, 과거 1년치를 한 번에 "
            "채울 수 있습니다."

        )

        vib_graph_file = st.file_uploader(

            "진동 그래프 엑셀",

            type=["xlsx"],

            key="vib_graph_uploader"

        )

        if vib_graph_file is not None:

            try:

                _graph_rows, _graph_unmatched, _graph_month_counts = parse_vibration_graph_workbook(

                    vib_graph_file.getvalue(),

                    ALL_PUMPS

                )

                st.success(

                    f"{len(_graph_rows)}건(측정점 기준)을 확인했습니다."

                )

                with st.expander(

                    "시트별 인식된 개월 수"

                ):

                    st.dataframe(

                        pd.DataFrame(

                            [

                                {"시트": k, "인식된 개월 수": v}

                                for k, v in _graph_month_counts.items()

                            ]

                        ),

                        use_container_width=True,

                        hide_index=True

                    )

                if _graph_unmatched:

                    st.warning(

                        "설비 매칭에 실패한 시트: "
                        +
                        ", ".join(sorted(_graph_unmatched))

                    )

                if is_read_only():

                    st.info(
                        "🔒 보기 전용 모드에서는 반영할 수 없습니다."
                    )

                elif st.button(

                    "📥 과거 이력 반영하기",

                    type="primary",

                    use_container_width=True,

                    key="vib_graph_apply_btn"

                ):

                    apply_vibration_measure_upload(
                        _graph_rows
                    )

                    st.success(

                        f"{len(_graph_rows)}건이 누적 DB에 저장되었습니다."

                    )

            except Exception as e:

                st.error(
                    f"엑셀을 읽는 중 문제가 발생했습니다: {e}"
                )

        st.markdown(
            "#### 2. 보고서 생성"
        )

        vib_site_list = sorted(

            set(
                p["site"] for p in ALL_PUMPS
            )

        )

        selected_sites = st.multiselect(

            "대상 사업장",

            vib_site_list,

            default=vib_site_list,

            key="vib_report_sites"

        )

        vmc1, vmc2 = st.columns(2)

        vib_today = datetime.now()

        def _sync_vib_work_period():

            # 보고 연/월이 바뀌면 "작업 기간 문구"도 자동으로
            # 그 달의 1일~말일로 맞춰준다. (예전엔 월을 바꿔도
            # 문구가 안 따라와서 8월로 그대로 남아있던 문제가 있었음)

            y = st.session_state["vib_report_year"]

            m = st.session_state["vib_report_month"]

            last_day = calendar.monthrange(
                y,
                m
            )[1]

            st.session_state["vib_work_period"] = (

                f"'{str(y)[2:]}.{m:02d}.01. ~ {last_day}."

            )

        vib_year = vmc1.selectbox(

            "보고 연도",

            list(range(vib_today.year - 2, vib_today.year + 1)),

            index=2,

            key="vib_report_year",

            on_change=_sync_vib_work_period

        )

        vib_month = vmc2.selectbox(

            "보고 월",

            list(range(1, 13)),

            index=vib_today.month - 1,

            key="vib_report_month",

            on_change=_sync_vib_work_period

        )

        vib_month_label = f"{vib_year}-{vib_month:02d}"

        if "vib_work_period" not in st.session_state:

            _last_day_init = calendar.monthrange(
                vib_year,
                vib_month
            )[1]

            st.session_state["vib_work_period"] = (

                f"'{str(vib_year)[2:]}.{vib_month:02d}.01. "
                f"~ {_last_day_init}."

            )

        work_period_text = st.text_input(

            "작업 기간 문구",

            key="vib_work_period"

        )

        work_members_text = st.text_input(

            "작업 인원 문구",

            value="조원기, 김광일, 정현철",

            key="vib_work_members"

        )

        if st.button(

            "📝 월간 진동측정 보고서 생성",

            type="primary",

            use_container_width=True,

            key="generate_vib_report_btn"

        ):

            with st.spinner(
                "보고서를 만드는 중입니다..."
            ):

                vib_docx_bytes = build_vibration_monthly_report_docx(

                    selected_sites,

                    vib_month_label,

                    ALL_PUMPS,

                    work_period_text,

                    work_members_text

                )

            st.session_state["_vib_report_bytes"] = vib_docx_bytes

            st.session_state["_vib_report_filename"] = (

                f"{vib_month_label}_월간진동측정보고서.docx"

            )

        if "_vib_report_bytes" in st.session_state:

            st.download_button(

                "⬇️ Word 보고서 다운로드",

                data=st.session_state["_vib_report_bytes"],

                file_name=st.session_state["_vib_report_filename"],

                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                ),

                type="primary",

                use_container_width=True,

                key="download_vib_report_btn"

            )

            if st.button(

                "🖨️ PDF로 변환해서 받기",

                use_container_width=True,

                key="convert_vib_report_pdf_btn"

            ):

                with st.spinner(
                    "PDF로 변환하는 중입니다..."
                ):

                    _vib_pdf_bytes = convert_docx_bytes_to_pdf_bytes(

                        st.session_state["_vib_report_bytes"]

                    )

                if _vib_pdf_bytes:

                    st.session_state["_vib_report_pdf_bytes"] = _vib_pdf_bytes

                else:

                    st.error(

                        "PDF 변환에 실패했습니다. 서버에 "
                        "LibreOffice가 설치돼 있는지 확인해주세요 "
                        "(packages.txt에 libreoffice-writer 필요)."

                    )

            if "_vib_report_pdf_bytes" in st.session_state:

                st.download_button(

                    "⬇️ PDF 보고서 다운로드",

                    data=st.session_state["_vib_report_pdf_bytes"],

                    file_name=st.session_state[

                        "_vib_report_filename"

                    ].replace(".docx", ".pdf"),

                    mime="application/pdf",

                    use_container_width=True,

                    key="download_vib_report_pdf_btn"

                )

    with report_tab3:

        st.markdown(
            "### 📐 연간 축정렬(얼라인먼트) 보고서"
        )

        st.caption(

            "축정렬 측정기 화면(또는 결과지) 사진 1장을 올리면, "
            "AI가 실제 보고서 양식(개요·측정데이터표·판정기준· "
            "추이그래프)에 맞춰 보고서를 만들어줍니다."

        )

        if not ALL_PUMPS:

            st.info(
                "등록된 설비가 없습니다."
            )

        else:

            _align_pump_sel = st.selectbox(

                "설비 선택",

                [p["equip"] for p in ALL_PUMPS],

                key="align_report_pump_select"

            )

            _align_pump = next(

                p for p in ALL_PUMPS

                if p["equip"] == _align_pump_sel

            )

            _align_image = st.file_uploader(

                "축정렬 측정 결과 사진 (1장)",

                type=["png", "jpg", "jpeg"],

                key="align_image_uploader"

            )

            if _align_image:

                st.image(

                    _align_image,

                    caption="업로드한 축정렬 측정 사진",

                    width=400

                )

            with st.expander(

                "✏️ AI 대신 직접 값 입력하기 "
                "(API 키가 없을 때)"

            ):

                st.caption(

                    "측정기 화면을 보고 값을 직접 입력하시면, "
                    "AI 분석 없이도 똑같은 양식의 보고서가 "
                    "만들어집니다."

                )

                st.markdown(

                    "**📊 엑셀로 한번에 채우기**"

                )

                st.download_button(

                    "📄 축정렬 입력 양식 다운로드",

                    data=build_align_excel_template(),

                    file_name="축정렬_입력양식.xlsx",

                    key="align_excel_template_download"

                )

                _align_excel_file = st.file_uploader(

                    "채운 엑셀 업로드",

                    type=["xlsx"],

                    key="align_excel_uploader"

                )

                if _align_excel_file and st.button(

                    "엑셀 값 반영하기",

                    key="align_excel_apply_btn"

                ):

                    try:

                        _align_from_excel = parse_align_excel(

                            _align_excel_file.getvalue()

                        )

                        _align_from_excel["점검일자"] = (

                            datetime.now().strftime("%Y-%m-%d")

                        )

                        st.session_state[
                            "_align_ai_data"
                        ] = _align_from_excel

                        st.session_state["_align_image_bytes"] = (

                            _align_image.getvalue()

                            if _align_image

                            else None

                        )

                        st.success(

                            "엑셀 값이 반영됐습니다. 아래에서 "
                            "보고서를 생성하세요."

                        )

                    except Exception as e:

                        st.error(
                            f"엑셀을 읽는 중 오류: {e}"
                        )

                st.markdown("---")

                st.markdown(

                    "**✏️ 또는 여기에 직접 입력**"

                )

                mc1, mc2 = st.columns(2)

                _m_angle_h = mc1.text_input(

                    "각도오차 수평 (mm/100mm)",

                    key="manual_align_angle_h"

                )

                _m_angle_v = mc2.text_input(

                    "각도오차 수직 (mm/100mm)",

                    key="manual_align_angle_v"

                )

                mc1, mc2 = st.columns(2)

                _m_offset_h = mc1.text_input(

                    "옵셋오차 수평 (mm)",

                    key="manual_align_offset_h"

                )

                _m_offset_v = mc2.text_input(

                    "옵셋오차 수직 (mm)",

                    key="manual_align_offset_v"

                )

                mc1, mc2 = st.columns(2)

                _m_front_h = mc1.text_input(

                    "펌프전면변위 수평 (mm)",

                    key="manual_align_front_h"

                )

                _m_front_v = mc2.text_input(

                    "펌프전면변위 수직 (mm)",

                    key="manual_align_front_v"

                )

                mc1, mc2 = st.columns(2)

                _m_back_h = mc1.text_input(

                    "펌프후면변위 수평 (mm)",

                    key="manual_align_back_h"

                )

                _m_back_v = mc2.text_input(

                    "펌프후면변위 수직 (mm)",

                    key="manual_align_back_v"

                )

                mc1, mc2 = st.columns(2)

                _m_judge_h = mc1.selectbox(

                    "판정 수평",

                    ["양호", "보통", "불량"],

                    key="manual_align_judge_h"

                )

                _m_judge_v = mc2.selectbox(

                    "판정 수직",

                    ["양호", "보통", "불량"],

                    key="manual_align_judge_v"

                )

                _m_equip_name = st.text_input(

                    "측정장비명 (선택)",

                    key="manual_align_equip"

                )

                _m_opinion = st.text_area(

                    "종합의견",

                    placeholder=(

                        "예: 전체 정렬 상태는 ISO/ANSI 회전기계 "
                        "정렬 허용 기준 내에 있으며 즉각적인 "
                        "재정렬이 필요한 수준은 아님"

                    ),

                    key="manual_align_opinion"

                )

                if st.button(

                    "이 값으로 보고서용 데이터 만들기",

                    key="manual_align_apply_btn"

                ):

                    st.session_state["_align_ai_data"] = {

                        "점검일자": datetime.now().strftime(
                            "%Y-%m-%d"
                        ),
                        "측정장비": _m_equip_name,
                        "각도오차_수평": _m_angle_h,
                        "각도오차_수직": _m_angle_v,
                        "옵셋오차_수평": _m_offset_h,
                        "옵셋오차_수직": _m_offset_v,
                        "전면변위_수평": _m_front_h,
                        "전면변위_수직": _m_front_v,
                        "후면변위_수평": _m_back_h,
                        "후면변위_수직": _m_back_v,
                        "판정_수평": _m_judge_h,
                        "판정_수직": _m_judge_v,
                        "종합의견": _m_opinion

                    }

                    st.session_state["_align_image_bytes"] = (

                        _align_image.getvalue()

                        if _align_image

                        else None

                    )

                    st.success(

                        "직접 입력한 값으로 준비됐습니다. "
                        "아래에서 보고서를 생성하세요."

                    )

            if st.button(

                "📷 사진에서 자동 읽기 (OCR, API 불필요)",

                key="align_ocr_btn",

                disabled=(_align_image is None)

            ):

                _align_img_bytes_for_ocr = _align_image.getvalue()

                with st.spinner(
                    "사진 속 숫자를 읽는 중..."
                ):

                    _ocr_result, _ocr_err = ocr_extract_alignment_values(

                        _align_img_bytes_for_ocr

                    )

                if _ocr_err:

                    st.error(
                        _ocr_err
                    )

                else:

                    st.session_state["_align_ocr_result"] = _ocr_result

                    st.session_state["_align_image_bytes"] = (

                        _align_img_bytes_for_ocr

                    )

                    st.success(

                        "숫자를 읽었습니다. 아래에서 맞는지 "
                        "확인·수정한 뒤 확정해주세요."

                    )

            if "_align_ocr_result" in st.session_state:

                _ocr = st.session_state["_align_ocr_result"]

                st.caption(
                    "🔍 OCR이 읽어낸 값 (틀린 부분은 고쳐주세요)"
                )

                oc1, oc2 = st.columns(2)

                _ocr_angle_v = oc1.text_input(
                    "각도오차 수직", value=_ocr.get("각도오차_수직", ""),
                    key="ocr_angle_v"
                )

                _ocr_angle_h = oc2.text_input(
                    "각도오차 수평", value=_ocr.get("각도오차_수평", ""),
                    key="ocr_angle_h"
                )

                oc1, oc2 = st.columns(2)

                _ocr_offset_v = oc1.text_input(
                    "옵셋오차 수직", value=_ocr.get("옵셋오차_수직", ""),
                    key="ocr_offset_v"
                )

                _ocr_offset_h = oc2.text_input(
                    "옵셋오차 수평", value=_ocr.get("옵셋오차_수평", ""),
                    key="ocr_offset_h"
                )

                oc1, oc2 = st.columns(2)

                _ocr_front_v = oc1.text_input(
                    "전면변위 수직", value=_ocr.get("전면변위_수직", ""),
                    key="ocr_front_v"
                )

                _ocr_front_h = oc2.text_input(
                    "전면변위 수평", value=_ocr.get("전면변위_수평", ""),
                    key="ocr_front_h"
                )

                oc1, oc2 = st.columns(2)

                _ocr_back_v = oc1.text_input(
                    "후면변위 수직", value=_ocr.get("후면변위_수직", ""),
                    key="ocr_back_v"
                )

                _ocr_back_h = oc2.text_input(
                    "후면변위 수평", value=_ocr.get("후면변위_수평", ""),
                    key="ocr_back_h"
                )

                oc1, oc2 = st.columns(2)

                _ocr_judge_v = oc1.selectbox(
                    "판정 수직", ["양호", "보통", "불량"],
                    key="ocr_judge_v"
                )

                _ocr_judge_h = oc2.selectbox(
                    "판정 수평", ["양호", "보통", "불량"],
                    key="ocr_judge_h"
                )

                _ocr_opinion = st.text_area(

                    "종합의견",

                    placeholder=(

                        "예: 전체 정렬 상태는 ISO/ANSI 회전기계 "
                        "정렬 허용 기준 내에 있으며 즉각적인 "
                        "재정렬이 필요한 수준은 아님"

                    ),

                    key="ocr_opinion"

                )

                if st.button(

                    "이 값으로 확정하기",

                    key="align_ocr_confirm_btn",

                    type="primary"

                ):

                    st.session_state["_align_ai_data"] = {

                        "점검일자": datetime.now().strftime(
                            "%Y-%m-%d"
                        ),
                        "측정장비": "",
                        "각도오차_수직": _ocr_angle_v,
                        "각도오차_수평": _ocr_angle_h,
                        "옵셋오차_수직": _ocr_offset_v,
                        "옵셋오차_수평": _ocr_offset_h,
                        "전면변위_수직": _ocr_front_v,
                        "전면변위_수평": _ocr_front_h,
                        "후면변위_수직": _ocr_back_v,
                        "후면변위_수평": _ocr_back_h,
                        "판정_수직": _ocr_judge_v,
                        "판정_수평": _ocr_judge_h,
                        "종합의견": _ocr_opinion

                    }

                    st.success(

                        "확정됐습니다. 아래에서 보고서를 "
                        "생성하세요."

                    )

            st.markdown("---")

            if st.button(

                "🤖 AI로 분석하기",

                key="align_ai_analyze_btn",

                type="primary",

                disabled=(_align_image is None)

            ):

                _align_img_bytes = _align_image.getvalue()

                with st.status(

                    "축정렬 사진을 분석하는 중...",

                    expanded=True

                ) as _align_status:

                    st.write("1/2 · 사진 속 수치 읽는 중...")

                    _align_raw, _align_err = call_claude_vision_analysis(

                        [(
                            "축정렬 측정 결과",
                            _align_img_bytes,
                            guess_image_mime(
                                _align_image.name
                            )
                        )],

                        "이 사진은 펌프-모터 축정렬(얼라인먼트) "
                        "측정기 화면 또는 결과지입니다. 다음 JSON "
                        "형식으로만 답하세요. 다른 설명 문장은 "
                        "붙이지 마세요.\n\n"
                        "{\n"
                        '  "점검일자": "화면에 보이는 날짜, 없으면 빈문자열",\n'
                        '  "측정장비": "화면에 보이는 장비명, 없으면 빈문자열",\n'
                        '  "각도오차_수평": "mm/100mm 단위 숫자와 부호",\n'
                        '  "각도오차_수직": "mm/100mm 단위 숫자와 부호",\n'
                        '  "옵셋오차_수평": "mm 단위 숫자와 부호",\n'
                        '  "옵셋오차_수직": "mm 단위 숫자와 부호",\n'
                        '  "전면변위_수평": "mm 단위 숫자와 부호",\n'
                        '  "전면변위_수직": "mm 단위 숫자와 부호",\n'
                        '  "후면변위_수평": "mm 단위 숫자와 부호",\n'
                        '  "후면변위_수직": "mm 단위 숫자와 부호",\n'
                        '  "판정_수평": "양호 또는 보통 또는 불량",\n'
                        '  "판정_수직": "양호 또는 보통 또는 불량",\n'
                        '  "거리_커플링센서": "펌프모터 커플링 중심↔센서 '
                        '거리(mm), 화면에 없으면 빈문자열",\n'
                        '  "거리_센서센서": "센서↔센서 거리(mm), 없으면 빈문자열",\n'
                        '  "거리_모터전센서": "모터 전Base중심→센서 '
                        '거리(mm), 없으면 빈문자열",\n'
                        '  "거리_모터전후": "모터 전Base중심↔모터 후Base중심 '
                        '거리(mm), 없으면 빈문자열",\n'
                        '  "종합의견": "이 설비의 정렬상태에 대한 '
                        '종합 소견 2~3문장. ISO/ANSI 회전기계 정렬 '
                        '허용기준과 비교하여 재정렬 필요 여부와 '
                        '운전 가능 여부를 판단해서 서술"\n'
                        "}"

                    )

                    if _align_err:

                        _align_status.update(

                            label="분석 실패",

                            state="error"

                        )

                    else:

                        _align_data = extract_json_from_ai_response(

                            _align_raw

                        )

                        if not _align_data:

                            _align_err = (

                                "AI 응답을 표 형식으로 해석하지 "
                                "못했습니다. 다시 시도해주세요.\n\n"
                                "AI가 실제로 보낸 답변(참고용):\n"
                                + _align_raw[-800:]

                            )

                            _align_status.update(

                                label="분석 실패",

                                state="error"

                            )

                        else:

                            st.write("2/2 · 분석 완료")

                            _align_status.update(

                                label="분석 완료",

                                state="complete"

                            )

                if _align_err:

                    st.error(
                        _align_err
                    )

                else:

                    st.session_state[
                        "_align_ai_data"
                    ] = _align_data

                    st.session_state[
                        "_align_image_bytes"
                    ] = _align_img_bytes

                    # 이력DB에 구조화 저장 (다음부터 추이그래프에 쌓임)

                    _today_str = datetime.now().strftime(
                        "%Y-%m-%d"
                    )

                    def _to_float(_v):

                        try:

                            return float(

                                str(_v).replace(
                                    "mm/100mm", ""
                                ).replace(
                                    "mm", ""
                                ).strip()

                            )

                        except Exception:

                            return 0.0

                    safe_append_row(

                        ALIGNMENT_DB_PATH,

                        "축정렬이력",

                        [
                            _today_str,
                            _align_pump["site"],
                            _align_pump["equip"],
                            st.session_state.user_name,
                            "",
                            abs(_to_float(
                                _align_data.get("옵셋오차_수직")
                            )),
                            abs(_to_float(
                                _align_data.get("옵셋오차_수평")
                            )),
                            abs(_to_float(
                                _align_data.get("각도오차_수직")
                            )),
                            abs(_to_float(
                                _align_data.get("각도오차_수평")
                            )),
                            _align_data.get("판정_수평", ""),
                            _align_data.get("종합의견", "")
                        ]

                    )

            if "_align_ai_data" in st.session_state:

                st.markdown(
                    "##### 🤖 AI 분석 결과"
                )

                st.json(
                    st.session_state["_align_ai_data"]
                )

                _align_date_label = datetime.now().strftime(
                    "%y.%m.%d"
                )

                if st.button(

                    "📄 축정렬 보고서 생성",

                    key="align_report_gen_btn",

                    type="primary"

                ):

                    # 이력DB에서 이 설비의 추이 데이터 조회

                    _df_align_hist = read_excel(

                        ALIGNMENT_DB_PATH,

                        "축정렬이력"

                    )

                    _history_rows = []

                    if (

                        not _df_align_hist.empty

                        and "설비명" in _df_align_hist.columns

                    ):

                        _hist_pump = _df_align_hist[

                            _df_align_hist["설비명"]
                            == _align_pump["equip"]

                        ]

                        for _, _r in _hist_pump.iterrows():

                            try:

                                _history_rows.append((

                                    str(_r.get("측정일자", "")),

                                    float(
                                        _r.get(
                                            "수직오프셋(mm)", 0
                                        ) or 0
                                    ),

                                    float(
                                        _r.get(
                                            "수평오프셋(mm)", 0
                                        ) or 0
                                    )

                                ))

                            except Exception:

                                continue

                    _align_docx = build_alignment_report_docx(

                        _align_pump,

                        _align_date_label,

                        st.session_state["_align_ai_data"],

                        st.session_state.get(
                            "_align_image_bytes"
                        ),

                        _history_rows

                    )

                    st.session_state["_align_report_bytes"] = _align_docx

                if "_align_report_bytes" in st.session_state:

                    st.download_button(

                        "⬇️ 축정렬 보고서 다운로드",

                        data=st.session_state["_align_report_bytes"],

                        file_name=(

                            f"{_align_pump['equip']}_축정렬보고서_"
                            f"{_align_date_label}.docx"

                        ),

                        use_container_width=True,

                        key="align_report_download_btn"

                    )

                    if st.button(

                        "🖨️ PDF로 변환해서 받기",

                        use_container_width=True,

                        key="convert_align_report_pdf_btn"

                    ):

                        with st.spinner(
                            "PDF로 변환하는 중입니다..."
                        ):

                            _align_pdf_bytes = convert_docx_bytes_to_pdf_bytes(

                                st.session_state["_align_report_bytes"]

                            )

                        if _align_pdf_bytes:

                            st.session_state[
                                "_align_report_pdf_bytes"
                            ] = _align_pdf_bytes

                        else:

                            st.error(

                                "PDF 변환에 실패했습니다. 서버에 "
                                "LibreOffice가 설치돼 있는지 확인해주세요."

                            )

                    if "_align_report_pdf_bytes" in st.session_state:

                        st.download_button(

                            "⬇️ PDF 보고서 다운로드",

                            data=st.session_state[
                                "_align_report_pdf_bytes"
                            ],

                            file_name=(

                                f"{_align_pump['equip']}_축정렬보고서_"
                                f"{_align_date_label}.pdf"

                            ),

                            mime="application/pdf",

                            use_container_width=True,

                            key="download_align_report_pdf_btn"

                        )

    with report_tab4:

        st.markdown(
            "### 🔬 연간 진동분석(스펙트럼) 보고서"
        )

        st.caption(

            "모터·펌프 각각 8장씩(스펙트럼 2종·시간파형·설비정보· "
            "주파수분석), 총 16장을 올리면 AI가 전부 읽고 종합분석 "
            "해서 보고서를 만들어줍니다. 없는 항목은 비워두셔도 "
            "됩니다."

        )

        VA_ATTACH_TYPES = [

            "반부하 3200Hz 수직수평축 스펙트럼 그래프",
            "부하 3200Hz 수직수평 스펙트럼 그래프",
            "반부하 320Hz 수직수평축 스펙트럼 그래프",
            "부하 320Hz 수직수평 스펙트럼 그래프",
            "반부하 수직수평축 시간파형 그래프",
            "부하 수직수평 시간파형 그래프",
            "설비정보",
            "주파수 분석"

        ]

        if not ALL_PUMPS:

            st.info(
                "등록된 설비가 없습니다."
            )

        else:

            _va_pump_sel = st.selectbox(

                "설비 선택",

                [p["equip"] for p in ALL_PUMPS],

                key="va_report_pump_select"

            )

            _va_pump = next(

                p for p in ALL_PUMPS

                if p["equip"] == _va_pump_sel

            )

            _va_uploaded_images = {}

            _va_col_motor, _va_col_pump = st.columns(2)

            with _va_col_motor:

                st.markdown("**⚙️ 모터 부분**")

                for _idx, _atype in enumerate(VA_ATTACH_TYPES):

                    _label = f"모터 {_atype}"

                    _va_file = st.file_uploader(

                        f"(첨부{_idx+1}) {_label}",

                        type=["png", "jpg", "jpeg"],

                        key=f"va_upload_motor_{_idx}"

                    )

                    if _va_file:

                        _va_uploaded_images[_label] = _va_file

            with _va_col_pump:

                st.markdown("**🔧 펌프 부분**")

                for _idx, _atype in enumerate(VA_ATTACH_TYPES):

                    _label = f"펌프 {_atype}"

                    _va_file = st.file_uploader(

                        f"(첨부{_idx+1}) {_label}",

                        type=["png", "jpg", "jpeg"],

                        key=f"va_upload_pump_{_idx}"

                    )

                    if _va_file:

                        _va_uploaded_images[_label] = _va_file

            _va_uploaded_count = len(_va_uploaded_images)

            st.caption(

                f"현재 {_va_uploaded_count}/16개 첨부가 "
                "올라와 있습니다."

            )

            if st.button(

                "📷 업로드된 사진에서 숫자 읽기 (OCR, API 불필요)",

                key="va_ocr_btn",

                disabled=(_va_uploaded_count == 0)

            ):

                with st.spinner(
                    "사진들 속 숫자를 읽는 중..."
                ):

                    _va_ocr_lines = []

                    for _label, _file in _va_uploaded_images.items():

                        _v_result, _v_err = ocr_extract_vibration_numbers(

                            _file.getvalue()

                        )

                        if _v_err or not _v_result:

                            continue

                        _overall = _v_result.get("총진동_후보", [])

                        _peak = _v_result.get("최대진폭_후보", [])

                        _freq = _v_result.get("주파수_후보", [])

                        if _overall or _peak or _freq:

                            _va_ocr_lines.append(

                                f"**{_label}**  ·  "
                                f"총진동 후보: {_overall or '없음'}  ·  "
                                f"최대진폭 후보: {_peak or '없음'}  ·  "
                                f"주파수 후보: {_freq or '없음'}"

                            )

                st.session_state["_va_ocr_lines"] = _va_ocr_lines

            if st.session_state.get("_va_ocr_lines"):

                st.caption(

                    "🔍 사진에서 읽은 숫자입니다. 그래프 곡선 "
                    "해석(결함 여부)은 자동으로 안 되니, 이 숫자를 "
                    "보고 아래 표에 직접 옮겨 적어주세요."

                )

                for _line in st.session_state["_va_ocr_lines"]:

                    st.markdown(_line)

                if not st.session_state["_va_ocr_lines"]:

                    st.info(

                        "사진에서 숫자를 찾지 못했습니다. 사진이 "
                        "너무 흐리거나 해상도가 낮을 수 있습니다."

                    )

            with st.expander(

                "✏️ AI 대신 직접 값 입력하기 "
                "(API 키가 없을 때)"

            ):

                st.caption(

                    "10개 측정점의 값을 표에서 직접 입력하시면, "
                    "AI 분석 없이도 요약표가 채워진 보고서가 "
                    "만들어집니다. (그룹별 상세 결함판정표·시간파형표는 "
                    "AI 분석에서만 자동으로 채워집니다.)"

                )

                st.markdown(

                    "**📊 엑셀로 한번에 채우기**"

                )

                st.download_button(

                    "📄 진동분석 입력 양식 다운로드",

                    data=build_va_excel_template(),

                    file_name="진동분석_입력양식.xlsx",

                    key="va_excel_template_download"

                )

                _va_excel_file = st.file_uploader(

                    "채운 엑셀 업로드",

                    type=["xlsx"],

                    key="va_excel_uploader"

                )

                _va_opinion_for_excel = st.text_input(

                    "종합의견 (엑셀 업로드시 같이 반영)",

                    key="va_excel_opinion"

                )

                if _va_excel_file and st.button(

                    "엑셀 값 반영하기",

                    key="va_excel_apply_btn"

                ):

                    try:

                        _va_points_from_excel = parse_va_excel(

                            _va_excel_file.getvalue()

                        )

                        _va_groups_from_excel = {

                            _gname: {

                                "세부의견": [],
                                "defect_table": [],
                                "waveform_table": [],
                                "해석": []

                            }

                            for _gname in VA_GROUP_ORDER

                        }

                        st.session_state["_va_ai_data"] = {

                            "측정장비": "",
                            "종합의견": _va_opinion_for_excel,
                            "points": _va_points_from_excel,
                            "freq_table": {},
                            "groups": _va_groups_from_excel,
                            "종합분석": [],
                            "향후대책": []

                        }

                        st.session_state["_va_images_bytes"] = {

                            _label: _file.getvalue()

                            for _label, _file
                            in _va_uploaded_images.items()

                        }

                        st.success(

                            "엑셀 값이 반영됐습니다. 아래에서 "
                            "보고서를 생성하세요."

                        )

                    except Exception as e:

                        st.error(
                            f"엑셀을 읽는 중 오류: {e}"
                        )

                st.markdown("---")

                st.markdown(

                    "**✏️ 또는 여기 표에 직접 입력**"

                )

                _va_manual_rows = [

                    {"측정점": "모터 부하 수직", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},
                    {"측정점": "모터 부하 수평", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},
                    {"측정점": "모터 반부하 수직", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},
                    {"측정점": "모터 반부하 수평", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},
                    {"측정점": "모터 반부하 축", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},
                    {"측정점": "펌프 부하 수직", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},
                    {"측정점": "펌프 부하 수평", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},
                    {"측정점": "펌프 반부하 수직", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},
                    {"측정점": "펌프 반부하 수평", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},
                    {"측정점": "펌프 반부하 축", "Overall(mm/s)": 0.0, "판정": "A(양호)", "pk-pk(µm)": 0.0, "주파수": "", "키워드": ""},

                ]

                _va_manual_df = st.data_editor(

                    pd.DataFrame(_va_manual_rows),

                    hide_index=True,

                    use_container_width=True,

                    key="manual_va_data_editor",

                    disabled=["측정점"]

                )

                _m_va_opinion = st.text_area(

                    "종합의견",

                    placeholder=(

                        "예: 펌프와 모터 모두 진동이 허용 기준 "
                        "안쪽이라 운전에 문제 없음."

                    ),

                    key="manual_va_opinion"

                )

                if st.button(

                    "이 값으로 보고서용 데이터 만들기",

                    key="manual_va_apply_btn"

                ):

                    _m_points = {}

                    for _, _r in _va_manual_df.iterrows():

                        _m_points[_r["측정점"]] = {

                            "overall_rms": _r["Overall(mm/s)"],
                            "판정": _r["판정"],
                            "pk_pk": _r["pk-pk(µm)"],
                            "주파수": _r["주파수"],
                            "키워드": _r["키워드"]

                        }

                    _m_groups = {}

                    for _gname in VA_GROUP_ORDER:

                        _m_groups[_gname] = {

                            "세부의견": [],
                            "defect_table": [],
                            "waveform_table": [],
                            "해석": []

                        }

                    st.session_state["_va_ai_data"] = {

                        "측정장비": "",
                        "종합의견": _m_va_opinion,
                        "points": _m_points,
                        "freq_table": {},
                        "groups": _m_groups,
                        "종합분석": [],
                        "향후대책": []

                    }

                    st.session_state["_va_images_bytes"] = {

                        _label: _file.getvalue()

                        for _label, _file in _va_uploaded_images.items()

                    }

                    st.success(

                        "직접 입력한 값으로 준비됐습니다. "
                        "아래에서 보고서를 생성하세요."

                    )

            if st.button(

                "🤖 AI로 분석하기",

                key="va_ai_analyze_btn",

                type="primary",

                disabled=(_va_uploaded_count == 0)

            ):

                _va_images_with_labels = [

                    (
                        _label,
                        _file.getvalue(),
                        guess_image_mime(_file.name)
                    )

                    for _label, _file in _va_uploaded_images.items()

                ]

                with st.status(

                    "진동 스펙트럼을 분석하는 중...",

                    expanded=True

                ) as _va_status:

                    st.write(

                        f"1/2 · {_va_uploaded_count}개 사진 "
                        "판독 중..."

                    )

                    _va_result, _va_err = call_claude_vision_analysis(

                        _va_images_with_labels,

                        "이 사진들은 펌프·모터 설비의 진동분석 "
                        "측정자료입니다. 각 사진 앞에 라벨이 붙어 "
                        "있는데, '반부하 3200Hz/320Hz 수직수평축 "
                        "스펙트럼'과 '부하 3200Hz/320Hz 수직수평 "
                        "스펙트럼' 사진은 한 장 안에 수직·수평·"
                        "(반부하의 경우)축 방향 그래프가 함께 들어 "
                        "있습니다(부하 상태는 축방향 측정이 없습니다). "
                        "'시간파형' 사진도 마찬가지로 여러 방향이 "
                        "한 장에 들어 있습니다. '설비정보' 사진은 "
                        "설비 제원 화면이고, '주파수 분석' 사진은 "
                        "불평형·축정렬불량·전기적불평형·베인이상·"
                        "베어링 결함주파수와 그 배수(nX)가 표시된 "
                        "설정창입니다.\n\n"
                        "다음 JSON 형식으로만 답하세요. 다른 설명 "
                        "문장은 붙이지 마세요.\n\n"
                        "{\n"
                        '  "측정장비": "사진에 보이면 장비명",\n'
                        '  "종합의견": "펌프와 모터 전체 진동상태 종합판정 1~2문장",\n'
                        '  "points": {\n'
                        '    "모터 부하 수직": {"overall_rms": 숫자, "판정":"A(양호) 등", "pk_pk":숫자, "주파수":"...", "키워드":"..."},\n'
                        '    "모터 부하 수평": {...}, "모터 반부하 수직": {...}, '
                        '"모터 반부하 수평": {...}, "모터 반부하 축": {...},\n'
                        '    "펌프 부하 수직": {...}, "펌프 부하 수평": {...}, '
                        '"펌프 반부하 수직": {...}, "펌프 반부하 수평": {...}, '
                        '"펌프 반부하 축": {...}\n'
                        "    (스펙트럼 사진 안의 각 방향 곡선을 읽어서 "
                        "10개 항목 모두 채워주세요. 사진이 없으면 "
                        "빈 값으로 두세요)\n"
                        "  },\n"
                        '  "freq_table": {\n'
                        '    "모터": [ {"구분":"불평형","주파수":"29.5","배수":"1X"}, '
                        '{"구분":"축정렬불량","주파수":"...","배수":"2X/3X~"}, '
                        '{"구분":"전기적불평형","주파수":"...","배수":"-"}, '
                        '{"구분":"베인이상","주파수":"...","배수":"5X"}, '
                        '{"구분":"베어링 결함","주파수":"...","배수":"..."} ],\n'
                        '    "펌프": [ 위와 동일 구조, 각 설비의 "주파수 분석" 사진에서 읽은 값 ]\n'
                        "  },\n"
                        '  "groups": {\n'
                        '    "펌프 반부하": {\n'
                        '      "세부의견": ["문장1", "문장2"],\n'
                        '      "defect_table": [\n'
                        '        {"item":"불평형","result":"X 또는 O","reason":"..."},\n'
                        '        {"item":"축정렬불량","result":"X 또는 O","reason":"..."},\n'
                        '        {"item":"전기적불평형","result":"X 또는 O","reason":"..."},\n'
                        '        {"item":"베인이상","result":"X 또는 O","reason":"..."},\n'
                        '        {"item":"베어링결함","result":"X 또는 O","reason":"..."}\n'
                        "      ],\n"
                        '      "waveform_table": [\n'
                        '        {"방향":"수직","형상특징":"...","시사점":"..."},\n'
                        '        {"방향":"수평","형상특징":"...","시사점":"..."},\n'
                        '        {"방향":"축","형상특징":"...","시사점":"..."}\n'
                        "      ],\n"
                        '      "해석": ["문장1", "문장2"]\n'
                        "    },\n"
                        '    "펌프 부하": { 위와 동일 구조(축 방향 없이 수직·수평만) },\n'
                        '    "모터 부하": { 위와 동일 구조(축 방향 없이 수직·수평만) },\n'
                        '    "모터 반부하": { 위와 동일 구조 }\n'
                        "  },\n"
                        '  "종합분석": ["문장1", "문장2"],\n'
                        '  "향후대책": ["문장1", "문장2"]\n'
                        "}"

                    )

                    if _va_err:

                        _va_status.update(

                            label="분석 실패",

                            state="error"

                        )

                    else:

                        _va_data = extract_json_from_ai_response(

                            _va_result

                        )

                        if not _va_data:

                            _va_err = (

                                "AI 응답을 표 형식으로 해석하지 "
                                "못했습니다. 다시 시도해주세요.\n\n"
                                "AI가 실제로 보낸 답변(참고용):\n"
                                + _va_result[-800:]

                            )

                            _va_status.update(

                                label="분석 실패",

                                state="error"

                            )

                        else:

                            st.write("2/2 · 분석 완료")

                            _va_status.update(

                                label="분석 완료",

                                state="complete"

                            )

                if _va_err:

                    st.error(
                        _va_err
                    )

                else:

                    st.session_state[
                        "_va_ai_data"
                    ] = _va_data

                    st.session_state[
                        "_va_images_bytes"
                    ] = {

                        _label: _file.getvalue()

                        for _label, _file in _va_uploaded_images.items()

                    }

            if "_va_ai_data" in st.session_state:

                st.markdown(
                    "##### 🤖 AI 분석 결과"
                )

                st.json(
                    st.session_state["_va_ai_data"]
                )

                _va_date_label = datetime.now().strftime(
                    "%y. %m. %d."
                )

                if st.button(

                    "📄 진동분석 보고서 생성",

                    key="va_report_gen_btn",

                    type="primary"

                ):

                    _va_docx = build_vib_analysis_report_docx(

                        _va_pump,

                        _va_date_label,

                        st.session_state["_va_ai_data"],

                        st.session_state.get(
                            "_va_images_bytes"
                        )

                    )

                    st.session_state["_va_report_bytes"] = _va_docx

                if "_va_report_bytes" in st.session_state:

                    st.download_button(

                        "⬇️ 진동분석 보고서 다운로드",

                        data=st.session_state["_va_report_bytes"],

                        file_name=(

                            f"{_va_pump['equip']}_진동분석보고서_"
                            f"{_va_date_label}.docx"

                        ),

                        use_container_width=True,

                        key="va_report_download_btn"

                    )

                    if st.button(

                        "🖨️ PDF로 변환해서 받기",

                        use_container_width=True,

                        key="convert_va_report_pdf_btn"

                    ):

                        with st.spinner(
                            "PDF로 변환하는 중입니다..."
                        ):

                            _va_pdf_bytes = convert_docx_bytes_to_pdf_bytes(

                                st.session_state["_va_report_bytes"]

                            )

                        if _va_pdf_bytes:

                            st.session_state[
                                "_va_report_pdf_bytes"
                            ] = _va_pdf_bytes

                        else:

                            st.error(

                                "PDF 변환에 실패했습니다. 서버에 "
                                "LibreOffice가 설치돼 있는지 확인해주세요."

                            )

                    if "_va_report_pdf_bytes" in st.session_state:

                        st.download_button(

                            "⬇️ PDF 보고서 다운로드",

                            data=st.session_state[
                                "_va_report_pdf_bytes"
                            ],

                            file_name=(

                                f"{_va_pump['equip']}_진동분석보고서_"
                                f"{_va_date_label}.pdf"

                            ),

                            mime="application/pdf",

                            use_container_width=True,

                            key="download_va_report_pdf_btn"

                        )

    with report_tab5:

        st.markdown(
            "### ⚡ 효율진단 보고서"
        )

        st.caption(

            "효율진단 데이터(성능곡선, 측정결과지 등) 사진 1장을 "
            "올리면, AI가 정격 대비 실측치를 읽고 원인·개선방향을 "
            "분석해서 보고서를 만들어줍니다."

        )

        if not ALL_PUMPS:

            st.info(
                "등록된 설비가 없습니다."
            )

        else:

            _eff_pump_sel = st.selectbox(

                "설비 선택",

                [p["equip"] for p in ALL_PUMPS],

                key="eff_report_pump_select"

            )

            _eff_pump = next(

                p for p in ALL_PUMPS

                if p["equip"] == _eff_pump_sel

            )

            _eff_image = st.file_uploader(

                "효율진단 데이터 사진 (1장)",

                type=["png", "jpg", "jpeg"],

                key="eff_image_uploader"

            )

            if _eff_image:

                st.image(

                    _eff_image,

                    caption="업로드한 효율진단 사진",

                    width=400

                )

            with st.expander(

                "✏️ AI 대신 직접 값 입력하기 "
                "(API 키가 없을 때)"

            ):

                st.caption(

                    "측정 데이터를 보고 값을 직접 입력하시면, "
                    "AI 분석 없이도 보고서가 만들어집니다."

                )

                st.markdown(

                    "**📊 엑셀로 한번에 채우기**"

                )

                st.download_button(

                    "📄 효율진단 입력 양식 다운로드",

                    data=build_eff_excel_template(),

                    file_name="효율진단_입력양식.xlsx",

                    key="eff_excel_template_download"

                )

                _eff_excel_file = st.file_uploader(

                    "채운 엑셀 업로드",

                    type=["xlsx"],

                    key="eff_excel_uploader"

                )

                if _eff_excel_file and st.button(

                    "엑셀 값 반영하기",

                    key="eff_excel_apply_btn"

                ):

                    try:

                        (
                            _e_rated, _e_measured,
                            _e_flow_rated, _e_flow_measured,
                            _e_opinion
                        ) = parse_eff_excel(

                            _eff_excel_file.getvalue()

                        )

                        _e_summary = (

                            f"정격효율 {_e_rated}%, "
                            f"실측효율 {_e_measured}%\n"
                            f"정격유량 {_e_flow_rated}m3/h, "
                            f"실측유량 {_e_flow_measured}m3/h\n"
                            f"{_e_opinion}"

                        )

                        st.session_state[
                            "_eff_ai_result"
                        ] = _e_summary

                        st.session_state["_eff_image_bytes"] = (

                            _eff_image.getvalue()

                            if _eff_image

                            else None

                        )

                        st.success(

                            "엑셀 값이 반영됐습니다. 아래에서 "
                            "보고서를 생성하세요."

                        )

                    except Exception as e:

                        st.error(
                            f"엑셀을 읽는 중 오류: {e}"
                        )

                st.markdown("---")

                st.markdown(

                    "**✏️ 또는 여기에 직접 입력**"

                )

                ec1, ec2 = st.columns(2)

                _m_eff_rated = ec1.text_input(

                    "정격효율(%)",

                    key="manual_eff_rated"

                )

                _m_eff_measured = ec2.text_input(

                    "실측효율(%)",

                    key="manual_eff_measured"

                )

                ec1, ec2 = st.columns(2)

                _m_flow_rated = ec1.text_input(

                    "정격유량(m3/h)",

                    key="manual_eff_flow_rated"

                )

                _m_flow_measured = ec2.text_input(

                    "실측유량(m3/h)",

                    key="manual_eff_flow_measured"

                )

                _m_eff_opinion = st.text_area(

                    "원인분석 및 개선권고",

                    placeholder=(

                        "예: 정격 대비 실측효율이 9%p 낮음. "
                        "임펠러 마모가 의심되며 정밀점검 및 "
                        "교체를 권고함."

                    ),

                    key="manual_eff_opinion"

                )

                if st.button(

                    "이 값으로 보고서용 데이터 만들기",

                    key="manual_eff_apply_btn"

                ):

                    _m_summary = (

                        f"정격효율 {_m_eff_rated}%, "
                        f"실측효율 {_m_eff_measured}%\n"
                        f"정격유량 {_m_flow_rated}m3/h, "
                        f"실측유량 {_m_flow_measured}m3/h\n"
                        f"{_m_eff_opinion}"

                    )

                    st.session_state[
                        "_eff_ai_result"
                    ] = _m_summary

                    st.session_state["_eff_image_bytes"] = (

                        _eff_image.getvalue()

                        if _eff_image

                        else None

                    )

                    st.success(

                        "직접 입력한 값으로 준비됐습니다. "
                        "아래에서 보고서를 생성하세요."

                    )

            if st.button(

                "🤖 AI로 분석하기",

                key="eff_ai_analyze_btn",

                type="primary",

                disabled=(_eff_image is None)

            ):

                _eff_img_bytes = _eff_image.getvalue()

                with st.status(

                    "효율진단 사진을 분석하는 중...",

                    expanded=True

                ) as _eff_status:

                    st.write("1/2 · 사진 속 수치 읽는 중...")

                    _eff_result, _eff_err = call_claude_vision_analysis(

                        [(
                            "효율진단 측정 결과",
                            _eff_img_bytes,
                            guess_image_mime(
                                _eff_image.name
                            )
                        )],

                        "이 사진은 펌프설비의 효율진단(성능시험) "
                        "결과입니다. 사진에서 정격효율과 실측효율, "
                        "정격유량과 실측유량, 정격양정과 실측양정, "
                        "전동기효율 등 읽을 수 있는 모든 수치를 "
                        "정리해주세요. 정격 대비 실측 효율이 얼마나 "
                        "차이나는지(%p) 계산하고, 효율이 떨어졌다면 "
                        "일반적으로 어떤 원인(임펠러 마모, 웨어링 "
                        "간극 증가, 이물질 등)일 가능성이 있는지, "
                        "개선을 위해 무엇을 해야 하는지 권고해주세요. "
                        "한국어로, 보고서에 바로 쓸 수 있는 간결한 "
                        "문장으로 작성해주세요."

                    )

                    if _eff_err:

                        _eff_status.update(

                            label="분석 실패",

                            state="error"

                        )

                    else:

                        st.write("2/2 · 분석 완료")

                        _eff_status.update(

                            label="분석 완료",

                            state="complete"

                        )

                if _eff_err:

                    st.error(
                        _eff_err
                    )

                else:

                    st.session_state[
                        "_eff_ai_result"
                    ] = _eff_result

                    st.session_state[
                        "_eff_image_bytes"
                    ] = _eff_img_bytes

            if "_eff_ai_result" in st.session_state:

                st.markdown(
                    "##### 🤖 AI 분석 결과"
                )

                st.info(
                    st.session_state["_eff_ai_result"]
                )

                _eff_month = datetime.now().strftime("%Y-%m")

                if st.button(

                    "📄 효율진단 보고서 생성",

                    key="eff_report_gen_btn",

                    type="primary"

                ):

                    _eff_docx = build_efficiency_report_docx(

                        _eff_pump,

                        _eff_month,

                        st.session_state["_eff_ai_result"],

                        st.session_state.get(
                            "_eff_image_bytes"
                        )

                    )

                    st.session_state["_eff_report_bytes"] = _eff_docx

                if "_eff_report_bytes" in st.session_state:

                    st.download_button(

                        "⬇️ 효율진단 보고서 다운로드",

                        data=st.session_state["_eff_report_bytes"],

                        file_name=(

                            f"{_eff_pump['equip']}_효율진단보고서_"
                            f"{_eff_month}.docx"

                        ),

                        use_container_width=True,

                        key="eff_report_download_btn"

                    )

                    if st.button(

                        "🖨️ PDF로 변환해서 받기",

                        use_container_width=True,

                        key="convert_eff_report_pdf_btn"

                    ):

                        with st.spinner(
                            "PDF로 변환하는 중입니다..."
                        ):

                            _eff_pdf_bytes = convert_docx_bytes_to_pdf_bytes(

                                st.session_state["_eff_report_bytes"]

                            )

                        if _eff_pdf_bytes:

                            st.session_state[
                                "_eff_report_pdf_bytes"
                            ] = _eff_pdf_bytes

                        else:

                            st.error(

                                "PDF 변환에 실패했습니다. 서버에 "
                                "LibreOffice가 설치돼 있는지 확인해주세요."

                            )

                    if "_eff_report_pdf_bytes" in st.session_state:

                        st.download_button(

                            "⬇️ PDF 보고서 다운로드",

                            data=st.session_state[
                                "_eff_report_pdf_bytes"
                            ],

                            file_name=(

                                f"{_eff_pump['equip']}_효율진단보고서_"
                                f"{_eff_month}.pdf"

                            ),

                            mime="application/pdf",

                            use_container_width=True,

                            key="download_eff_report_pdf_btn"

                        )



# ============================================================
# 26. 데이터 관리 / QR 비용
# ============================================================

elif st.session_state.page == "백업":

    st.markdown(
        """
        <div class="section-title">
        💾 데이터 및 사업관리
        </div>

        <div class="section-caption">
        설비 DB 백업과 QR 구축·부착 비용,
        플랫폼 운영비용을 관리합니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    st.info(
        "현재 PFM 연동 단계 : 1단계 (수동 입력형) · "
        "독자 DB 구축·운영 중 — "
        "2단계(엑셀 Batch 연계) 및 "
        "3단계(API 연계)는 향후 추진 예정"
    )

    _databack_tab1, _databack_tab2, _databack_tab3, _databack_tab4, _databack_tab5, _databack_tab6 = st.tabs(
        [
            "📥 설비 데이터 관리",
            "📱 QR 관리",
            "🔔 알림 · 보안",
            "💾 백업 · 리포트",
            "💰 비용 참고",
            "🗑️ 휴지통",
        ]
    )

    with _databack_tab1:

        st.markdown(
            "### 🧾 통합 업로드 양식 (설비+정밀진단+오버홀 한번에)"
        )

        st.caption(

            "엑셀 파일 하나에 설비정보·정밀진단·오버홀이력을 "
            "전부 담아서 한 번에 올리는 양식입니다. "
            "빈 양식을 받아서 채운 뒤 그대로 다시 올리시면 됩니다."

        )

        template_bytes = build_unified_import_template_bytes()

        st.download_button(

            "📄 통합 업로드 양식 다운로드",

            data=template_bytes,

            file_name="통합업로드_양식.xlsx",

            mime=(
                "application/vnd.openxmlformats-"
                "officedocument.spreadsheetml.sheet"
            ),

            use_container_width=True,

            key="download_unified_template_btn"

        )

        unified_file = st.file_uploader(

            "작성한 통합 양식 업로드",

            type=["xlsx"],

            key="unified_import_uploader"

        )

        if unified_file is not None:

            try:

                _parsed = parse_unified_import_workbook(

                    unified_file.getvalue()

                )

                st.success(

                    f"설비 {len(_parsed['equip_rows'])}건 · "
                    f"정밀진단 {len(_parsed['diag_rows'])}건 · "
                    f"오버홀 {len(_parsed['overhaul_rows'])}건을 "
                    "확인했습니다."

                )

                if _parsed["diag_skipped"]:

                    st.warning(

                        "정밀진단 시트에서 설비정보와 매칭 안 되어 "
                        "제외된 항목: "
                        +
                        ", ".join(_parsed["diag_skipped"])

                    )

                if _parsed["overhaul_skipped"]:

                    st.warning(

                        f"오버홀이력 시트에서 {_parsed['overhaul_skipped']}건이 "
                        "설비 매칭 실패 또는 필수값 누락으로 제외되었습니다."

                    )

                if _parsed["equip_rows"]:

                    with st.expander(
                        f"설비정보 미리보기 ({len(_parsed['equip_rows'])}건)"
                    ):

                        st.dataframe(

                            pd.DataFrame(_parsed["equip_rows"]),

                            use_container_width=True,

                            hide_index=True

                        )

                if _parsed["diag_rows"]:

                    with st.expander(
                        f"정밀진단 미리보기 ({len(_parsed['diag_rows'])}건)"
                    ):

                        _diag_preview_cols = (

                            ["점검일","사업장","설비명","제조사","모델명",
                             "마력","양정","준공일","점검자","종합점수","최종등급"]

                            +
                            [item[1] for item in EVAL_ITEMS]

                            +
                            ["효율측정값(%)","진동측정값(mm/s)",
                             "온도측정값(°C)","전류측정값(A)"]

                        )

                        st.dataframe(

                            pd.DataFrame(
                                _parsed["diag_rows"],
                                columns=_diag_preview_cols
                            ),

                            use_container_width=True,

                            hide_index=True

                        )

                if _parsed["overhaul_rows"]:

                    with st.expander(
                        f"오버홀이력 미리보기 ({len(_parsed['overhaul_rows'])}건)"
                    ):

                        st.dataframe(

                            pd.DataFrame(

                                _parsed["overhaul_rows"],

                                columns=[
                                    "작업일자","사업장","설비명","공정단계",
                                    "작업자","작업내용","사진파일명",
                                    "전후효율","전후진동"
                                ]

                            ),

                            use_container_width=True,

                            hide_index=True

                        )

                if is_read_only():

                    st.info(
                        "🔒 보기 전용 모드에서는 반영할 수 없습니다."
                    )

                elif st.button(

                    "📥 통합 업로드 반영하기",

                    type="primary",

                    use_container_width=True,

                    key="unified_import_open_btn"

                ):

                    st.session_state["_show_unified_import_confirm"] = True

                if st.session_state.get("_show_unified_import_confirm"):

                    @st.dialog("통합 업로드 반영 확인")
                    def _confirm_unified_import():

                        st.write(

                            f"설비 {len(_parsed['equip_rows'])}건, "
                            f"정밀진단 {len(_parsed['diag_rows'])}건, "
                            f"오버홀 {len(_parsed['overhaul_rows'])}건을 "
                            "반영하시겠습니까?"

                        )

                        st.caption(

                            "이미 등록된 설비명은 건드리지 않고, "
                            "새로운 설비명만 추가됩니다. "
                            "정밀진단·오버홀은 항상 새 기록으로 "
                            "추가됩니다(기존 이력은 지워지지 않습니다)."

                        )

                        uc1, uc2 = st.columns(2)

                        with uc1:

                            if st.button(

                                "예, 반영합니다",

                                type="primary",

                                use_container_width=True,

                                key="unified_import_confirm_yes"

                            ):

                                apply_unified_import(
                                    _parsed
                                )

                                st.session_state[
                                    "_show_unified_import_confirm"
                                ] = False

                                st.session_state[
                                    "_unified_import_done"
                                ] = True

                                st.rerun()

                        with uc2:

                            if st.button(

                                "아니오",

                                use_container_width=True,

                                key="unified_import_confirm_no"

                            ):

                                st.session_state[
                                    "_show_unified_import_confirm"
                                ] = False

                                st.rerun()

                    _confirm_unified_import()

                if st.session_state.get("_unified_import_done"):

                    st.success(

                        "통합 업로드가 반영되었습니다. "
                        "메뉴를 다시 열면 확인할 수 있습니다."

                    )

            except Exception as e:

                st.error(

                    f"엑셀을 읽는 중 문제가 발생했습니다: {e}"

                )


        st.markdown(
            "### 📥 실제 설비 데이터 일괄 업로드"
        )

        st.caption(

            "사업소에서 받은 '설비 기본정보' 엑셀과 '오버홀 실적' "
            "엑셀을 그대로 올리면 설비마스터와 오버홀이력이 한 번에 "
            "채워집니다. 원본에 없는 값(제조사·모델·정격출력 등)은 "
            "억지로 채우지 않고 빈칸으로 둡니다. "
            "⚠️ 기존에 등록된 설비 목록은 전부 지워지고 "
            "새로 올리는 목록으로 교체됩니다."

        )

        bulk_c1, bulk_c2 = st.columns(2)

        equip_excel_file = bulk_c1.file_uploader(

            "설비 기본정보 엑셀",

            type=["xlsx"],

            key="bulk_equip_excel"

        )

        overhaul_excel_file = bulk_c2.file_uploader(

            "오버홀 실적 엑셀",

            type=["xlsx"],

            key="bulk_overhaul_excel"

        )

        if equip_excel_file is not None:

            try:

                _df1 = pd.read_excel(equip_excel_file)

                _equip_rows, _equip_map, _site_map = parse_equipment_import_excel(
                    _df1
                )

                st.success(

                    f"설비 {len(_equip_rows)}건을 확인했습니다."

                )

                with st.expander(

                    f"설비 미리보기 ({len(_equip_rows)}건)"

                ):

                    st.dataframe(

                        pd.DataFrame(_equip_rows),

                        use_container_width=True,

                        hide_index=True

                    )

                _overhaul_rows = []

                if overhaul_excel_file is not None:

                    _df2 = pd.read_excel(
                        overhaul_excel_file
                    )

                    _overhaul_rows, _skipped = parse_overhaul_import_excel(

                        _df2,

                        _equip_map,

                        _site_map

                    )

                    st.success(

                        f"오버홀 실적 {len(_overhaul_rows)}건을 "
                        f"확인했습니다"
                        +
                        (
                            f" ({_skipped}건은 설비 매칭 실패 또는 "
                            f"날짜 오류로 제외)"

                            if _skipped

                            else ""
                        )

                    )

                    with st.expander(

                        f"오버홀 실적 미리보기 ({len(_overhaul_rows)}건)"

                    ):

                        st.dataframe(

                            pd.DataFrame(_overhaul_rows),

                            use_container_width=True,

                            hide_index=True

                        )

                if is_read_only():

                    st.info(
                        "🔒 보기 전용 모드에서는 반영할 수 없습니다."
                    )

                elif st.button(

                    "📥 일괄 업로드 확인 팝업 열기",

                    type="primary",

                    use_container_width=True,

                    key="bulk_import_open_btn"

                ):

                    st.session_state["_show_bulk_import_confirm"] = True

                if st.session_state.get("_show_bulk_import_confirm"):

                    @st.dialog("일괄 업로드 확인")
                    def _confirm_bulk_import():

                        st.write(

                            f"설비 {len(_equip_rows)}건으로 "
                            "**기존 설비 목록을 전부 교체**하고, "
                            f"오버홀 실적 {len(_overhaul_rows)}건을 "
                            "추가하시겠습니까?"

                        )

                        st.warning(

                            "기존에 등록돼 있던 설비는 삭제되고 "
                            "새 목록으로 바뀝니다. 기존 오버홀·"
                            "정밀진단 이력은 그대로 남지만, "
                            "삭제된 설비 이름과 연결이 끊깁니다."

                        )

                        dc1, dc2 = st.columns(2)

                        with dc1:

                            if st.button(

                                "예, 진행합니다",

                                type="primary",

                                use_container_width=True,

                                key="bulk_import_confirm_yes"

                            ):

                                replace_all_equipment(
                                    _equip_rows
                                )

                                if _overhaul_rows:

                                    bulk_append_overhaul(
                                        _overhaul_rows
                                    )

                                st.session_state[
                                    "_show_bulk_import_confirm"
                                ] = False

                                st.session_state[
                                    "_bulk_import_done"
                                ] = True

                                st.rerun()

                        with dc2:

                            if st.button(

                                "아니오",

                                use_container_width=True,

                                key="bulk_import_confirm_no"

                            ):

                                st.session_state[
                                    "_show_bulk_import_confirm"
                                ] = False

                                st.rerun()

                    _confirm_bulk_import()

                if st.session_state.get("_bulk_import_done"):

                    st.success(

                        "설비 목록과 오버홀 실적이 반영되었습니다. "
                        "메뉴를 다시 열면 확인할 수 있습니다."

                    )

            except Exception as e:

                st.error(

                    f"엑셀을 읽는 중 문제가 발생했습니다: {e}"

                )


        st.markdown(
            "### ⚙️ 설비 마스터 관리"
        )

        st.caption(
            "설비 목록이 코드에 고정되어 있지 않고 여기서 "
            "추가·삭제할 수 있습니다. 설비별로 진동·효율 기준을 "
            "다르게 두고 싶으면 기준값도 함께 입력하세요 "
            "(비워두면 공통 기준 적용)."
        )

        _sample_names_now = {
            p["equip"] for p in DEFAULT_PUMPS
        } & {
            p["equip"] for p in get_all_pumps()
        }

        if _sample_names_now:

            st.warning(

                f"코드에 예시로 들어있던 샘플(가짜) 설비가 "
                f"{len(_sample_names_now)}대 남아있습니다: "
                +
                ", ".join(sorted(_sample_names_now))

            )

            if is_read_only():

                st.info(
                    "🔒 보기 전용 모드에서는 삭제할 수 없습니다."
                )

            elif st.button(

                "🗑️ 샘플(가짜) 설비 전체 삭제",

                type="primary",

                use_container_width=True,

                key="delete_sample_btn"

            ):

                st.session_state["_show_delete_sample_confirm"] = True

            if st.session_state.get("_show_delete_sample_confirm"):

                @st.dialog("샘플 설비 삭제 확인")
                def _confirm_delete_sample():

                    st.write(

                        f"코드 예시용 샘플 설비 {len(_sample_names_now)}대를 "
                        "삭제하시겠습니까?"

                    )

                    st.caption(

                        "실제로 업로드해서 등록한 설비는 이름이 달라서 "
                        "영향받지 않습니다."

                    )

                    sc1, sc2 = st.columns(2)

                    with sc1:

                        if st.button(

                            "예, 삭제합니다",

                            type="primary",

                            use_container_width=True,

                            key="delete_sample_confirm_yes"

                        ):

                            deleted_count = delete_sample_default_equipment()

                            st.session_state[
                                "_show_delete_sample_confirm"
                            ] = False

                            st.session_state[
                                "_sample_delete_done"
                            ] = deleted_count

                            st.rerun()

                    with sc2:

                        if st.button(

                            "아니오",

                            use_container_width=True,

                            key="delete_sample_confirm_no"

                        ):

                            st.session_state[
                                "_show_delete_sample_confirm"
                            ] = False

                            st.rerun()

                _confirm_delete_sample()

            if st.session_state.get("_sample_delete_done"):

                st.success(

                    f"샘플 설비 {st.session_state['_sample_delete_done']}대가 "
                    "삭제되었습니다."

                )

        else:

            st.caption(
                "✅ 코드 예시용 샘플(가짜) 설비는 남아있지 않습니다."
            )

        with st.expander(
            "➕ 새 설비 추가"
        ):

            nc1, nc2 = st.columns(2)

            new_site = nc1.text_input(
                "사업장",
                "밀양정수장",
                key="new_equip_site"
            )

            new_equip = nc2.text_input(
                "설비명",
                key="new_equip_name"
            )

            new_maker = nc1.text_input(
                "제조사",
                key="new_equip_maker"
            )

            new_model = nc2.text_input(
                "모델명",
                key="new_equip_model"
            )

            new_hp = nc1.number_input(
                "정격출력(HP)",
                min_value=0.0,
                value=150.0,
                key="new_equip_hp"
            )

            new_head = nc2.number_input(
                "정격양정(m)",
                min_value=0.0,
                value=45.0,
                key="new_equip_head"
            )

            new_flow = nc1.number_input(
                "정격유량(m³/h)",
                min_value=0.0,
                value=1200.0,
                key="new_equip_flow"
            )

            new_rpm = nc2.number_input(
                "회전수(RPM)",
                min_value=0.0,
                value=1780.0,
                key="new_equip_rpm"
            )

            new_build = nc1.text_input(
                "준공일 (YYYY-MM-DD)",
                "2024-01-01",
                key="new_equip_build"
            )

            new_hours = nc2.number_input(
                "누적 운전시간",
                min_value=0.0,
                value=0.0,
                key="new_equip_hours"
            )

            nc1, nc2 = st.columns(2)

            new_vib_limit = nc1.number_input(
                "기준진동(mm/s) — 선택사항, 0이면 공통기준",
                min_value=0.0,
                value=0.0,
                key="new_equip_vib"
            )

            new_eff_target = nc2.number_input(
                "기준효율(%) — 선택사항, 0이면 공통기준",
                min_value=0.0,
                max_value=100.0,
                value=0.0,
                key="new_equip_eff"
            )

            if is_read_only():

                st.info(
                    "🔒 보기 전용 모드에서는 설비를 추가할 수 없습니다."
                )

            elif st.button(
                "설비 추가",
                type="primary",
                key="add_equip_btn"
            ):

                existing_names = [

                    p["equip"].strip()

                    for p in ALL_PUMPS

                ]

                if not new_equip.strip():

                    st.error(
                        "설비명을 입력해주세요."
                    )

                elif new_equip.strip() in existing_names:

                    st.error(

                        f"'{new_equip.strip()}'는 이미 등록된 설비명입니다. "
                        "같은 이름이 있으면 QR포털·정밀진단 등에서 "
                        "어느 설비인지 구분할 수 없으니, "
                        "다른 이름(예: 번호를 다르게)으로 등록해주세요."

                    )

                else:

                    add_equipment(

                        {
                            "site": new_site,
                            "equip": new_equip.strip(),
                            "maker": new_maker,
                            "model": new_model,
                            "hp": new_hp,
                            "head": new_head,
                            "flow": new_flow,
                            "rpm": new_rpm,
                            "build_date": new_build,
                            "op_hours": new_hours,
                            "기준진동": new_vib_limit if new_vib_limit > 0 else None,
                            "기준효율": new_eff_target if new_eff_target > 0 else None
                        }

                    )

                    st.success(
                        f"{new_equip} 설비가 추가되었습니다. "
                        "메뉴를 다시 열면 목록에 반영됩니다."
                    )

        with st.expander(
            "🗑️ 설비 삭제"
        ):

            del_target = st.selectbox(

                "삭제할 설비",

                [
                    p["equip"]
                    for p in ALL_PUMPS
                ],

                key="del_equip_select"

            )

            if is_read_only():

                st.info(
                    "🔒 보기 전용 모드에서는 삭제할 수 없습니다."
                )

            elif st.button(
                "선택한 설비 삭제",
                key="del_equip_btn"
            ):

                delete_equipment(
                    del_target
                )

                st.success(
                    f"{del_target} 설비가 삭제되었습니다. "
                    "정밀진단·오버홀 이력은 그대로 남아있습니다."
                )


        st.markdown(
            "### 📦 소모품 재고관리"
        )

        st.caption(

            "그랜드패킹·베어링·축슬리브 같은 소모품의 현재고를 "
            "관리합니다. 안전재고보다 부족하면 경고가 뜹니다."

        )

        df_consumables = get_consumables()

        if not df_consumables.empty:

            st.dataframe(

                df_consumables,

                use_container_width=True,

                hide_index=True

            )

            _low_stock = get_low_stock_consumables()

            if not _low_stock.empty:

                st.warning(

                    "🔴 안전재고 미달 소모품: "
                    +
                    ", ".join(

                        f"{r['소모품명']}(현재고 {r['현재고']}/"
                        f"안전재고 {r['안전재고']})"

                        for _, r in _low_stock.iterrows()

                    )

                )

        else:

            st.caption(
                "등록된 소모품이 없습니다."
            )

        with st.expander(
            "➕ 소모품 등록"
        ):

            cc1, cc2 = st.columns(2)

            new_cons_name = cc1.text_input(

                "소모품명",

                key="new_cons_name"

            )

            new_cons_spec = cc2.text_input(

                "규격",

                key="new_cons_spec"

            )

            cc3, cc4 = st.columns(2)

            new_cons_qty = cc3.number_input(

                "현재고",

                min_value=0,

                step=1,

                key="new_cons_qty"

            )

            new_cons_safety = cc4.number_input(

                "안전재고",

                min_value=0,

                step=1,

                key="new_cons_safety"

            )

            new_cons_date = st.text_input(

                "최근입고일 (YYYY-MM-DD)",

                value=datetime.now().strftime("%Y-%m-%d"),

                key="new_cons_date"

            )

            if is_read_only():

                st.info(
                    "🔒 보기 전용 모드에서는 등록할 수 없습니다."
                )

            elif st.button(

                "소모품 등록",

                key="add_consumable_btn"

            ):

                if not new_cons_name.strip():

                    st.error(
                        "소모품명을 입력해주세요."
                    )

                else:

                    add_consumable(

                        {
                            "소모품명": new_cons_name.strip(),
                            "규격": new_cons_spec.strip(),
                            "현재고": int(new_cons_qty),
                            "안전재고": int(new_cons_safety),
                            "최근입고일": new_cons_date
                        }

                    )

                    st.success(

                        f"{new_cons_name} 등록되었습니다."

                    )

                    st.rerun()

        if not df_consumables.empty:

            with st.expander(
                "✏️ 재고 수량 변경 / 삭제"
            ):

                cons_names = df_consumables["소모품명"].tolist()

                target_cons = st.selectbox(

                    "대상 소모품",

                    cons_names,

                    key="target_consumable_select"

                )

                target_row = df_consumables[

                    df_consumables["소모품명"] == target_cons

                ].iloc[0]

                new_qty_val = st.number_input(

                    "새 현재고",

                    min_value=0,

                    step=1,

                    value=int(target_row["현재고"]),

                    key="update_cons_qty"

                )

                uc1, uc2 = st.columns(2)

                if is_read_only():

                    st.info(
                        "🔒 보기 전용 모드에서는 변경할 수 없습니다."
                    )

                else:

                    with uc1:

                        if st.button(

                            "재고 수량 저장",

                            key="update_consumable_btn"

                        ):

                            update_consumable_stock(

                                target_cons,

                                int(new_qty_val),

                                datetime.now().strftime("%Y-%m-%d")

                            )

                            st.success(
                                "재고가 변경되었습니다."
                            )

                            st.rerun()

                    with uc2:

                        if st.button(

                            "🗑️ 이 소모품 삭제",

                            key="delete_consumable_btn"

                        ):

                            delete_consumable(
                                target_cons
                            )

                            st.success(
                                "삭제되었습니다."
                            )

                            st.rerun()


    with _databack_tab2:

        st.markdown(
            "### 🌐 배포 주소 설정 (QR 코드용)"
        )

        st.caption(
            "QR 포털의 QR 이미지가 가리키는 주소입니다. "
            "실제로 이 앱을 배포한 주소와 다르면 QR을 스캔해도 "
            "연결되지 않으니, 배포 후 반드시 실제 주소로 바꿔주세요."
        )

        current_url = get_app_base_url()

        new_url = st.text_input(

            "배포 주소 (예: https://kwatertech-pump.streamlit.app)",

            value=current_url,

            key="app_base_url_input",

            disabled=is_read_only()

        )

        if is_read_only():

            st.info(
                "🔒 보기 전용 모드에서는 변경할 수 없습니다."
            )

        elif st.button(
            "배포 주소 저장",
            key="save_app_url_btn"
        ):

            set_app_base_url(
                new_url
            )

            st.success(
                "배포 주소가 저장되었습니다. "
                "QR 포털 화면을 다시 열면 반영됩니다."
            )


        st.markdown(
            "### 🏷️ QR 라벨 인쇄용 PDF"
        )

        st.caption(

            "등록된 설비의 QR코드+설비명+사업장을 A4 한 장에 "
            "여러 개씩 격자로 배치해서, 현장에 붙일 라벨을 한 번에 "
            "인쇄할 수 있게 만듭니다. 점선을 따라 잘라 쓰시면 됩니다."

        )

        if not ALL_PUMPS:

            st.caption(
                "등록된 설비가 없어 라벨을 만들 수 없습니다."
            )

        else:

            label_site_list = sorted(

                set(
                    p["site"] for p in ALL_PUMPS
                )

            )

            label_site_filter = st.multiselect(

                "라벨 만들 사업장 선택 (비워두면 전체)",

                label_site_list,

                default=label_site_list,

                key="qr_label_site_filter"

            )

            label_cols_per_row = st.selectbox(

                "한 줄에 라벨 개수",

                [2, 3, 4],

                index=1,

                key="qr_label_cols"

            )

            target_pumps_for_label = [

                p for p in ALL_PUMPS

                if p["site"] in label_site_filter

            ]

            st.caption(

                f"대상 설비 {len(target_pumps_for_label)}개 "
                f"(A4 {label_cols_per_row}열×6행 기준 "
                f"{-(-len(target_pumps_for_label) // (label_cols_per_row * 6))}페이지)"

            )

            if target_pumps_for_label and st.button(

                "🏷️ QR 라벨 PDF 생성",

                type="primary",

                use_container_width=True,

                key="generate_qr_label_pdf_btn"

            ):

                with st.spinner(
                    "라벨 PDF를 만드는 중입니다..."
                ):

                    label_pdf_bytes = build_qr_label_sheet_pdf(

                        target_pumps_for_label,

                        cols=label_cols_per_row,

                        rows=6

                    )

                st.session_state["_qr_label_pdf_bytes"] = label_pdf_bytes

            if "_qr_label_pdf_bytes" in st.session_state:

                st.download_button(

                    "⬇️ QR 라벨 PDF 다운로드",

                    data=st.session_state["_qr_label_pdf_bytes"],

                    file_name="QR_설비라벨.pdf",

                    mime="application/pdf",

                    use_container_width=True,

                    key="download_qr_label_pdf_btn"

                )


        st.markdown(
            "### 📱 QR 구축비용"
        )

        qr_count = st.number_input(
            "QR 부착 설비 수",
            min_value=1,
            value=10
        )

        unit_cost, total_cost = calculate_qr_cost(
            qr_count
        )

        c1, c2 = st.columns(2)

        c1.metric(
            "QR 1개당 구축비",
            f"{unit_cost:,}원"
        )

        c2.metric(
            f"총 {qr_count}개 구축비",
            f"{total_cost:,}원"
        )

        qr_df = pd.DataFrame(

            [

                [
                    "QR 라벨 제작",
                    QR_UNIT_COST["QR 라벨 제작"]
                ],

                [
                    "표면 세척/준비",
                    QR_UNIT_COST["표면 세척/준비"]
                ],

                [
                    "QR 부착 작업",
                    QR_UNIT_COST["QR 부착 작업"]
                ],

                [
                    "설비 등록 및 검수",
                    QR_UNIT_COST["설비 등록 및 검수"]
                ]

            ],

            columns=[
                "항목",
                "원가"
            ]

        )

        st.dataframe(
            qr_df,
            use_container_width=True,
            hide_index=True
        )

        st.info(
            "권장 현장 부착시간 : "
            "설비 1대당 약 10~15분"
        )


        st.markdown(
            "### 🏷️ QR 부착 방법"
        )

        st.markdown(
            """
            **① 설치 위치 선정**

            펌프 본체 또는 제어반 등
            작업자가 쉽게 확인할 수 있는 위치 선정

            **② 표면 세척**

            먼지·유분 제거

            **③ QR 라벨 부착**

            평탄한 면에 부착 후
            모서리 및 접착면 압착

            **④ 설비번호 확인**

            QR과 설비번호가 일치하는지 확인

            **⑤ 스마트폰 테스트**

            실제 현장에서 스캔하여
            해당 설비 페이지 연결 확인

            **⑥ DB 등록 완료**

            설비 Master DB와 QR ID 연결
            """,
            unsafe_allow_html=True
        )


    with _databack_tab3:

        st.markdown(
            "### 🚨 알림 기준값 설정"
        )

        st.caption(

            "AI 이상징후·정비권고사항·홈 화면에서 쓰는 "
            "'관찰/주의' 보조 알림 기준입니다. "
            "정밀진단 17개 항목의 A~E 등급 계산 기준(ISO 10816-3 등)은 "
            "여기서 바뀌지 않고 그대로 유지됩니다."

        )

        _cur_th = get_alert_thresholds()

        th_c1, th_c2 = st.columns(2)

        new_vib_watch = th_c1.number_input(

            "진동 관찰 기준(mm/s)",

            value=float(_cur_th["vib_watch"]),

            key="th_vib_watch",

            disabled=is_read_only()

        )

        new_vib_danger = th_c2.number_input(

            "진동 주의 기준(mm/s)",

            value=float(_cur_th["vib_danger"]),

            key="th_vib_danger",

            disabled=is_read_only()

        )

        new_eff_watch = th_c1.number_input(

            "효율 관찰 기준(%)",

            value=float(_cur_th["eff_watch"]),

            key="th_eff_watch",

            disabled=is_read_only()

        )

        new_eff_danger = th_c2.number_input(

            "효율 주의 기준(%)",

            value=float(_cur_th["eff_danger"]),

            key="th_eff_danger",

            disabled=is_read_only()

        )

        new_temp_watch = th_c1.number_input(

            "온도 관찰 기준(°C)",

            value=float(_cur_th["temp_watch"]),

            key="th_temp_watch",

            disabled=is_read_only()

        )

        new_temp_danger = th_c2.number_input(

            "온도 주의 기준(°C)",

            value=float(_cur_th["temp_danger"]),

            key="th_temp_danger",

            disabled=is_read_only()

        )

        if is_read_only():

            st.info(
                "🔒 보기 전용 모드에서는 변경할 수 없습니다."
            )

        elif st.button(
            "알림 기준값 저장",
            key="save_thresholds_btn"
        ):

            save_alert_thresholds(

                {
                    "vib_watch": new_vib_watch,
                    "vib_danger": new_vib_danger,
                    "eff_watch": new_eff_watch,
                    "eff_danger": new_eff_danger,
                    "temp_watch": new_temp_watch,
                    "temp_danger": new_temp_danger
                }

            )

            log_audit(

                "알림 기준값 변경",

                "전체",

                f"진동({new_vib_watch}/{new_vib_danger}) "
                f"효율({new_eff_watch}/{new_eff_danger}) "
                f"온도({new_temp_watch}/{new_temp_danger})"

            )

            st.success(
                "알림 기준값이 저장되었습니다."
            )


        st.markdown(
            "### 🔔 알림 웹훅 설정 (Slack 등)"
        )

        st.caption(

            "정밀진단 결과가 D/E등급으로 저장되면 아래 웹훅 주소로 "
            "자동 알림을 보냅니다. Slack의 'Incoming Webhook' 주소를 "
            "그대로 붙여넣으면 됩니다. 비워두면 알림을 보내지 않습니다. "
            "⚠️ 실제 알림이 도착하는지는 워크스페이스에 웹훅을 "
            "등록하신 뒤 직접 확인해주세요."

        )

        current_webhook = get_webhook_url()

        new_webhook = st.text_input(

            "웹훅 URL",

            value=current_webhook,

            key="webhook_url_input",

            type="password",

            disabled=is_read_only()

        )

        if is_read_only():

            st.info(
                "🔒 보기 전용 모드에서는 변경할 수 없습니다."
            )

        elif st.button(
            "웹훅 주소 저장",
            key="save_webhook_btn"
        ):

            set_webhook_url(
                new_webhook
            )

            log_audit(
                "웹훅 설정 변경",
                "알림설정"
            )

            st.success(
                "웹훅 주소가 저장되었습니다."
            )


        st.markdown(
            "### 🕵️ 감사 로그 (누가 언제 무엇을 했는지)"
        )

        df_audit = read_excel(

            AUDIT_LOG_PATH,

            "감사로그"

        )

        if not df_audit.empty:

            with st.expander(

                f"최근 감사 로그 (총 {len(df_audit)}건)"

            ):

                st.dataframe(

                    df_audit.tail(30).iloc[::-1],

                    use_container_width=True,

                    hide_index=True

                )

        else:

            st.caption(
                "아직 기록된 감사 로그가 없습니다."
            )


    with _databack_tab4:

        st.markdown(
            "### ☁️ 구글시트 백업/복원"
        )

        st.caption(

            "서버가 재배포되면 로컬 데이터가 사라질 수 있는 문제의 "
            "안전장치입니다. 구글 서비스계정 연결이 되어있어야 "
            "동작합니다 (Streamlit Secrets에 gcp_service_account, "
            "GSHEET_BACKUP_URL 등록 필요)."

        )

        gcol1, gcol2 = st.columns(2)

        with gcol1:

            if st.button(

                "☁️ 지금 구글시트로 백업",

                use_container_width=True,

                key="gsheet_backup_btn"

            ):

                with st.status(

                    "구글시트로 백업 중...",

                    expanded=True

                ) as status_box:

                    st.write(

                        f"총 {len(GSHEET_BACKUP_TARGETS)}개 시트 "
                        "대상입니다."

                    )

                    st.write("구글 인증 확인 중...")

                    ok, msg = backup_all_to_gsheet()

                    if ok:

                        status_box.update(

                            label="백업 완료",

                            state="complete"

                        )

                    else:

                        status_box.update(

                            label="백업 실패",

                            state="error"

                        )

                if ok:

                    st.success(
                        msg
                    )

                else:

                    st.error(
                        msg
                    )

        with gcol2:

            if is_read_only():

                st.info(
                    "🔒 보기 전용 모드에서는 복원할 수 없습니다."
                )

            elif st.button(

                "⬇️ 구글시트에서 복원",

                use_container_width=True,

                key="gsheet_restore_btn"

            ):

                st.session_state["_confirm_gsheet_restore"] = True

            if st.session_state.get(

                "_confirm_gsheet_restore"

            ):

                st.warning(

                    "⚠️ 지금 로컬 데이터를 구글시트 내용으로 "
                    "덮어씁니다. 계속할까요?"

                )

                ccol1, ccol2 = st.columns(2)

                if ccol1.button(

                    "예, 복원합니다",

                    key="confirm_restore_yes"

                ):

                    with st.spinner(
                        "복원하는 중입니다..."
                    ):

                        ok, msg = restore_all_from_gsheet()

                    st.session_state["_confirm_gsheet_restore"] = False

                    if ok:

                        st.success(
                            msg
                        )

                    else:

                        st.error(
                            msg
                        )

                    st.rerun()

                if ccol2.button(

                    "취소",

                    key="confirm_restore_no"

                ):

                    st.session_state["_confirm_gsheet_restore"] = False

                    st.rerun()

        st.write("")


        st.markdown(
            "### 📦 DB 백업"
        )

        st.warning(
            "⚠️ 이 데이터는 서버 로컬 파일로 저장됩니다. "
            "배포 환경이 재시작·재배포되면 초기화될 수 있으니 "
            "정기적으로 아래에서 전체 백업을 받아두는 것을 권장합니다."
        )

        _backup_files = [

            (DB_FILE_PATH, "진단 DB"),
            (OVERHAUL_DB_PATH, "오버홀 DB"),
            (KNOWHOW_DB_PATH, "노하우 DB"),
            (EQUIP_DB_PATH, "설비마스터 DB"),
            (KPI_DB_PATH, "KPI DB")

        ]

        import zipfile

        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(
            zip_buffer,
            "w",
            zipfile.ZIP_DEFLATED
        ) as zf:

            for path, label in _backup_files:

                if os.path.exists(path):

                    zf.write(
                        path,
                        arcname=os.path.basename(path)
                    )

            # 오버홀 작업사진도 백업에 포함시킨다.
            # (예전에는 사진 파일만 빠져있어서 서버가
            #  재시작되면 사진만 조용히 사라졌었다)

            if os.path.exists(PHOTO_DIR):

                for fname in os.listdir(PHOTO_DIR):

                    fpath = os.path.join(
                        PHOTO_DIR,
                        fname
                    )

                    if os.path.isfile(fpath):

                        zf.write(

                            fpath,

                            arcname=os.path.join(
                                "overhaul_photos",
                                fname
                            )

                        )

        st.download_button(

            "📥 전체 DB 한번에 백업 (zip)",

            data=zip_buffer.getvalue(),

            file_name=

            f"kwatertech_backup_"

            +
            datetime.now().strftime("%Y%m%d_%H%M")

            +
            ".zip",

            mime="application/zip",

            type="primary",

            use_container_width=True,

            on_click=_record_last_backup_time

        )

        for path, label in _backup_files:

            if os.path.exists(path):

                with open(
                    path,
                    "rb"
                ) as f:

                    st.download_button(

                        f"📥 {label} 다운로드",

                        data=f.read(),

                        file_name=os.path.basename(
                            path
                        ),

                        key=f"download_{path}"

                    )

        render_excel_export_section(

            "backup",

            "설비마스터_목록.xlsx",

            lambda: pd.DataFrame(ALL_PUMPS)

        )

        st.markdown(
            "### 📊 종합 대시보드 PDF"
        )

        st.caption(

            "설비 하나씩이 아니라, 이번 달 전체 사업장 현황을 "
            "한 장으로 정리한 경영보고용 요약 PDF를 만듭니다."

        )

        if st.button(

            "📊 종합 대시보드 PDF 생성",

            type="primary",

            use_container_width=True,

            key="generate_dashboard_pdf_btn"

        ):

            with st.spinner(
                "대시보드를 만드는 중입니다..."
            ):

                _dash_site_list = sorted(

                    set(
                        p["site"] for p in ALL_PUMPS
                    )

                )

                dashboard_pdf_bytes = build_dashboard_summary_pdf(

                    ALL_PUMPS,

                    df_history,

                    _dash_site_list

                )

            st.session_state[

                "_dashboard_pdf_bytes"

            ] = dashboard_pdf_bytes

        if "_dashboard_pdf_bytes" in st.session_state:

            st.download_button(

                "⬇️ 대시보드 PDF 다운로드",

                data=st.session_state["_dashboard_pdf_bytes"],

                file_name=(

                    f"{datetime.now().strftime('%Y%m')}_"
                    "종합대시보드.pdf"

                ),

                mime="application/pdf",

                use_container_width=True,

                key="download_dashboard_pdf_btn"

            )

        st.write("")


        st.markdown(
            "### 📧 이메일 월간요약 발송"
        )

        st.caption(

            "지금 버튼을 누르면 즉시 발송됩니다. 매달 자동으로 "
            "보내려면 이 기능을 외부 스케줄러(예: GitHub Actions)가 "
            "주기적으로 호출해야 합니다 — 앱 자체는 접속 없이 "
            "혼자 실행되지 않습니다. (Streamlit Secrets에 "
            "SMTP_USER, SMTP_PASSWORD 등록 필요)"

        )

        email_recipients_text = st.text_input(

            "받는 사람 이메일 (쉼표로 구분)",

            key="email_recipients_input"

        )

        if st.button(

            "📧 지금 이메일 발송",

            use_container_width=True,

            key="send_email_now_btn"

        ):

            _recipients = [

                e.strip()

                for e in email_recipients_text.split(",")

                if e.strip()

            ]

            if not _recipients:

                st.error(
                    "받는 사람 이메일을 입력해주세요."
                )

            else:

                _body = build_monthly_summary_email_body(

                    ALL_PUMPS,

                    df_history

                )

                ok, msg = send_monthly_summary_email(

                    _recipients,

                    f"[K-water tech] "
                    f"{datetime.now().strftime('%Y년 %m월')} "
                    "설비관리 월간요약",

                    _body

                )

                if ok:

                    st.success(
                        msg
                    )

                else:

                    st.error(
                        msg
                    )

        st.write("")


    with _databack_tab5:

        st.markdown(
            "### 💰 플랫폼 사업화 예상비용"
        )

        cost_df = pd.DataFrame(

            [

                [
                    "1단계 PoC",
                    "자체개발",
                    "약 150~300만원",
                    "기존 용역 내 실증"
                ],

                [
                    "2단계 플랫폼화",
                    "고도화",
                    "약 500~800만원/년",
                    "유지보수·서버·운영"
                ],

                [
                    "3단계 외부사업",
                    "지자체",
                    "별도 견적",
                    "설비 수량별 산정"
                ]

            ],

            columns=[
                "단계",
                "구축방식",
                "예상비용",
                "사업모델"
            ]

        )

        st.dataframe(
            cost_df,
            use_container_width=True,
            hide_index=True
        )

        st.write("")

    with _databack_tab6:

        st.markdown(
            "### 🗑️ 휴지통"
        )

        st.caption(

            f"삭제한 항목이 실수였다면 여기서 되돌릴 수 있습니다. "
            f"{TRASH_RETENTION_DAYS}일이 지나면 자동으로 완전히 "
            "삭제됩니다."

        )

        _trash_df = get_trash_items()

        if _trash_df.empty:

            st.info(
                "휴지통이 비어있습니다."
            )

        else:

            for _t_idx, _t_row in _trash_df.iterrows():

                _t_col1, _t_col2, _t_col3 = st.columns(

                    [3, 1, 1]

                )

                with _t_col1:

                    st.write(

                        f"**[{_t_row['종류']}]** "
                        f"{_t_row['이름']} · "
                        f"{_t_row['삭제일시']} 삭제 "
                        f"({_t_row['삭제자']})"

                    )

                with _t_col2:

                    if st.button(

                        "↩️ 복원",

                        key=f"trash_restore_{_t_idx}",

                        use_container_width=True,

                        disabled=is_read_only()

                    ):

                        import json as _json_trash

                        restore_from_trash(

                            _t_idx,

                            _t_row["종류"],

                            _t_row["이름"],

                            _t_row["원본시트"],

                            _json_trash.loads(
                                _t_row["데이터JSON"]
                            )

                        )

                        st.success(

                            f"{_t_row['이름']} 복원되었습니다."

                        )

                        st.rerun()

                with _t_col3:

                    if st.button(

                        "🗑️ 완전삭제",

                        key=f"trash_purge_{_t_idx}",

                        use_container_width=True,

                        disabled=is_read_only()

                    ):

                        permanently_delete_from_trash(

                            _t_row["종류"],

                            _t_row["이름"]

                        )

                        st.success(
                            "완전히 삭제되었습니다."
                        )

                        st.rerun()



# ============================================================
# 25-1. 통합검색
# ============================================================

elif st.session_state.page == "통합검색":

    render_back_to_home_button("search")

    st.markdown(
        """
        <div class="section-title">
        🔎 통합검색
        </div>

        <div class="section-caption">
        설비마스터·정밀진단·오버홀이력·소모품재고·일위대가·
        기술노하우를 한 번에 검색합니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    search_keyword = st.text_input(

        "검색어",

        placeholder="예: 가압펌프 #3, 최진욱, 볼밸브 ...",

        key="unified_search_keyword"

    )

    with st.expander(
        "🔧 검색 설정"
    ):

        similarity_threshold = st.slider(

            "AI 비슷한 사례 추천 민감도",

            min_value=0.05,

            max_value=0.40,

            value=0.12,

            step=0.01,

            key="search_similarity_threshold",

            help="낮을수록 느슨하게(더 많이) 추천하고, "
            "높을수록 확실히 비슷한 것만 추천합니다."

        )

    if "_recent_searches" not in st.session_state:

        st.session_state["_recent_searches"] = []

    if search_keyword.strip():

        _recent = st.session_state["_recent_searches"]

        if search_keyword.strip() in _recent:

            _recent.remove(
                search_keyword.strip()
            )

        _recent.insert(
            0,
            search_keyword.strip()
        )

        st.session_state["_recent_searches"] = _recent[:5]

    if st.session_state["_recent_searches"]:

        st.caption(
            "🕘 최근 검색어"
        )

        recent_cols = st.columns(

            len(st.session_state["_recent_searches"])

        )

        for i, kw in enumerate(

            st.session_state["_recent_searches"]

        ):

            if recent_cols[i].button(

                kw,

                key=f"recent_search_{i}",

                use_container_width=True

            ):

                st.session_state["unified_search_keyword"] = kw

                st.rerun()

    if search_keyword.strip():

        search_results = run_unified_search(

            search_keyword,

            ALL_PUMPS,

            similarity_threshold=similarity_threshold

        )

        if not search_results:

            st.info(
                "검색 결과가 없습니다."
            )

        else:

            total_hits = sum(

                len(v)

                for v in search_results.values()

            )

            st.caption(

                f"'{search_keyword}' 검색 결과: 총 "
                f"{total_hits}건"

            )

            if "설비" in search_results:

                st.markdown(
                    "##### 🏭 설비"
                )

                for p in search_results["설비"]:

                    scol1, scol2 = st.columns(

                        [4, 1]

                    )

                    scol1.write(

                        f"**{p['equip']}** ({p['site']})"

                    )

                    if scol2.button(

                        "열기",

                        key=f"search_goto_{p['equip']}"

                    ):

                        st.session_state.page = "QR"

                        st.query_params["page"] = "QR"

                        st.query_params["equip"] = p["equip"]

                        st.rerun()

            for label, df_result in search_results.items():

                if label == "설비":

                    continue

                st.markdown(
                    f"##### {label} ({len(df_result)}건)"
                )

                if "AI 추천" in label or "비슷한 사례" in label:

                    # 의미기반 추천은 정확한 단어가 안 겹칠 수
                    # 있어서 하이라이트를 강제로 하지 않는다.

                    st.dataframe(

                        df_result,

                        use_container_width=True,

                        hide_index=True

                    )

                else:

                    st.markdown(

                        render_highlighted_table_html(

                            df_result,

                            search_keyword

                        ),

                        unsafe_allow_html=True

                    )

    else:

        st.info(
            "검색어를 입력해주세요."
        )


# ============================================================
# 26-0. 설계적산
# ============================================================

# ============================================================
# 26-2. 업데이트 내역
# ============================================================

elif st.session_state.page == "업데이트내역":

    render_back_to_home_button("changelog")

    st.markdown(
        """
        <div class="section-title">
        📋 업데이트 내역
        </div>

        <div class="section-caption">
        이 플랫폼이 어떻게 발전해왔는지 정리한 기록입니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    CHANGELOG = [

        (
            "v1.9",
            "2026-09-01",
            "보고서 자동화 (사진 → AI/OCR 분석 → 완성)",
            [
                "축정렬·진동분석·효율진단 3종 보고서, 실제 사업소 양식 그대로 재현",
                "사진 업로드 → AI가 직접 읽고 판정문까지 작성(API 연동)",
                "API 없이도 OCR 자동읽기·엑셀업로드·수동입력으로 동일 양식 완성",
                "휴지통(30일 복원), 백업 알림 배너, 앱 내 매뉴얼 뷰어"
            ]
        ),

        (
            "v1.8",
            "2026-08-28",
            "메뉴 재구성 · 정비 캘린더 · 즐겨찾기 · 체크리스트",
            [
                "사이드바 메뉴를 업무흐름 기준 5개 그룹으로 재정리",
                "오버홀 예정일정 캘린더 뷰 추가",
                "즐겨찾기·최근 본 설비 영구 저장(파일 기반)",
                "설비 종류별 정비 체크리스트 템플릿"
            ]
        ),

        (
            "v1.7",
            "2026-08-23",
            "UI/UX 전면 개선",
            [
                "회사 공식 마스코트 TEKI 적용 (5개 페이지)",
                "다크모드, 모바일 반응형 레이아웃 대응",
                "카드 애니메이션·여백·색상 통일"
            ]
        ),

        (
            "v1.6",
            "2026-08-18",
            "데이터 안전망 강화",
            [
                "구글시트 백업/복원 (행개수 자동 검증 포함)",
                "이메일 월간요약 발송 (재시도 로직 포함)",
                "종합 대시보드 PDF 자동생성"
            ]
        ),

        (
            "v1.5",
            "2026-08-11",
            "통합검색 · 보안 강화",
            [
                "전체 데이터 통합검색 + TF-IDF 의미기반 유사사례 추천",
                "PIN 로그인 5회 실패시 잠금, 세션 자동 로그아웃",
                "긴급 차단 스위치, 로그인 감사로그"
            ]
        ),

        (
            "v1.4",
            "2026-08-04",
            "설계적산 모듈",
            [
                "일위대가 라이브러리 (실제 산출내역서 기반)",
                "원본 양식 그대로 엑셀 재현·다운로드",
                "새 공사 만들기, 유사 항목명 경고"
            ]
        ),

        (
            "v1.3",
            "2026-07-29",
            "AI 예측 · LLM 종합의견",
            [
                "진동 추세 회귀분석 기반 고장예측 (계절성 반영)",
                "예측 사후검증 (지난 예측 vs 실제 비교)",
                "Claude API 기반 실시간 종합의견 생성"
            ]
        ),

        (
            "v1.2",
            "2026-07-22",
            "정비효과 · 성과관리",
            [
                "설비별 총소유비용(TCO) 추적",
                "종합 성과 요약 (Before/After)",
                "PoC 단계 KPI 관리"
            ]
        ),

        (
            "v1.1",
            "2026-07-16",
            "데이터 관리 체계화",
            [
                "설비 마스터를 코드 고정 → 엑셀 DB 기반으로 전환",
                "통합 업로드 양식, 소모품 재고관리",
                "QR 라벨 인쇄, 감사로그"
            ]
        ),

        (
            "v1.0",
            "2026-07-10",
            "기본 CBM 플랫폼",
            [
                "QR 기반 설비 포털",
                "17개 항목 정밀진단 · CBM Score 산출",
                "오버홀 관리, 월간 보고서 자동생성"
            ]
        )

    ]

    for version, ver_date, title, items in CHANGELOG:

        st.markdown(

            f"""
            <div class="platform-card">
            <div style="display:flex; align-items:center; gap:10px;
            flex-wrap:wrap;">
                <span style="background:#087ea4; color:white;
                font-weight:800; font-size:0.78rem;
                padding:3px 10px; border-radius:20px;">
                {version}
                </span>
                <span style="font-weight:800; font-size:1.02rem;
                color:#0f3552;">
                {title}
                </span>
                <span style="color:#94a3b8; font-size:0.78rem;
                margin-left:auto;">
                {ver_date}
                </span>
            </div>
            <ul style="margin:10px 0 0 0; padding-left:20px;
            color:#334155; font-size:0.88rem; line-height:1.8;">
            {"".join(f"<li>{item}</li>" for item in items)}
            </ul>
            </div>
            """,

            unsafe_allow_html=True

        )


elif st.session_state.page == "설계적산":

    render_back_to_home_button("design")

    st.markdown(
        """
        <div class="section-title">
        📐 설계적산
        </div>

        <div class="section-caption">
        일위대가 라이브러리를 쌓아두고, 새 공사를 만들 때
        자유롭게 검색해서 담으면 자동으로 금액이 계산됩니다.
        </div>
        """,
        unsafe_allow_html=True
    )

    design_tab1, design_tab2, design_tab3 = st.tabs(

        [
            "📦 일위대가 라이브러리",
            "🏗️ 새 공사 만들기",
            "📋 저장된 공사"
        ]

    )

    with design_tab1:

        st.markdown(
            "##### 산출내역서 엑셀 업로드"
        )

        st.caption(

            "실제 설계·시공 산출내역서(\"일위대가표\" 시트가 "
            "있는 엑셀)를 올리면 자동으로 항목을 뽑아 "
            "라이브러리에 추가합니다. 오래된 엑셀이라 파일이 "
            "손상돼 있어도(정의된 이름 오류 등) 자동으로 "
            "복구를 시도합니다."

        )

        design_upload = st.file_uploader(

            "산출내역서 엑셀",

            type=["xlsx"],

            key="design_unitprice_uploader"

        )

        if design_upload is not None:

            try:

                _raw_bytes = design_upload.getvalue()

                _df_up = safe_read_excel_any(

                    _raw_bytes,

                    sheet_name="일위대가표"

                )

                _parsed_items = parse_unitprice_sheet(
                    _df_up
                )

                if not _parsed_items:

                    st.warning(

                        "\"일위대가표\" 시트에서 항목을 "
                        "찾지 못했습니다. 시트 이름과 양식을 "
                        "확인해주세요."

                    )

                else:

                    st.success(

                        f"{len(_parsed_items)}개 항목을 "
                        "찾았습니다."

                    )

                    with st.expander(
                        "찾은 항목 미리보기"
                    ):

                        st.dataframe(

                            pd.DataFrame(

                                [
                                    {
                                        "항목명": it["항목명"],
                                        "단위": it["단위"],
                                        "합계단가": it["합계단가"]
                                    }
                                    for it in _parsed_items
                                ]

                            ),

                            use_container_width=True,

                            hide_index=True

                        )

                    _existing_names_for_check = set(

                        get_unitprice_library()["항목명"].tolist()

                        if not get_unitprice_library().empty

                        else []

                    )

                    _similar_warnings = []

                    for it in _parsed_items:

                        _sims = find_similar_item_names(

                            it["항목명"],

                            _existing_names_for_check

                        )

                        if _sims:

                            _similar_warnings.append(

                                f"'{it['항목명']}' ↔ "
                                f"'{_sims[0][0]}' "
                                f"(유사도 {_sims[0][1]*100:.0f}%)"

                            )

                    if _similar_warnings:

                        st.warning(

                            "⚠️ 기존 라이브러리에 비슷한 이름이 "
                            "이미 있습니다. 오타나 표기 차이로 "
                            "중복 등록되는 건 아닌지 확인해주세요:\n"
                            +
                            "\n".join(

                                f"- {w}"

                                for w in _similar_warnings

                            )

                        )

                    if is_read_only():

                        st.info(
                            "🔒 보기 전용 모드에서는 "
                            "반영할 수 없습니다."
                        )

                    elif st.button(

                        "📥 라이브러리에 추가하기",

                        type="primary",

                        use_container_width=True,

                        key="design_import_btn"

                    ):

                        _added = import_unitprice_items_to_library(

                            _parsed_items,

                            source_name=design_upload.name

                        )

                        st.success(

                            f"{_added}건이 라이브러리에 "
                            "새로 추가되었습니다 (이미 있던 "
                            "이름은 건너뛰었습니다)."

                        )

                        st.rerun()

            except Exception as e:

                st.error(

                    f"엑셀을 읽는 중 문제가 발생했습니다: {e}"

                )

        st.write("---")

        st.markdown(
            "##### 라이브러리 (자주 쓰는 순)"
        )

        _lib_df = get_unitprice_library()

        if _lib_df.empty:

            st.info(
                "아직 등록된 일위대가가 없습니다. "
                "위에서 엑셀을 올려 시작해보세요."
            )

        else:

            _lib_search = st.text_input(

                "🔍 항목명 검색",

                key="design_lib_search"

            )

            _lib_view = _lib_df

            if _lib_search.strip():

                _lib_view = _lib_df[

                    _lib_df["항목명"].str.contains(

                        _lib_search.strip(),

                        case=False,

                        na=False

                    )

                ]

            st.dataframe(

                _lib_view[

                    [
                        "항목명",
                        "단위",
                        "합계단가",
                        "사용횟수",
                        "출처공사"
                    ]

                ],

                use_container_width=True,

                hide_index=True

            )

            st.download_button(

                "⬇️ 이 목록 원본 양식 엑셀로 다운로드",

                data=build_unitprice_excel_bytes(

                    _lib_view.to_dict("records")

                ),

                file_name="일위대가표.xlsx",

                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.spreadsheetml.sheet"
                ),

                use_container_width=True,

                key="download_unitprice_excel_btn"

            )

    with design_tab2:

        st.markdown(
            "##### 새 공사 정보"
        )

        _lib_df2 = get_unitprice_library()

        if _lib_df2.empty:

            st.info(

                "먼저 \"일위대가 라이브러리\" 탭에서 "
                "산출내역서를 올려 항목을 등록해주세요."

            )

        else:

            pc1, pc2 = st.columns(2)

            new_project_name = pc1.text_input(

                "공사명",

                placeholder="예: 부곡가압장 배관 재설치",

                key="new_project_name"

            )

            new_project_gongjong = pc2.selectbox(

                "공사 유형",

                DESIGN_GONGJONG_LIST + ["기타"],

                key="new_project_gongjong"

            )

            pc3, pc4 = st.columns(2)

            new_project_site = pc3.selectbox(

                "사업장",

                sorted(

                    set(
                        p["site"] for p in ALL_PUMPS
                    )

                ) or ["미지정"],

                key="new_project_site"

            )

            new_project_author = pc4.text_input(

                "작성자",

                value=st.session_state.user_name,

                key="new_project_author"

            )

            _guideline_names = get_guideline_items(

                new_project_gongjong

            )

            if _guideline_names:

                st.info(

                    f"💡 \"{new_project_gongjong}\" 공사에 "
                    "보통 쓰는 항목: "
                    +
                    ", ".join(_guideline_names)

                )

            st.write("---")

            st.markdown(
                "##### 담을 항목 고르기"
            )

            st.caption(

                "목록에서 필요한 항목을 체크(선택)하세요. "
                "타이핑하면 검색도 됩니다."

            )

            _cart_key = "design_cart_items"

            if _cart_key not in st.session_state:

                st.session_state[_cart_key] = []

            _item_options = [

                f"{row['항목명']} ({row['단위']}, "
                f"{row['합계단가']:,.0f}원)"

                for _, row in _lib_df2.iterrows()

            ]

            _name_by_option = {

                f"{row['항목명']} ({row['단위']}, "
                f"{row['합계단가']:,.0f}원)": row["항목명"]

                for _, row in _lib_df2.iterrows()

            }

            _default_options = [

                opt for opt, name in _name_by_option.items()

                if name in _guideline_names

            ]

            selected_options = st.multiselect(

                "일위대가 목록",

                _item_options,

                default=_default_options,

                key="design_item_multiselect"

            )

            if is_read_only():

                st.info(
                    "🔒 보기 전용 모드에서는 담을 수 없습니다."
                )

            elif st.button(

                "📥 선택한 항목 담기",

                type="primary",

                use_container_width=True,

                key="add_selected_items_btn"

            ):

                _cart_names = {

                    c["항목명"]

                    for c in st.session_state[_cart_key]

                }

                _added_count = 0

                for opt in selected_options:

                    _name = _name_by_option[opt]

                    if _name in _cart_names:

                        continue

                    _row = _lib_df2[

                        _lib_df2["항목명"] == _name

                    ].iloc[0]

                    st.session_state[_cart_key].append(

                        {
                            "항목명": _row["항목명"],
                            "단위": _row["단위"],
                            "단가": _row["합계단가"],
                            "규격": "",
                            "수량": 1.0
                        }

                    )

                    _added_count += 1

                if _added_count:

                    st.rerun()

            if st.session_state[_cart_key]:

                st.write("---")

                st.markdown(
                    "##### 담은 항목 (규격·수량 입력)"
                )

                _total_amount = 0

                for idx, cart_item in enumerate(

                    st.session_state[_cart_key]

                ):

                    ccol1, ccol2, ccol3, ccol4 = st.columns(

                        [2, 2, 1, 1]

                    )

                    ccol1.write(
                        f"**{cart_item['항목명']}**"
                    )

                    cart_item["규격"] = ccol2.text_input(

                        "규격",

                        value=cart_item["규격"],

                        key=f"cart_spec_{idx}",

                        label_visibility="collapsed"

                    )

                    cart_item["수량"] = ccol3.number_input(

                        f"수량({cart_item['단위']})",

                        min_value=0.0,

                        value=cart_item["수량"],

                        key=f"cart_qty_{idx}",

                        label_visibility="collapsed"

                    )

                    amount = (

                        cart_item["수량"]
                        *
                        cart_item["단가"]

                    )

                    _total_amount += amount

                    ccol4.write(
                        f"{amount:,.0f}원"
                    )

                st.markdown(

                    f"### 합계: {_total_amount:,.0f}원"

                )

                if st.button(

                    "🗑️ 담은 항목 전체 비우기",

                    key="clear_cart_btn"

                ):

                    st.session_state[_cart_key] = []

                    st.rerun()

                if is_read_only():

                    st.info(
                        "🔒 보기 전용 모드에서는 "
                        "저장할 수 없습니다."
                    )

                elif st.button(

                    "💾 이 공사로 저장하기",

                    type="primary",

                    use_container_width=True,

                    key="save_project_btn"

                ):

                    if not new_project_name.strip():

                        st.error(
                            "공사명을 입력해주세요."
                        )

                    else:

                        _pid = create_design_project(

                            new_project_name.strip(),

                            new_project_gongjong,

                            new_project_site,

                            new_project_author

                        )

                        for cart_item in st.session_state[

                            _cart_key

                        ]:

                            add_item_to_project(

                                _pid,

                                cart_item["항목명"],

                                cart_item["규격"],

                                cart_item["수량"],

                                cart_item["단위"],

                                cart_item["단가"]

                            )

                        st.session_state[_cart_key] = []

                        st.success(

                            f"\"{new_project_name}\" "
                            "공사가 저장되었습니다."

                        )

                        st.rerun()

    with design_tab3:

        st.markdown(
            "##### 저장된 공사 목록"
        )

        _proj_df = read_excel(

            DESIGN_PROJECT_DB_PATH,

            "설계공사"

        )

        if _proj_df.empty:

            st.info(
                "아직 저장된 공사가 없습니다. "
                "'새 공사 만들기' 탭에서 첫 공사를 만들어보세요."
            )

        else:

            st.dataframe(

                _proj_df[

                    [
                        "공사명",
                        "공사유형",
                        "사업장",
                        "작성자",
                        "작성일"
                    ]

                ],

                use_container_width=True,

                hide_index=True

            )

            _sel_project_name = st.selectbox(

                "상세 보기",

                _proj_df["공사명"].tolist(),

                key="view_project_select"

            )

            _sel_pid = _proj_df[

                _proj_df["공사명"] == _sel_project_name

            ]["공사ID"].iloc[0]

            _items_df = get_project_items(
                _sel_pid
            )

            if not _items_df.empty:

                st.dataframe(

                    _items_df[

                        [
                            "항목명",
                            "규격",
                            "수량",
                            "단위",
                            "적용단가",
                            "금액"
                        ]

                    ],

                    use_container_width=True,

                    hide_index=True

                )

                st.markdown(

                    f"### 총 공사금액: "
                    f"{_items_df['금액'].sum():,.0f}원"

                )


# ============================================================
# 26-1. 부서별 관리 (개발예정)
# ============================================================

elif st.session_state.page in (

    "토목",
    "전기",
    "전자",
    "행정",
    "안전"

):

    _dept_icon = {

        "토목": "🏗️",
        "전기": "⚡",
        "전자": "💻",
        "행정": "📋",
        "안전": "🦺"

    }[st.session_state.page]

    st.markdown(

        f"""
        <div class="section-title">
        {_dept_icon} {st.session_state.page} 관리
        </div>

        <div class="section-caption">
        곧 추가될 예정인 메뉴입니다.
        </div>
        """,

        unsafe_allow_html=True

    )

    st.info(

        f"🚧 {st.session_state.page} 관리 메뉴는 "
        "현재 개발 준비 중입니다. 필요한 기능이 있으면 "
        "말씀해 주세요."

    )


# ============================================================
# 27. QR 비용 별도 상세 계산
# ============================================================

st.write("")


# ============================================================
# 28. 하단 플랫폼 정보
# ============================================================

st.markdown(
    """
    <div style="
    margin-top:35px;
    padding:14px;
    border-top:1px solid #dce8ef;
    text-align:center;
    color:#94a3b8;
    font-size:0.7rem;
    ">

    K-water tech · 설비관리 통합 플랫폼

    <br>

    QR 기반 설비정보 · 상태진단 · CBM · 오버홀 ·
    데이터 기반 유지관리

    </div>
    """,
    unsafe_allow_html=True
)
